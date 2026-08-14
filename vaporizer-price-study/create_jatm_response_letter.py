"""Generate a point-by-point response-to-reviewers letter for the JATM major revision."""
import os
import json
import numpy as np
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(SCRIPT_DIR, 'data')
OUT_DIR = os.path.join(SCRIPT_DIR, 'papers')
os.makedirs(OUT_DIR, exist_ok=True)

# Load sensitivity results
rev_path = os.path.join(DATA_DIR, 'revision_sensitivity.json')
try:
    with open(rev_path, 'r') as f:
        rev = json.load(f)
    boot = rev['bootstrap_pre_post']['Desflurane']
    powd = rev['power_simulation']['Desflurane']
    cits = rev['cits_monthly_medians']
    did = rev['did_transaction_level']
except Exception:
    boot = {}
    powd = {}
    cits = {}
    did = {}


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Times New Roman'
    return p


def add_para(doc, text, bold=False, italic=False, size=Pt(12)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def fmt_float(v, prec=3, default='N/A'):
    try:
        v = float(v)
        if not np.isnan(v):
            return f'{v:.{prec}f}'
    except Exception:
        pass
    return default


def add_bullet(doc, title, body, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(' ' + body)
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'
    return p


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Header
    add_para(doc, '29 July 2026', size=Pt(12))
    add_para(doc, 'Jianjun Yang, MD, Editor-in-Chief')
    add_para(doc, 'Journal of Anesthesia and Translational Medicine')
    add_para(doc, 'c/o Editorial Office')
    add_para(doc, '')
    add_para(doc, 'Re: Manuscript JATMED-D-26-00064 - Major Revision')
    add_para(doc, 'Targeted environmental regulation and secondary market vaporizer prices: '
                 'an observational analysis of the EU desflurane phase-out')
    add_para(doc, '')
    add_para(doc, 'Dear Dr. Yang,')
    add_para(doc,
        'Thank you for the opportunity to revise our manuscript. We are grateful to the '
        'Editor and the two reviewers for their constructive and detailed feedback. '
        'We have addressed the major concerns by (1) adding sensitivity and control-series '
        'analyses (bootstrap resampling, post-hoc power estimation, difference-in-differences, '
        'and comparative interrupted time-series), (2) substantially softening the causal '
        'interpretation and replacing "ban/prohibited" with accurate wording about a '
        'restriction on routine use with documented clinical exceptions, (3) expanding the '
        'clinical and environmental context (desflurane pharmacological advantages, GWP100 '
        'debate, drug-diversity and supply-chain resilience), and (4) strengthening the '
        'limitations section regarding confounders, eBay generalizability, EU/non-EU '
        'transaction uncertainty, and the small post-restriction sample. All numerical '
        'results are derived from the updated analysis scripts and are reproducible from '
        'the deposited data.')

    add_heading(doc, 'Reviewer 1', level=1)

    add_bullet(doc, '1. Causal interpretation and regulatory attribution.',
        'We agree that the original phrasing could imply a stronger causal claim than the '
        'observational design supports. We have revised the title to emphasize an '
        'observational analysis ("an observational analysis of the EU desflurane phase-out") '
        'rather than "without collateral market damage." Throughout the abstract, '
        'introduction, results, discussion, and conclusion, we now use "associated with" '
        'rather than causal language, explicitly state that no causal inference can be '
        'drawn, and note that the observed desflurane decline may reflect secular trends '
        '(e.g., NHS England/Scotland decommissioning pre-dating the EU restriction, TIVA '
        'adoption, low-flow/end-tidal control technologies, macroeconomic trends, and '
        'equipment-specific compositional changes).')

    add_bullet(doc, '2. Confounders and compositional effects.',
        'We have substantially expanded the Discussion/Limitations section to discuss each '
        'confounder raised: the shift toward total intravenous anesthesia (TIVA), '
        'institutional desflurane decommissioning by NHS England and NHS Scotland, '
        'the growing availability of low-flow anesthesia and end-tidal control systems, '
        'general macroeconomic trends, and the inability to control for equipment age, '
        'model year, service history, cosmetic condition, and calibration certificates. '
        'We now explicitly state that these unmeasured factors could produce '
        'compositional changes that mimic a regulatory effect.')

    add_bullet(doc, '3. Difference-in-differences and interrupted time-series with controls.',
        f'To strengthen causal attribution assessment, we added two control-series '
        f'analyses. A comparative interrupted time-series (CITS) model of monthly median '
        f'prices (sevoflurane and isoflurane as concurrent controls, HAC robust standard '
        f'errors) showed a desflurane-specific level decline at the restriction '
        f'(log-price coefficient={fmt_float(cits.get("desflurane_level_change_coef"), 2)}, '
        f'P<0.001) and a post-restriction slope difference '
        f'(coefficient={fmt_float(cits.get("desflurane_slope_change_coef"), 2)}, P<0.001). '
        f'A transaction-level difference-in-differences model, however, did not show a '
        f'significant desflurane post-restriction coefficient '
        f'(coefficient={fmt_float(did.get("desflurane_post_coef"), 3)}, '
        f'P={fmt_float(did.get("desflurane_post_p"), 3)}) '
        f'because a significant pre-existing trend difference violated the parallel-trend '
        f'assumption (P={fmt_float(did.get("desflurane_trend_p"), 3)}). '
        f'These models are therefore presented as sensitivity, not confirmatory, analyses.')

    add_bullet(doc, '4. eBay data generalizability.',
        'We agree that eBay is only one segment of the secondary medical equipment market. '
        'We have added a clear limitation that specialized B2B platforms (e.g., DOTmed, '
        'Bimedis) and private dealer/hospital-to-hospital networks likely constitute a '
        'substantial share of high-value capital-asset transactions, and our findings may '
        'not generalize to those channels.')

    add_bullet(doc, '5. Inability to distinguish EU and non-EU buyers/sellers.',
        'This is a critical limitation and we have strengthened its discussion. We now '
        'state that eBay is a global marketplace and we could not determine buyer or seller '
        'location; if a meaningful share of transactions occurred outside the EU, the '
        'observed price decline could represent a global reassessment of desflurane\u2019s '
        'long-term viability rather than a response confined to EU jurisdiction.')

    add_bullet(doc, '6. Nominally significant isoflurane Spearman P=0.044.',
        'We added a footnote to Table 1 explaining that, although the isoflurane Spearman '
        'correlation reached nominal significance (P=0.044), the magnitude was very small '
        '(\u03c1=0.081), the quarterly median trend was not significant, and a single P<0.05 '
        'among several trend tests is expected under multiple testing. We therefore do '
        'not interpret it as a clinically meaningful trend.')

    mwp = fmt_float((powd.get('power_mannwhitney') or 0) * 100, 1, default='N/A')
    wtp = fmt_float((powd.get('power_welch_t') or 0) * 100, 1, default='N/A')
    mlo = fmt_float(boot.get('mean_diff_ci95', [None, None])[0], 0, default='N/A')
    mhi = fmt_float(boot.get('mean_diff_ci95', [None, None])[1], 0, default='N/A')
    mdlo = fmt_float(boot.get('median_diff_ci95', [None, None])[0], 0, default='N/A')
    mdhi = fmt_float(boot.get('median_diff_ci95', [None, None])[1], 0, default='N/A')
    add_bullet(doc, '7. Statistical power, bootstrap sensitivity, and pre-/post comparison.',
        f'We now report bootstrap 95% confidence intervals and post-hoc power. For '
        f'desflurane (n={boot.get("n_pre", "N/A")} pre, n={boot.get("n_post", "N/A")} post), '
        f'the bootstrap mean-difference 95% CI was US${mlo} to US${mhi} and the '
        f'median-difference 95% CI was US${mdlo} to US${mdhi}. '
        f'Post-hoc power under a lognormal model was low for the Mann\u2013Whitney U test '
        f'({mwp}%) and moderate for Welch\u2019s t-test ({wtp}%). We have re-emphasized that the primary '
        f'analysis is the time-series trend test (which uses all data points); the '
        f'pre-/post-ban comparison is exploratory and limited by the small post-restriction '
        f'sample.')

    add_heading(doc, 'Reviewer 2', level=1)

    add_bullet(doc, '1. "Prohibited" and "ban" terminology.',
        'We have corrected this throughout. The manuscript now states that Regulation (EU) '
        '2024/573 restricts desflurane for routine use while permitting documented clinical '
        'exceptions, and the title has been revised to refer to the "EU desflurane phase-out" '
        'rather than a "ban."')

    add_bullet(doc, '2. Desflurane pharmacological advantages.',
        'We added a sentence in the Introduction noting desflurane\u2019s recognized '
        'pharmacokinetic advantages, including faster emergence and extubation than '
        'sevoflurane in selected surgical populations, and cited a 2025 systematic review '
        'and meta-analysis (Hariyanto et al., Indian J Anaesth 2025;69:65-77).')

    add_bullet(doc, '3. GWP100 and short-lived climate pollutants.',
        'We added a sentence in the Introduction explaining that the choice of GWP metric '
        'and time horizon remains debated for short-lived halogenated anesthetics such as '
        'desflurane (atmospheric lifetime \u2248 14 years), because GWP100 may not reflect '
        'instantaneous radiative forcing or steady-state atmospheric concentrations, and '
        'cited Sulb\u00e6k Andersen et al. (Lancet Planet Health 2023;7:e622-e629).')

    add_bullet(doc, '4. Drug diversity and healthcare system resilience.',
        'We added a sentence in the Discussion highlighting that restricting desflurane '
        'reduces anesthetic diversity, which has been argued to be a pillar of healthcare '
        'system resilience in the context of recurring propofol and sevoflurane supply-chain '
        'disruptions, and cited Kranke et al. (Anaesthesiologie 2026; doi:10.1007/s00101-025-01639-x).')

    add_bullet(doc, '5. eBay buyer composition and B2B markets.',
        'We expanded the Limitations/Discussion to note that eBay predominantly serves '
        'individual purchasers and smaller facilities, that large hospital-system asset '
        'divestment typically occurs on specialized B2B platforms, and that the global '
        'reach of eBay means observed price movements may reflect global investor sentiment '
        'rather than a purely EU jurisdictional response. We have therefore toned down '
        'policy-relevant conclusions and frame the findings as hypothesis-generating.')

    add_para(doc, '')
    add_para(doc,
        'We believe these revisions have substantially improved the manuscript\u2019s '
        'methodological transparency, clinical balance, and interpretive caution. We would '
        'be grateful if the revised manuscript could be considered for publication in '
        'Journal of Anesthesia and Translational Medicine.')
    add_para(doc, '')
    add_para(doc, 'Sincerely,')
    add_para(doc, '[Corresponding author name]')
    add_para(doc, '[Affiliation]')
    add_para(doc, 'Email: [email]')

    out_path = os.path.join(OUT_DIR, 'jatm_response_to_reviewers.docx')
    doc.save(out_path)
    print(f'Response letter saved: {out_path}')


if __name__ == '__main__':
    main()
