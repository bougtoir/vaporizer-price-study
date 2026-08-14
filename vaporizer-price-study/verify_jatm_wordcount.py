"""Verify JATM manuscript word counts and citation completeness."""
import re
from docx import Document
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
doc_path = os.path.join(SCRIPT_DIR, 'papers', 'jatm_manuscript_english.docx')
doc = Document(doc_path)

# Extract all text by section
sections = {}
current_section = 'preamble'
current_text = []

for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    # Check if heading
    if para.style.name.startswith('Heading'):
        if current_text or current_section not in sections:
            sections[current_section] = ' '.join(current_text)
        current_section = text
        current_text = []
    else:
        current_text.append(text)

if current_text:
    sections[current_section] = ' '.join(current_text)

# Word counts
print("=" * 60)
print("JATM MANUSCRIPT WORD COUNT VERIFICATION")
print("=" * 60)

# Abstract
if 'Abstract' in sections:
    abstract_text = sections['Abstract']
    abstract_words = len(abstract_text.split())
    print(f"\nAbstract: {abstract_words} words (max 300)")
    if abstract_words > 300:
        print("  *** EXCEEDS LIMIT ***")
    else:
        print(f"  OK ({300 - abstract_words} words under limit)")

    # Check abstract sections
    for label in ['Background:', 'Methods:', 'Results:', 'Conclusions:']:
        if label in abstract_text:
            print(f"  {label} present")
        else:
            print(f"  *** MISSING: {label} ***")

# Introduction
if 'Introduction' in sections:
    intro_words = len(sections['Introduction'].split())
    print(f"\nIntroduction: {intro_words} words (max 500)")
    if intro_words > 500:
        print("  *** EXCEEDS LIMIT ***")
    else:
        print(f"  OK ({500 - intro_words} words under limit)")

# Discussion
if 'Discussion' in sections:
    disc_words = len(sections['Discussion'].split())
    print(f"\nDiscussion: {disc_words} words (max 1500)")
    if disc_words > 1500:
        print("  *** EXCEEDS LIMIT ***")
    else:
        print(f"  OK ({1500 - disc_words} words under limit)")

# Combined
if 'Introduction' in sections and 'Discussion' in sections:
    combined = intro_words + disc_words
    print(f"\nIntro + Discussion combined: {combined} words (max 2000)")
    if combined > 2000:
        print("  *** EXCEEDS LIMIT ***")
    else:
        print(f"  OK ({2000 - combined} words under limit)")

# Figure citations
print("\n" + "=" * 60)
print("FIGURE AND TABLE CITATION CHECK")
print("=" * 60)

full_text = '\n'.join(para.text for para in doc.paragraphs)

for i in range(1, 7):
    pattern = f'Fig. {i}'
    count = full_text.count(pattern)
    if count > 0:
        print(f"Fig. {i}: cited {count} time(s)")
    else:
        print(f"Fig. {i}: *** NOT CITED ***")

for i in range(1, 3):
    pattern = f'Table {i}'
    count = full_text.count(pattern)
    if count > 0:
        print(f"Table {i}: cited {count} time(s)")
    else:
        print(f"Table {i}: *** NOT CITED ***")

# Check Table S1
if 'Table S1' in full_text:
    print(f"Table S1: cited")
else:
    print(f"Table S1: *** NOT CITED ***")

# Reference count
print("\n" + "=" * 60)
print("REFERENCE CHECK")
print("=" * 60)

ref_section_found = False
ref_count = 0
for para in doc.paragraphs:
    text = para.text.strip()
    if text == 'References':
        ref_section_found = True
        continue
    if ref_section_found and text == 'Supplementary material':
        break
    if ref_section_found and text:
        # Check if starts with superscript number
        for run in para.runs:
            if run.font.superscript and run.text.strip().isdigit():
                ref_count += 1
                break

print(f"Total references: {ref_count}")

# Check citation numbers in text
max_cite = 0
cite_pattern = re.compile(r'\{(\d+(?:[\u2013,-]\d+)*)\}')
for m in cite_pattern.finditer(full_text):
    nums = re.findall(r'\d+', m.group(1))
    for n in nums:
        max_cite = max(max_cite, int(n))

# Also check for superscript references in runs
for para in doc.paragraphs:
    for run in para.runs:
        if run.font.superscript and run.text.strip():
            nums = re.findall(r'\d+', run.text)
            for n in nums:
                try:
                    max_cite = max(max_cite, int(n))
                except ValueError:
                    pass

print(f"Highest citation number in text: {max_cite}")
if max_cite > ref_count:
    print(f"  *** Citation {max_cite} exceeds reference count {ref_count} ***")
elif max_cite == ref_count:
    print(f"  OK (matches reference count)")

# Section structure check
print("\n" + "=" * 60)
print("SECTION STRUCTURE CHECK (JATM)")
print("=" * 60)
expected_sections = [
    'Abstract', 'Introduction', 'Materials and Methods',
    'Results', 'Discussion', 'Conclusion',
    'CRediT authorship contribution statement', 'Disclosure statement',
    'Ethical statement', 'Funding', 'Data availability',
    'Declaration of competing interest', 'Acknowledgments', 'References'
]
for s in expected_sections:
    if s in sections:
        print(f"  {s}: present")
    else:
        print(f"  {s}: *** MISSING ***")

# Total main text word count
print("\n" + "=" * 60)
print("TOTAL MAIN TEXT WORD COUNT")
print("=" * 60)
main_sections = ['Introduction', 'Materials and Methods', 'Ethics',
                 'Reporting guidelines', 'Study design and data source',
                 'Eligibility criteria', 'Variables', 'Statistical analysis',
                 'Results', 'Discussion', 'Conclusion']
total_main = 0
for s in main_sections:
    if s in sections:
        wc = len(sections[s].split())
        total_main += wc
        print(f"  {s}: {wc} words")
print(f"\n  Total main text: {total_main} words")
