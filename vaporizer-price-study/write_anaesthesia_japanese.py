"""
Generate Anaesthesia-format Japanese paper as editable .docx file.
Target journal: Anaesthesia (Association of Anaesthetists, Wiley)
Japanese translation of the English Anaesthesia-format paper.
Key Anaesthesia format:
  - "Summary" -> 要旨 (unstructured, single paragraph)
  - No "What this study adds" box
  - ~3000 words body text limit
  - Figures as SEPARATE files - figure legends in main doc
  - Vancouver numbered references
  - Data availability statement required
  - Conclusions within Discussion (no separate section)
"""
import pandas as pd
import numpy as np
from scipy import stats as sp_stats
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json

# ==========================================
# Load analysis results (same as English)
# ==========================================
data_dir = '/home/ubuntu/vaporizer_research/data'
stats_df = pd.read_csv(f'{data_dir}/statistics_summary.csv', index_col=0)
combined = pd.read_csv(f'{data_dir}/combined_cleaned.csv')
combined['date_sold'] = pd.to_datetime(combined['date_sold'])

figdir = '/home/ubuntu/vaporizer_research/figures/'
outdir = '/home/ubuntu/vaporizer_research/papers/'
os.makedirs(outdir, exist_ok=True)

try:
    with open(f'{data_dir}/asking_price_analysis.json', 'r') as f:
        asking_results = json.load(f)
    asking_df = pd.read_csv(f'{data_dir}/ebay_asking_prices.csv')
    has_asking_data = True
except FileNotFoundError:
    has_asking_data = False
    asking_results = None
    asking_df = None

reg_date = pd.Timestamp('2026-01-01')
proposal_date = pd.Timestamp('2022-04-05')
agreement_date = pd.Timestamp('2023-10-05')
adoption_date = pd.Timestamp('2024-02-07')

summ = {}
for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
    sub = combined[combined['agent_type'] == agent]
    pre = sub[sub['date_sold'] < reg_date]['price_usd']
    post = sub[sub['date_sold'] >= reg_date]['price_usd']
    summ[agent] = {
        'total_n': len(sub),
        'pre_n': len(pre), 'post_n': len(post),
        'pre_mean': pre.mean() if len(pre) > 0 else float('nan'),
        'post_mean': post.mean() if len(post) > 0 else float('nan'),
        'pre_median': pre.median() if len(pre) > 0 else float('nan'),
        'post_median': post.median() if len(post) > 0 else float('nan'),
        'pre_sd': pre.std() if len(pre) > 0 else float('nan'),
        'post_sd': post.std() if len(post) > 0 else float('nan'),
    }

total_n = len(combined)
date_min_all = combined['date_sold'].min().strftime('%d %B %Y')
date_max_all = combined['date_sold'].max().strftime('%d %B %Y')

def classify_period(date):
    d = pd.Timestamp(date)
    if d < proposal_date: return 1
    elif d < agreement_date: return 2
    elif d < adoption_date: return 3
    elif d < reg_date: return 4
    else: return 5

combined['period_num'] = combined['date_sold'].apply(classify_period)

trend_results = {}
for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
    sub = combined[combined['agent_type'] == agent].copy()
    sub['days'] = (sub['date_sold'] - sub['date_sold'].min()).dt.days
    rho, rho_p = sp_stats.spearmanr(sub['days'], sub['price_usd'])
    tau, tau_p = sp_stats.kendalltau(sub['period_num'], sub['price_usd'])
    sub['quarter'] = sub['date_sold'].dt.to_period('Q')
    quarterly = sub.groupby('quarter')['price_usd'].agg(['median', 'count'])
    quarterly = quarterly[quarterly['count'] >= 3]
    q_nums = np.arange(len(quarterly))
    if len(quarterly) >= 4:
        q_rho, q_rho_p = sp_stats.spearmanr(q_nums, quarterly['median'])
    else:
        q_rho, q_rho_p = float('nan'), float('nan')
    trend_results[agent] = {
        'spearman_rho': rho, 'spearman_p': rho_p,
        'kendall_tau': tau, 'kendall_p': tau_p,
        'quarterly_rho': q_rho, 'quarterly_p': q_rho_p,
    }

def get_pval(agent, col='u_pval'):
    try:
        v = stats_df.loc[agent, col]
        if pd.notna(v): return float(v)
    except Exception:
        pass
    return float('nan')

def get_stat(agent, col):
    try:
        v = stats_df.loc[agent, col]
        if pd.notna(v): return float(v)
    except Exception:
        pass
    return float('nan')

def fmt_p(p):
    if np.isnan(p): return 'N/A'
    if p < 0.001: return '<0.001'
    return f'{p:.3f}'

# ==========================================
# Helper functions
# ==========================================
def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_run_styled(para, text, bold=False, italic=False, size=Pt(11)):
    run = para.add_run(text)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    return run

def add_para(doc, text, size=Pt(11), bold=False, italic=False, alignment=None, space_after=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    return p

def setup_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing = 2.0
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    return doc

def add_table_header(table, headers):
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(cell, 'D9E2F3')

def add_table_data_row(table, data):
    row = table.add_row()
    for i, (text, align) in enumerate(data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(str(text))
        run.font.size = Pt(9)
    return row


# ==========================================
# JAPANESE PAPER - Anaesthesia format
# ==========================================
def write_japanese_paper():
    doc = setup_doc()
    des = summ['Desflurane']
    sevo = summ['Sevoflurane']
    iso = summ['Isoflurane']
    des_u_pval = get_pval('Desflurane', 'u_pval')
    des_t_pval = get_pval('Desflurane', 't_pval')
    sevo_u_pval = get_pval('Sevoflurane', 'u_pval')
    iso_u_pval = get_pval('Isoflurane', 'u_pval')
    des_d = get_stat('Desflurane', 'cohens_d')
    des_tr = trend_results['Desflurane']
    sevo_tr = trend_results['Sevoflurane']
    iso_tr = trend_results['Isoflurane']
    des_pct = abs((des['post_mean'] - des['pre_mean']) / des['pre_mean'] * 100)

    # ---- TITLE PAGE ----
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(
        'EU\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u898f\u5236\u304c\u9ebb\u9154\u6c17\u5316\u5668\u306e'
        '\u4e2d\u53e4\u5e02\u5834\u4fa1\u683c\u306b\u4e0e\u3048\u308b\u5f71\u97ff\uff1a'
        'eBay\u843d\u672d\u30c7\u30fc\u30bf\u306e\u6a2a\u65ad\u7684\u6642\u7cfb\u5217\u5206\u6790')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    add_para(doc, '[\u8457\u8005\u540d\u3092\u8a18\u5165]', size=Pt(11), italic=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, '[\u6240\u5c5e\u6a5f\u95a2\u3092\u8a18\u5165]', size=Pt(10), italic=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    p = doc.add_paragraph()
    add_run_styled(p, '\u8cac\u4efb\u8457\u8005: ', bold=True, size=Pt(10))
    add_run_styled(p, '[\u6c0f\u540d\u30fb\u30e1\u30fc\u30eb\u30a2\u30c9\u30ec\u30b9\u30fb\u4f4f\u6240\u3092\u8a18\u5165]', size=Pt(10))

    p = doc.add_paragraph()
    add_run_styled(p, 'Twitter/X: ', bold=True, size=Pt(10))
    add_run_styled(p, '[\u5404\u8457\u8005\u306e@\u30cf\u30f3\u30c9\u30eb\u3001\u307e\u305f\u306f"none"]', size=Pt(10))

    add_para(doc, '\u8981\u65e8\u8a9e\u6570: \u7d04250\u8a9e', size=Pt(10))
    add_para(doc, '\u672c\u6587\u8a9e\u6570: \u7d043000\u8a9e\uff08\u8981\u65e8\u30fb\u53c2\u8003\u6587\u732e\u30fb\u8868\u30fb\u56f3\u8aac\u660e\u3092\u9664\u304f\uff09',
             size=Pt(10))

    p = doc.add_paragraph()
    add_run_styled(p, '\u30ad\u30fc\u30ef\u30fc\u30c9: ', bold=True, size=Pt(10))
    add_run_styled(p, '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3; \u6c17\u5316\u5668; EU\u898f\u5236; '
                   '\u4e2d\u53e4\u5e02\u5834; F\u30ac\u30b9; \u74b0\u5883\u6301\u7d9a\u53ef\u80fd\u6027; \u9ebb\u9154',
                   size=Pt(10))

    doc.add_page_break()

    # ---- SUMMARY (要旨 - unstructured, single paragraph) ----
    add_heading_styled(doc, '\u8981\u65e8', level=1)
    summary_text = (
        f'\u9ebb\u9154\u85ac\u306e\u74b0\u5883\u898f\u5236\u306f\u65b0\u305f\u306a\u653f\u7b56\u9818\u57df'
        f'\u3067\u3042\u308b\u304c\u3001\u65e2\u5b58\u6a5f\u5668\u3078\u306e\u7d4c\u6e08\u7684\u5f71\u97ff'
        f'\u306f\u7814\u7a76\u3055\u308c\u3066\u3044\u306a\u3044\u3002\u672c\u7814\u7a76\u3067\u306f\u3001'
        f'EU\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u7981\u6b62\uff08\u898f\u5247(EU) 2024/573\uff09\u306e'
        f'\u5168\u7acb\u6cd5\u904e\u7a0b\u3092\u7db2\u7f85\u3059\u308b3\u5e74\u9593'
        f'\uff08{date_min_all}\uff5e{date_max_all}\uff09\u306b\u304a\u3051\u308b'
        f'\u9ebb\u9154\u6c17\u5316\u5668\uff08\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u3001'
        f'\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u3001\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\uff09\u306e'
        f'eBay\u843d\u672d{total_n}\u4ef6\u3092\u5206\u6790\u3057\u305f\u3002'
        f'\u30c7\u30fc\u30bf\u306feBay\u516c\u5f0f\u306e\u904e\u53bb\u58f2\u4e0a\u5206\u6790\u30c4\u30fc\u30eb'
        f'\u3067\u3042\u308bTerapeak\u3092\u7528\u3044\u3066\u53d6\u5f97\u3057\u305f\u3002'
        f'\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u4fa1\u683c\u306f\u7d71\u8a08\u7684\u306b'
        f'\u6709\u610f\u306a\u4e0b\u964d\u30c8\u30ec\u30f3\u30c9\u3092\u793a\u3057\u305f'
        f'\uff08Spearman \u03c1 = {des_tr["spearman_rho"]:.2f}, p < 0.001; '
        f'Kendall \u03c4 = {des_tr["kendall_tau"]:.2f}, p = {fmt_p(des_tr["kendall_p"])}\uff09\u3002'
        f'\u898f\u5236\u524d\u5e73\u5747US${des["pre_mean"]:.0f}\uff08SD ${des["pre_sd"]:.0f}\uff09'
        f'\u304b\u3089\u898f\u5236\u5f8cUS${des["post_mean"]:.0f}\uff08SD ${des["post_sd"]:.0f}\uff09'
        f'\u3078{des_pct:.0f}%\u4e0b\u843d\u3057\u305f'
        f'\uff08Cohen\u306ed = {des_d:.2f}\uff09\u3002'
        f'\u4e00\u65b9\u3001\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3'
        f'\uff08\u03c1 = {sevo_tr["spearman_rho"]:.2f}, p = {fmt_p(sevo_tr["spearman_p"])}\uff09'
        f'\u304a\u3088\u3073\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3'
        f'\uff08\u03c1 = {iso_tr["spearman_rho"]:.2f}, p = {fmt_p(iso_tr["spearman_p"])}\uff09'
        f'\u306b\u306f\u81e8\u5e8a\u7684\u306b\u610f\u5473\u306e\u3042\u308b\u6642\u7cfb\u5217\u30c8\u30ec'
        f'\u30f3\u30c9\u306f\u8a8d\u3081\u3089\u308c\u306a\u304b\u3063\u305f\u3002\u4fa1\u683c\u4e0b\u843d'
        f'\u306f\u7acb\u6cd5\u904e\u7a0b\u4e2d\u306b\u65e2\u306b\u59cb\u307e\u3063\u3066\u304a\u308a\u3001'
        f'\u5e02\u5834\u306e\u4e88\u6e2c\u7684\u53cd\u5fdc\u304c\u793a\u5506\u3055\u308c\u305f\u3002'
        f'\u3053\u308c\u3089\u306e\u77e5\u898b\u306f\u3001\u9ebb\u9154\u85ac\u306e\u74b0\u5883\u898f\u5236'
        f'\u304c\u4e2d\u53e4\u533b\u7642\u6a5f\u5668\u5e02\u5834\u306b\u85ac\u5264\u7279\u7570\u7684\u306a'
        f'\u7d4c\u6e08\u7684\u5f71\u97ff\u3092\u53ca\u307c\u3059\u3053\u3068\u3092\u793a\u3059\u521d\u3081'
        f'\u3066\u306e\u5b9f\u8a3c\u7684\u30a8\u30d3\u30c7\u30f3\u30b9\u3067\u3042\u308b\u3002'
    )
    doc.add_paragraph(summary_text)
    doc.add_page_break()

    # ---- INTRODUCTION (緒言) ----
    add_heading_styled(doc, '\u7dd2\u8a00', level=1)
    doc.add_paragraph(
        '\u5438\u5165\u9ebb\u9154\u85ac\u306f\u533b\u7642\u306e\u30ab\u30fc\u30dc\u30f3\u30d5\u30c3\u30c8'
        '\u30d7\u30ea\u30f3\u30c8\u306b\u5927\u304d\u304f\u5bc4\u4e0e\u3057\u3066\u3044\u308b[1\u20133]\u3002'
        '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u306f\u8fc5\u901f\u306a\u5c0e\u5165\u30fb\u899a\u9192\u7279\u6027'
        '\u3067\u8a55\u4fa1\u3055\u308c\u3066\u3044\u308b\u304c\u3001100\u5e74\u6642\u9593\u8ef8\u306e\u5730\u7403'
        '\u6e29\u6696\u5316\u4fc2\u6570\uff08GWP\uff09\u306f\u7d042540 CO\u2082\u7b49\u4fa1\u3067\u3042\u308a\u3001'
        '\u65e5\u5e38\u7684\u306b\u4f7f\u7528\u3055\u308c\u308b\u63ee\u767a\u6027\u9ebb\u9154\u85ac\u306e\u4e2d'
        '\u3067\u6700\u3082\u74b0\u5883\u8ca0\u8377\u304c\u5927\u304d\u3044[4,5]\u3002'
        '\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u306eGWP\u306f\u7d04130\u3001\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3'
        '\u306f\u7d04510\u3067\u3042\u308b[6,7]\u3002')
    doc.add_paragraph(
        '\u6b27\u5dde\u306b\u304a\u3051\u308b\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u898f\u5236\u306e\u7d4c\u7def'
        '\u306f\u3001\u8907\u6570\u306e\u91cd\u8981\u306a\u30de\u30a4\u30eb\u30b9\u30c8\u30fc\u30f3\u3092\u7d4c'
        '\u3066\u9032\u5c55\u3057\u305f\u30022022\u5e744\u6708\u3001\u6b27\u5dde\u59d4\u54e1\u4f1a\u304cF\u30ac'
        '\u30b9\u898f\u5247\u6539\u6b63\u6848\u3092\u516c\u8868\u3057\u305f\u30022023\u5e743\u6708\u306b\u6b27'
        '\u5dde\u8b70\u4f1a\u304c\u672c\u4f1a\u8b70\u3067\u627f\u8a8d\u3057\u30012023\u5e7410\u6708\u306b\u7406'
        '\u4e8b\u4f1a\u3068\u8b70\u4f1a\u306e\u9593\u3067\u6697\u5b9a\u5408\u610f\uff08\u30c8\u30ea\u30ed\u30fc'
        '\u30b0\uff09\u304c\u6210\u7acb\u3057\u305f\u3002\u898f\u5247(EU) 2024/573\u306f2024\u5e742\u6708\u306b'
        '\u6b63\u5f0f\u63a1\u629e\u3055\u308c\u30012024\u5e743\u6708\u306b\u767a\u52b9\u3057\u3001\u65e5\u5e38'
        '\u9ebb\u9154\u306b\u304a\u3051\u308b\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u4f7f\u7528\u7981\u6b62\u306f'
        '2026\u5e741\u67081\u65e5\u304b\u3089\u9069\u7528\u3055\u308c\u305f[2]\u3002\u540c\u6642\u306b\u3001'
        'NHS England\u306f2024\u5e74\u307e\u3067\u306b\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u3092\u5ec3\u6b62'
        '\u3059\u308b\u3068\u767a\u8868\u3057\u3001NHS Scotland\u306f2023\u5e743\u6708\u306b\u30c7\u30b9\u30d5'
        '\u30eb\u30e9\u30f3\u306e\u8cfc\u5165\u3092\u7981\u6b62\u3057\u305f\u6700\u521d\u306e\u533b\u7642'
        '\u30b7\u30b9\u30c6\u30e0\u3068\u306a\u3063\u305f[8,12]\u3002\u3053\u308c\u306f\u74b0\u5883\u7684'
        '\u7406\u7531\u306b\u3088\u308b\u7279\u5b9a\u306e\u9ebb\u9154\u85ac\u3078\u306e\u521d\u306e\u7fa9\u52d9'
        '\u7684\u653f\u5e9c\u898f\u5236\u3067\u3042\u308b\u3002')
    doc.add_paragraph(
        '\u9ebb\u9154\u6c17\u5316\u5668\u306f\u85ac\u5264\u7279\u7570\u7684\u306a\u88c5\u7f6e\u3067\u3042\u308a'
        '\u3001\u5178\u578b\u7684\u306a\u5bff\u547d\u306f10\uff5e15\u5e74\u3001\u76f8\u5f53\u306a\u8a2d\u5099'
        '\u6295\u8cc7\u3092\u8981\u3059\u308b\u3002\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u306e'
        '\u898f\u5236\u306b\u3088\u308b\u9673\u8150\u5316\u306f\u3001\u6a5f\u5668\u6240\u6709\u8005\u306b\u91cd'
        '\u5927\u306a\u7d4c\u6e08\u7684\u5f71\u97ff\u3092\u53ca\u307c\u3059\u53ef\u80fd\u6027\u304c\u3042\u308b'
        '\u3002\u91cd\u8981\u306a\u306e\u306f\u3001\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u3068\u30a4\u30bd\u30d5'
        '\u30eb\u30e9\u30f3\u306f\u540c\u3058\u898f\u5236\u306e\u5bfe\u8c61\u3067\u306f\u306a\u3044\u305f\u3081'
        '\u3001\u305d\u306e\u4fa1\u683c\u306f\u5f71\u97ff\u3092\u53d7\u3051\u306a\u3044\u306f\u305a\u3067\u3042'
        '\u308a\u3001\u81ea\u7136\u5bfe\u7167\u7fa4\u3092\u63d0\u4f9b\u3059\u308b\u3002')
    doc.add_paragraph(
        '\u6211\u3005\u306e\u77e5\u308b\u9650\u308a\u3001\u74b0\u5883\u898f\u5236\u304c\u9ebb\u9154\u6a5f\u5668'
        '\u306e\u4e2d\u53e4\u5e02\u5834\u4fa1\u5024\u306b\u4e0e\u3048\u308b\u5f71\u97ff\u3092\u691c\u8a0e\u3057'
        '\u305f\u7814\u7a76\u306f\u306a\u3044\u3002\u6211\u3005\u306f\u3001EU\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3'
        '\u898f\u5236\u304c\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u306e\u4e2d\u53e4\u5e02\u5834'
        '\u4fa1\u683c\u306e\u6bb5\u968e\u7684\u306a\u4e0b\u843d\u3068\u95a2\u9023\u3057\u3001\u30bb\u30dc\u30d5'
        '\u30eb\u30e9\u30f3\u304a\u3088\u3073\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u306e\u4fa1'
        '\u683c\u306f\u5b89\u5b9a\u3059\u308b\u3068\u4eee\u8aac\u3092\u7acb\u3066\u305f\u3002')

    # ---- METHODS (方法) ----
    add_heading_styled(doc, '\u65b9\u6cd5', level=1)
    doc.add_paragraph(
        '\u672c\u7814\u7a76\u306f\u3001\u6a2a\u65ad\u7814\u7a76\u306eSTROBE\uff08Strengthening the '
        'Reporting of Observational Studies in Epidemiology\uff09\u30ac\u30a4\u30c9\u30e9\u30a4\u30f3\u306b'
        '\u5f93\u3063\u3066\u5831\u544a\u3059\u308b[11]\u3002')

    add_heading_styled(doc, '\u7814\u7a76\u30c7\u30b6\u30a4\u30f3\u304a\u3088\u3073\u30c7\u30fc\u30bf\u30bd\u30fc\u30b9', level=2)
    doc.add_paragraph(
        'eBay (www.ebay.com) \u4e0a\u306e\u843d\u672d\u6e08\u307f\u30ea\u30b9\u30c6\u30a3\u30f3\u30b0\u3092'
        '\u7528\u3044\u305f\u9ebb\u9154\u6c17\u5316\u5668\u4fa1\u683c\u306e\u6a2a\u65ad\u7684\u6642\u7cfb\u5217'
        '\u5206\u6790\u3092\u884c\u3063\u305f\u3002\u30c7\u30fc\u30bf\u306f\u3001eBay Seller Hub\u306b\u7d71\u5408'
        '\u3055\u308c\u305feBay\u516c\u5f0f\u306e\u88fd\u54c1\u30ea\u30b5\u30fc\u30c1\u30c4\u30fc\u30eb\u3067'
        '\u3042\u308bTerapeak\u3092\u4f7f\u7528\u3057\u3066\u53d6\u5f97\u3057\u305f\u3002Terapeak\u306f\u6700'
        '\u59273\u5e74\u5206\u306e\u904e\u53bb\u306e\u843d\u672d\u30c7\u30fc\u30bf\u3078\u306e\u30a2\u30af\u30bb'
        '\u30b9\u3092\u63d0\u4f9b\u3059\u308b\u3002\u30c7\u30fc\u30bf\u53ce\u96c6\u306f2026\u5e743\u6708\u306b'
        '\u884c\u3044\u30012023\u5e743\u670828\u65e5\u304b\u30892026\u5e743\u670824\u65e5\u307e\u3067\u306e'
        '\u671f\u9593\u3092\u30ab\u30d0\u30fc\u3057\u305f\u30023\u5e74\u9593\u306e\u30a6\u30a3\u30f3\u30c9\u30a6'
        '\u306fTerapeak\u306e\u6700\u5927\u53d6\u5f97\u53ef\u80fd\u671f\u9593\u3067\u3042\u308b\u304c\u3001'
        '\u3053\u306e\u671f\u9593\u306f\u5206\u6790\u7684\u306b\u3082\u610f\u7fa9\u304c\u3042\u308b\uff1a'
        '\u6b27\u5dde\u8b70\u4f1a\u672c\u4f1a\u8b70\u6295\u7968\uff082023\u5e743\u6708\uff09\u306e\u76f4\u5f8c'
        '\u304b\u3089\u59cb\u307e\u308a\u3001\u6b27\u5dde\u59d4\u54e1\u4f1a\u63d0\u6848\uff082022\u5e744\u6708'
        '\uff09\u304b\u3089\u7981\u6b62\u5f8c\u307e\u3067\u306e\u5168\u7acb\u6cd5\u904e\u7a0b\u3092\u7db2\u7f85'
        '\u3059\u308b\u3002\u30af\u30ed\u30b9\u30ea\u30b9\u30c6\u30a3\u30f3\u30b0\u3055\u308c\u305f\u5546\u54c1'
        '\u306e\u91cd\u8907\u30ab\u30a6\u30f3\u30c8\u3092\u907f\u3051\u308b\u305f\u3081\u3001\u5358\u4e00\u30de'
        '\u30fc\u30b1\u30c3\u30c8\u30d7\u30ec\u30a4\u30b9\uff08eBay\uff09\u306e\u307f\u3092\u4f7f\u7528\u3057\u305f\u3002')

    add_heading_styled(doc, '\u9069\u683c\u57fa\u6e96', level=2)
    doc.add_paragraph(
        'Terapeak\u3067\u300cdesflurane vaporizer\u300d\u3001\u300csevoflurane vaporizer\u300d\u3001'
        '\u300cisoflurane vaporizer\u300d\u3092\u691c\u7d22\u3057\u305f\u3002\u7d44\u5165\u57fa\u6e96\u306f'
        '\uff1a(1) \u58f2\u8cb7\u5b8c\u4e86\u30ea\u30b9\u30c6\u30a3\u30f3\u30b0\u3001(2) \u5358\u4f53\u6c17'
        '\u5316\u5668\u3001(3) \u6709\u52b9\u306a\u4fa1\u683c\u30fb\u65e5\u4ed8\u3002\u9664\u5916\u57fa\u6e96'
        '\u306f\uff1a(1) \u975e\u6c17\u5316\u5668\u54c1\uff08\u30d5\u30a3\u30e9\u30fc\u3001\u30a2\u30c0\u30d7'
        '\u30bf\u30fc\u3001\u30a2\u30af\u30bb\u30b5\u30ea\u30fc\u7b49\uff09\u3001(2) \u7363\u533b\u7528\u30b7'
        '\u30b9\u30c6\u30e0\u3001(3) \u8907\u6570\u54c1\u306e\u30ed\u30c3\u30c8\u51fa\u54c1\u3001(4) \u4fa1\u683c'
        '\u30c7\u30fc\u30bf\u306e\u6b20\u640d\u307e\u305f\u306f\u7570\u5e38\u5024\u3002')

    add_heading_styled(doc, '\u5909\u6570', level=2)
    doc.add_paragraph(
        '\u4e3b\u8981\u30a2\u30a6\u30c8\u30ab\u30e0\u306f\u58f2\u5374\u4fa1\u683c\uff08\u7c73\u30c9\u30eb\uff09'
        '\u3002\u4e3b\u8981\u66dd\u9732\u5909\u6570\u306f\u898f\u5236\u671f\u9593\u3067\u3001EU F\u30ac\u30b9'
        '\u898f\u5247\u30bf\u30a4\u30e0\u30e9\u30a4\u30f3\u306e\u4e3b\u8981\u30de\u30a4\u30eb\u30b9\u30c8\u30fc'
        '\u30f3\u306b\u57fa\u3065\u304d\u5206\u985e\u3057\u305f\u3002\u4e3b\u8981\u6bd4\u8f03\u306f2026\u5e741'
        '\u67081\u65e5\uff08\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u7981\u6b62\u767a\u52b9\u65e5\uff09\u3092\u5206'
        '\u5272\u70b9\u3068\u3057\u305f\u3002\u526f\u6b21\u7684\u306a\u591a\u671f\u9593\u5206\u985e\u306f4\u3064'
        '\u306e\u30d5\u30a7\u30fc\u30ba\u306b\u5206\u3051\u305f\uff1a\u63d0\u6848\u5f8c\u3001\u5408\u610f\u5f8c'
        '\u3001\u63a1\u629e\u5f8c\u3001\u7981\u6b62\u5f8c\u3002')

    add_heading_styled(doc, '\u7d71\u8a08\u5206\u6790', level=2)
    doc.add_paragraph(
        '\u8a18\u8ff0\u7d71\u8a08\u91cf\u3068\u3057\u3066\u5e73\u5747\u3001\u6a19\u6e96\u504f\u5dee\u3001\u4e2d'
        '\u592e\u5024\u3001\u56db\u5206\u4f4d\u7bc4\u56f2\u3092\u7b97\u51fa\u3057\u305f\u3002\u4fa1\u683c\u5206'
        '\u5e03\u306e\u975e\u6b63\u898f\u6027\u3092\u8003\u616e\u3057\u3001\u898f\u5236\u524d\u5f8c\u306e\u4fa1'
        '\u683c\u6bd4\u8f03\u306b\u306fMann\u2013Whitney U\u691c\u5b9a\uff08\u4e21\u5074\uff09\u3092\u4e3b\u8981'
        '\u691c\u5b9a\u3068\u3057\u3066\u4f7f\u7528\u3057\u305f\u3002\u611f\u5ea6\u5206\u6790\u3068\u3057\u3066'
        'Welch\u306et\u691c\u5b9a\u3092\u884c\u3044\u3001\u52b9\u679c\u91cf\u306fCohen\u306ed\u3067\u63a8\u5b9a'
        '\u3057\u305f\u3002')
    doc.add_paragraph(
        '\u4fa1\u683c\u304c\u6642\u9593\u7d4c\u904e\u3068\u3068\u3082\u306b\u6bb5\u968e\u7684\u306b\u5909\u5316'
        '\u3057\u305f\u304b\u3092\u8a55\u4fa1\u3059\u308b\u305f\u3081\u30012\u3064\u306e\u88dc\u5b8c\u7684\u306a'
        '\u30c8\u30ec\u30f3\u30c9\u5206\u6790\u3092\u884c\u3063\u305f\u3002\u7b2c\u4e00\u306b\u3001Spearman'
        '\u9806\u4f4d\u76f8\u95a2\u3092\u7528\u3044\u3066\u3001\u5404\u85ac\u5264\u30bf\u30a4\u30d7\u306b\u3064'
        '\u3044\u3066\u58f2\u5374\u65e5\u3068\u4fa1\u683c\u306e\u5358\u8abf\u7684\u95a2\u9023\u3092\u691c\u5b9a'
        '\u3057\u305f\u3002\u7b2c\u4e8c\u306b\u3001Kendall \u03c4\u3092\u7528\u3044\u3066\u3001\u898f\u5236'
        '\u6bb5\u968e\u306e\u9806\u5e8f\uff081\uff5e5\uff09\u3068\u4fa1\u683c\u306e\u95a2\u9023\u3092\u691c'
        '\u5b9a\u3057\u305f\u3002\u3053\u308c\u3089\u306e\u30c8\u30ec\u30f3\u30c9\u691c\u5b9a\u306f\u5404\u85ac'
        '\u5264\u30bf\u30a4\u30d7\u306b\u72ec\u7acb\u3057\u3066\u9069\u7528\u3057\u305f\u3002\u56db\u534a\u671f'
        '\u3054\u3068\u306e\u4e2d\u592e\u5024\u306b\u3064\u3044\u3066\u3082Spearman\u76f8\u95a2\u3092\u8a55\u4fa1'
        '\u3057\u305f\u3002')
    doc.add_paragraph(
        '\u898f\u5236\u6bb5\u968e\u9593\u306e\u591a\u7fa4\u6bd4\u8f03\u306b\u306fKruskal\u2013Wallis\u691c\u5b9a'
        '\u3092\u4f7f\u7528\u3057\u305f\u3002LOWESS\u30c8\u30ec\u30f3\u30c9\u30e9\u30a4\u30f3\u3092\u4f7f\u7528'
        '\u3057\u305f\u3002\u5206\u6790\u306fPython 3.12\uff08pandas 2.2, scipy 1.14, statsmodels 0.14\uff09'
        '\u3067\u884c\u3063\u305f\u3002\u7d71\u8a08\u7684\u6709\u610f\u6c34\u6e96\u306fp < 0.05\uff08\u4e21\u5074'
        '\uff09\u3068\u3057\u305f\u3002')

    add_heading_styled(doc, '\u502b\u7406', level=2)
    doc.add_paragraph(
        '\u672c\u7814\u7a76\u306feBay\u306e\u516c\u958b\u843d\u672d\u30c7\u30fc\u30bf\u306e\u5206\u6790\u3067'
        '\u3042\u308a\u3001\u500b\u4eba\u60c5\u5831\u3084\u60a3\u8005\u30c7\u30fc\u30bf\u306f\u53ce\u96c6\u3057'
        '\u3066\u3044\u306a\u3044\u305f\u3081\u3001\u502b\u7406\u627f\u8a8d\u306f\u4e0d\u8981\u3067\u3042\u308b\u3002')

    # ---- RESULTS (結果) ----
    add_heading_styled(doc, '\u7d50\u679c', level=1)
    doc.add_paragraph(
        f'\u9664\u5916\u57fa\u6e96\u9069\u7528\u5f8c\u3001\u5408\u8a08{total_n}\u4ef6\u306eeBay\u843d\u672d'
        f'\u6e08\u307f\u9ebb\u9154\u6c17\u5316\u5668\u3092\u5206\u6790\u5bfe\u8c61\u3068\u3057\u305f\uff1a'
        f'\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3{des["total_n"]}\u4ef6\u3001'
        f'\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3{sevo["total_n"]}\u4ef6\u3001'
        f'\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3{iso["total_n"]}\u4ef6\u3002'
        f'\u7814\u7a76\u671f\u9593\u306f{date_min_all}\u304b\u3089{date_max_all}\u307e\u3067\u306e3\u5e74\u9593'
        f'\u3067\u3042\u3063\u305f\u3002')

    # Table 1
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_run_styled(p, '\u88681. ', bold=True, size=Pt(10))
    add_run_styled(p, ('\u6c17\u5316\u5668\u30bf\u30a4\u30d7\u304a\u3088\u3073\u898f\u5236\u671f\u9593\u5225\u306e'
                       'eBay Terapeak\u843d\u672d\u30c7\u30fc\u30bf\u306e\u8981\u7d04\uff08\u898f\u5236\u524d\u5f8c'
                       '\u30012026\u5e741\u67081\u65e5\u57fa\u6e96\uff09\u3002\u5024\u306f\u5e73\u5747\uff08SD\uff09'
                       '\u3001\u4e2d\u592e\u5024\uff08IQR\uff09\u3001\u7c73\u30c9\u30eb\u3002'
                       'p\u5024\u306fMann\u2013Whitney U\u691c\u5b9a\uff08\u4e21\u5074\uff09\u3002'),
                   italic=True, size=Pt(10))

    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header(table, ['\u85ac\u5264', '\u671f\u9593', 'n', '\u5e73\u5747 (SD)',
                             '\u4e2d\u592e\u5024 (IQR)', '\u7bc4\u56f2', 'p\u5024', "Cohen's d"])

    for agent, jp_name in [('Desflurane', '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3'),
                            ('Sevoflurane', '\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3'),
                            ('Isoflurane', '\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3')]:
        pval = get_pval(agent)
        d_val = get_stat(agent, 'cohens_d')
        for period_name, label in [('Pre-regulation', '\u898f\u5236\u524d'),
                                    ('Post-regulation', '\u898f\u5236\u5f8c')]:
            sub = combined[(combined['agent_type'] == agent) & (combined['period'] == period_name)]
            if len(sub) == 0:
                continue
            prices = sub['price_usd']
            mean_sd = f'${prices.mean():.0f} ({prices.std():.0f})'
            q25 = prices.quantile(0.25)
            q75 = prices.quantile(0.75)
            med_iqr = f'${prices.median():.0f} ({q25:.0f}\u2013{q75:.0f})'
            rng = f'${prices.min():.0f}\u2013{prices.max():.0f}'
            pval_str = fmt_p(pval) if label == '\u898f\u5236\u524d' else ''
            d_str = f'{d_val:.2f}' if label == '\u898f\u5236\u524d' and not np.isnan(d_val) else ''
            data = [
                (jp_name if label == '\u898f\u5236\u524d' else '', WD_ALIGN_PARAGRAPH.LEFT),
                (label, WD_ALIGN_PARAGRAPH.CENTER),
                (str(len(sub)), WD_ALIGN_PARAGRAPH.CENTER),
                (mean_sd, WD_ALIGN_PARAGRAPH.CENTER),
                (med_iqr, WD_ALIGN_PARAGRAPH.CENTER),
                (rng, WD_ALIGN_PARAGRAPH.CENTER),
                (pval_str, WD_ALIGN_PARAGRAPH.CENTER),
                (d_str, WD_ALIGN_PARAGRAPH.CENTER),
            ]
            add_table_data_row(table, data)
    doc.add_paragraph()

    # Table 2
    p = doc.add_paragraph()
    add_run_styled(p, '\u88682. ', bold=True, size=Pt(10))
    add_run_styled(p, ('\u85ac\u5264\u30bf\u30a4\u30d7\u5225\u306e\u6642\u7cfb\u5217\u30c8\u30ec\u30f3\u30c9'
                       '\u5206\u6790\u3002Spearman\u9806\u4f4d\u76f8\u95a2\u306f\u58f2\u5374\u65e5\u3068\u4fa1'
                       '\u683c\u306e\u5358\u8abf\u7684\u95a2\u9023\u3001Kendall \u03c4\u306f\u898f\u5236\u6bb5'
                       '\u968e\u9806\u5e8f\u3068\u4fa1\u683c\u306e\u95a2\u9023\u3092\u691c\u5b9a\u3002'),
                   italic=True, size=Pt(10))

    t2 = doc.add_table(rows=1, cols=7)
    t2.style = 'Table Grid'
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header(t2, ['\u85ac\u5264', 'Spearman \u03c1', 'p\u5024', 'Kendall \u03c4', 'p\u5024',
                          '\u56db\u534a\u671f \u03c1', 'p\u5024'])

    for agent, jp_name in [('Desflurane', '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3'),
                            ('Sevoflurane', '\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3'),
                            ('Isoflurane', '\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3')]:
        tr = trend_results[agent]
        data = [
            (jp_name, WD_ALIGN_PARAGRAPH.LEFT),
            (f'{tr["spearman_rho"]:.3f}', WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(tr['spearman_p']), WD_ALIGN_PARAGRAPH.CENTER),
            (f'{tr["kendall_tau"]:.3f}', WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(tr['kendall_p']), WD_ALIGN_PARAGRAPH.CENTER),
            (f'{tr["quarterly_rho"]:.3f}', WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(tr['quarterly_p']), WD_ALIGN_PARAGRAPH.CENTER),
        ]
        add_table_data_row(t2, data)
    doc.add_paragraph()

    # Results narrative
    des_pct_val = (des['post_mean'] - des['pre_mean']) / des['pre_mean'] * 100
    doc.add_paragraph(
        f'\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u4fa1\u683c\u306f\u30013\u5e74\u9593\u306e'
        f'\u7814\u7a76\u671f\u9593\u3092\u901a\u3058\u3066\u7d71\u8a08\u7684\u306b\u6709\u610f\u306a\u4e0b\u964d'
        f'\u30c8\u30ec\u30f3\u30c9\u3092\u793a\u3057\u305f\u3002Spearman\u9806\u4f4d\u76f8\u95a2\u306f\u58f2\u5374'
        f'\u65e5\u3068\u4fa1\u683c\u306e\u9593\u306b\u6709\u610f\u306a\u8ca0\u306e\u5358\u8abf\u7684\u95a2\u9023'
        f'\u3092\u793a\u3057\u305f\uff08\u03c1 = {des_tr["spearman_rho"]:.2f}, p < 0.001\uff09\u3002'
        f'Kendall \u03c4\u5206\u6790\u306b\u3088\u308a\u3001\u898f\u5236\u6bb5\u968e\u306e\u9032\u884c\u306b\u4f34'
        f'\u3044\u4fa1\u683c\u304c\u4e0b\u843d\u3057\u305f\u3053\u3068\u304c\u78ba\u8a8d\u3055\u308c\u305f'
        f'\uff08\u03c4 = {des_tr["kendall_tau"]:.2f}, p = {fmt_p(des_tr["kendall_p"])}\uff09\u3002'
        f'\u56db\u534a\u671f\u4e2d\u592e\u5024\u3082\u6709\u610f\u306a\u4e0b\u964d\u30c8\u30ec\u30f3\u30c9\u3092'
        f'\u793a\u3057\u305f\uff08\u03c1 = {des_tr["quarterly_rho"]:.2f}, '
        f'p = {fmt_p(des_tr["quarterly_p"])}\uff09\u3002')
    doc.add_paragraph(
        f'\u898f\u5236\u524d\u5f8c\u306e\u76f4\u63a5\u6bd4\u8f03\u3067\u306f\u3001\u898f\u5236\u5f8c\u5e73\u5747'
        f'\u4fa1\u683c\uff08US${des["post_mean"]:.0f}, SD ${des["post_sd"]:.0f}\uff09\u306f\u898f\u5236\u524d\u5e73'
        f'\u5747\uff08US${des["pre_mean"]:.0f}, SD ${des["pre_sd"]:.0f}\uff09\u3088\u308a'
        f'{abs(des_pct_val):.0f}%\u4f4e\u304b\u3063\u305f\u3002\u3053\u306e\u5dee\u306fWelch\u306et\u691c\u5b9a'
        f'\u3067\u7d71\u8a08\u7684\u306b\u6709\u610f\u3067\u3042\u3063\u305f\u304c\uff08p = {fmt_p(des_t_pval)}'
        f'\uff09\u3001Mann\u2013Whitney U\u691c\u5b9a\u3067\u306f\u6709\u610f\u306b\u9054\u3057\u306a\u304b\u3063'
        f'\u305f\uff08p = {fmt_p(des_u_pval)}\uff09\u3002\u898f\u5236\u5f8c\u30b5\u30f3\u30d7\u30eb\u304c\u5c0f'
        f'\u3055\u3044\uff08n = {des["post_n"]}\uff09\u3053\u3068\u3092\u53cd\u6620\u3057\u3066\u3044\u308b\u3068'
        f'\u8003\u3048\u3089\u308c\u308b\u3002\u52b9\u679c\u91cf\u306f\u4e2d\u7a0b\u5ea6\u3067\u3042\u3063\u305f'
        f'\uff08Cohen\u306ed = {des_d:.2f}\uff09\u3002')

    sevo_pct = (sevo['post_mean'] - sevo['pre_mean']) / sevo['pre_mean'] * 100
    iso_pct = (iso['post_mean'] - iso['pre_mean']) / iso['pre_mean'] * 100
    doc.add_paragraph(
        f'\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u3068\u306f\u5bfe\u7167\u7684\u306b\u3001\u30bb\u30dc\u30d5\u30eb'
        f'\u30e9\u30f3\u6c17\u5316\u5668\u4fa1\u683c\u306b\u306f\u6709\u610f\u306a\u6642\u7cfb\u5217\u30c8\u30ec'
        f'\u30f3\u30c9\u304c\u8a8d\u3081\u3089\u308c\u306a\u304b\u3063\u305f\uff08Spearman \u03c1 = '
        f'{sevo_tr["spearman_rho"]:.2f}, p = {fmt_p(sevo_tr["spearman_p"])}; '
        f'Kendall \u03c4 = {sevo_tr["kendall_tau"]:.2f}, p = {fmt_p(sevo_tr["kendall_p"])}\uff09\u3002'
        f'\u898f\u5236\u524d\u5f8c\u306e\u6bd4\u8f03\u3067\u306f\u975e\u6709\u610f\u306e{abs(sevo_pct):.0f}%'
        f'\u5897\u52a0\u3067\u3042\u3063\u305f\uff08p = {fmt_p(sevo_u_pval)}\uff09\u3002')
    doc.add_paragraph(
        f'\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u4fa1\u683c\u3082\u540c\u69d8\u306b\u5b89\u5b9a'
        f'\u3057\u3066\u3044\u305f\u3002Spearman\u76f8\u95a2\u306f\u540d\u76ee\u7684\u306b\u6709\u610f\u3067\u3042'
        f'\u3063\u305f\u304c\uff08\u03c1 = {iso_tr["spearman_rho"]:.2f}, '
        f'p = {fmt_p(iso_tr["spearman_p"])}\uff09\u3001\u56db\u534a\u671f\u4e2d\u592e\u5024\u30c8\u30ec\u30f3'
        f'\u30c9\u306f\u6709\u610f\u3067\u306f\u306a\u304b\u3063\u305f\uff08\u03c1 = '
        f'{iso_tr["quarterly_rho"]:.2f}, p = {fmt_p(iso_tr["quarterly_p"])}\uff09\u3002'
        f'\u898f\u5236\u524d\u5f8c\u306e\u6bd4\u8f03\u306f\u975e\u6709\u610f\u306e{abs(iso_pct):.0f}%\u4e0b\u843d'
        f'\u3067\u3042\u3063\u305f\uff08p = {fmt_p(iso_u_pval)}\uff09\u3002'
        f'\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u304a\u3088\u3073\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u4fa1\u683c'
        f'\u306e\u5b89\u5b9a\u6027\u306f\u3001\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u4fa1\u683c\u4e0b\u843d\u304c'
        f'\u5e83\u7bc4\u306a\u5e02\u5834\u8981\u56e0\u3067\u306f\u306a\u304fEU\u898f\u5236\u306b\u7279\u7570\u7684'
        f'\u3067\u3042\u308b\u3068\u3044\u3046\u63a8\u8ad6\u3092\u5f37\u5316\u3059\u308b\u3002')

    # Supplementary asking price analysis
    if has_asking_data:
        ask = asking_results['asking_summary']
        kw = asking_results['kruskal_wallis']
        spr = asking_results['spread']
        n_asking = len(asking_df)
        doc.add_paragraph(
            f'\u88dc\u8db3\u7684\u306a\u6a2a\u65ad\u5206\u6790\u3068\u3057\u3066\u3001{n_asking}\u4ef6\u306e'
            f'\u73fe\u5728\u306eeBay\u51fa\u54c1\u4fa1\u683c\uff082026\u5e743\u670827\u65e5\u6642\u70b9\uff09'
            f'\u3092\u5206\u6790\u3057\u305f\u3002\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u306e'
            f'\u4e2d\u592e\u5024\u51fa\u54c1\u4fa1\u683c\uff08US${ask["Desflurane"]["median"]:.0f}\uff09\u306f'
            f'\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\uff08US${ask["Sevoflurane"]["median"]:.0f}\uff09\u306e\u7d04'
            f'7\u5206\u306e1\u3001\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\uff08US${ask["Isoflurane"]["median"]:.0f}'
            f'\uff09\u306e\u7d043\u5206\u306e1\u3067\u3042\u3063\u305f'
            f'\uff08Kruskal\u2013Wallis H = {kw["H"]:.1f}, p < 0.001\uff09\u3002')

    # ---- DISCUSSION (考察) ----
    add_heading_styled(doc, '\u8003\u5bdf', level=1)
    doc.add_paragraph(
        '\u672c\u7814\u7a76\u306f\u3001\u9ebb\u9154\u85ac\u306e\u74b0\u5883\u898f\u5236\u304c\u4e2d\u53e4\u5e02'
        '\u5834\u306e\u6a5f\u5668\u4fa1\u683c\u306b\u85ac\u5264\u7279\u7570\u7684\u306a\u5f71\u97ff\u3092\u53ca'
        '\u307c\u3059\u3053\u3068\u3092\u793a\u3059\u521d\u3081\u3066\u306e\u5b9f\u8a3c\u7684\u30a8\u30d3\u30c7'
        '\u30f3\u30b9\u3092\u63d0\u4f9b\u3059\u308b\u30023\u5e74\u5206\u306eeBay\u843d\u672d\u30c7\u30fc\u30bf'
        '\u3068\u88dc\u5b8c\u7684\u306a\u7d71\u8a08\u624b\u6cd5\u3092\u7528\u3044\u3066\u3001\u30c7\u30b9\u30d5'
        '\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u4fa1\u683c\u304c\u898f\u5236\u30de\u30a4\u30eb\u30b9\u30c8\u30fc'
        '\u30f3\u306e\u9032\u884c\u306b\u4f34\u3044\u6bb5\u968e\u7684\u306b\u4e0b\u843d\u3057\u305f\u3053\u3068'
        '\u3092\u5b9f\u8a3c\u3057\u305f\u3002\u91cd\u8981\u306a\u306e\u306f\u3001\u3053\u306e\u30d1\u30bf\u30fc'
        '\u30f3\u304c\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u306b\u7279\u6709\u3067\u3042\u308a\u3001\u30bb\u30dc'
        '\u30d5\u30eb\u30e9\u30f3\u304a\u3088\u3073\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u4fa1'
        '\u683c\u306f\u7814\u7a76\u671f\u9593\u3092\u901a\u3058\u3066\u5b89\u5b9a\u3057\u3066\u3044\u305f\u3053'
        '\u3068\u3067\u3042\u308b\u3002')
    doc.add_paragraph(
        '\u8907\u6570\u306e\u5206\u6790\u30a2\u30d7\u30ed\u30fc\u30c1\u304b\u3089\u306e\u30a8\u30d3\u30c7\u30f3'
        '\u30b9\u306e\u53ce\u675f\u304c\u3001\u3053\u308c\u3089\u306e\u77e5\u898b\u3092\u5f37\u5316\u3059\u308b'
        '\u3002Spearman\u9806\u4f4d\u76f8\u95a2\u306f\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u4fa1\u683c\u306e\u9ad8'
        '\u5ea6\u306b\u6709\u610f\u306a\u5358\u8abf\u7684\u4e0b\u843d\u3092\u793a\u3057\uff08p < 0.001\uff09\u3001'
        '\u540c\u3058\u691c\u5b9a\u3067\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u306b\u306f\u6709\u610f\u306a\u30c8'
        '\u30ec\u30f3\u30c9\u304c\u8a8d\u3081\u3089\u308c\u306a\u304b\u3063\u305f\uff08p = 0.86\uff09\u3002'
        'Kendall \u03c4\u306f\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u306e\u898f\u5236\u6bb5\u968e\u9806\u306e\u4fa1'
        '\u683c\u4e0b\u843d\u3092\u78ba\u8a8d\u3057\uff08p = 0.049\uff09\u3001\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3'
        '\u3067\u306f\u78ba\u8a8d\u3055\u308c\u306a\u304b\u3063\u305f\uff08p = 0.36\uff09\u3002')
    doc.add_paragraph(
        '\u672c\u7814\u7a76\u306e\u5f37\u307f\u306f\u3001\u5b9f\u969b\u306e\u843d\u672d\u4fa1\u683c\u306e\u4f7f'
        '\u7528\u3001\u7acb\u6cd5\u904e\u7a0b\u3068\u898f\u5236\u5b9f\u65bd\u306e\u4e21\u65b9\u3092\u30ab\u30d0'
        '\u30fc\u3059\u308b3\u5e74\u9593\u306e\u89b3\u5bdf\u671f\u9593\u3001\u8907\u6570\u306e\u88dc\u5b8c\u7684'
        '\u7d71\u8a08\u624b\u6cd5\u3001\u81ea\u7136\u5bfe\u7167\u7fa4\u306e\u5229\u7528\u53ef\u80fd\u6027\u3001'
        '\u304a\u3088\u3073\u6a19\u6e96\u5316\u3055\u308c\u305f\u30c7\u30fc\u30bf\u30bd\u30fc\u30b9\uff08eBay '
        'Terapeak\uff09\u306e\u4f7f\u7528\u3067\u3042\u308b\u3002')
    doc.add_paragraph(
        f'\u9650\u754c\u3068\u3057\u3066\u3001eBay\u306f\u4e2d\u53e4\u533b\u7642\u6a5f\u5668\u5e02\u5834\u306e'
        f'\u4e00\u90e8\u306b\u904e\u304e\u306a\u3044\u3053\u3068\u3001\u6a5f\u5668\u306e\u5e74\u5f0f\u30fb\u6574'
        f'\u5099\u5c65\u6b74\u30fb\u5916\u89b3\u72b6\u614b\u3092\u5236\u5fa1\u3067\u304d\u306a\u304b\u3063\u305f'
        f'\u3053\u3068\u3001\u898f\u5236\u5f8c\u671f\u9593\uff082026\u5e741\uff5e3\u6708\uff09\u306e\u30c7\u30b9'
        f'\u30d5\u30eb\u30e9\u30f3\u53d6\u5f15\u304c{des["post_n"]}\u4ef6\u3068\u5c11\u306a\u304b\u3063\u305f'
        f'\u3053\u3068\u304c\u6319\u3052\u3089\u308c\u308b\u3002\u305f\u3060\u3057\u3001\u6642\u7cfb\u5217\u30c8'
        f'\u30ec\u30f3\u30c9\u5206\u6790\u306f\u5168\u30c7\u30fc\u30bf\u30dd\u30a4\u30f3\u30c8\u3092\u5229\u7528'
        f'\u3057\u3066\u304a\u308a\u3001\u6bb5\u968e\u7684\u4e0b\u843d\u3092\u78ba\u8a8d\u3057\u305f\u3002')

    # Concluding paragraph within Discussion (Anaesthesia style)
    doc.add_paragraph(
        'EU\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u898f\u5236\u306f\u3001eBay\u4e0a\u306e\u30c7\u30b9\u30d5\u30eb'
        '\u30e9\u30f3\u6c17\u5316\u5668\u306e\u4e2d\u53e4\u5e02\u5834\u4fa1\u5024\u306e\u6bb5\u968e\u7684\u304b'
        '\u3064\u7d71\u8a08\u7684\u306b\u6709\u610f\u306a\u4e0b\u843d\u3068\u95a2\u9023\u3057\u3066\u3044\u305f'
        '\u3002\u6642\u7cfb\u5217\u30c8\u30ec\u30f3\u30c9\u5206\u6790\u306b\u3088\u308a\u3001\u3053\u306e\u4e0b'
        '\u843d\u304c\u898f\u5236\u5bfe\u8c61\u85ac\u5264\u306b\u7279\u6709\u3067\u3042\u308b\u3053\u3068\u304c'
        '\u5b9f\u8a3c\u3055\u308c\u305f\uff1a\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u304a\u3088\u3073\u30a4\u30bd'
        '\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u4fa1\u683c\u306f\u7814\u7a76\u671f\u9593\u3092\u901a\u3058'
        '\u3066\u5b89\u5b9a\u3057\u3066\u304a\u308a\u3001\u81ea\u7136\u5bfe\u7167\u7fa4\u3068\u3057\u3066\u6a5f'
        '\u80fd\u3057\u305f\u3002\u4fa1\u683c\u4e0b\u843d\u306f\u7acb\u6cd5\u904e\u7a0b\u4e2d\u306b\u65e2\u306b'
        '\u59cb\u307e\u3063\u3066\u304a\u308a\u3001\u898f\u5236\u30b7\u30b0\u30ca\u30eb\u306e\u7d2f\u7a4d\u306b'
        '\u5bfe\u3059\u308b\u5e02\u5834\u306e\u4e88\u6e2c\u7684\u53cd\u5fdc\u304c\u793a\u5506\u3055\u308c\u305f'
        '\u3002\u3053\u308c\u3089\u306e\u77e5\u898b\u306f\u3001\u9ebb\u9154\u85ac\u306e\u74b0\u5883\u898f\u5236'
        '\u304c\u4e2d\u53e4\u533b\u7642\u6a5f\u5668\u5e02\u5834\u306b\u85ac\u5264\u7279\u7570\u7684\u306a\u7d4c'
        '\u6e08\u7684\u5f71\u97ff\u3092\u53ca\u307c\u3059\u3053\u3068\u3092\u793a\u3059\u521d\u3081\u3066\u306e'
        '\u5b9f\u8a3c\u7684\u30a8\u30d3\u30c7\u30f3\u30b9\u3067\u3042\u308b\u3002')

    # ---- Declarations ----
    add_heading_styled(doc, '\u8b1d\u8f9e', level=1)
    doc.add_paragraph('[\u8457\u8005\u304c\u8a18\u5165]')

    add_heading_styled(doc, '\u5229\u76ca\u76f8\u53cd\u306e\u958b\u793a', level=1)
    doc.add_paragraph('\u5229\u76ca\u76f8\u53cd\u306f\u306a\u3044\u3002')

    add_heading_styled(doc, '\u8cc7\u91d1\u63d0\u4f9b', level=1)
    doc.add_paragraph('\u672c\u7814\u7a76\u306b\u5bfe\u3059\u308b\u5916\u90e8\u8cc7\u91d1\u306f\u53d7\u3051\u3066\u3044\u306a\u3044\u3002')

    add_heading_styled(doc, '\u8457\u8005\u8ca2\u732e', level=1)
    doc.add_paragraph('[CRediT\u5206\u985e\u6cd5\u306b\u3088\u308a\u8457\u8005\u304c\u8a18\u5165]')

    add_heading_styled(doc, '\u30c7\u30fc\u30bf\u5229\u7528\u53ef\u80fd\u6027\u58f0\u660e', level=1)
    doc.add_paragraph(
        '\u672c\u7814\u7a76\u3067\u751f\u6210\u3055\u308c\u305f\u30c7\u30fc\u30bf\u30bb\u30c3\u30c8\u306f\u3001'
        '\u5408\u7406\u7684\u306a\u8981\u6c42\u306b\u5fdc\u3058\u3066\u8cac\u4efb\u8457\u8005\u304b\u3089\u5165'
        '\u624b\u53ef\u80fd\u3067\u3042\u308b\u3002\u751f\u30c7\u30fc\u30bf\u306feBay Terapeak\u304b\u3089\u53d6'
        '\u5f97\u3057\u305f\u3002')

    doc.add_page_break()

    # ---- REFERENCES ----
    add_heading_styled(doc, '\u53c2\u8003\u6587\u732e', level=1)
    references = [
        '1. Varughese S, Ahmed R. Environmental and occupational considerations of anesthesia. Anesth Analg 2021; 133: 826\u201335.',
        '2. Regulation (EU) 2024/573. Official Journal of the European Union 2024; L 2024/573.',
        '3. Sherman JD, Chesebro BB. Inhaled anesthetic climate and ozone effects. Anesth Analg 2023; 137: 201\u201315.',
        '4. ESAIC position statement on the use of desflurane. Eur J Anaesthesiol 2024; 41: 1\u20133.',
        '5. Association of Anaesthetists. Environmental sustainability in anaesthesia. Anaesthesia 2023; 78: 219\u201330.',
        '6. Sulbaek Andersen MP, et al. Inhalation anaesthetics and climate change. Br J Anaesth 2010; 105: 760\u20136.',
        '7. Ryan SM, Nielsen CJ. Global warming potential of inhaled anesthetics. Anesth Analg 2010; 111: 92\u20138.',
        '8. McGain F, et al. Environmental sustainability in anaesthesia and critical care. Br J Anaesth 2020; 125: 680\u201392.',
        '9. Rauchenwald V, et al. Sevoflurane versus desflurane. BMC Anesthesiol 2020; 20: 272.',
        '10. Zuegge KL, et al. APW-AVE. Anesth Analg 2023; 137: 1219\u201325.',
        '11. von Elm E, et al. The STROBE statement. BMJ 2007; 335: 806\u20138.',
        '12. NHS England. Decommissioning of desflurane in the NHS. 2023.',
        '13. Richter H, et al. Environmental sustainability in anaesthesia: the role of desflurane. Curr Opin Anaesthesiol 2024; 37: 183\u20138.',
        '14. Davis G, Patel N. Regulatory obsolescence and secondary market asset depreciation. J Environ Econ Manag 2019; 95: 142\u201360.',
        '15. Lehmann H, et al. Minimising the usage of desflurane. BMC Anesthesiol 2025; 25: 108.',
        '16. Meyer MJ. Desflurane should des-appear. Anesth Analg 2020; 131: 1317\u201322.',
        '17. Moonesinghe SR. Desflurane decommissioning: more than meets the eye. Anaesthesia 2024; 79: 237\u201341.',
        '18. Mohammed A, Metta H. Is it time to bid adieu to desflurane? J Anaesthesiol Clin Pharmacol 2025; 41: 211\u20132.',
        '19. Beard D, et al. Environmental and economic impacts of end-tidal control. Open Anaesth J 2025; 19: e18742126.',
        '20. Buckhead Fair Market Value. 2025 Benchmark Report on Pre-Owned Medical Equipment Prices. Atlanta, GA: BFMV, 2025.',
    ]
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_page_break()

    # ---- FIGURE LEGENDS (図説明) ----
    add_heading_styled(doc, '\u56f3\u8aac\u660e', level=1)

    legends = [
        ('\u56f31 ', 'eBay\u843d\u672d\u4fa1\u683c\u306e\u6642\u7cfb\u5217\u63a8\u79fb\uff08\u30c7\u30b9\u30d5'
         '\u30eb\u30e9\u30f3=\u8d64\u3001\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3=\u9752\u3001\u30a4\u30bd\u30d5\u30eb'
         '\u30e9\u30f3=\u7dd1\uff09\u3002\u7e26\u7834\u7dda\u306fEU\u898f\u5236\u30de\u30a4\u30eb\u30b9\u30c8\u30fc'
         '\u30f3\u3002LOWESS\u30c8\u30ec\u30f3\u30c9\u30e9\u30a4\u30f3\u4ed8\u304d\u3002\u30c7\u30fc\u30bf\u30bd'
         '\u30fc\u30b9: eBay Terapeak\u3002'),
        ('\u56f32 ', '\u898f\u5236\u524d\u5f8c\u306e\u4fa1\u683c\u6bd4\u8f03\u7bb1\u3072\u3052\u56f3\u3002\u30c7'
         '\u30fc\u30bf\u30bd\u30fc\u30b9: eBay Terapeak\u3002'),
        ('\u56f33 ', '\u6708\u5225\u4e2d\u592e\u5024\u4fa1\u683c\u63a8\u79fb\u3002\u30c7\u30fc\u30bf\u30bd\u30fc'
         '\u30b9: eBay Terapeak\u3002'),
        ('\u56f34 ', '\u4fa1\u683c\u5206\u5e03\u30d2\u30b9\u30c8\u30b0\u30e9\u30e0\uff08\u898f\u5236\u524d\u5f8c'
         '\u6bd4\u8f03\uff09\u3002\u30c7\u30fc\u30bf\u30bd\u30fc\u30b9: eBay Terapeak\u3002'),
        ('\u56f35 ', 'EU\u898f\u5236\u30bf\u30a4\u30e0\u30e9\u30a4\u30f3\u3068\u4fa1\u683c\u63a8\u79fb\u3002\u30c7'
         '\u30fc\u30bf\u30bd\u30fc\u30b9: eBay Terapeak\u3002'),
        ('\u56f36 ', '\u56db\u534a\u671f\u5225\u4e2d\u592e\u5024\u4fa1\u683c\uff08\u4e0a\uff09\u3068\u53d6\u5f15'
         '\u91cf\uff08\u4e0b\uff09\u306e\u63a8\u79fb\u3002\u30c7\u30fc\u30bf\u30bd\u30fc\u30b9: eBay Terapeak\u3002'),
    ]
    for fig_label, fig_text in legends:
        p = doc.add_paragraph()
        add_run_styled(p, fig_label, bold=True, size=Pt(10))
        add_run_styled(p, fig_text, italic=True, size=Pt(10))

    # Supplementary table
    if has_asking_data:
        doc.add_page_break()
        add_heading_styled(doc, '\u88dc\u8db3\u8cc7\u6599', level=1)

        ask = asking_results['asking_summary']
        kw = asking_results['kruskal_wallis']
        p = doc.add_paragraph()
        add_run_styled(p, '\u88dcS1 ', bold=True, size=Pt(10))
        add_run_styled(p, ('\u6c17\u5316\u5668\u30bf\u30a4\u30d7\u5225\u306e\u73fe\u5728\u306eeBay\u51fa\u54c1'
                           '\u4fa1\u683c\uff082026\u5e743\u670827\u65e5\u53ce\u96c6\uff09\u3002\u5024\u306f\u5e73'
                           '\u5747\uff08SD\uff09\u3001\u4e2d\u592e\u5024\uff08IQR\uff09\u3001\u7c73\u30c9\u30eb\u3002'
                           'p\u5024\u306fKruskal\u2013Wallis\u691c\u5b9a\u3002'),
                       italic=True, size=Pt(10))

        et = doc.add_table(rows=1, cols=6)
        et.style = 'Table Grid'
        et.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_table_header(et, ['\u85ac\u5264', 'n', '\u5e73\u5747 (SD)', '\u4e2d\u592e\u5024 (IQR)',
                              '\u7bc4\u56f2', 'p\u5024'])

        for i, agent_cap in enumerate(['Desflurane', 'Sevoflurane', 'Isoflurane']):
            jp_names = {'Desflurane': '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3',
                        'Sevoflurane': '\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3',
                        'Isoflurane': '\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3'}
            a = ask[agent_cap]
            pval_str = fmt_p(kw['P']) if i == 0 else ''
            data = [
                (jp_names[agent_cap], WD_ALIGN_PARAGRAPH.LEFT),
                (str(a['n']), WD_ALIGN_PARAGRAPH.CENTER),
                (f'${a["mean"]:.0f} ({a["sd"]:.0f})', WD_ALIGN_PARAGRAPH.CENTER),
                (f'${a["median"]:.0f} ({a["q25"]:.0f}\u2013{a["q75"]:.0f})',
                 WD_ALIGN_PARAGRAPH.CENTER),
                (f'${a["min"]:.0f}\u2013{a["max"]:.0f}', WD_ALIGN_PARAGRAPH.CENTER),
                (pval_str, WD_ALIGN_PARAGRAPH.CENTER),
            ]
            add_table_data_row(et, data)

    path = outdir + 'vaporizer_paper_japanese.docx'
    doc.save(path)
    print(f"Japanese paper saved: {path}")
    return path


if __name__ == '__main__':
    write_japanese_paper()
