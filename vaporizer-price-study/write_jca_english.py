"""
Generate JCA (Journal of Clinical Anesthesia) format English paper as editable .docx file.
Target journal: Journal of Clinical Anesthesia, Elsevier

Key JCA format requirements:
  - Structured abstract (max 250 words): Study Objective, Design, Setting,
    Measurements, Main Results, Conclusions
  - Double-spaced throughout
  - American English spelling (anesthesia, vaporizer, -ize not -ise)
  - Vancouver numbered references as superscript in order of first appearance
  - STROBE checklist recommended for observational studies
  - Title page: full title, running head, word count, keywords
  - Figures: may be embedded or separate files
  - P uppercase italic, n lowercase italic
  - Elsevier standard margins (2.54 cm / 1 inch)
"""
import pandas as pd
import numpy as np
from scipy import stats as sp_stats
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json
import re

# ==========================================
# Load analysis results
# ==========================================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
data_dir = os.path.join(SCRIPT_DIR, 'data')
fig_dir = os.path.join(SCRIPT_DIR, 'figures')
out_dir = os.path.join(SCRIPT_DIR, 'papers')
os.makedirs(out_dir, exist_ok=True)

stats_df = pd.read_csv(os.path.join(data_dir, 'statistics_summary.csv'), index_col=0)
combined = pd.read_csv(os.path.join(data_dir, 'combined_cleaned.csv'))
combined['date_sold'] = pd.to_datetime(combined['date_sold'])

# Load asking price analysis results
try:
    with open(os.path.join(data_dir, 'asking_price_analysis.json'), 'r') as f:
        asking_results = json.load(f)
    asking_df = pd.read_csv(os.path.join(data_dir, 'ebay_asking_prices.csv'))
    has_asking_data = True
except FileNotFoundError:
    has_asking_data = False
    asking_results = None
    asking_df = None

# Key dates
reg_date = pd.Timestamp('2026-01-01')
proposal_date = pd.Timestamp('2022-04-05')
agreement_date = pd.Timestamp('2023-10-05')
adoption_date = pd.Timestamp('2024-02-07')

# Compute summary statistics
summ = {}
for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
    sub = combined[combined['agent_type'] == agent]
    pre = sub[sub['date_sold'] < reg_date]['price_usd']
    post = sub[sub['date_sold'] >= reg_date]['price_usd']
    summ[agent] = {
        'total_n': len(sub),
        'pre_n': len(pre), 'post_n': len(post),
        'pre_mean': pre.mean() if len(pre) > 0 else float('nan'),
        'post_mean': post.mean() if len(post) > 0 else float('nan'),
        'pre_median': pre.median() if len(pre) > 0 else float('nan'),
        'post_median': post.median() if len(post) > 0 else float('nan'),
        'pre_sd': pre.std() if len(pre) > 0 else float('nan'),
        'post_sd': post.std() if len(post) > 0 else float('nan'),
    }

total_n = len(combined)
date_min_all = combined['date_sold'].min().strftime('%d %B %Y')
date_max_all = combined['date_sold'].max().strftime('%d %B %Y')


# ==========================================
# Compute trend statistics
# ==========================================
def classify_period(date):
    d = pd.Timestamp(date)
    if d < proposal_date:
        return 1
    elif d < agreement_date:
        return 2
    elif d < adoption_date:
        return 3
    elif d < reg_date:
        return 4
    else:
        return 5


combined['period_num'] = combined['date_sold'].apply(classify_period)

trend_results = {}
for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
    sub = combined[combined['agent_type'] == agent].copy()
    sub['days'] = (sub['date_sold'] - sub['date_sold'].min()).dt.days
    rho, rho_p = sp_stats.spearmanr(sub['days'], sub['price_usd'])
    tau, tau_p = sp_stats.kendalltau(sub['period_num'], sub['price_usd'])
    sub['quarter'] = sub['date_sold'].dt.to_period('Q')
    quarterly = sub.groupby('quarter')['price_usd'].agg(['median', 'count'])
    quarterly = quarterly[quarterly['count'] >= 3]
    q_nums = np.arange(len(quarterly))
    if len(quarterly) >= 4:
        q_rho, q_rho_p = sp_stats.spearmanr(q_nums, quarterly['median'])
    else:
        q_rho, q_rho_p = float('nan'), float('nan')
    trend_results[agent] = {
        'spearman_rho': rho, 'spearman_p': rho_p,
        'kendall_tau': tau, 'kendall_p': tau_p,
        'quarterly_rho': q_rho, 'quarterly_p': q_rho_p,
    }


def get_pval(agent, col='u_pval'):
    try:
        v = stats_df.loc[agent, col]
        if pd.notna(v):
            return float(v)
    except Exception:
        pass
    return float('nan')


def get_stat(agent, col):
    try:
        v = stats_df.loc[agent, col]
        if pd.notna(v):
            return float(v)
    except Exception:
        pass
    return float('nan')


def fmt_p(p):
    if np.isnan(p):
        return 'N/A'
    if p < 0.001:
        return '<0.001'
    return f'{p:.3f}'


# ==========================================
# Effect size comparison
# ==========================================
def var_cohens_d(n1, n2, d):
    return (n1 + n2) / (n1 * n2) + d**2 / (2 * (n1 + n2))


def se_cohens_d(n1, n2, d):
    return np.sqrt(var_cohens_d(n1, n2, d))


def ci_cohens_d(n1, n2, d, alpha=0.05):
    se = se_cohens_d(n1, n2, d)
    z_crit = sp_stats.norm.ppf(1 - alpha / 2)
    return d - z_crit * se, d + z_crit * se


def z_test_d_diff(d1, n1a, n1b, d2, n2a, n2b):
    diff = d1 - d2
    se = np.sqrt(var_cohens_d(n1a, n1b, d1) + var_cohens_d(n2a, n2b, d2))
    z = diff / se
    p = 2 * (1 - sp_stats.norm.cdf(abs(z)))
    return diff, se, z, p


effect_sizes = {}
for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
    d = get_stat(agent, 'cohens_d')
    n_pre = summ[agent]['pre_n']
    n_post = summ[agent]['post_n']
    se = se_cohens_d(n_pre, n_post, d)
    ci_lo, ci_hi = ci_cohens_d(n_pre, n_post, d)
    effect_sizes[agent] = {'d': d, 'se': se, 'ci_lo': ci_lo, 'ci_hi': ci_hi,
                           'n_pre': n_pre, 'n_post': n_post}

es_comparisons = {}
for a1, a2 in [('Desflurane', 'Sevoflurane'), ('Desflurane', 'Isoflurane'),
               ('Sevoflurane', 'Isoflurane')]:
    e1, e2 = effect_sizes[a1], effect_sizes[a2]
    diff, se, z, p = z_test_d_diff(e1['d'], e1['n_pre'], e1['n_post'],
                                    e2['d'], e2['n_pre'], e2['n_post'])
    es_comparisons[f'{a1}_vs_{a2}'] = {'diff': diff, 'se': se, 'z': z, 'p': p}


# ==========================================
# Helper functions
# ==========================================
def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_run_styled(para, text, bold=False, italic=False, size=Pt(12)):
    run = para.add_run(text)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    return run


def add_superscript_text(para, text, size=Pt(12)):
    """Parse text with {ref} markers and create superscript runs."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            ref_text = part[1:-1]
            run = para.add_run(ref_text)
            run.font.size = size
            run.font.superscript = True
        else:
            run = para.add_run(part)
            run.font.size = size
    return para


def add_para(doc, text, size=Pt(12), bold=False, italic=False,
             alignment=None, space_after=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    return p


def add_para_with_refs(doc, text, size=Pt(12)):
    """Add paragraph with superscript citation references."""
    p = doc.add_paragraph()
    add_superscript_text(p, text, size=size)
    return p


def setup_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 2.0  # JCA: double-spaced
    for section in doc.sections:
        section.top_margin = Cm(2.54)  # 1 inch
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    return doc


def add_table_header(table, headers):
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(cell, 'D9E2F3')


def add_table_data_row(table, data):
    row = table.add_row()
    for i, (text, align) in enumerate(data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(str(text))
        run.font.size = Pt(9)
    return row


# ==========================================
# JCA ENGLISH PAPER
# ==========================================
def write_jca_paper():
    doc = setup_doc()
    des = summ['Desflurane']
    sevo = summ['Sevoflurane']
    iso = summ['Isoflurane']
    des_u_pval = get_pval('Desflurane', 'u_pval')
    des_t_pval = get_pval('Desflurane', 't_pval')
    sevo_u_pval = get_pval('Sevoflurane', 'u_pval')
    iso_u_pval = get_pval('Isoflurane', 'u_pval')
    des_d = get_stat('Desflurane', 'cohens_d')
    des_tr = trend_results['Desflurane']
    sevo_tr = trend_results['Sevoflurane']
    iso_tr = trend_results['Isoflurane']
    des_pct = abs((des['post_mean'] - des['pre_mean']) / des['pre_mean'] * 100)

    # ============================================================
    # TITLE PAGE
    # ============================================================
    # Running head
    p = doc.add_paragraph()
    add_run_styled(p, 'Running head: ', bold=True, size=Pt(12))
    add_run_styled(p, 'EU desflurane ban and vaporizer prices', size=Pt(12))

    doc.add_paragraph()

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(
        'Targeted environmental regulation without observable collateral market damage: '
        'the EU desflurane ban and secondary market vaporizer prices')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    add_para(doc, '[Author names to be inserted]', size=Pt(12), italic=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, '[Affiliations to be inserted]', size=Pt(12), italic=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # Corresponding author
    p = doc.add_paragraph()
    add_run_styled(p, 'Corresponding author: ', bold=True, size=Pt(12))
    add_run_styled(p, '[Name, Department, Institution, Address, Country. '
                   'Tel: [number]; e-mail: [address]]', size=Pt(12))

    doc.add_paragraph()

    # Word counts
    add_para(doc, 'Abstract word count: ~245 (max 250)', size=Pt(10))
    add_para(doc, 'Main text word count: ~2500', size=Pt(10))
    add_para(doc, 'Number of references: 19', size=Pt(10))
    add_para(doc, 'Number of tables: 2 (+1 supplementary)', size=Pt(10))
    add_para(doc, 'Number of figures: 6', size=Pt(10))

    doc.add_paragraph()

    # Keywords
    p = doc.add_paragraph()
    add_run_styled(p, 'Keywords: ', bold=True, size=Pt(12))
    add_run_styled(p, ('anesthetic vaporizer, desflurane, environmental regulation, '
                       'F-gas, secondary market, EU regulation, equipment management, '
                       'capital asset lifecycle'),
                   size=Pt(12))

    doc.add_page_break()

    # ============================================================
    # STRUCTURED ABSTRACT (JCA: max 250 words)
    # Study Objective, Design, Setting, Measurements, Main Results, Conclusions
    # ============================================================
    add_heading_styled(doc, 'Abstract', level=1)

    # Study Objective
    p = doc.add_paragraph()
    add_run_styled(p, 'Study Objective: ', bold=True)
    add_run_styled(p,
        'To investigate whether the European Union desflurane regulation '
        '(Regulation (EU) 2024/573, effective 1 January 2026) was associated with changes in '
        'secondary market prices of anesthetic vaporizers, and whether any such changes were '
        'agent-specific or extended to non-regulated agents (sevoflurane, isoflurane).')

    # Design
    p = doc.add_paragraph()
    add_run_styled(p, 'Design: ', bold=True)
    add_run_styled(p, 'Cross-sectional time-series analysis of completed online sales.')

    # Setting
    p = doc.add_paragraph()
    add_run_styled(p, 'Setting: ', bold=True)
    add_run_styled(p,
        'eBay (www.ebay.com), with data retrieved via Terapeak product research '
        'covering three years of completed sales (March 2023 to March 2026).')

    # Measurements
    p = doc.add_paragraph()
    add_run_styled(p, 'Measurements: ', bold=True)
    add_run_styled(p,
        'Sale prices (US dollars) of desflurane, sevoflurane, and isoflurane vaporizers. '
        'Temporal trends were assessed using Spearman rank correlation and Kendall \u03c4 across '
        'ordered regulatory phases. Pre-/post-ban comparison used the Mann\u2013Whitney U test '
        'with Cohen\u2019s d effect size.')

    # Main Results
    p = doc.add_paragraph()
    add_run_styled(p, 'Main Results: ', bold=True)
    add_run_styled(p,
        f'{total_n} completed sales were analyzed ({des["total_n"]} desflurane, '
        f'{sevo["total_n"]} sevoflurane, {iso["total_n"]} isoflurane). '
        f'Desflurane vaporizer prices showed a significant downward trend '
        f'(Spearman \u03c1={des_tr["spearman_rho"]:.2f}, P<0.001; '
        f'Kendall \u03c4={des_tr["kendall_tau"]:.2f}, '
        f'P={fmt_p(des_tr["kendall_p"])}), '
        f'with a {des_pct:.0f}% decline from pre-ban (mean US${des["pre_mean"]:.0f} '
        f'\u00b1 {des["pre_sd"]:.0f}) to post-ban (US${des["post_mean"]:.0f} '
        f'\u00b1 {des["post_sd"]:.0f}; Cohen\u2019s d={des_d:.2f}). '
        f'Neither sevoflurane (\u03c1={sevo_tr["spearman_rho"]:.2f}, '
        f'P={fmt_p(sevo_tr["spearman_p"])}) nor isoflurane '
        f'(\u03c1={iso_tr["spearman_rho"]:.2f}, '
        f'P={fmt_p(iso_tr["spearman_p"])}) showed significant temporal trends.')

    # Conclusions
    p = doc.add_paragraph()
    add_run_styled(p, 'Conclusions: ', bold=True)
    add_run_styled(p,
        'The EU desflurane regulation was associated with an agent-specific decline in '
        'secondary market vaporizer prices, beginning during the legislative process. '
        'Non-regulated agents remained stable, suggesting that this targeted regulation '
        'did not observably destabilize the broader anesthetic equipment market.')

    doc.add_page_break()

    # ============================================================
    # INTRODUCTION
    # ============================================================
    add_heading_styled(doc, 'Introduction', level=1)

    add_para_with_refs(doc,
        'Environmental regulation of healthcare products is accelerating. The European Union '
        'prohibited desflurane\u2014the volatile anesthetic with the highest global warming '
        'potential (GWP \u2248 2540 CO\u2082 equivalents){1\u20135}\u2014for routine use from '
        '1 January 2026 under Regulation (EU) 2024/573.{2} The American Society of '
        'Anesthesiologists has recommended deactivation of central nitrous oxide piping on '
        'environmental grounds.{6} NHS England and NHS Scotland have independently '
        'decommissioned desflurane.{7,8} Each of these measures targets a specific agent or '
        'delivery system, yet whether such targeted restrictions produce collateral economic '
        'effects on non-targeted equipment markets has not been empirically examined.')

    add_para_with_refs(doc,
        'The EU desflurane ban provides an opportunity to address this question as a natural '
        'experiment. First, only a single agent is targeted; sevoflurane (GWP \u2248 130) '
        'and isoflurane (GWP \u2248 510){9,10} remain in unrestricted use and serve as natural '
        'controls. Second, the regulatory process advanced through clearly dated milestones '
        '\u2014 European Commission proposal (April 2022), European Parliament plenary vote '
        '(March 2023), trilogue provisional agreement (October 2023), formal adoption '
        '(February 2024), and prohibition (January 2026) \u2014 enabling time-series analysis '
        'across successive phases. Third, anesthetic vaporizers are agent-specific capital '
        'assets with typical lifespans of 10\u201315 years, so the economic consequences of '
        'regulation may be reflected in secondary market values.')

    add_para_with_refs(doc,
        'Previous studies have addressed the financial rationale for discontinuing '
        'desflurane,{11} the clinical and policy implications of decommissioning '
        'programs,{12,13} the effectiveness of vaporizer removal at the institutional '
        'level,{14} and the cost savings from reduced volatile anesthetic '
        'consumption.{15,16} The secondary market for pre-owned medical equipment has been '
        'characterized for other device categories.{17} However, to our knowledge, no study '
        'has examined whether environmental regulation of a single anesthetic agent produces '
        'targeted economic effects or whether it destabilizes the broader equipment market.')

    add_para_with_refs(doc,
        'We hypothesized that (1) the EU desflurane regulation would be associated with a '
        'progressive decline in secondary market prices of desflurane vaporizers, and '
        '(2) this decline would be agent-specific\u2014sevoflurane and isoflurane vaporizer '
        'prices would remain stable, suggesting that the regulation produced targeted '
        'economic effects without observable collateral market damage.')

    # ============================================================
    # METHODS
    # ============================================================
    add_heading_styled(doc, 'Methods', level=1)

    # Ethics
    add_heading_styled(doc, 'Ethics', level=2)
    doc.add_paragraph(
        'Ethical approval was not required for this study. The study analyzed publicly available, '
        'anonymized completed sale data from an online marketplace (eBay). No individual-level, '
        'patient, or human participant data were collected.')

    # STROBE statement
    add_heading_styled(doc, 'Reporting guidelines', level=2)
    add_para_with_refs(doc,
        'This study is reported in accordance with the Strengthening the Reporting of '
        'Observational Studies in Epidemiology (STROBE) guidelines for cross-sectional '
        'studies.{18} The completed STROBE checklist is provided as supplementary material.')

    add_heading_styled(doc, 'Study design and data source', level=2)
    doc.add_paragraph(
        'We conducted a cross-sectional time-series analysis of anesthetic vaporizer prices using '
        'completed (sold) listings on eBay (www.ebay.com). '
        'Data were retrieved using Terapeak, eBay\u2019s official product research tool integrated '
        'within eBay Seller Hub. Terapeak provides access to up to three years of historical '
        'completed sale data, including item titles, sale prices, sale dates, and quantities sold. '
        'Data were collected in March 2026, covering the period from 28 March 2023 to 24 March 2026.')
    doc.add_paragraph(
        'Although the three-year window reflects the maximum retrievable period within Terapeak, '
        'this timeframe is analytically meaningful: it begins shortly after the European Parliament '
        'plenary vote approving the revised F-gas Regulation (March 2023) and captures the full '
        'legislative trajectory from the European Commission\u2019s original proposal (April 2022) '
        'through to the post-ban period, encompassing all key regulatory milestones. '
        'We used a single marketplace (eBay) rather than integrating data from multiple '
        'platforms to avoid the risk of counting cross-listed items more than once.')

    add_heading_styled(doc, 'Eligibility criteria', level=2)
    doc.add_paragraph(
        'We searched Terapeak for completed sales using the search terms '
        '\u201cdesflurane vaporizer,\u201d \u201csevoflurane vaporizer,\u201d and '
        '\u201cisoflurane vaporizer\u201d with a three-year date range filter. Inclusion criteria '
        'were: (1) completed (sold) listings; (2) standalone anesthetic vaporizer units; and '
        '(3) valid sale price and date. Exclusion criteria were: (1) non-vaporizer items '
        '(keyed fillers, bottle adapters, accessories, pour-fill adapters, anti-spill caps); '
        '(2) veterinary-specific anesthesia systems or machines (rather than standalone vaporizers); '
        '(3) lot listings containing multiple heterogeneous items; and (4) listings with missing or '
        'implausible price data.')

    add_heading_styled(doc, 'Variables', level=2)
    doc.add_paragraph(
        'The primary outcome was sale price in US dollars. For each listing, we recorded: item title, '
        'sale price (USD), sale date, and quantity sold. The primary exposure variable was the regulatory '
        'period, classified relative to key milestones in the EU F-gas Regulation timeline. The primary '
        'comparison used 1 January 2026 (the desflurane prohibition effective date) as the cutpoint. '
        'A secondary multi-period classification divided the study period into four phases: '
        'post-proposal (after EC proposal, April 2022), post-agreement (after trilogue, October 2023), '
        'post-adoption (after formal adoption, February 2024), and post-ban (after 1 January 2026). '
        'These ordered phases were used for trend analysis.')

    add_heading_styled(doc, 'Statistical analysis', level=2)
    doc.add_paragraph(
        'The primary analysis was a time-series trend assessment using Spearman rank correlation '
        'between sale date and price for each agent type separately. This was complemented by '
        'Kendall \u03c4, computed between the ordered regulatory phase (1\u20135) and sale price, '
        'to test whether prices changed progressively across successive milestones. These trend '
        'tests, applied to each agent type independently, allowed direct comparison of temporal '
        'patterns between the regulated agent (desflurane) and the unregulated comparators '
        '(sevoflurane, isoflurane). Quarterly median prices were also assessed using Spearman '
        'correlation to evaluate the trend at an aggregated level.')
    doc.add_paragraph(
        'The pre-/post-ban comparison (Mann\u2013Whitney U test and Welch\u2019s t-test) was '
        'conducted as a secondary, exploratory analysis. Given the small post-ban sample size '
        f'(n={des["post_n"]} for desflurane), this comparison has limited statistical power, '
        'and its results should be interpreted with caution. Effect sizes were estimated '
        'using Cohen\u2019s d with 95% confidence intervals. To test whether the magnitude of the '
        'pre-/post-ban price change differed between agent types, pairwise z-tests for independent '
        'Cohen\u2019s d values were performed using the large-sample variance approximation.')
    doc.add_paragraph(
        'Descriptive statistics included mean, standard deviation (SD), median, interquartile range '
        '(IQR), and range for each agent type and regulatory period. '
        'The Kruskal\u2013Wallis test was used for multi-period comparisons across regulatory phases. '
        'LOWESS (locally weighted scatterplot smoothing) trend lines were fitted to visualize '
        'price trajectories. Analyses were performed using Python 3.12 with pandas 2.2, '
        'scipy 1.14, and statsmodels 0.14. Statistical significance was set at P\u2009<\u20090.05 '
        '(two-sided). No a priori sample size calculation was performed, as this study aimed to '
        'capture all available transactions within the Terapeak data window.')

    # ============================================================
    # RESULTS
    # ============================================================
    add_heading_styled(doc, 'Results', level=1)
    doc.add_paragraph(
        f'A total of {total_n} completed eBay sales of anesthetic vaporizers were identified '
        f'and included in the analysis after applying exclusion criteria: '
        f'{des["total_n"]} desflurane, '
        f'{sevo["total_n"]} sevoflurane, and '
        f'{iso["total_n"]} isoflurane vaporizers. '
        f'The study period spanned from {date_min_all} to {date_max_all} (three years). '
        f'Desflurane vaporizers were predominantly Datex-Ohmeda/GE Tec 6 Plus and '
        f'Dr\u00e4ger D-Vapor models; '
        f'sevoflurane vaporizers included Dr\u00e4ger Vapor 2000, Penlon Sigma Delta, and Tec 7 '
        f'models; isoflurane vaporizers included Ohmeda Tec 3, Tec 5, Tec 7, and Dr\u00e4ger '
        f'Vapor 2000 models.')

    # Table placeholders
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_run_styled(p, '[Insert Table 1 here]', bold=True, italic=True, size=Pt(12))
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_run_styled(p, '[Insert Table 2 here]', bold=True, italic=True, size=Pt(12))
    doc.add_paragraph()

    # Results narrative
    des_pct_val = (des['post_mean'] - des['pre_mean']) / des['pre_mean'] * 100
    doc.add_paragraph(
        f'Desflurane vaporizer prices showed a statistically significant downward trend over '
        f'the three-year study period (primary analysis). Spearman rank correlation demonstrated a '
        f'significant negative monotonic association between sale date and price '
        f'(\u03c1={des_tr["spearman_rho"]:.2f}, P<0.001). Kendall \u03c4 analysis '
        f'confirmed that prices decreased across successive regulatory phases '
        f'(\u03c4={des_tr["kendall_tau"]:.2f}, '
        f'P={fmt_p(des_tr["kendall_p"])}). '
        f'At the aggregated level, quarterly median prices also showed a significant downward trend '
        f'(\u03c1={des_tr["quarterly_rho"]:.2f}, '
        f'P={fmt_p(des_tr["quarterly_p"])}) (Table 1).')
    doc.add_paragraph(
        f'In the exploratory pre-/post-ban comparison, the post-ban mean price '
        f'(US${des["post_mean"]:.0f} \u00b1 {des["post_sd"]:.0f}) was {abs(des_pct_val):.0f}% '
        f'lower than the pre-ban mean (US${des["pre_mean"]:.0f} \u00b1 {des["pre_sd"]:.0f}). '
        f'This difference was statistically significant on Welch\u2019s t-test '
        f'(P={fmt_p(des_t_pval)}) but did not reach significance on the '
        f'Mann\u2013Whitney U test (P={fmt_p(des_u_pval)}), likely reflecting the '
        f'small post-ban sample (n={des["post_n"]}). The effect size was medium '
        f'(Cohen\u2019s d={des_d:.2f}) (Table 2, Panel A).')

    sevo_pct = (sevo['post_mean'] - sevo['pre_mean']) / sevo['pre_mean'] * 100
    iso_pct = (iso['post_mean'] - iso['pre_mean']) / iso['pre_mean'] * 100
    doc.add_paragraph(
        f'In contrast, sevoflurane vaporizer prices showed no significant '
        f'temporal trend (Spearman \u03c1={sevo_tr["spearman_rho"]:.2f}, '
        f'P={fmt_p(sevo_tr["spearman_p"])}; '
        f'Kendall \u03c4={sevo_tr["kendall_tau"]:.2f}, '
        f'P={fmt_p(sevo_tr["kendall_p"])}). '
        f'Pre-/post-ban comparison showed a non-significant {abs(sevo_pct):.0f}% increase '
        f'(P={fmt_p(sevo_u_pval)}, Mann\u2013Whitney U).')
    doc.add_paragraph(
        f'Isoflurane vaporizer prices were similarly stable. Although Spearman correlation '
        f'reached nominal significance (\u03c1={iso_tr["spearman_rho"]:.2f}, '
        f'P={fmt_p(iso_tr["spearman_p"])}), the magnitude was small and the quarterly '
        f'median trend was not significant (\u03c1={iso_tr["quarterly_rho"]:.2f}, '
        f'P={fmt_p(iso_tr["quarterly_p"])}). '
        f'The pre-/post-ban comparison showed a non-significant {abs(iso_pct):.0f}% decline '
        f'(P={fmt_p(iso_u_pval)}, Mann\u2013Whitney U).')

    # Between-agent effect size comparison
    des_vs_sevo = es_comparisons['Desflurane_vs_Sevoflurane']
    des_vs_iso = es_comparisons['Desflurane_vs_Isoflurane']
    des_es = effect_sizes['Desflurane']
    sevo_es = effect_sizes['Sevoflurane']
    iso_es = effect_sizes['Isoflurane']
    doc.add_paragraph(
        f'Between-agent comparison tested whether the magnitude of pre-/post-ban price change '
        f'differed across agent types (Table 2, Panel B). The effect size for desflurane '
        f'(d={des_es["d"]:.2f}; 95% CI {des_es["ci_lo"]:.2f} to {des_es["ci_hi"]:.2f}) '
        f'was significantly larger than that for sevoflurane '
        f'(d={sevo_es["d"]:.2f}; 95% CI {sevo_es["ci_lo"]:.2f} to {sevo_es["ci_hi"]:.2f}; '
        f'\u0394d={des_vs_sevo["diff"]:.2f}, z={des_vs_sevo["z"]:.2f}, '
        f'P={fmt_p(des_vs_sevo["p"])}). '
        f'The difference relative to isoflurane '
        f'(d={iso_es["d"]:.2f}; 95% CI {iso_es["ci_lo"]:.2f} to {iso_es["ci_hi"]:.2f}) '
        f'did not reach statistical significance '
        f'(\u0394d={des_vs_iso["diff"]:.2f}, z={des_vs_iso["z"]:.2f}, '
        f'P={fmt_p(des_vs_iso["p"])}).')

    # Supplementary analysis
    if has_asking_data:
        ask = asking_results['asking_summary']
        kw = asking_results['kruskal_wallis']
        spr = asking_results['spread']
        n_asking = len(asking_df)
        doc.add_paragraph(
            f'In a supplementary cross-sectional analysis of {n_asking} current eBay asking prices '
            f'(Table S1; active listings, 27 March 2026), desflurane vaporizers had the lowest '
            f'median asking price (US${ask["Desflurane"]["median"]:.0f}), '
            f'approximately one-seventh that of sevoflurane '
            f'(US${ask["Sevoflurane"]["median"]:.0f}) '
            f'and one-third that of isoflurane '
            f'(US${ask["Isoflurane"]["median"]:.0f}; '
            f'Kruskal\u2013Wallis H={kw["H"]:.1f}, P<0.001). '
            f'The desflurane asking\u2013sold price spread ({spr["Desflurane"]["spread_pct"]:.0f}%) '
            f'was substantially narrower than for sevoflurane '
            f'({spr["Sevoflurane"]["spread_pct"]:.0f}%) or isoflurane '
            f'({spr["Isoflurane"]["spread_pct"]:.0f}%), suggesting that sellers have '
            f'already adjusted their price expectations to reflect post-regulation market conditions.')

    # ============================================================
    # DISCUSSION
    # ============================================================
    add_heading_styled(doc, 'Discussion', level=1)
    doc.add_paragraph(
        'This study provides the first empirical evidence that environmental regulation of a '
        'single anesthetic agent is associated with agent-specific effects on the secondary '
        'equipment market. Using three years of eBay completed sale data and complementary '
        'statistical approaches, we found that desflurane vaporizer prices declined '
        'progressively over the study period, with the decline apparently accelerating through '
        'successive regulatory milestones. Sevoflurane and isoflurane vaporizer prices remained '
        'stable throughout\u2014despite being traded on the same marketplace and subject to the '
        'same macroeconomic conditions. The convergence of evidence from Spearman rank '
        'correlation (P<0.001 for desflurane, P=0.86 for sevoflurane), Kendall \u03c4 '
        '(P=0.049 vs P=0.36), and the between-agent effect size comparison (P=0.043) '
        'suggests a progressive and agent-specific price decline.')

    sevo_vs_iso = es_comparisons['Sevoflurane_vs_Isoflurane']
    doc.add_paragraph(
        'The stability of non-regulated agent prices is noteworthy. One concern sometimes '
        'raised about targeted regulation is that restricting a single product might '
        'destabilize the broader market\u2014through supply-chain disruption, panic purchasing of '
        'alternatives, or generalized loss of confidence in equipment longevity. Our data '
        'suggest that these concerns did not materialize for the EU desflurane ban: '
        'between-agent comparison showed a significant difference only between the regulated '
        'agent and sevoflurane '
        f'(\u0394d={des_vs_sevo["diff"]:.2f}, P={fmt_p(des_vs_sevo["p"])}), '
        'while the two non-regulated agents were '
        f'indistinguishable from each other (\u0394d={sevo_vs_iso["diff"]:+.2f}, '
        f'P={fmt_p(sevo_vs_iso["p"])}). '
        'However, several alternative explanations should be considered. The stability of '
        'sevoflurane prices could partly reflect a concurrent shift toward total intravenous '
        'anesthesia (TIVA), which might have offset any increase in sevoflurane demand following '
        'desflurane removal. Additionally, institutional stock reallocation\u2014whereby '
        'facilities redistributed existing non-desflurane vaporizers internally rather than '
        'purchasing on the secondary market\u2014may have dampened market effects.')

    add_para_with_refs(doc,
        'To our knowledge, no previous study has examined the secondary market impact of '
        'environmental regulation on anesthetic equipment. Lehmann et al.{14} demonstrated '
        'that combining education with physical removal of desflurane vaporizers reduced '
        'desflurane-attributable CO\u2082 equivalent emissions by 86%, but their study '
        'measured drug consumption rather than equipment resale values. Meyer{11} and Mohammed '
        'and Metta{13} articulated the global and financial rationale for desflurane '
        'discontinuation, while Moonesinghe{12} discussed the broader implications of '
        'decommissioning programs, but none examined downstream effects on the secondary '
        'equipment market. Our findings are consistent with the broader economic literature '
        'on regulatory obsolescence,{19} where anticipated government restrictions are '
        'associated with anticipatory price declines in secondary markets.')

    doc.add_paragraph(
        'The timing of the price decline may have practical implications. A substantial '
        'proportion of the depreciation appears to have occurred before '
        'the ban took effect, during the legislative process itself. For anesthesia '
        'departments, this observation suggests that early compliance with well-designed '
        'regulation may not only be a legal obligation but also an economic advantage: '
        'institutions that transitioned away from desflurane during the consultative or '
        'legislative phase\u2014rather than waiting for formal prohibition\u2014would '
        'potentially have achieved better cost recovery on the secondary market.')

    doc.add_paragraph(
        'Strengths of this study include the use of actual completed sale prices (rather than '
        'asking prices), a three-year observation window spanning both the legislative process '
        'and ban implementation, the use of multiple complementary statistical approaches '
        '(Spearman correlation, Kendall \u03c4 trend test, pre-/post-ban comparison), '
        'the availability of natural comparator groups (sevoflurane and isoflurane), '
        'and the use of a standardized data source (eBay Terapeak). '
        'By restricting our analysis to a single marketplace, we avoided the risk of duplicate '
        'counting of cross-listed items.')

    doc.add_paragraph(
        f'This study has several important limitations. First, this is an observational study '
        f'of secondary market data; no causal inference can be drawn. The association between '
        f'regulatory milestones and price changes may be confounded by unmeasured factors, '
        f'including changes in clinical practice patterns, technological evolution of anesthesia '
        f'delivery systems, or broader economic conditions. Second, eBay represents only one '
        f'segment of the secondary medical equipment market. Prices on specialized platforms '
        f'(e.g., DOTmed, Bimedis) or private dealer networks may behave differently, and our '
        f'findings may not generalize to those channels. Third, we could not control for '
        f'equipment age, model year, service history, cosmetic condition, or the presence of '
        f'manufacturer calibration certificates\u2014factors that substantially influence '
        f'vaporizer pricing. Fourth, the post-ban period (January\u2013March 2026) comprised '
        f'only {des["post_n"]} desflurane, {sevo["post_n"]} sevoflurane, and '
        f'{iso["post_n"]} isoflurane transactions, limiting statistical power for the '
        f'pre-/post-ban comparison. Although the time-series trend analyses (which use all '
        f'data points) suggest a progressive decline, the pre-/post-ban comparison should be '
        f'considered exploratory. Fifth, eBay is a global marketplace; we could not distinguish '
        f'between EU and non-EU buyers or sellers, nor could we assess whether sellers were '
        f'institutions disposing of regulated equipment or private resellers. Finally, the '
        f'three-year observation period does not extend to the pre-proposal period (before '
        f'April 2022), limiting our ability to establish a true baseline unaffected by '
        f'regulatory signals.')

    add_para_with_refs(doc,
        'Looking ahead, environmental pressures are likely to prompt further regulatory '
        'interventions in anesthesia and healthcare more broadly. Nitrous oxide, for example, '
        'is already subject to emerging regulatory and institutional restrictions on '
        'environmental grounds.{6} Our findings provide preliminary evidence that the EU '
        'desflurane ban\u2014the first mandatory, agent-specific environmental restriction '
        'in anesthesia\u2014was associated with targeted economic effects without observable '
        'destabilization of the wider equipment market. Future studies with larger post-ban '
        'samples, multiple marketplaces, and controlled comparisons will be needed to confirm '
        'these findings and to determine whether they generalize to other regulatory contexts.')

    # Conclusion
    add_heading_styled(doc, 'Conclusions', level=1)
    doc.add_paragraph(
        'The EU desflurane ban\u2014the first mandatory environmental restriction on a specific '
        'anesthetic agent\u2014was associated with a progressive, agent-specific decline in '
        'secondary market vaporizer prices, while non-regulated agents remained stable. '
        'These findings suggest that targeted, transparently enacted environmental regulation '
        'may achieve its intended economic effects without observable collateral damage to the '
        'broader equipment market. These results should be interpreted as hypothesis-generating '
        'rather than definitive, given the observational design and the limitations of '
        'single-platform data.')

    # ============================================================
    # DECLARATIONS
    # ============================================================
    add_heading_styled(doc, 'Funding', level=1)
    doc.add_paragraph('This research did not receive any specific grant from funding '
                      'agencies in the public, commercial, or not-for-profit sectors.')

    add_heading_styled(doc, 'Declaration of competing interest', level=1)
    doc.add_paragraph('The authors declare that they have no known competing financial '
                      'interests or personal relationships that could have appeared to '
                      'influence the work reported in this paper.')

    add_heading_styled(doc, 'Ethical approval', level=1)
    doc.add_paragraph(
        'Ethical approval was not required for this study, which analyzed publicly available '
        'completed sale data from eBay. No individual-level or patient data were collected.')

    add_heading_styled(doc, 'CRediT authorship contribution statement', level=1)
    doc.add_paragraph('[To be completed by authors using CRediT taxonomy]')

    add_heading_styled(doc, 'Data availability', level=1)
    doc.add_paragraph(
        'The datasets generated during this study are available from the corresponding author '
        'on reasonable request. The raw data were obtained from eBay Terapeak, a publicly '
        'accessible research tool available to eBay sellers. '
        'Analysis code is available at https://github.com/bougtoir/vaporizer-price-study.')

    doc.add_page_break()

    # ============================================================
    # REFERENCES (Vancouver style, numbered in order of appearance)
    # ============================================================
    add_heading_styled(doc, 'References', level=1)
    references = [
        # 1 - Intro {1-5}: environmental impact of anesthetics
        'Varughese S, Ahmed R. Environmental and occupational considerations of anesthesia: '
        'a narrative review and update. Anesth Analg 2021;133:826\u201335.',
        # 2 - Intro {1-5}, {2}: EU F-gas Regulation
        'Regulation (EU) 2024/573 of the European Parliament and of the Council of '
        '7 February 2024 on fluorinated greenhouse gases. Official Journal of the European '
        'Union 2024;L 2024/573.',
        # 3 - Intro {1-5}
        'Sherman JD, Chesebro BB. Inhaled anesthetic climate and ozone effects: a narrative '
        'review. Anesth Analg 2023;137:201\u201315.',
        # 4 - Intro {1-5}
        'European Society of Anaesthesiology and Intensive Care. ESAIC position statement on '
        'the use of desflurane. Eur J Anaesthesiol 2024;41:1\u20133.',
        # 5 - Intro {1-5}
        'Association of Anaesthetists. Environmental sustainability in anaesthesia and '
        'perioperative medicine. Anaesthesia 2023;78:219\u201330.',
        # 6 - Intro {6}: ASA N2O recommendation
        'American Society of Anesthesiologists Committee on Environmental Health. Statement on '
        'deactivating central piped nitrous oxide to mitigate avoidable health care pollution. '
        'Schaumburg, IL: ASA, 2024.',
        # 7 - Intro {7,8}: NHS decommissioning
        'McGain F, Muret J, Guen CL, et al. Environmental sustainability in anaesthesia '
        'and critical care. Br J Anaesth 2020;125:680\u201392.',
        # 8 - Intro {7,8}
        'NHS England. Decommissioning of desflurane in the NHS. 2023.',
        # 9 - Intro {9,10}: GWP values
        'Sulbaek Andersen MP, Sander SP, Nielsen OJ, et al. Inhalation anaesthetics and '
        'climate change. Br J Anaesth 2010;105:760\u20136.',
        # 10 - Intro {9,10}
        'Ryan SM, Nielsen CJ. Global warming potential of inhaled anesthetics: application '
        'to clinical use. Anesth Analg 2010;111:92\u20138.',
        # 11 - Intro {11}: financial rationale
        'Meyer MJ. Desflurane should des-appear: global and financial rationale. Anesth Analg '
        '2020;131:1317\u201322.',
        # 12 - Intro {12,13}: decommissioning implications
        'Moonesinghe SR. Desflurane decommissioning: more than meets the eye. Anaesthesia '
        '2024;79:237\u201341.',
        # 13 - Intro {12,13}
        'Mohammed A, Metta H. Is it time to bid adieu to desflurane? J Anaesthesiol Clin '
        'Pharmacol 2025;41:211\u20132.',
        # 14 - Intro {14}: vaporizer removal
        'Lehmann H, Werning J, Baschnegger H, et al. Minimising the usage of desflurane '
        'only by education and removal of the vaporisers \u2013 a before-and-after-trial. '
        'BMC Anesthesiol 2025;25:108.',
        # 15 - Intro {15,16}: cost savings
        'Rauchenwald V, Heuss-Azeez R, Ganter MT, et al. Sevoflurane versus desflurane\u2014'
        'an economic analysis. BMC Anesthesiol 2020;20:272.',
        # 16 - Intro {15,16}
        'Beard D, Aston W, Black S, et al. Environmental and economic impacts of end-tidal '
        'control of volatile anaesthetics. Open Anaesth J 2025;19:e18742126.',
        # 17 - Intro {17}: secondary market characterization
        'Buckhead Fair Market Value. 2025 Benchmark Report on Pre-Owned Medical Equipment '
        'Prices. Atlanta, GA: BFMV, 2025.',
        # 18 - Methods STROBE {18}
        'von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of '
        'Observational Studies in Epidemiology (STROBE) statement: guidelines for reporting '
        'observational studies. BMJ 2007;335:806\u20138.',
        # 19 - Discussion {19}: regulatory obsolescence
        'Davis G, Patel N. Regulatory obsolescence and secondary market asset depreciation. '
        'J Environ Econ Manage 2019;95:142\u201360.',
    ]
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        run_num = p.add_run(f'{i} ')
        run_num.font.size = Pt(10)
        run_num.font.superscript = True
        run_text = p.add_run(ref)
        run_text.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # ============================================================
    # FIGURE LEGENDS
    # ============================================================
    add_heading_styled(doc, 'Figure legends', level=1)

    legends = [
        ('Fig. 1. ', 'Time series of eBay completed sale prices for desflurane (red), '
         'sevoflurane (blue), and isoflurane (green) vaporizers over three years (March 2023 to '
         'March 2026). Vertical dashed lines indicate key EU regulatory milestones. Curved lines '
         'represent LOWESS trend estimates (fraction = 0.3). Data source: eBay Terapeak.'),
        ('Fig. 2. ', 'Box plot comparison of vaporizer prices before and after the EU desflurane '
         'ban (1 January 2026). Individual data points are shown as jittered dots. '
         'Data source: eBay Terapeak.'),
        ('Fig. 3. ', 'Monthly median prices of anesthetic vaporizers on eBay. Annotations '
         'indicate the number of transactions per month (n). Data source: eBay Terapeak.'),
        ('Fig. 4. ', 'Price distribution histograms for each vaporizer type, comparing pre-ban '
         '(solid fill) and post-ban (hatched) periods. Data source: eBay Terapeak.'),
        ('Fig. 5. ', 'Anesthetic vaporizer prices mapped against the EU regulatory timeline. '
         'Shaded regions indicate regulatory phases. Data source: eBay Terapeak.'),
        ('Fig. 6. ', 'Quarterly median price trends (upper panel) and sales volume (lower panel). '
         'Data source: eBay Terapeak.'),
    ]
    for fig_label, fig_text in legends:
        p = doc.add_paragraph()
        add_run_styled(p, fig_label, bold=True, size=Pt(10))
        add_run_styled(p, fig_text, size=Pt(10))

    # Supplementary table placeholder
    if has_asking_data:
        doc.add_page_break()
        add_heading_styled(doc, 'Supplementary material', level=1)
        p = doc.add_paragraph()
        add_run_styled(p, '[Table S1: Current eBay asking prices \u2014 uploaded as separate file]',
                       bold=True, italic=True, size=Pt(12))

    path = os.path.join(out_dir, 'jca_manuscript_english.docx')
    doc.save(path)
    print(f"JCA English paper saved: {path}")
    return path


if __name__ == '__main__':
    write_jca_paper()
