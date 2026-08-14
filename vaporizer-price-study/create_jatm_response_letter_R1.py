"""Generate a point-by-point response-to-reviewers letter for the JATM R1 minor revision."""
import os
import json
import numpy as np
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(SCRIPT_DIR, 'data')
OUT_DIR = os.path.join(SCRIPT_DIR, 'papers')
os.makedirs(OUT_DIR, exist_ok=True)

# Load sensitivity results
rev_path = os.path.join(DATA_DIR, 'revision_sensitivity.json')
try:
    with open(rev_path, 'r') as f:
        rev = json.load(f)
    boot = rev['bootstrap_pre_post']['Desflurane']
    cits = rev['cits_monthly_medians']
    did = rev['did_transaction_level']
    des_n_post = rev['bootstrap_pre_post']['Desflurane']['n_post']
except Exception:
    boot = {}
    cits = {}
    did = {}
    des_n_post = 'N/A'


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = 'Times New Roman'
    return p


def add_para(doc, text, bold=False, italic=False, size=Pt(12)):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = size
    run.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def fmt_float(v, prec=3, default='N/A'):
    try:
        v = float(v)
        if not np.isnan(v):
            return f'{v:.{prec}f}'
    except Exception:
        pass
    return default


def add_bullet(doc, title, body, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run2 = p.add_run(' ' + body)
    run2.font.size = Pt(12)
    run2.font.name = 'Times New Roman'
    return p


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Header
    add_para(doc, '13 August 2026', size=Pt(12))
    add_para(doc, 'Jianjun Yang, MD, Editor-in-Chief')
    add_para(doc, 'Journal of Anesthesia and Translational Medicine')
    add_para(doc, 'c/o Editorial Office')
    add_para(doc, '')
    add_para(doc, 'Re: Manuscript JATMED-D-26-00064R1 - Minor Revision')
    add_para(doc,
             'Association between the EU desflurane phase-out and secondary market '
             'vaporizer prices: an observational time-series analysis')
    add_para(doc, '')
    add_para(doc, 'Dear Dr. Yang,')
    add_para(doc,
        'Thank you for the opportunity to revise our manuscript. We are grateful to the '
        'Editor and the two reviewers for their thorough and constructive feedback. '
        'We have addressed each of Reviewer 1\u2019s remaining minor concerns and have '
        'further softened the causal language throughout the title, abstract, and body text. '
        'All numerical results remain derived from the updated analysis scripts and are '
        'reproducible from the deposited data.')

    add_heading(doc, 'Reviewer 1', level=1)

    add_bullet(doc, '1. Title and abstract causal language.',
        'We agree that the title and abstract still implied a stronger causal claim than the '
        'observational design supports. The title has been changed to "Association between the '
        'EU desflurane phase-out and secondary market vaporizer prices: an observational '
        'time-series analysis." The abstract now states that desflurane prices "were associated '
        'with" an agent-specific decline, that the findings are "consistent with, but do not '
        'prove" an effect of the EU phase-out, and that they should be interpreted as '
        'hypothesis-generating. Similar "associated with" wording has been applied throughout '
        'the manuscript.')

    add_bullet(doc, '2. Post-restriction sample size and longer follow-up.',
        f'The post-restriction period contains only {des_n_post} desflurane transactions. '
        'We have added an explicit statement in the Study design and data source section noting '
        'that this constraint reflects Terapeak\u2019s three-year historical window, which ends '
        'at the data-extraction date (March 2026). Consequently, the post-restriction '
        'observation window is limited to approximately three months after the 1 January 2026 '
        'effective date; extending follow-up would require either prospective data collection '
        'or a later extraction, neither of which was possible in this retrospective study. '
        'This limitation is also reiterated in the Discussion/Limitations section.')

    add_bullet(doc, '3. Interpretation of the DiD null finding.',
        f'We have expanded the Discussion to explicitly interpret the transaction-level '
        f'difference-in-differences (DiD) result. The desflurane post-restriction coefficient '
        f'was not significant (log-price coefficient={fmt_float(did.get("desflurane_post_coef"), 3)}, '
        f'P={fmt_float(did.get("desflurane_post_p"), 3)}). This null finding should not be '
        f'interpreted as evidence that the regulation had no effect. Instead, the DiD '
        f'parallel-trend assumption was violated (P={fmt_float(did.get("desflurane_trend_p"), 3)} '
        f'for the pre-restriction trend difference), indicating that desflurane prices were already '
        f'diverging from the control agents before 1 January 2026. Because the regulation was '
        f'proposed in 2022 and adopted in 2024, market participants may have anticipated the '
        f'2026 effective date; the DiD violation is consistent with a gradual, '
        f'anticipation-driven decline rather than a discrete post-restriction shock. The DiD '
        f'and CITS results are therefore presented as sensitivity, not causal, analyses.')

    add_bullet(doc, '4. Multiple testing across trend tests.',
        'We have added an explicit multiple-testing acknowledgment in the Statistical analysis '
        'section and in the Discussion. Because trend tests (Spearman, Kendall \u03c4, and '
        'quarterly Spearman) were applied to each of three agent types, the probability of at '
        'least one nominally significant P value arising under the null is inflated. No formal '
        'family-wise correction was applied; P values were interpreted alongside effect magnitude, '
        'direction, and consistency across tests. The isolated isoflurane Spearman P=0.044 is '
        'therefore not interpreted as evidence of a true trend, and the desflurane finding is '
        'supported by consistency across Spearman, Kendall \u03c4, and quarterly median analyses.')

    add_heading(doc, 'Reviewer 2', level=1)

    add_bullet(doc, 'Satisfaction with revisions.',
        'We thank Reviewer 2 for the positive assessment. The major-revision changes '
        '(control-series sensitivity analyses, expanded confounder discussion, softened '
        'causal language, and transparent limitations) have been retained and further refined '
        'in this minor revision. No additional substantive concerns were raised.')

    add_para(doc, '')
    add_para(doc,
        'We believe these revisions have further strengthened the manuscript\u2019s '
        'interpretive caution and transparency. We would be grateful if the revised '
        'manuscript could be considered for publication in Journal of Anesthesia and '
        'Translational Medicine.')
    add_para(doc, '')
    add_para(doc, 'Sincerely,')
    add_para(doc, '[Corresponding author name]')
    add_para(doc, '[Affiliation]')
    add_para(doc, 'Email: [email]')

    out_path = os.path.join(OUT_DIR, 'jatm_response_to_reviewers_R1.docx')
    doc.save(out_path)
    print(f'R1 response letter saved: {out_path}')


if __name__ == '__main__':
    main()
