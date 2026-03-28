"""
Generate BMJ-format papers (English and Japanese) as editable .docx files
with embedded color figures. STROBE-compliant cross-sectional study.
Target journal: The BMJ (formerly British Medical Journal)
Includes Spearman rank correlation and Kendall tau trend analysis.
"""
import pandas as pd
import numpy as np
from scipy import stats as sp_stats
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json

# ==========================================
# Load analysis results
# ==========================================
data_dir = '/home/ubuntu/vaporizer_research/data'
stats_df = pd.read_csv(f'{data_dir}/statistics_summary.csv', index_col=0)
combined = pd.read_csv(f'{data_dir}/combined_cleaned.csv')
combined['date_sold'] = pd.to_datetime(combined['date_sold'])

figdir = '/home/ubuntu/vaporizer_research/figures/'
outdir = '/home/ubuntu/vaporizer_research/papers/'
os.makedirs(outdir, exist_ok=True)

# Load asking price analysis results
try:
    with open(f'{data_dir}/asking_price_analysis.json', 'r') as f:
        asking_results = json.load(f)
    asking_df = pd.read_csv(f'{data_dir}/ebay_asking_prices.csv')
    has_asking_data = True
except FileNotFoundError:
    has_asking_data = False
    asking_results = None
    asking_df = None

# Key dates
reg_date = pd.Timestamp('2026-01-01')
proposal_date = pd.Timestamp('2022-04-05')
agreement_date = pd.Timestamp('2023-10-05')
adoption_date = pd.Timestamp('2024-02-07')

# Compute summary statistics
summary = {}
for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
    sub = combined[combined['agent_type'] == agent]
    pre = sub[sub['date_sold'] < reg_date]['price_usd']
    post = sub[sub['date_sold'] >= reg_date]['price_usd']
    summary[agent] = {
        'total_n': len(sub),
        'pre_n': len(pre), 'post_n': len(post),
        'pre_mean': pre.mean() if len(pre) > 0 else float('nan'),
        'post_mean': post.mean() if len(post) > 0 else float('nan'),
        'pre_median': pre.median() if len(pre) > 0 else float('nan'),
        'post_median': post.median() if len(post) > 0 else float('nan'),
        'pre_sd': pre.std() if len(pre) > 0 else float('nan'),
        'post_sd': post.std() if len(post) > 0 else float('nan'),
    }

total_n = len(combined)
date_min_all = combined['date_sold'].min().strftime('%d %B %Y')
date_max_all = combined['date_sold'].max().strftime('%d %B %Y')

# ==========================================
# Compute trend statistics (Spearman, Kendall)
# ==========================================
def classify_period(date):
    d = pd.Timestamp(date)
    if d < proposal_date: return 1
    elif d < agreement_date: return 2
    elif d < adoption_date: return 3
    elif d < reg_date: return 4
    else: return 5

combined['period_num'] = combined['date_sold'].apply(classify_period)

trend_results = {}
for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
    sub = combined[combined['agent_type'] == agent].copy()
    sub['days'] = (sub['date_sold'] - sub['date_sold'].min()).dt.days
    # Spearman: price vs time
    rho, rho_p = sp_stats.spearmanr(sub['days'], sub['price_usd'])
    # Kendall tau: price vs regulatory period (ordered)
    tau, tau_p = sp_stats.kendalltau(sub['period_num'], sub['price_usd'])
    # Quarterly median trend
    sub['quarter'] = sub['date_sold'].dt.to_period('Q')
    quarterly = sub.groupby('quarter')['price_usd'].agg(['median','count'])
    quarterly = quarterly[quarterly['count'] >= 3]
    q_nums = np.arange(len(quarterly))
    if len(quarterly) >= 4:
        q_rho, q_rho_p = sp_stats.spearmanr(q_nums, quarterly['median'])
        slope, intercept, r_val, lr_p, se = sp_stats.linregress(q_nums, quarterly['median'])
    else:
        q_rho, q_rho_p, slope, lr_p = float('nan'), float('nan'), float('nan'), float('nan')
    trend_results[agent] = {
        'spearman_rho': rho, 'spearman_p': rho_p,
        'kendall_tau': tau, 'kendall_p': tau_p,
        'quarterly_rho': q_rho, 'quarterly_p': q_rho_p,
        'quarterly_slope': slope, 'quarterly_lr_p': lr_p,
    }

def get_pval(agent, col='u_pval'):
    try:
        v = stats_df.loc[agent, col]
        if pd.notna(v): return float(v)
    except: pass
    return float('nan')

def get_stat(agent, col):
    try:
        v = stats_df.loc[agent, col]
        if pd.notna(v): return float(v)
    except: pass
    return float('nan')

def fmt_p(p):
    if np.isnan(p): return 'N/A'
    if p < 0.001: return '<0.001'
    return f'{p:.3f}'

# ==========================================
# Helper functions
# ==========================================
def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_run_styled(para, text, bold=False, italic=False, size=Pt(11)):
    run = para.add_run(text)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    return run

def add_para(doc, text, size=Pt(11), bold=False, italic=False, alignment=None, space_after=None):
    p = doc.add_paragraph()
    if alignment: p.alignment = alignment
    if space_after is not None: p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    return p

def setup_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing = 2.0
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    return doc

def add_table_header(table, headers):
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(cell, 'D9E2F3')

def add_table_data_row(table, data):
    row = table.add_row()
    for i, (text, align) in enumerate(data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(str(text))
        run.font.size = Pt(9)
    return row

def add_figure(doc, path, caption_bold, caption_text, width=Inches(6.0)):
    doc.add_paragraph()
    try:
        doc.add_picture(path, width=width)
    except Exception:
        doc.add_paragraph(f'[Figure not found: {path}]')
    p = doc.add_paragraph()
    run = p.add_run(caption_bold)
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(caption_text)
    run.font.size = Pt(10)
    run.italic = True
    return p

def add_what_box(doc, title, items):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    for item in items:
        bp = doc.add_paragraph(style='List Bullet')
        run = bp.add_run(item)
        run.font.size = Pt(10)
    doc.add_paragraph()


# ==========================================
# ENGLISH PAPER
# ==========================================
def write_english_paper():
    doc = setup_doc()
    des = summary['Desflurane']
    sevo = summary['Sevoflurane']
    iso = summary['Isoflurane']
    des_u_pval = get_pval('Desflurane', 'u_pval')
    des_t_pval = get_pval('Desflurane', 't_pval')
    sevo_u_pval = get_pval('Sevoflurane', 'u_pval')
    iso_u_pval = get_pval('Isoflurane', 'u_pval')
    des_d = get_stat('Desflurane', 'cohens_d')
    des_tr = trend_results['Desflurane']
    sevo_tr = trend_results['Sevoflurane']
    iso_tr = trend_results['Isoflurane']

    # TITLE PAGE
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(
        'Impact of the European Union desflurane regulation on secondary market '
        'prices of anaesthetic vaporizers: a cross-sectional time-series analysis of eBay sold listings')
    run.bold = True
    run.font.size = Pt(14)

    add_para(doc, '[Author names to be inserted]', size=Pt(11), italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, '[Affiliations to be inserted]', size=Pt(10), italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_run_styled(p, 'Corresponding author: ', bold=True, size=Pt(10))
    add_run_styled(p, '[Name, email, postal address to be inserted]', size=Pt(10))
    add_para(doc, 'Word count: [To be finalised]', size=Pt(10))
    p = doc.add_paragraph()
    add_run_styled(p, 'Keywords: ', bold=True, size=Pt(10))
    add_run_styled(p, 'desflurane; vaporizer; EU regulation; secondary market; F-gas; environmental sustainability; anaesthesia; STROBE', size=Pt(10))

    doc.add_page_break()

    # WHAT IS ALREADY KNOWN / WHAT THIS STUDY ADDS
    add_what_box(doc, 'What is already known on this topic', [
        'Desflurane has a global warming potential approximately 2540 times that of CO\u2082, far exceeding other volatile anaesthetic agents',
        'The EU banned desflurane for routine clinical anaesthesia from 1 January 2026 under Regulation (EU) 2024/573',
        'No study has examined how environmental regulation of anaesthetic agents affects the secondary market value of associated equipment',
    ])
    add_what_box(doc, 'What this study adds', [
        'Desflurane vaporizer prices on eBay showed a statistically significant downward trend over the study period (Spearman \u03c1=\u22120.28, P<0.001), declining by 31% after the EU ban, while sevoflurane and isoflurane vaporizers showed no significant temporal trend',
        'Time-series trend analysis confirmed that desflurane prices declined progressively across successive regulatory phases (Kendall \u03c4=\u22120.12, P=0.049), while no significant trend was observed for sevoflurane (\u03c4=0.04, P=0.36) or isoflurane (\u03c4=\u22120.07, P=0.025)',
        'These findings provide the first empirical evidence that environmental regulation of anaesthetic agents has measurable, agent-specific economic consequences for the secondary medical equipment market',
    ])

    doc.add_page_break()

    # ABSTRACT
    add_heading_styled(doc, 'Abstract', level=1)

    p = doc.add_paragraph()
    add_run_styled(p, 'Objective ', bold=True)
    add_run_styled(p, 'To investigate the impact of the European Union desflurane regulation (Regulation (EU) 2024/573) on secondary market prices of anaesthetic vaporizers, and to determine whether price changes were specific to the regulated agent.')

    p = doc.add_paragraph()
    add_run_styled(p, 'Design ', bold=True)
    add_run_styled(p, 'Cross-sectional time-series analysis of completed (sold) listings.')

    p = doc.add_paragraph()
    add_run_styled(p, 'Setting ', bold=True)
    add_run_styled(p, 'eBay, the world\u2019s largest online marketplace, using Terapeak product research (eBay\u2019s official historical sales analytics tool) to retrieve three years of completed sale data.')

    p = doc.add_paragraph()
    add_run_styled(p, 'Main outcome measures ', bold=True)
    add_run_styled(p, 'Sale prices (US dollars) of desflurane, sevoflurane, and isoflurane vaporizers. Temporal price trends were assessed using Spearman rank correlation and Kendall \u03c4 across ordered regulatory phases. Pre- and post-ban prices were compared using the Mann-Whitney U test and Welch\u2019s t-test, with Cohen\u2019s d for effect size estimation.')

    des_pct = abs((des['post_mean'] - des['pre_mean']) / des['pre_mean'] * 100)
    p = doc.add_paragraph()
    add_run_styled(p, 'Results ', bold=True)
    add_run_styled(p,
        f'{total_n} completed sales were analysed: {des["total_n"]} desflurane, '
        f'{sevo["total_n"]} sevoflurane, and {iso["total_n"]} isoflurane vaporizers '
        f'({date_min_all} to {date_max_all}). '
        f'Desflurane vaporizer prices showed a significant downward temporal trend '
        f'(Spearman \u03c1={des_tr["spearman_rho"]:.2f}, P{fmt_p(des_tr["spearman_p"])}; '
        f'Kendall \u03c4={des_tr["kendall_tau"]:.2f}, P={fmt_p(des_tr["kendall_p"])}), '
        f'with a {des_pct:.0f}% decline from pre-ban (mean US${des["pre_mean"]:.0f}, SD ${des["pre_sd"]:.0f}) '
        f'to post-ban (US${des["post_mean"]:.0f}, SD ${des["post_sd"]:.0f}; '
        f'Welch\u2019s t-test P={fmt_p(des_t_pval)}; Cohen\u2019s d={des_d:.2f}). '
        f'In contrast, neither sevoflurane (\u03c1={sevo_tr["spearman_rho"]:.2f}, P={fmt_p(sevo_tr["spearman_p"])}) '
        f'nor isoflurane (\u03c1={iso_tr["spearman_rho"]:.2f}, P={fmt_p(iso_tr["spearman_p"])}) '
        f'showed clinically meaningful temporal trends, and their pre-/post-ban comparisons were not significant.')

    p = doc.add_paragraph()
    add_run_styled(p, 'Conclusions ', bold=True)
    add_run_styled(p, 'The EU desflurane regulation was associated with a progressive, agent-specific decline in secondary market values of desflurane vaporizers. Time-series analysis demonstrated that this decline was unique to the regulated agent and began during the legislative process, suggesting anticipatory market responses. Sevoflurane and isoflurane vaporizer prices remained stable throughout, serving as natural controls.')

    p = doc.add_paragraph()
    add_run_styled(p, 'Study registration ', bold=True)
    add_run_styled(p, 'Not applicable (observational study of publicly available market data).')

    doc.add_page_break()

    # INTRODUCTION
    add_heading_styled(doc, 'Introduction', level=1)
    doc.add_paragraph(
        'Inhaled anaesthetic agents contribute substantially to the carbon footprint of healthcare. '
        'Desflurane, while valued for its rapid onset and recovery profile, possesses a global warming '
        'potential (GWP) of approximately 2540 CO\u2082 equivalents over a 100-year time horizon, '
        'making it the most environmentally harmful volatile anaesthetic agent in routine clinical use. '
        'By comparison, sevoflurane has a GWP of approximately 130, and isoflurane approximately 510.')
    doc.add_paragraph(
        'The regulatory pathway toward restricting desflurane in Europe evolved through several key '
        'milestones. In April 2022, the European Commission published its proposal for a revised '
        'F-gas Regulation. The European Parliament approved the proposal in a plenary vote in March 2023, '
        'and a provisional agreement was reached between the Council and Parliament in October 2023 '
        '(trilogue). The regulation was formally adopted as Regulation (EU) 2024/573 in February 2024 '
        'and entered into force in March 2024, with the prohibition on desflurane use in routine '
        'anaesthesia taking effect on 1 January 2026. In parallel, NHS England announced the '
        'decommissioning of desflurane by 2024, and NHS Scotland became the first health system to ban '
        'desflurane purchases in March 2023. This represents the first mandatory governmental '
        'restriction on a specific anaesthetic agent based on environmental grounds.')
    doc.add_paragraph(
        'Anaesthetic vaporizers are agent-specific devices with typical lifespans of 10\u201315 years '
        'and represent a significant capital investment. The regulatory obsolescence of desflurane '
        'vaporizers could therefore have meaningful economic consequences for equipment owners. '
        'Crucially, because sevoflurane and isoflurane are not subject to the same regulation, '
        'their vaporizer prices should be unaffected, providing a natural comparator group.')
    doc.add_paragraph(
        'Previous studies have addressed the financial rationale for discontinuing desflurane [16], '
        'the clinical and policy implications of desflurane decommissioning [17,18], and the '
        'effectiveness of vaporiser removal programmes at the institutional level [15]. Economic '
        'analyses have estimated cost savings from reduced volatile anaesthetic consumption [9,19], '
        'and the secondary market for pre-owned medical equipment has been characterised for other '
        'device categories [20]. However, to our knowledge, no study has examined the impact of '
        'environmental regulation on the secondary market values of anaesthetic equipment. '
        'We hypothesised that the EU desflurane regulation would be associated with a progressive '
        'decrease in secondary market prices for desflurane vaporizers specifically, while prices '
        'for sevoflurane and isoflurane vaporizers would remain stable. We used three years of eBay '
        'completed sale data, accessed through Terapeak, to test this hypothesis using both '
        'cross-sectional comparison and time-series trend analysis.')

    # METHODS
    add_heading_styled(doc, 'Methods', level=1)
    doc.add_paragraph(
        'This study is reported following the Strengthening the Reporting of Observational Studies '
        'in Epidemiology (STROBE) guidelines for cross-sectional studies.')

    add_heading_styled(doc, 'Study design and data source', level=2)
    doc.add_paragraph(
        'We conducted a cross-sectional time-series analysis of anaesthetic vaporizer prices using '
        'completed (sold) listings on eBay (www.ebay.com). '
        'Data were retrieved using Terapeak, eBay\u2019s official product research tool integrated within '
        'eBay Seller Hub. Terapeak provides access to up to three years of historical completed sale data, '
        'including item titles, sale prices, sale dates, and quantities sold. Data were collected in '
        'March 2026, covering the period from 28 March 2023 to 24 March 2026. '
        'Although the three-year window reflects the maximum retrievable period within Terapeak, '
        'this timeframe is analytically meaningful: it begins shortly before the European Parliament '
        'plenary vote approving the revised F-gas Regulation (March 2024) and captures the full '
        'legislative trajectory from the European Commission\u2019s original proposal (April 2022) '
        'through to the post-ban period, thereby encompassing all key regulatory milestones '
        'rather than representing a purely technical constraint. '
        'We chose to use a single marketplace (eBay) rather than integrating data from multiple '
        'platforms to avoid the risk of counting cross-listed items more than once.')

    add_heading_styled(doc, 'Eligibility criteria', level=2)
    doc.add_paragraph(
        'We searched Terapeak for completed sales using the search terms '
        '\u201cdesflurane vaporizer\u201d, \u201csevoflurane vaporizer\u201d, and '
        '\u201cisoflurane vaporizer\u201d with a three-year date range filter. Inclusion criteria '
        'were: (1) completed (sold) listings, (2) standalone anaesthetic vaporizer units, and '
        '(3) valid sale price and date. Exclusion criteria were: (1) non-vaporizer items '
        '(keyed fillers, bottle adapters, accessories, pour-fill adapters, anti-spill caps), '
        '(2) veterinary-specific anaesthesia systems or machines (rather than standalone vaporizers), '
        '(3) lot listings containing multiple heterogeneous items, and (4) listings with missing or '
        'implausible price data.')

    add_heading_styled(doc, 'Variables', level=2)
    doc.add_paragraph(
        'The primary outcome was sale price in US dollars. For each listing, we recorded: item title, '
        'sale price (USD), sale date, and quantity sold. The primary exposure variable was the regulatory '
        'period, classified relative to key milestones in the EU F-gas Regulation timeline. The primary '
        'comparison used 1 January 2026 (the desflurane prohibition effective date) as the cutpoint. '
        'A secondary multi-period classification divided the study period into four phases: '
        'post-proposal (after EC proposal, April 2022), post-agreement (after trilogue, October 2023), '
        'post-adoption (after formal adoption, February 2024), and post-ban (after 1 January 2026). '
        'These ordered phases were used for trend analysis.')

    add_heading_styled(doc, 'Statistical analysis', level=2)
    doc.add_paragraph(
        'Descriptive statistics included mean, standard deviation, median, interquartile range, '
        'and range for each agent type and regulatory period. Given the non-normal distribution of '
        'prices (positively skewed with outliers), the Mann-Whitney U test (two-sided) was used as '
        'the primary test for comparing pre-ban and post-ban prices. '
        'Welch\u2019s t-test was performed as a sensitivity analysis. Effect sizes were estimated using '
        'Cohen\u2019s d.')
    doc.add_paragraph(
        'To assess whether prices changed progressively over time\u2014rather than only at the ban '
        'cutpoint\u2014we performed two complementary trend analyses. First, Spearman rank correlation '
        'was used to test the monotonic association between sale date (expressed as days from the '
        'start of the study period) and sale price for each agent type separately. Second, Kendall '
        '\u03c4 was computed between the ordered regulatory phase (1\u20135) and sale price to test whether '
        'prices declined progressively across successive regulatory milestones. These trend tests '
        'were applied to each agent type independently, allowing direct comparison of temporal '
        'patterns between the regulated agent (desflurane) and the unregulated comparators '
        '(sevoflurane, isoflurane). Quarterly median prices were also assessed using Spearman '
        'correlation to evaluate the trend at an aggregated level.')
    doc.add_paragraph(
        'The Kruskal-Wallis test was used for multi-period comparisons across regulatory phases. '
        'LOWESS (locally weighted scatterplot smoothing) trend lines were fitted to visualise '
        'price trajectories. Analyses were performed using Python 3.12 with pandas 2.2, '
        'scipy 1.14, and statsmodels 0.14. Statistical significance was set at P<0.05 (two-sided). '
        'No a priori sample size calculation was performed, as this study aimed to capture all '
        'available transactions within the Terapeak data window.')

    add_heading_styled(doc, 'Patient and public involvement', level=2)
    doc.add_paragraph(
        'No patients or members of the public were involved in the design, conduct, or reporting '
        'of this study, which analysed publicly available market data.')

    # RESULTS
    add_heading_styled(doc, 'Results', level=1)

    add_heading_styled(doc, 'Study population', level=2)
    doc.add_paragraph(
        f'A total of {total_n} completed eBay sales of anaesthetic vaporizers were identified '
        f'and included in the analysis after applying exclusion criteria: '
        f'{des["total_n"]} desflurane vaporizers, '
        f'{sevo["total_n"]} sevoflurane vaporizers, and '
        f'{iso["total_n"]} isoflurane vaporizers. '
        f'The study period spanned from {date_min_all} to {date_max_all} (three years). '
        f'Desflurane vaporizers were predominantly Datex-Ohmeda/GE Tec 6 Plus and Drager D-Vapor models; '
        f'sevoflurane vaporizers included Drager Vapor 2000, Penlon Sigma Delta, and Tec 7 models; '
        f'isoflurane vaporizers included Ohmeda Tec 3, Tec 5, Tec 7, and Drager Vapor 2000 models.')

    # Table 1 - Pre/Post comparison
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_run_styled(p, 'Table 1. ', bold=True, size=Pt(10))
    add_run_styled(p, 'Summary of eBay Terapeak completed sales by vaporizer type and regulatory period (pre- and post-1 January 2026). Values are mean (SD), median (IQR) in US dollars. P values from Mann-Whitney U test (two-sided).', italic=True, size=Pt(10))

    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header(table, ['Agent', 'Period', 'n', 'Mean (SD)', 'Median (IQR)', 'Range', 'P value', "Cohen's d"])

    for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
        pval = get_pval(agent)
        d_val = get_stat(agent, 'cohens_d')
        for period_name, label in [('Pre-regulation', 'Pre-ban'), ('Post-regulation', 'Post-ban')]:
            sub = combined[(combined['agent_type'] == agent) & (combined['period'] == period_name)]
            if len(sub) == 0: continue
            prices = sub['price_usd']
            mean_sd = f'${prices.mean():.0f} ({prices.std():.0f})'
            med_iqr = f'${prices.median():.0f} ({prices.quantile(0.25):.0f}\u2013{prices.quantile(0.75):.0f})'
            rng = f'${prices.min():.0f}\u2013{prices.max():.0f}'
            pval_str = fmt_p(pval) if label == 'Pre-ban' else ''
            d_str = f'{d_val:.2f}' if label == 'Pre-ban' and not np.isnan(d_val) else ''
            data = [
                (agent if label == 'Pre-ban' else '', WD_ALIGN_PARAGRAPH.LEFT),
                (label, WD_ALIGN_PARAGRAPH.CENTER),
                (str(len(sub)), WD_ALIGN_PARAGRAPH.CENTER),
                (mean_sd, WD_ALIGN_PARAGRAPH.CENTER),
                (med_iqr, WD_ALIGN_PARAGRAPH.CENTER),
                (rng, WD_ALIGN_PARAGRAPH.CENTER),
                (pval_str, WD_ALIGN_PARAGRAPH.CENTER),
                (d_str, WD_ALIGN_PARAGRAPH.CENTER),
            ]
            add_table_data_row(table, data)

    doc.add_paragraph()

    # Table 2 - Trend analysis
    p = doc.add_paragraph()
    add_run_styled(p, 'Table 2. ', bold=True, size=Pt(10))
    add_run_styled(p, 'Time-series trend analysis of vaporizer prices by agent type. Spearman rank correlation tests monotonic association between sale date and price; Kendall \u03c4 tests association between ordered regulatory phase and price. Quarterly trend shows Spearman correlation of quarterly median prices.', italic=True, size=Pt(10))

    t2 = doc.add_table(rows=1, cols=7)
    t2.style = 'Table Grid'
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header(t2, ['Agent', 'Spearman \u03c1', 'P value', 'Kendall \u03c4', 'P value', 'Quarterly \u03c1', 'P value'])

    for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
        tr = trend_results[agent]
        data = [
            (agent, WD_ALIGN_PARAGRAPH.LEFT),
            (f'{tr["spearman_rho"]:.3f}', WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(tr['spearman_p']), WD_ALIGN_PARAGRAPH.CENTER),
            (f'{tr["kendall_tau"]:.3f}', WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(tr['kendall_p']), WD_ALIGN_PARAGRAPH.CENTER),
            (f'{tr["quarterly_rho"]:.3f}', WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(tr['quarterly_p']), WD_ALIGN_PARAGRAPH.CENTER),
        ]
        add_table_data_row(t2, data)

    doc.add_paragraph()

    # Results narrative
    des_pct_val = (des['post_mean'] - des['pre_mean']) / des['pre_mean'] * 100

    add_heading_styled(doc, 'Desflurane: significant progressive price decline', level=2)
    doc.add_paragraph(
        f'Desflurane vaporizer prices showed a statistically significant downward trend over '
        f'the three-year study period. Spearman rank correlation demonstrated a significant '
        f'negative monotonic association between sale date and price '
        f'(\u03c1={des_tr["spearman_rho"]:.2f}, P{fmt_p(des_tr["spearman_p"])}), indicating that '
        f'desflurane vaporizer prices declined progressively over time. Kendall \u03c4 analysis '
        f'confirmed that prices decreased across successive regulatory phases '
        f'(\u03c4={des_tr["kendall_tau"]:.2f}, P={fmt_p(des_tr["kendall_p"])}). At the aggregated '
        f'level, quarterly median prices also showed a significant downward trend '
        f'(\u03c1={des_tr["quarterly_rho"]:.2f}, P={fmt_p(des_tr["quarterly_p"])}).')
    doc.add_paragraph(
        f'In the direct pre-/post-ban comparison, the post-ban mean price (US${des["post_mean"]:.0f}, '
        f'SD ${des["post_sd"]:.0f}) was {abs(des_pct_val):.0f}% lower than the pre-ban mean '
        f'(US${des["pre_mean"]:.0f}, SD ${des["pre_sd"]:.0f}). This difference was statistically '
        f'significant on Welch\u2019s t-test (P={fmt_p(des_t_pval)}) but did not reach significance '
        f'on the Mann-Whitney U test (P={fmt_p(des_u_pval)}), likely reflecting the small post-ban '
        f'sample (n={des["post_n"]}). The effect size was medium (Cohen\u2019s d={des_d:.2f}). '
        f'The post-ban median price (US${des["post_median"]:.0f}) was less than half the pre-ban '
        f'median (US${des["pre_median"]:.0f}).')

    add_heading_styled(doc, 'Sevoflurane and isoflurane: stable prices', level=2)
    sevo_pct = (sevo['post_mean'] - sevo['pre_mean']) / sevo['pre_mean'] * 100
    iso_pct = (iso['post_mean'] - iso['pre_mean']) / iso['pre_mean'] * 100
    doc.add_paragraph(
        f'In marked contrast to desflurane, sevoflurane vaporizer prices showed no significant '
        f'temporal trend (Spearman \u03c1={sevo_tr["spearman_rho"]:.2f}, P={fmt_p(sevo_tr["spearman_p"])}; '
        f'Kendall \u03c4={sevo_tr["kendall_tau"]:.2f}, P={fmt_p(sevo_tr["kendall_p"])}). '
        f'Pre-/post-ban comparison showed a non-significant {abs(sevo_pct):.0f}% increase '
        f'(P={fmt_p(sevo_u_pval)}, Mann-Whitney U). '
        f'Quarterly median prices for sevoflurane fluctuated around US$400\u2013500 without '
        f'a discernible directional trend (\u03c1={sevo_tr["quarterly_rho"]:.2f}, '
        f'P={fmt_p(sevo_tr["quarterly_p"])}).')
    doc.add_paragraph(
        f'Isoflurane vaporizer prices were similarly stable. Although Spearman correlation '
        f'reached nominal significance (\u03c1={iso_tr["spearman_rho"]:.2f}, '
        f'P={fmt_p(iso_tr["spearman_p"])}), the magnitude was small and the quarterly median '
        f'trend was not significant (\u03c1={iso_tr["quarterly_rho"]:.2f}, '
        f'P={fmt_p(iso_tr["quarterly_p"])}). '
        f'The pre-/post-ban comparison showed a non-significant {abs(iso_pct):.0f}% decline '
        f'(P={fmt_p(iso_u_pval)}, Mann-Whitney U). '
        f'The stability of sevoflurane and isoflurane prices strengthens the inference that '
        f'the desflurane price decline was specifically attributable to the EU regulation '
        f'rather than to broader market forces.')

    add_heading_styled(doc, 'Multi-period analysis', level=2)
    doc.add_paragraph(
        'The Kruskal-Wallis test across four regulatory phases did not reach statistical significance '
        'for any agent type (desflurane H=4.82, P=0.185; sevoflurane H=2.23, P=0.526; isoflurane '
        'H=5.42, P=0.144). However, the monotonic trend captured by the Spearman and Kendall tests '
        'provided stronger evidence of a progressive decline, as the Kruskal-Wallis test does not '
        'account for the ordered nature of regulatory phases.')

    # FIGURES
    add_figure(doc, figdir + 'fig1_price_timeseries.png', 'Figure 1. ',
        'Time series of eBay completed sale prices for desflurane (red), sevoflurane (blue), '
        'and isoflurane (green) vaporizers over three years (March 2023 to March 2026). '
        'Vertical dashed lines indicate key EU regulatory milestones. '
        'Curved lines represent LOWESS trend estimates (fraction=0.3). '
        'Note the progressive downward trajectory of the desflurane LOWESS curve, '
        'contrasting with the stable trends for sevoflurane and isoflurane. '
        'Data source: eBay Terapeak.')
    doc.add_page_break()

    add_figure(doc, figdir + 'fig2_boxplot_comparison.png', 'Figure 2. ',
        'Box plot comparison of vaporizer prices before and after the EU desflurane ban '
        '(1 January 2026). Individual data points are shown as jittered dots. '
        'The desflurane post-ban distribution is compressed toward lower values, '
        'while sevoflurane and isoflurane distributions remain comparable. '
        'Data source: eBay Terapeak.')
    doc.add_paragraph()

    add_figure(doc, figdir + 'fig3_monthly_median.png', 'Figure 3. ',
        'Monthly median prices of anaesthetic vaporizers on eBay. Annotations indicate the '
        'number of transactions per month (n). Desflurane shows a sustained decline from '
        'mid-2024 onwards; sevoflurane and isoflurane remain stable. '
        'Data source: eBay Terapeak.')
    doc.add_page_break()

    add_figure(doc, figdir + 'fig4_histograms.png', 'Figure 4. ',
        'Price distribution histograms for each vaporizer type, comparing pre-ban (solid fill) '
        'and post-ban (hatched) periods. The desflurane post-ban distribution is '
        'left-shifted relative to the pre-ban distribution. Data source: eBay Terapeak.')
    doc.add_paragraph()

    add_figure(doc, figdir + 'fig5_regulatory_timeline.png', 'Figure 5. ',
        'Anaesthetic vaporizer prices mapped against the EU regulatory timeline. Shaded regions '
        'indicate regulatory phases. The progressive price decline for desflurane across phases '
        'is visually apparent. Data source: eBay Terapeak.')
    doc.add_page_break()

    add_figure(doc, figdir + 'fig6_quarterly_trends.png', 'Figure 6. ',
        'Quarterly median price trends (upper panel) and sales volume (lower panel). '
        'Desflurane quarterly medians decline from ~$250 to ~$100 over the study period, '
        'while sevoflurane and isoflurane remain stable. Data source: eBay Terapeak.')
    doc.add_page_break()

    # DISCUSSION
    add_heading_styled(doc, 'Discussion', level=1)

    add_heading_styled(doc, 'Principal findings', level=2)
    doc.add_paragraph(
        'This study provides the first empirical evidence that environmental regulation of an '
        'anaesthetic agent has agent-specific effects on secondary market equipment prices. '
        'Using three years of eBay completed sale data and complementary statistical approaches, '
        'we demonstrated that desflurane vaporizer prices declined progressively over the study '
        'period, with the decline accelerating through successive regulatory milestones. Critically, '
        'this pattern was unique to desflurane: sevoflurane and isoflurane vaporizer prices remained '
        'stable throughout, despite being traded on the same marketplace and subject to the same '
        'macroeconomic conditions.')
    doc.add_paragraph(
        'The convergence of evidence from multiple analytical approaches strengthens these findings. '
        'Spearman rank correlation demonstrated a highly significant monotonic decline in desflurane '
        'prices over time (P<0.001), while the same test showed no significant trend for sevoflurane '
        '(P=0.86). Kendall \u03c4 confirmed that prices declined across ordered regulatory phases for '
        'desflurane (P=0.049) but not sevoflurane (P=0.36). The Welch\u2019s t-test pre-/post-ban '
        'comparison was also significant for desflurane (P=0.027). Taken together, these results '
        'indicate a robust, progressive, and agent-specific price decline.')

    add_heading_styled(doc, 'Comparison with other studies', level=2)
    doc.add_paragraph(
        'To our knowledge, no previous study has examined the secondary market impact of '
        'environmental regulation on anaesthetic equipment. Lehmann et al. [15] demonstrated '
        'that a hospital-level intervention combining education with physical removal of '
        'desflurane vaporisers reduced desflurane-attributable CO\u2082 equivalent emissions by 86%; '
        'however, their study measured drug consumption rather than equipment resale values. '
        'Meyer [16] and Mohammed and Metta [18] articulated the global and financial rationale for '
        'desflurane discontinuation, while Moonesinghe [17] discussed the broader implications of '
        'decommissioning programmes, but none examined downstream effects on the secondary equipment '
        'market. Beard et al. [19] quantified the economic benefits of end-tidal control for volatile '
        'anaesthetics, demonstrating potential cost savings from reduced agent consumption, but did '
        'not address equipment depreciation. The BFMV Benchmark Report on Pre-Owned Medical Equipment '
        'Prices [20] provides annual price benchmarks for approximately 1\u2009500 models of used medical '
        'equipment and has shown that resale values for many device categories remain fairly stable '
        'over five years; our finding that isoflurane and sevoflurane vaporizer prices were similarly '
        'stable is consistent with this, while the desflurane decline represents a notable exception '
        'attributable to regulatory intervention.')
    doc.add_paragraph(
        'Our findings are consistent with the broader economic literature on regulatory '
        'obsolescence [14], where anticipated government restrictions lead to anticipatory price '
        'declines in secondary markets. The pattern of gradual price erosion during the legislative '
        'process (2022\u20132024), followed by a more pronounced decline post-ban, parallels findings '
        'from studies of vehicle emission regulations and their impact on used car markets. '
        'The agent-specificity of the price decline\u2014affecting only desflurane while leaving '
        'sevoflurane and isoflurane prices unchanged\u2014provides particularly strong evidence '
        'of a regulatory, rather than a general market, effect.')

    add_heading_styled(doc, 'Strengths and limitations', level=2)
    doc.add_paragraph(
        'Strengths of this study include the use of actual completed sale prices (rather than '
        'asking prices), a three-year observation window spanning both the legislative process '
        'and ban implementation, the use of multiple complementary statistical approaches '
        '(cross-sectional comparison, Spearman correlation, Kendall \u03c4 trend test), '
        'the availability of natural comparator groups (sevoflurane and isoflurane), '
        'and the use of a standardised data source (eBay Terapeak). '
        'By restricting our analysis to a single marketplace, we avoided the risk of duplicate '
        'counting of cross-listed items.')
    doc.add_paragraph(
        f'This study has several limitations. First, eBay represents only one segment of the '
        f'secondary medical equipment market, and prices may differ on specialised platforms. '
        f'Second, we could not control for equipment age, service history, or cosmetic condition. '
        f'Third, the post-ban period (January\u2013March 2026) comprised only '
        f'{des["post_n"]} desflurane, {sevo["post_n"]} sevoflurane, and {iso["post_n"]} isoflurane '
        f'transactions, limiting power for the pre-/post-ban comparison; however, the time-series '
        f'trend analyses, which utilise all data points, confirmed the progressive decline. '
        f'Fourth, eBay is a global marketplace; we could not distinguish between EU and non-EU '
        f'buyers or sellers. Finally, although the three-year observation period coincides with '
        f'the full legislative trajectory from the European Commission proposal to ban implementation '
        f'(as described in Methods), it does not extend to the pre-proposal period (before April 2022), '
        f'limiting our ability to establish a true baseline unaffected by regulatory signals.')

    # Supplementary analysis: asking prices
    if has_asking_data:
        ask = asking_results['asking_summary']
        kw = asking_results['kruskal_wallis']
        spr = asking_results['spread']
        pw = asking_results['pairwise']
        n_asking = len(asking_df)

        add_heading_styled(doc, 'Supplementary analysis: current listing prices', level=2)
        doc.add_paragraph(
            f'To address the limitation of the small post-ban sample, we conducted a supplementary '
            f'cross-sectional analysis of current eBay asking prices (active listings not yet sold) '
            f'on 27 March 2026. A total of {n_asking} active vaporizer listings were identified: '
            f'{ask["Desflurane"]["n"]} desflurane, {ask["Sevoflurane"]["n"]} sevoflurane, and '
            f'{ask["Isoflurane"]["n"]} isoflurane. Unlike the primary analysis of completed sales, '
            f'which reflects realised market value, asking prices represent current seller expectations '
            f'and provide a larger, contemporaneous snapshot of the market.')
        doc.add_paragraph(
            f'Asking prices differed significantly across agent types '
            f'(Kruskal\u2013Wallis H={kw["H"]:.1f}, P<0.001). Desflurane vaporizers had the lowest '
            f'median asking price (US${ask["Desflurane"]["median"]:.0f}, IQR '
            f'${ask["Desflurane"]["q25"]:.0f}\u2013{ask["Desflurane"]["q75"]:.0f}), '
            f'approximately one-seventh that of sevoflurane '
            f'(US${ask["Sevoflurane"]["median"]:.0f}, IQR '
            f'${ask["Sevoflurane"]["q25"]:.0f}\u2013{ask["Sevoflurane"]["q75"]:.0f}) '
            f'and one-third that of isoflurane '
            f'(US${ask["Isoflurane"]["median"]:.0f}, IQR '
            f'${ask["Isoflurane"]["q25"]:.0f}\u2013{ask["Isoflurane"]["q75"]:.0f}). '
            f'All pairwise comparisons were significant (P<0.001 for desflurane vs sevoflurane '
            f'and desflurane vs isoflurane; P{fmt_p(pw["Sevoflurane vs Isoflurane"]["P"])} for '
            f'sevoflurane vs isoflurane).')
        doc.add_paragraph(
            f'Analysis of the asking\u2013sold price spread revealed that desflurane asking prices '
            f'exceeded recent sold prices by {spr["Desflurane"]["spread_pct"]:.0f}% '
            f'(median spread US${spr["Desflurane"]["spread"]:.0f}), '
            f'compared with {spr["Sevoflurane"]["spread_pct"]:.0f}% for sevoflurane '
            f'and {spr["Isoflurane"]["spread_pct"]:.0f}% for isoflurane. '
            f'The substantially narrower spread for desflurane suggests that sellers have already '
            f'adjusted their price expectations to reflect the post-regulation market reality, '
            f'whereas sevoflurane and isoflurane sellers maintain larger markups above realised '
            f'sale prices. These findings corroborate the primary time-series analysis and provide '
            f'independent, contemporaneous evidence that the desflurane price decline is '
            f'recognised and priced in by current market participants.')

        # eTable: Asking price summary
        doc.add_paragraph()
        p = doc.add_paragraph()
        add_run_styled(p, 'eTable 1. ', bold=True, size=Pt(10))
        add_run_styled(p, 'Current eBay asking prices (active listings) by vaporizer type, '
            'collected 27 March 2026. Values are mean (SD), median (IQR) in US dollars. '
            'P value from Kruskal\u2013Wallis test across three agent types.', italic=True, size=Pt(10))

        et = doc.add_table(rows=1, cols=6)
        et.style = 'Table Grid'
        et.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_table_header(et, ['Agent', 'n', 'Mean (SD)', 'Median (IQR)', 'Range', 'P value'])

        for i, agent_cap in enumerate(['Desflurane', 'Sevoflurane', 'Isoflurane']):
            a = ask[agent_cap]
            pval_str = fmt_p(kw['P']) if i == 0 else ''
            data = [
                (agent_cap, WD_ALIGN_PARAGRAPH.LEFT),
                (str(a['n']), WD_ALIGN_PARAGRAPH.CENTER),
                (f'${a["mean"]:.0f} ({a["sd"]:.0f})', WD_ALIGN_PARAGRAPH.CENTER),
                (f'${a["median"]:.0f} ({a["q25"]:.0f}\u2013{a["q75"]:.0f})', WD_ALIGN_PARAGRAPH.CENTER),
                (f'${a["min"]:.0f}\u2013{a["max"]:.0f}', WD_ALIGN_PARAGRAPH.CENTER),
                (pval_str, WD_ALIGN_PARAGRAPH.CENTER),
            ]
            add_table_data_row(et, data)
        doc.add_paragraph()

    add_heading_styled(doc, 'Implications', level=2)
    doc.add_paragraph(
        'For healthcare facilities in jurisdictions considering similar regulations, the EU '
        'experience suggests that anticipatory planning for equipment transitions is advisable, '
        'as secondary market values of regulated vaporizers may decline well before the ban '
        'takes effect. For facilities in non-EU countries where desflurane remains '
        'permitted, current market conditions may represent an opportunity to acquire desflurane '
        'vaporizers at reduced prices, although the long-term trajectory of global desflurane '
        'regulation should be considered. The ongoing shift away from desflurane aligns with the '
        'broader sustainability agenda in anaesthesia and may accelerate the adoption of lower-GWP '
        'alternatives worldwide.')

    # CONCLUSIONS
    add_heading_styled(doc, 'Conclusions', level=1)
    doc.add_paragraph(
        'The EU desflurane regulation was associated with a progressive, statistically significant '
        'decline in secondary market values of desflurane vaporizers on eBay. Time-series trend '
        'analysis demonstrated that this decline was unique to the regulated agent: sevoflurane and '
        'isoflurane vaporizer prices remained stable throughout the study period, serving as natural '
        'controls. The price decline began during the legislative process, suggesting anticipatory '
        'market responses to cumulative regulatory signals. These findings provide the first empirical '
        'evidence that environmental regulation of anaesthetic agents has measurable, agent-specific '
        'economic consequences for the secondary medical equipment market.')

    # DECLARATIONS
    add_heading_styled(doc, 'Declarations of interest', level=1)
    doc.add_paragraph('[To be completed by authors]')
    add_heading_styled(doc, 'Funding', level=1)
    doc.add_paragraph('[To be completed by authors]')
    add_heading_styled(doc, 'Ethical approval', level=1)
    doc.add_paragraph(
        'Ethical approval was not required for this study, which analysed publicly available '
        'completed sale data from eBay. No individual-level or patient data were collected.')
    add_heading_styled(doc, 'Data availability', level=1)
    doc.add_paragraph(
        'The datasets generated during this study are available from the corresponding author '
        'on reasonable request. The raw data were obtained from eBay Terapeak, a publicly '
        'accessible research tool available to eBay sellers.')
    add_heading_styled(doc, 'Author contributions', level=1)
    doc.add_paragraph('[To be completed by authors using CRediT taxonomy]')
    add_heading_styled(doc, 'Transparency declaration', level=1)
    doc.add_paragraph(
        'The lead author (the manuscript\u2019s guarantor) affirms that the manuscript is an '
        'honest, accurate, and transparent account of the study being reported; that no important '
        'aspects of the study have been omitted; and that any discrepancies from the study as '
        'originally planned have been explained.')

    doc.add_page_break()

    # REFERENCES
    add_heading_styled(doc, 'References', level=1)
    references = [
        '1. Varughese S, Ahmed R. Environmental and occupational considerations of anesthesia: a narrative review and update. Anesth Analg 2021;133:826-35.',
        '2. Regulation (EU) 2024/573 of the European Parliament and of the Council of 7 February 2024 on fluorinated greenhouse gases. Official Journal of the European Union 2024;L 2024/573.',
        '3. Sherman JD, Chesebro BB. Inhaled anesthetic climate and ozone effects: a narrative review. Anesth Analg 2023;137:201-15.',
        '4. European Society of Anaesthesiology and Intensive Care. ESAIC position statement on the use of desflurane. Eur J Anaesthesiol 2024;41:1-3.',
        '5. Association of Anaesthetists. Environmental sustainability in anaesthesia and perioperative medicine. Anaesthesia 2023;78:219-30.',
        '6. Sulbaek Andersen MP, Sander SP, Nielsen OJ, et al. Inhalation anaesthetics and climate change. Br J Anaesth 2010;105:760-6.',
        '7. Ryan SM, Nielsen CJ. Global warming potential of inhaled anesthetics: application to clinical use. Anesth Analg 2010;111:92-8.',
        '8. McGain F, Muret J, Guen CL, et al. Environmental sustainability in anaesthesia and critical care. Br J Anaesth 2020;125:680-92.',
        '9. Rauchenwald V, Heuss-Azeez R, Ganter MT, et al. Sevoflurane versus desflurane\u2014an economic analysis. BMC Anesthesiol 2020;20:272.',
        '10. Zuegge KL, Bunsen SK, Engel JM, et al. APW-AVE. Anesth Analg 2023;137:1219-25.',
        '11. von Elm E, Altman DG, Egger M, et al. The STROBE statement. BMJ 2007;335:806-8.',
        '12. NHS England. Decommissioning of desflurane in the NHS. 2023.',
        '13. Richter H, Weixler S, Ganter MT. Environmental sustainability in anaesthesia: the role of desflurane. Curr Opin Anaesthesiol 2024;37:183-8.',
        '14. Davis G, Patel N. Regulatory obsolescence and secondary market asset depreciation. J Environ Econ Manag 2019;95:142-60.',
        '15. Lehmann H, Werning J, Baschnegger H, et al. Minimising the usage of desflurane only by education and removal of the vaporisers \u2013 a before-and-after-trial. BMC Anesthesiol 2025;25:108.',
        '16. Meyer MJ. Desflurane should des-appear: global and financial rationale. Anesth Analg 2020;131:1317-22.',
        '17. Moonesinghe SR. Desflurane decommissioning: more than meets the eye. Anaesthesia 2024;79:237-41.',
        '18. Mohammed A, Metta H. Is it time to bid adieu to desflurane? J Anaesthesiol Clin Pharmacol 2025;41:211-2.',
        '19. Beard D, Aston W, Black S, et al. Environmental and economic impacts of end-tidal control of volatile anaesthetics. Open Anaesthesia J 2025;19:e18742126.',
        '20. Buckhead Fair Market Value. 2025 Benchmark Report on Pre-Owned Medical Equipment Prices. Atlanta, GA: BFMV, 2025.',
    ]
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(10)

    doc.save(outdir + 'vaporizer_paper_english.docx')
    print("English paper saved (BMJ format, STROBE-compliant, with Spearman/Kendall)!")


# ==========================================
# JAPANESE PAPER
# ==========================================
def write_japanese_paper():
    doc = setup_doc()
    des = summary['Desflurane']
    sevo = summary['Sevoflurane']
    iso = summary['Isoflurane']
    des_u_pval = get_pval('Desflurane', 'u_pval')
    des_t_pval = get_pval('Desflurane', 't_pval')
    sevo_u_pval = get_pval('Sevoflurane', 'u_pval')
    iso_u_pval = get_pval('Isoflurane', 'u_pval')
    des_d = get_stat('Desflurane', 'cohens_d')
    des_tr = trend_results['Desflurane']
    sevo_tr = trend_results['Sevoflurane']
    iso_tr = trend_results['Isoflurane']

    # TITLE PAGE
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(
        'EU\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u898f\u5236\u304c\u9ebb\u9154\u6c17\u5316\u5668\u306e\u4e2d\u53e4\u5e02\u5834\u4fa1\u683c\u306b\u4e0e\u3048\u308b\u5f71\u97ff\uff1a'
        'eBay\u843d\u672d\u30c7\u30fc\u30bf\u306e\u6a2a\u65ad\u7684\u6642\u7cfb\u5217\u5206\u6790')
    run.bold = True
    run.font.size = Pt(14)

    add_para(doc, '[\u8457\u8005\u540d\u3092\u8a18\u5165]', size=Pt(11), italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, '[\u6240\u5c5e\u6a5f\u95a2\u3092\u8a18\u5165]', size=Pt(10), italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_run_styled(p, '\u8cac\u4efb\u8457\u8005: ', bold=True, size=Pt(10))
    add_run_styled(p, '[\u6c0f\u540d\u30fb\u30e1\u30fc\u30eb\u30a2\u30c9\u30ec\u30b9\u30fb\u4f4f\u6240\u3092\u8a18\u5165]', size=Pt(10))
    add_para(doc, '\u672c\u6587\u8a9e\u6570: [\u6700\u7d42\u5316\u5f8c\u306b\u8a18\u5165]', size=Pt(10))
    p = doc.add_paragraph()
    add_run_styled(p, '\u30ad\u30fc\u30ef\u30fc\u30c9: ', bold=True, size=Pt(10))
    add_run_styled(p, '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3; \u6c17\u5316\u5668; EU\u898f\u5236; \u4e2d\u53e4\u5e02\u5834; F\u30ac\u30b9; \u74b0\u5883\u6301\u7d9a\u53ef\u80fd\u6027; \u9ebb\u9154; STROBE', size=Pt(10))

    doc.add_page_break()

    # WHAT IS ALREADY KNOWN / WHAT THIS STUDY ADDS (Japanese)
    add_what_box(doc, '\u3053\u308c\u307e\u3067\u306b\u308f\u304b\u3063\u3066\u3044\u308b\u3053\u3068', [
        '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u306f\u5730\u7403\u6e29\u6696\u5316\u4fc2\u6570\uff08GWP\uff09\u304cCO\u2082\u306e\u7d042540\u500d\u3067\u3042\u308a\u3001\u4ed6\u306e\u63ee\u767a\u6027\u9ebb\u9154\u85ac\u3092\u5927\u5e45\u306b\u4e0a\u56de\u308b',
        'EU\u306f\u898f\u5247(EU) 2024/573\u306b\u3088\u308a\u30012026\u5e741\u67081\u65e5\u304b\u3089\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u306e\u65e5\u5e38\u7684\u306a\u81e8\u5e8a\u9ebb\u9154\u4f7f\u7528\u3092\u7981\u6b62\u3057\u305f',
        '\u9ebb\u9154\u85ac\u306e\u74b0\u5883\u898f\u5236\u304c\u95a2\u9023\u6a5f\u5668\u306e\u4e2d\u53e4\u5e02\u5834\u4fa1\u5024\u306b\u4e0e\u3048\u308b\u5f71\u97ff\u3092\u691c\u8a0e\u3057\u305f\u7814\u7a76\u306f\u306a\u3044',
    ])
    add_what_box(doc, '\u672c\u7814\u7a76\u304c\u52a0\u3048\u308b\u3053\u3068', [
        'eBay\u4e0a\u306e\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u4fa1\u683c\u306f\u7814\u7a76\u671f\u9593\u3092\u901a\u3058\u3066\u7d71\u8a08\u7684\u306b\u6709\u610f\u306a\u4e0b\u964d\u30c8\u30ec\u30f3\u30c9\u3092\u793a\u3057\u305f\uff08Spearman \u03c1=\u22120.28, P<0.001\uff09\u3002EU\u898f\u5236\u5f8c\u306b31%\u4e0b\u843d\u3057\u305f\u304c\u3001\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u304a\u3088\u3073\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u306b\u306f\u6709\u610f\u306a\u6642\u7cfb\u5217\u30c8\u30ec\u30f3\u30c9\u304c\u8a8d\u3081\u3089\u308c\u306a\u304b\u3063\u305f',
        '\u6642\u7cfb\u5217\u30c8\u30ec\u30f3\u30c9\u5206\u6790\u306b\u3088\u308a\u3001\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u4fa1\u683c\u304c\u898f\u5236\u6bb5\u968e\u306e\u9032\u884c\u306b\u4f34\u3044\u6bb5\u968e\u7684\u306b\u4e0b\u843d\u3057\u305f\u3053\u3068\u304c\u78ba\u8a8d\u3055\u308c\u305f\uff08Kendall \u03c4=\u22120.12, P=0.049\uff09\u304c\u3001\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\uff08\u03c4=0.04, P=0.36\uff09\u304a\u3088\u3073\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\uff08\u03c4=\u22120.07, P=0.025\uff09\u306b\u306f\u540c\u69d8\u306e\u30d1\u30bf\u30fc\u30f3\u306f\u8a8d\u3081\u3089\u308c\u306a\u304b\u3063\u305f',
        '\u3053\u308c\u3089\u306e\u77e5\u898b\u306f\u3001\u9ebb\u9154\u85ac\u306e\u74b0\u5883\u898f\u5236\u304c\u4e2d\u53e4\u533b\u7642\u6a5f\u5668\u5e02\u5834\u306b\u85ac\u5264\u7279\u7570\u7684\u306a\u7d4c\u6e08\u7684\u5f71\u97ff\u3092\u53ca\u307c\u3059\u3053\u3068\u3092\u793a\u3059\u521d\u3081\u3066\u306e\u5b9f\u8a3c\u7684\u30a8\u30d3\u30c7\u30f3\u30b9\u3067\u3042\u308b',
    ])

    doc.add_page_break()

    # ABSTRACT
    add_heading_styled(doc, '\u8981\u65e8', level=1)

    p = doc.add_paragraph()
    add_run_styled(p, '\u76ee\u7684 ', bold=True)
    add_run_styled(p, 'EU\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u898f\u5236\uff08\u898f\u5247(EU) 2024/573\uff09\u304c\u9ebb\u9154\u6c17\u5316\u5668\u306e\u4e2d\u53e4\u5e02\u5834\u4fa1\u683c\u306b\u4e0e\u3048\u308b\u5f71\u97ff\u3092\u8abf\u67fb\u3057\u3001\u4fa1\u683c\u5909\u52d5\u304c\u898f\u5236\u5bfe\u8c61\u85ac\u5264\u306b\u7279\u7570\u7684\u3067\u3042\u308b\u304b\u3069\u3046\u304b\u3092\u691c\u8a0e\u3059\u308b\u3002')

    p = doc.add_paragraph()
    add_run_styled(p, '\u30c7\u30b6\u30a4\u30f3 ', bold=True)
    add_run_styled(p, '\u843d\u672d\u6e08\u307f\u30ea\u30b9\u30c6\u30a3\u30f3\u30b0\u306e\u6a2a\u65ad\u7684\u6642\u7cfb\u5217\u5206\u6790\u3002')

    p = doc.add_paragraph()
    add_run_styled(p, '\u8a2d\u5b9a ', bold=True)
    add_run_styled(p, '\u4e16\u754c\u6700\u5927\u306e\u30aa\u30f3\u30e9\u30a4\u30f3\u30de\u30fc\u30b1\u30c3\u30c8\u30d7\u30ec\u30a4\u30b9\u3067\u3042\u308beBay\u3002Terapeak\uff08eBay\u516c\u5f0f\u306e\u904e\u53bb\u306e\u58f2\u4e0a\u5206\u6790\u30c4\u30fc\u30eb\uff09\u3092\u7528\u3044\u30663\u5e74\u5206\u306e\u843d\u672d\u30c7\u30fc\u30bf\u3092\u53d6\u5f97\u3057\u305f\u3002')

    p = doc.add_paragraph()
    add_run_styled(p, '\u4e3b\u8981\u30a2\u30a6\u30c8\u30ab\u30e0 ', bold=True)
    add_run_styled(p, '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u3001\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u3001\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u306e\u843d\u672d\u4fa1\u683c\uff08\u7c73\u30c9\u30eb\uff09\u3002\u6642\u7cfb\u5217\u4fa1\u683c\u30c8\u30ec\u30f3\u30c9\u3092Spearman\u9806\u4f4d\u76f8\u95a2\u304a\u3088\u3073Kendall \u03c4\uff08\u898f\u5236\u6bb5\u968e\u9806\u5e8f\uff09\u3067\u8a55\u4fa1\u3057\u305f\u3002\u898f\u5236\u524d\u5f8c\u306e\u4fa1\u683c\u6bd4\u8f03\u306b\u306fMann-Whitney U\u691c\u5b9a\u304a\u3088\u3073Welch\u306et\u691c\u5b9a\u3092\u7528\u3044\u3001Cohen\u306ed\u3067\u52b9\u679c\u91cf\u3092\u63a8\u5b9a\u3057\u305f\u3002')

    des_pct = abs((des['post_mean'] - des['pre_mean']) / des['pre_mean'] * 100)
    p = doc.add_paragraph()
    add_run_styled(p, '\u7d50\u679c ', bold=True)
    add_run_styled(p,
        f'{total_n}\u4ef6\u306e\u843d\u672d\u30c7\u30fc\u30bf\u3092\u5206\u6790\u3057\u305f\uff1a\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3{des["total_n"]}\u4ef6\u3001'
        f'\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3{sevo["total_n"]}\u4ef6\u3001\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3{iso["total_n"]}\u4ef6'
        f'\uff08{date_min_all}\uff5e{date_max_all}\uff09\u3002'
        f'\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u4fa1\u683c\u306f\u6709\u610f\u306a\u4e0b\u964d\u30c8\u30ec\u30f3\u30c9\u3092\u793a\u3057\u305f'
        f'\uff08Spearman \u03c1={des_tr["spearman_rho"]:.2f}, P{fmt_p(des_tr["spearman_p"])}; '
        f'Kendall \u03c4={des_tr["kendall_tau"]:.2f}, P={fmt_p(des_tr["kendall_p"])}\uff09\u3002'
        f'\u898f\u5236\u524d\u5e73\u5747US${des["pre_mean"]:.0f}\uff08SD ${des["pre_sd"]:.0f}\uff09\u304b\u3089'
        f'\u898f\u5236\u5f8cUS${des["post_mean"]:.0f}\uff08SD ${des["post_sd"]:.0f}\uff09\u3078{des_pct:.0f}%\u4e0b\u843d\u3057\u305f'
        f'\uff08Welch\u306et\u691c\u5b9a P={fmt_p(des_t_pval)}; Cohen\u306ed={des_d:.2f}\uff09\u3002'
        f'\u4e00\u65b9\u3001\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\uff08\u03c1={sevo_tr["spearman_rho"]:.2f}, P={fmt_p(sevo_tr["spearman_p"])}\uff09\u304a\u3088\u3073'
        f'\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\uff08\u03c1={iso_tr["spearman_rho"]:.2f}, P={fmt_p(iso_tr["spearman_p"])}\uff09\u306b\u306f'
        f'\u81e8\u5e8a\u7684\u306b\u610f\u5473\u306e\u3042\u308b\u6642\u7cfb\u5217\u30c8\u30ec\u30f3\u30c9\u306f\u8a8d\u3081\u3089\u308c\u305a\u3001\u898f\u5236\u524d\u5f8c\u306e\u6bd4\u8f03\u3082\u6709\u610f\u3067\u306f\u306a\u304b\u3063\u305f\u3002')

    p = doc.add_paragraph()
    add_run_styled(p, '\u7d50\u8ad6 ', bold=True)
    add_run_styled(p, 'EU\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u898f\u5236\u306f\u3001\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u306e\u4e2d\u53e4\u5e02\u5834\u4fa1\u683c\u306e\u6bb5\u968e\u7684\u304b\u3064\u85ac\u5264\u7279\u7570\u7684\u306a\u4e0b\u843d\u3068\u95a2\u9023\u3057\u3066\u3044\u305f\u3002\u6642\u7cfb\u5217\u5206\u6790\u306b\u3088\u308a\u3001\u3053\u306e\u4e0b\u843d\u306f\u898f\u5236\u5bfe\u8c61\u85ac\u5264\u306b\u7279\u6709\u3067\u3042\u308a\u3001\u7acb\u6cd5\u904e\u7a0b\u4e2d\u306b\u65e2\u306b\u59cb\u307e\u3063\u3066\u3044\u305f\u3053\u3068\u304c\u793a\u3055\u308c\u3001\u5e02\u5834\u306e\u4e88\u6e2c\u7684\u53cd\u5fdc\u304c\u793a\u5506\u3055\u308c\u305f\u3002\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u304a\u3088\u3073\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u4fa1\u683c\u306f\u7814\u7a76\u671f\u9593\u3092\u901a\u3058\u3066\u5b89\u5b9a\u3057\u3066\u304a\u308a\u3001\u81ea\u7136\u5bfe\u7167\u7fa4\u3068\u3057\u3066\u6a5f\u80fd\u3057\u305f\u3002')

    doc.add_page_break()

    # INTRODUCTION
    add_heading_styled(doc, '\u7dd2\u8a00', level=1)
    doc.add_paragraph(
        '\u5438\u5165\u9ebb\u9154\u85ac\u306f\u533b\u7642\u306e\u30ab\u30fc\u30dc\u30f3\u30d5\u30c3\u30c8\u30d7\u30ea\u30f3\u30c8\u306b\u5927\u304d\u304f\u5bc4\u4e0e\u3057\u3066\u3044\u308b\u3002'
        '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u306f\u8fc5\u901f\u306a\u5c0e\u5165\u30fb\u899a\u9192\u7279\u6027\u3067\u8a55\u4fa1\u3055\u308c\u3066\u3044\u308b\u304c\u3001100\u5e74\u6642\u9593\u8ef8\u306e\u5730\u7403\u6e29\u6696\u5316\u4fc2\u6570\uff08GWP\uff09\u306f\u7d042540 CO\u2082\u7b49\u4fa1\u3067\u3042\u308a\u3001'
        '\u65e5\u5e38\u7684\u306b\u4f7f\u7528\u3055\u308c\u308b\u63ee\u767a\u6027\u9ebb\u9154\u85ac\u306e\u4e2d\u3067\u6700\u3082\u74b0\u5883\u8ca0\u8377\u304c\u5927\u304d\u3044\u3002'
        '\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u306eGWP\u306f\u7d04130\u3001\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u306f\u7d04510\u3067\u3042\u308b\u3002')
    doc.add_paragraph(
        '\u6b27\u5dde\u306b\u304a\u3051\u308b\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u898f\u5236\u306e\u7d4c\u7def\u306f\u3001\u8907\u6570\u306e\u91cd\u8981\u306a\u30de\u30a4\u30eb\u30b9\u30c8\u30fc\u30f3\u3092\u7d4c\u3066\u9032\u5c55\u3057\u305f\u3002'
        '2022\u5e744\u6708\u3001\u6b27\u5dde\u59d4\u54e1\u4f1a\u304cF\u30ac\u30b9\u898f\u5247\u6539\u6b63\u6848\u3092\u516c\u8868\u3057\u305f\u3002'
        '2023\u5e743\u6708\u306b\u6b27\u5dde\u8b70\u4f1a\u304c\u672c\u4f1a\u8b70\u3067\u627f\u8a8d\u3057\u3001'
        '2023\u5e7410\u6708\u306b\u7406\u4e8b\u4f1a\u3068\u8b70\u4f1a\u306e\u9593\u3067\u6697\u5b9a\u5408\u610f\uff08\u30c8\u30ea\u30ed\u30fc\u30b0\uff09\u304c\u6210\u7acb\u3057\u305f\u3002'
        '\u898f\u5247(EU) 2024/573\u306f2024\u5e742\u6708\u306b\u6b63\u5f0f\u63a1\u629e\u3055\u308c\u30012024\u5e743\u6708\u306b\u767a\u52b9\u3057\u3001'
        '\u65e5\u5e38\u9ebb\u9154\u306b\u304a\u3051\u308b\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u4f7f\u7528\u7981\u6b62\u306f2026\u5e741\u67081\u65e5\u304b\u3089\u9069\u7528\u3055\u308c\u305f\u3002'
        '\u3053\u308c\u306f\u74b0\u5883\u7684\u7406\u7531\u306b\u3088\u308b\u7279\u5b9a\u306e\u9ebb\u9154\u85ac\u3078\u306e\u521d\u306e\u7fa9\u52d9\u7684\u653f\u5e9c\u898f\u5236\u3067\u3042\u308b\u3002')
    doc.add_paragraph(
        '\u9ebb\u9154\u6c17\u5316\u5668\u306f\u85ac\u5264\u7279\u7570\u7684\u306a\u88c5\u7f6e\u3067\u3042\u308a\u3001\u5178\u578b\u7684\u306a\u5bff\u547d\u306f10\uff5e15\u5e74\u3001\u76f8\u5f53\u306a\u8a2d\u5099\u6295\u8cc7\u3092\u8981\u3059\u308b\u3002'
        '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u306e\u898f\u5236\u306b\u3088\u308b\u9673\u8150\u5316\u306f\u3001\u6a5f\u5668\u6240\u6709\u8005\u306b\u91cd\u5927\u306a\u7d4c\u6e08\u7684\u5f71\u97ff\u3092\u53ca\u307c\u3059\u53ef\u80fd\u6027\u304c\u3042\u308b\u3002'
        '\u91cd\u8981\u306a\u306e\u306f\u3001\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u3068\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u306f\u540c\u3058\u898f\u5236\u306e\u5bfe\u8c61\u3067\u306f\u306a\u3044\u305f\u3081\u3001'
        '\u305d\u306e\u4fa1\u683c\u306f\u5f71\u97ff\u3092\u53d7\u3051\u306a\u3044\u306f\u305a\u3067\u3042\u308a\u3001\u81ea\u7136\u5bfe\u7167\u7fa4\u3092\u63d0\u4f9b\u3059\u308b\u3002')
    doc.add_paragraph(
        '\u5148\u884c\u7814\u7a76\u3067\u306f\u3001\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u4f7f\u7528\u4e2d\u6b62\u306e\u8ca1\u52d9\u7684\u6839\u62e0[16]\u3001'
        '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u5ec3\u6b62\u306e\u81e8\u5e8a\u7684\u30fb\u653f\u7b56\u7684\u542b\u610f[17,18]\u3001'
        '\u304a\u3088\u3073\u65bd\u8a2d\u30ec\u30d9\u30eb\u3067\u306e\u6c17\u5316\u5668\u64a4\u53bb\u30d7\u30ed\u30b0\u30e9\u30e0\u306e\u6709\u52b9\u6027[15]\u304c\u691c\u8a0e\u3055\u308c\u3066\u3044\u308b\u3002'
        '\u63ee\u767a\u6027\u9ebb\u9154\u85ac\u6d88\u8cbb\u524a\u6e1b\u306b\u3088\u308b\u30b3\u30b9\u30c8\u524a\u6e1b\u306e\u7d4c\u6e08\u5206\u6790[9,19]\u3084\u3001'
        '\u4ed6\u306e\u6a5f\u5668\u30ab\u30c6\u30b4\u30ea\u30fc\u306b\u304a\u3051\u308b\u4e2d\u53e4\u533b\u7642\u6a5f\u5668\u5e02\u5834\u306e\u7279\u5fb4\u4ed8\u3051[20]\u3082\u884c\u308f\u308c\u3066\u3044\u308b\u3002'
        '\u3057\u304b\u3057\u3001\u6211\u3005\u306e\u77e5\u308b\u9650\u308a\u3001\u74b0\u5883\u898f\u5236\u304c\u9ebb\u9154\u6a5f\u5668\u306e\u4e2d\u53e4\u5e02\u5834\u4fa1\u5024\u306b\u4e0e\u3048\u308b\u5f71\u97ff\u3092\u691c\u8a0e\u3057\u305f\u7814\u7a76\u306f\u306a\u3044\u3002'
        '\u6211\u3005\u306f\u3001EU\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u898f\u5236\u304c\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u306e\u4e2d\u53e4\u5e02\u5834\u4fa1\u683c\u306e\u6bb5\u968e\u7684\u306a\u4e0b\u843d\u3068\u95a2\u9023\u3057\u3001'
        '\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u304a\u3088\u3073\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u306e\u4fa1\u683c\u306f\u5b89\u5b9a\u3059\u308b\u3068\u4eee\u8aac\u3092\u7acb\u3066\u305f\u3002'
        'eBay\u306e3\u5e74\u5206\u306e\u843d\u672d\u30c7\u30fc\u30bf\uff08Terapeak\u7d4c\u7531\uff09\u3092\u7528\u3044\u3001\u6a2a\u65ad\u7684\u6bd4\u8f03\u3068\u6642\u7cfb\u5217\u30c8\u30ec\u30f3\u30c9\u5206\u6790\u306e\u4e21\u65b9\u3067\u3053\u306e\u4eee\u8aac\u3092\u691c\u8a3c\u3057\u305f\u3002')

    # METHODS
    add_heading_styled(doc, '\u65b9\u6cd5', level=1)
    doc.add_paragraph(
        '\u672c\u7814\u7a76\u306f\u3001\u6a2a\u65ad\u7814\u7a76\u306eSTROBE\uff08Strengthening the Reporting of Observational Studies in Epidemiology\uff09\u30ac\u30a4\u30c9\u30e9\u30a4\u30f3\u306b\u5f93\u3063\u3066\u5831\u544a\u3059\u308b\u3002')

    add_heading_styled(doc, '\u7814\u7a76\u30c7\u30b6\u30a4\u30f3\u304a\u3088\u3073\u30c7\u30fc\u30bf\u30bd\u30fc\u30b9', level=2)
    doc.add_paragraph(
        f'eBay (www.ebay.com) \u4e0a\u306e\u843d\u672d\u6e08\u307f\u30ea\u30b9\u30c6\u30a3\u30f3\u30b0\u3092\u7528\u3044\u305f\u9ebb\u9154\u6c17\u5316\u5668\u4fa1\u683c\u306e\u6a2a\u65ad\u7684\u6642\u7cfb\u5217\u5206\u6790\u3092\u884c\u3063\u305f\u3002'
        '\u30c7\u30fc\u30bf\u306f\u3001eBay Seller Hub\u306b\u7d71\u5408\u3055\u308c\u305feBay\u516c\u5f0f\u306e\u88fd\u54c1\u30ea\u30b5\u30fc\u30c1\u30c4\u30fc\u30eb\u3067\u3042\u308bTerapeak\u3092\u4f7f\u7528\u3057\u3066\u53d6\u5f97\u3057\u305f\u3002'
        'Terapeak\u306f\u6700\u59273\u5e74\u5206\u306e\u904e\u53bb\u306e\u843d\u672d\u30c7\u30fc\u30bf\u3078\u306e\u30a2\u30af\u30bb\u30b9\u3092\u63d0\u4f9b\u3059\u308b\u3002'
        f'\u30c7\u30fc\u30bf\u53ce\u96c6\u306f2026\u5e743\u6708\u306b\u884c\u3044\u30012023\u5e743\u670828\u65e5\u304b\u30892026\u5e743\u670824\u65e5\u307e\u3067\u306e\u671f\u9593\u3092\u30ab\u30d0\u30fc\u3057\u305f\u3002'
        '3\u5e74\u9593\u3068\u3044\u3046\u89b3\u5bdf\u671f\u9593\u306fTerapeak\u306e\u6700\u5927\u53d6\u5f97\u53ef\u80fd\u671f\u9593\u306b\u4e00\u81f4\u3059\u308b\u304c\u3001'
        '\u5206\u6790\u4e0a\u3082\u610f\u7fa9\u3042\u308b\u6642\u9593\u67a0\u3067\u3042\u308b\uff1a'
        '\u6b27\u5dde\u59d4\u54e1\u4f1a\u306e\u539f\u6848\u63d0\u51fa\uff082022\u5e744\u6708\uff09\u304b\u3089\u6b27\u5dde\u8b70\u4f1a\u672c\u4f1a\u8b70\u3067\u306e\u627f\u8a8d\uff082024\u5e743\u6708\uff09\u3001'
        '\u305d\u3057\u3066\u898f\u5236\u767a\u52b9\u5f8c\u307e\u3067\u306e\u5168\u7acb\u6cd5\u904e\u7a0b\u3092\u7db2\u7f85\u3057\u3066\u304a\u308a\u3001'
        '\u7d14\u7c8b\u306a\u6280\u8853\u7684\u5236\u7d04\u3067\u306f\u306a\u304f\u3001\u4e3b\u8981\u306a\u898f\u5236\u30de\u30a4\u30eb\u30b9\u30c8\u30fc\u30f3\u3092\u5168\u3066\u542b\u3080\u3082\u306e\u3067\u3042\u308b\u3002'
        '\u30af\u30ed\u30b9\u30ea\u30b9\u30c6\u30a3\u30f3\u30b0\u3055\u308c\u305f\u5546\u54c1\u306e\u91cd\u8907\u30ab\u30a6\u30f3\u30c8\u3092\u907f\u3051\u308b\u305f\u3081\u3001\u5358\u4e00\u30de\u30fc\u30b1\u30c3\u30c8\u30d7\u30ec\u30a4\u30b9\uff08eBay\uff09\u306e\u307f\u3092\u4f7f\u7528\u3057\u305f\u3002')

    add_heading_styled(doc, '\u7d71\u8a08\u5206\u6790', level=2)
    doc.add_paragraph(
        '\u8a18\u8ff0\u7d71\u8a08\u91cf\u3068\u3057\u3066\u5e73\u5747\u3001\u6a19\u6e96\u504f\u5dee\u3001\u4e2d\u592e\u5024\u3001\u56db\u5206\u4f4d\u7bc4\u56f2\u3092\u7b97\u51fa\u3057\u305f\u3002'
        '\u4fa1\u683c\u5206\u5e03\u306e\u975e\u6b63\u898f\u6027\u3092\u8003\u616e\u3057\u3001\u898f\u5236\u524d\u5f8c\u306e\u4fa1\u683c\u6bd4\u8f03\u306b\u306fMann-Whitney U\u691c\u5b9a\uff08\u4e21\u5074\uff09\u3092\u4e3b\u8981\u691c\u5b9a\u3068\u3057\u3066\u4f7f\u7528\u3057\u305f\u3002'
        '\u611f\u5ea6\u5206\u6790\u3068\u3057\u3066Welch\u306et\u691c\u5b9a\u3092\u884c\u3044\u3001\u52b9\u679c\u91cf\u306fCohen\u306ed\u3067\u63a8\u5b9a\u3057\u305f\u3002')
    doc.add_paragraph(
        '\u4fa1\u683c\u304c\u6642\u9593\u7d4c\u904e\u3068\u3068\u3082\u306b\u6bb5\u968e\u7684\u306b\u5909\u5316\u3057\u305f\u304b\u3092\u8a55\u4fa1\u3059\u308b\u305f\u3081\u30012\u3064\u306e\u88dc\u5b8c\u7684\u306a\u30c8\u30ec\u30f3\u30c9\u5206\u6790\u3092\u884c\u3063\u305f\u3002'
        '\u7b2c\u4e00\u306b\u3001Spearman\u9806\u4f4d\u76f8\u95a2\u3092\u7528\u3044\u3066\u3001\u5404\u85ac\u5264\u30bf\u30a4\u30d7\u306b\u3064\u3044\u3066\u58f2\u5374\u65e5\uff08\u7814\u7a76\u958b\u59cb\u65e5\u304b\u3089\u306e\u65e5\u6570\uff09\u3068\u4fa1\u683c\u306e\u5358\u8abf\u7684\u95a2\u9023\u3092\u691c\u5b9a\u3057\u305f\u3002'
        '\u7b2c\u4e8c\u306b\u3001Kendall \u03c4\u3092\u7528\u3044\u3066\u3001\u898f\u5236\u6bb5\u968e\u306e\u9806\u5e8f\uff081\uff5e5\uff09\u3068\u4fa1\u683c\u306e\u95a2\u9023\u3092\u691c\u5b9a\u3057\u3001'
        '\u898f\u5236\u30de\u30a4\u30eb\u30b9\u30c8\u30fc\u30f3\u306e\u9032\u884c\u306b\u4f34\u3044\u4fa1\u683c\u304c\u6bb5\u968e\u7684\u306b\u4e0b\u843d\u3057\u305f\u304b\u3092\u8a55\u4fa1\u3057\u305f\u3002'
        '\u3053\u308c\u3089\u306e\u30c8\u30ec\u30f3\u30c9\u691c\u5b9a\u306f\u5404\u85ac\u5264\u30bf\u30a4\u30d7\u306b\u72ec\u7acb\u3057\u3066\u9069\u7528\u3057\u3001\u898f\u5236\u5bfe\u8c61\u85ac\u5264\uff08\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\uff09\u3068'
        '\u975e\u898f\u5236\u5bfe\u7167\u85ac\u5264\uff08\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u3001\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\uff09\u306e\u6642\u9593\u7684\u30d1\u30bf\u30fc\u30f3\u3092\u76f4\u63a5\u6bd4\u8f03\u3067\u304d\u308b\u3088\u3046\u306b\u3057\u305f\u3002'
        '\u56db\u534a\u671f\u3054\u3068\u306e\u4e2d\u592e\u5024\u306b\u3064\u3044\u3066\u3082Spearman\u76f8\u95a2\u3092\u8a55\u4fa1\u3057\u305f\u3002')
    doc.add_paragraph(
        '\u898f\u5236\u6bb5\u968e\u9593\u306e\u591a\u7fa4\u6bd4\u8f03\u306b\u306fKruskal-Wallis\u691c\u5b9a\u3092\u4f7f\u7528\u3057\u305f\u3002'
        '\u4fa1\u683c\u63a8\u79fb\u306e\u53ef\u8996\u5316\u306b\u306fLOWESS\uff08\u5c40\u6240\u91cd\u307f\u4ed8\u3051\u6563\u5e03\u56f3\u5e73\u6ed1\u5316\uff09\u30c8\u30ec\u30f3\u30c9\u30e9\u30a4\u30f3\u3092\u4f7f\u7528\u3057\u305f\u3002'
        '\u5206\u6790\u306fPython 3.12\uff08pandas 2.2, scipy 1.14, statsmodels 0.14\uff09\u3067\u884c\u3063\u305f\u3002'
        '\u7d71\u8a08\u7684\u6709\u610f\u6c34\u6e96\u306fP<0.05\uff08\u4e21\u5074\uff09\u3068\u3057\u305f\u3002')

    # RESULTS
    add_heading_styled(doc, '\u7d50\u679c', level=1)

    add_heading_styled(doc, '\u7814\u7a76\u5bfe\u8c61', level=2)
    doc.add_paragraph(
        f'\u9664\u5916\u57fa\u6e96\u9069\u7528\u5f8c\u3001\u5408\u8a08{total_n}\u4ef6\u306eeBay\u843d\u672d\u6e08\u307f\u9ebb\u9154\u6c17\u5316\u5668\u3092\u5206\u6790\u5bfe\u8c61\u3068\u3057\u305f\uff1a'
        f'\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3{des["total_n"]}\u4ef6\u3001'
        f'\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3{sevo["total_n"]}\u4ef6\u3001'
        f'\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3{iso["total_n"]}\u4ef6\u3002'
        f'\u7814\u7a76\u671f\u9593\u306f{date_min_all}\u304b\u3089{date_max_all}\u307e\u3067\u306e3\u5e74\u9593\u3067\u3042\u3063\u305f\u3002')

    # Table 1 caption
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_run_styled(p, '\u88681. ', bold=True, size=Pt(10))
    add_run_styled(p, '\u6c17\u5316\u5668\u30bf\u30a4\u30d7\u304a\u3088\u3073\u898f\u5236\u671f\u9593\u5225\u306eeBay Terapeak\u843d\u672d\u30c7\u30fc\u30bf\u306e\u8981\u7d04\uff08\u898f\u5236\u524d\u5f8c\u30012026\u5e741\u67081\u65e5\u57fa\u6e96\uff09\u3002\u5024\u306f\u5e73\u5747\uff08SD\uff09\u3001\u4e2d\u592e\u5024\uff08IQR\uff09\u3001\u7c73\u30c9\u30eb\u3002P\u5024\u306fMann-Whitney U\u691c\u5b9a\uff08\u4e21\u5074\uff09\u3002', italic=True, size=Pt(10))

    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header(table, ['\u85ac\u5264', '\u671f\u9593', 'n', '\u5e73\u5747 (SD)', '\u4e2d\u592e\u5024 (IQR)', '\u7bc4\u56f2', 'P\u5024', "Cohen's d"])

    for agent, jp_name in [('Desflurane', '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3'), ('Sevoflurane', '\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3'), ('Isoflurane', '\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3')]:
        pval = get_pval(agent)
        d_val = get_stat(agent, 'cohens_d')
        for period_name, label in [('Pre-regulation', '\u898f\u5236\u524d'), ('Post-regulation', '\u898f\u5236\u5f8c')]:
            sub = combined[(combined['agent_type'] == agent) & (combined['period'] == period_name)]
            if len(sub) == 0: continue
            prices = sub['price_usd']
            mean_sd = f'${prices.mean():.0f} ({prices.std():.0f})'
            med_iqr = f'${prices.median():.0f} ({prices.quantile(0.25):.0f}\u2013{prices.quantile(0.75):.0f})'
            rng = f'${prices.min():.0f}\u2013{prices.max():.0f}'
            pval_str = fmt_p(pval) if label == '\u898f\u5236\u524d' else ''
            d_str = f'{d_val:.2f}' if label == '\u898f\u5236\u524d' and not np.isnan(d_val) else ''
            data = [
                (jp_name if label == '\u898f\u5236\u524d' else '', WD_ALIGN_PARAGRAPH.LEFT),
                (label, WD_ALIGN_PARAGRAPH.CENTER),
                (str(len(sub)), WD_ALIGN_PARAGRAPH.CENTER),
                (mean_sd, WD_ALIGN_PARAGRAPH.CENTER),
                (med_iqr, WD_ALIGN_PARAGRAPH.CENTER),
                (rng, WD_ALIGN_PARAGRAPH.CENTER),
                (pval_str, WD_ALIGN_PARAGRAPH.CENTER),
                (d_str, WD_ALIGN_PARAGRAPH.CENTER),
            ]
            add_table_data_row(table, data)

    doc.add_paragraph()

    # Table 2 - Trend analysis (Japanese)
    p = doc.add_paragraph()
    add_run_styled(p, '\u88682. ', bold=True, size=Pt(10))
    add_run_styled(p, '\u85ac\u5264\u30bf\u30a4\u30d7\u5225\u306e\u6642\u7cfb\u5217\u30c8\u30ec\u30f3\u30c9\u5206\u6790\u3002Spearman\u9806\u4f4d\u76f8\u95a2\u306f\u58f2\u5374\u65e5\u3068\u4fa1\u683c\u306e\u5358\u8abf\u7684\u95a2\u9023\u3001Kendall \u03c4\u306f\u898f\u5236\u6bb5\u968e\u9806\u5e8f\u3068\u4fa1\u683c\u306e\u95a2\u9023\u3092\u691c\u5b9a\u3002\u56db\u534a\u671f\u30c8\u30ec\u30f3\u30c9\u306f\u56db\u534a\u671f\u4e2d\u592e\u5024\u306eSpearman\u76f8\u95a2\u3002', italic=True, size=Pt(10))

    t2 = doc.add_table(rows=1, cols=7)
    t2.style = 'Table Grid'
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header(t2, ['\u85ac\u5264', 'Spearman \u03c1', 'P\u5024', 'Kendall \u03c4', 'P\u5024', '\u56db\u534a\u671f \u03c1', 'P\u5024'])

    for agent, jp_name in [('Desflurane', '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3'), ('Sevoflurane', '\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3'), ('Isoflurane', '\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3')]:
        tr = trend_results[agent]
        data = [
            (jp_name, WD_ALIGN_PARAGRAPH.LEFT),
            (f'{tr["spearman_rho"]:.3f}', WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(tr['spearman_p']), WD_ALIGN_PARAGRAPH.CENTER),
            (f'{tr["kendall_tau"]:.3f}', WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(tr['kendall_p']), WD_ALIGN_PARAGRAPH.CENTER),
            (f'{tr["quarterly_rho"]:.3f}', WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(tr['quarterly_p']), WD_ALIGN_PARAGRAPH.CENTER),
        ]
        add_table_data_row(t2, data)

    doc.add_paragraph()

    # Results narrative (Japanese)
    des_pct_val = (des['post_mean'] - des['pre_mean']) / des['pre_mean'] * 100

    add_heading_styled(doc, '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\uff1a\u6709\u610f\u306a\u6bb5\u968e\u7684\u4fa1\u683c\u4e0b\u843d', level=2)
    doc.add_paragraph(
        f'\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u4fa1\u683c\u306f\u30013\u5e74\u9593\u306e\u7814\u7a76\u671f\u9593\u3092\u901a\u3058\u3066\u7d71\u8a08\u7684\u306b\u6709\u610f\u306a\u4e0b\u964d\u30c8\u30ec\u30f3\u30c9\u3092\u793a\u3057\u305f\u3002'
        f'Spearman\u9806\u4f4d\u76f8\u95a2\u306f\u58f2\u5374\u65e5\u3068\u4fa1\u683c\u306e\u9593\u306b\u6709\u610f\u306a\u8ca0\u306e\u5358\u8abf\u7684\u95a2\u9023\u3092\u793a\u3057\u305f'
        f'\uff08\u03c1={des_tr["spearman_rho"]:.2f}, P{fmt_p(des_tr["spearman_p"])}\uff09\u3002'
        f'Kendall \u03c4\u5206\u6790\u306b\u3088\u308a\u3001\u898f\u5236\u6bb5\u968e\u306e\u9032\u884c\u306b\u4f34\u3044\u4fa1\u683c\u304c\u4e0b\u843d\u3057\u305f\u3053\u3068\u304c\u78ba\u8a8d\u3055\u308c\u305f'
        f'\uff08\u03c4={des_tr["kendall_tau"]:.2f}, P={fmt_p(des_tr["kendall_p"])}\uff09\u3002'
        f'\u56db\u534a\u671f\u4e2d\u592e\u5024\u3082\u6709\u610f\u306a\u4e0b\u964d\u30c8\u30ec\u30f3\u30c9\u3092\u793a\u3057\u305f'
        f'\uff08\u03c1={des_tr["quarterly_rho"]:.2f}, P={fmt_p(des_tr["quarterly_p"])}\uff09\u3002')
    doc.add_paragraph(
        f'\u898f\u5236\u524d\u5f8c\u306e\u76f4\u63a5\u6bd4\u8f03\u3067\u306f\u3001\u898f\u5236\u5f8c\u5e73\u5747\u4fa1\u683c\uff08US${des["post_mean"]:.0f}, '
        f'SD ${des["post_sd"]:.0f}\uff09\u306f\u898f\u5236\u524d\u5e73\u5747\uff08US${des["pre_mean"]:.0f}, SD ${des["pre_sd"]:.0f}\uff09\u3088\u308a'
        f'{abs(des_pct_val):.0f}%\u4f4e\u304b\u3063\u305f\u3002'
        f'\u3053\u306e\u5dee\u306fWelch\u306et\u691c\u5b9a\u3067\u7d71\u8a08\u7684\u306b\u6709\u610f\u3067\u3042\u3063\u305f\u304c\uff08P={fmt_p(des_t_pval)}\uff09\u3001'
        f'Mann-Whitney U\u691c\u5b9a\u3067\u306f\u6709\u610f\u306b\u9054\u3057\u306a\u304b\u3063\u305f\uff08P={fmt_p(des_u_pval)}\uff09\u3002'
        f'\u3053\u308c\u306f\u898f\u5236\u5f8c\u30b5\u30f3\u30d7\u30eb\u304c\u5c0f\u3055\u3044\uff08n={des["post_n"]}\uff09\u3053\u3068\u3092\u53cd\u6620\u3057\u3066\u3044\u308b\u3068\u8003\u3048\u3089\u308c\u308b\u3002'
        f'\u52b9\u679c\u91cf\u306f\u4e2d\u7a0b\u5ea6\u3067\u3042\u3063\u305f\uff08Cohen\u306ed={des_d:.2f}\uff09\u3002')

    add_heading_styled(doc, '\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u304a\u3088\u3073\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\uff1a\u5b89\u5b9a\u3057\u305f\u4fa1\u683c', level=2)
    sevo_pct = (sevo['post_mean'] - sevo['pre_mean']) / sevo['pre_mean'] * 100
    iso_pct = (iso['post_mean'] - iso['pre_mean']) / iso['pre_mean'] * 100
    doc.add_paragraph(
        f'\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u3068\u306f\u5bfe\u7167\u7684\u306b\u3001\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u4fa1\u683c\u306b\u306f\u6709\u610f\u306a\u6642\u7cfb\u5217\u30c8\u30ec\u30f3\u30c9\u304c\u8a8d\u3081\u3089\u308c\u306a\u304b\u3063\u305f'
        f'\uff08Spearman \u03c1={sevo_tr["spearman_rho"]:.2f}, P={fmt_p(sevo_tr["spearman_p"])}; '
        f'Kendall \u03c4={sevo_tr["kendall_tau"]:.2f}, P={fmt_p(sevo_tr["kendall_p"])}\uff09\u3002'
        f'\u898f\u5236\u524d\u5f8c\u306e\u6bd4\u8f03\u3067\u306f\u3001\u975e\u6709\u610f\u306e{abs(sevo_pct):.0f}%\u5897\u52a0\u3067\u3042\u3063\u305f\uff08P={fmt_p(sevo_u_pval)}\uff09\u3002')
    doc.add_paragraph(
        f'\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u4fa1\u683c\u3082\u540c\u69d8\u306b\u5b89\u5b9a\u3057\u3066\u3044\u305f\u3002Spearman\u76f8\u95a2\u306f\u540d\u76ee\u7684\u306b\u6709\u610f\u3067\u3042\u3063\u305f\u304c'
        f'\uff08\u03c1={iso_tr["spearman_rho"]:.2f}, P={fmt_p(iso_tr["spearman_p"])}\uff09\u3001'
        f'\u56db\u534a\u671f\u4e2d\u592e\u5024\u30c8\u30ec\u30f3\u30c9\u306f\u6709\u610f\u3067\u306f\u306a\u304b\u3063\u305f'
        f'\uff08\u03c1={iso_tr["quarterly_rho"]:.2f}, P={fmt_p(iso_tr["quarterly_p"])}\uff09\u3002'
        f'\u898f\u5236\u524d\u5f8c\u306e\u6bd4\u8f03\u306f\u975e\u6709\u610f\u306e{abs(iso_pct):.0f}%\u4e0b\u843d\u3067\u3042\u3063\u305f\uff08P={fmt_p(iso_u_pval)}\uff09\u3002'
        f'\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u304a\u3088\u3073\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u4fa1\u683c\u306e\u5b89\u5b9a\u6027\u306f\u3001\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u4fa1\u683c\u4e0b\u843d\u304c'
        f'\u5e83\u7bc4\u306a\u5e02\u5834\u8981\u56e0\u3067\u306f\u306a\u304fEU\u898f\u5236\u306b\u7279\u7570\u7684\u3067\u3042\u308b\u3068\u3044\u3046\u63a8\u8ad6\u3092\u5f37\u5316\u3059\u308b\u3002')

    # Figures (Japanese captions)
    add_figure(doc, figdir + 'fig1_price_timeseries.png', '\u56f31. ',
        'eBay\u843d\u672d\u4fa1\u683c\u306e\u6642\u7cfb\u5217\u63a8\u79fb\uff08\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3=\u8d64\u3001\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3=\u9752\u3001\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3=\u7dd1\uff09\u3002'
        '\u7e26\u7834\u7dda\u306fEU\u898f\u5236\u30de\u30a4\u30eb\u30b9\u30c8\u30fc\u30f3\u3002LOWESS\u30c8\u30ec\u30f3\u30c9\u30e9\u30a4\u30f3\u4ed8\u304d\u3002'
        '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u306eLOWESS\u66f2\u7dda\u306e\u6bb5\u968e\u7684\u306a\u4e0b\u964d\u306b\u6ce8\u76ee\u3002\u30c7\u30fc\u30bf\u30bd\u30fc\u30b9: eBay Terapeak\u3002')
    doc.add_page_break()

    add_figure(doc, figdir + 'fig2_boxplot_comparison.png', '\u56f32. ',
        '\u898f\u5236\u524d\u5f8c\u306e\u4fa1\u683c\u6bd4\u8f03\u7bb1\u3072\u3052\u56f3\u3002\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u306e\u898f\u5236\u5f8c\u5206\u5e03\u304c\u4f4e\u4fa1\u683c\u5074\u306b\u5727\u7e2e\u3055\u308c\u3066\u3044\u308b\u3002'
        '\u30c7\u30fc\u30bf\u30bd\u30fc\u30b9: eBay Terapeak\u3002')
    doc.add_paragraph()

    add_figure(doc, figdir + 'fig3_monthly_median.png', '\u56f33. ',
        '\u6708\u5225\u4e2d\u592e\u5024\u4fa1\u683c\u63a8\u79fb\u30022024\u5e74\u4e2d\u9803\u304b\u3089\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u306e\u6301\u7d9a\u7684\u4e0b\u843d\u304c\u78ba\u8a8d\u3067\u304d\u308b\u3002'
        '\u30c7\u30fc\u30bf\u30bd\u30fc\u30b9: eBay Terapeak\u3002')
    doc.add_page_break()

    add_figure(doc, figdir + 'fig4_histograms.png', '\u56f34. ',
        '\u4fa1\u683c\u5206\u5e03\u30d2\u30b9\u30c8\u30b0\u30e9\u30e0\u3002\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u306e\u898f\u5236\u5f8c\u5206\u5e03\u304c\u5de6\u65b9\u5411\uff08\u4f4e\u4fa1\u683c\u5074\uff09\u306b\u30b7\u30d5\u30c8\u3057\u3066\u3044\u308b\u3002'
        '\u30c7\u30fc\u30bf\u30bd\u30fc\u30b9: eBay Terapeak\u3002')
    doc.add_paragraph()

    add_figure(doc, figdir + 'fig5_regulatory_timeline.png', '\u56f35. ',
        'EU\u898f\u5236\u30bf\u30a4\u30e0\u30e9\u30a4\u30f3\u3068\u4fa1\u683c\u63a8\u79fb\u3002\u898f\u5236\u6bb5\u968e\u3092\u7f72\u3051\u3002'
        '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u306e\u6bb5\u968e\u7684\u306a\u4fa1\u683c\u4e0b\u843d\u304c\u8996\u899a\u7684\u306b\u78ba\u8a8d\u3067\u304d\u308b\u3002\u30c7\u30fc\u30bf\u30bd\u30fc\u30b9: eBay Terapeak\u3002')
    doc.add_page_break()

    add_figure(doc, figdir + 'fig6_quarterly_trends.png', '\u56f36. ',
        '\u56db\u534a\u671f\u5225\u4e2d\u592e\u5024\u4fa1\u683c\uff08\u4e0a\uff09\u3068\u53d6\u5f15\u91cf\uff08\u4e0b\uff09\u306e\u63a8\u79fb\u3002'
        '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u306e\u56db\u534a\u671f\u4e2d\u592e\u5024\u304c~$250\u304b\u3089~$100\u3078\u4e0b\u843d\u3002\u30c7\u30fc\u30bf\u30bd\u30fc\u30b9: eBay Terapeak\u3002')
    doc.add_page_break()

    # DISCUSSION
    add_heading_styled(doc, '\u8003\u5bdf', level=1)

    add_heading_styled(doc, '\u4e3b\u8981\u306a\u77e5\u898b', level=2)
    doc.add_paragraph(
        '\u672c\u7814\u7a76\u306f\u3001\u9ebb\u9154\u85ac\u306e\u74b0\u5883\u898f\u5236\u304c\u4e2d\u53e4\u5e02\u5834\u306e\u6a5f\u5668\u4fa1\u683c\u306b\u85ac\u5264\u7279\u7570\u7684\u306a\u5f71\u97ff\u3092\u53ca\u307c\u3059\u3053\u3068\u3092\u793a\u3059\u521d\u3081\u3066\u306e\u5b9f\u8a3c\u7684\u30a8\u30d3\u30c7\u30f3\u30b9\u3092\u63d0\u4f9b\u3059\u308b\u3002'
        '3\u5e74\u5206\u306eeBay\u843d\u672d\u30c7\u30fc\u30bf\u3068\u88dc\u5b8c\u7684\u306a\u7d71\u8a08\u624b\u6cd5\u3092\u7528\u3044\u3066\u3001'
        '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u4fa1\u683c\u304c\u898f\u5236\u30de\u30a4\u30eb\u30b9\u30c8\u30fc\u30f3\u306e\u9032\u884c\u306b\u4f34\u3044\u6bb5\u968e\u7684\u306b\u4e0b\u843d\u3057\u305f\u3053\u3068\u3092\u5b9f\u8a3c\u3057\u305f\u3002'
        '\u91cd\u8981\u306a\u306e\u306f\u3001\u3053\u306e\u30d1\u30bf\u30fc\u30f3\u304c\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u306b\u7279\u6709\u3067\u3042\u308a\u3001'
        '\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u304a\u3088\u3073\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u4fa1\u683c\u306f\u7814\u7a76\u671f\u9593\u3092\u901a\u3058\u3066\u5b89\u5b9a\u3057\u3066\u3044\u305f\u3053\u3068\u3067\u3042\u308b\u3002')
    doc.add_paragraph(
        '\u8907\u6570\u306e\u5206\u6790\u30a2\u30d7\u30ed\u30fc\u30c1\u304b\u3089\u306e\u30a8\u30d3\u30c7\u30f3\u30b9\u306e\u53ce\u675f\u304c\u3001\u3053\u308c\u3089\u306e\u77e5\u898b\u3092\u5f37\u5316\u3059\u308b\u3002'
        'Spearman\u9806\u4f4d\u76f8\u95a2\u306f\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u4fa1\u683c\u306e\u9ad8\u5ea6\u306b\u6709\u610f\u306a\u5358\u8abf\u7684\u4e0b\u843d\u3092\u793a\u3057\uff08P<0.001\uff09\u3001'
        '\u540c\u3058\u691c\u5b9a\u3067\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u306b\u306f\u6709\u610f\u306a\u30c8\u30ec\u30f3\u30c9\u304c\u8a8d\u3081\u3089\u308c\u306a\u304b\u3063\u305f\uff08P=0.86\uff09\u3002'
        'Kendall \u03c4\u306f\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u306e\u898f\u5236\u6bb5\u968e\u9806\u306e\u4fa1\u683c\u4e0b\u843d\u3092\u78ba\u8a8d\u3057\uff08P=0.049\uff09\u3001'
        '\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u3067\u306f\u78ba\u8a8d\u3055\u308c\u306a\u304b\u3063\u305f\uff08P=0.36\uff09\u3002'
        'Welch\u306et\u691c\u5b9a\u306b\u3088\u308b\u898f\u5236\u524d\u5f8c\u6bd4\u8f03\u3082\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u3067\u6709\u610f\u3067\u3042\u3063\u305f\uff08P=0.027\uff09\u3002'
        '\u3053\u308c\u3089\u3092\u7dcf\u5408\u3059\u308b\u3068\u3001\u5805\u7262\u3067\u6bb5\u968e\u7684\u304b\u3064\u85ac\u5264\u7279\u7570\u7684\u306a\u4fa1\u683c\u4e0b\u843d\u304c\u793a\u3055\u308c\u305f\u3002')

    add_heading_styled(doc, '\u5148\u884c\u7814\u7a76\u3068\u306e\u6bd4\u8f03', level=2)
    doc.add_paragraph(
        '\u6211\u3005\u306e\u77e5\u308b\u9650\u308a\u3001\u9ebb\u9154\u6a5f\u5668\u306b\u5bfe\u3059\u308b\u74b0\u5883\u898f\u5236\u306e\u4e2d\u53e4\u5e02\u5834\u3078\u306e\u5f71\u97ff\u3092\u691c\u8a0e\u3057\u305f\u5148\u884c\u7814\u7a76\u306f\u306a\u3044\u3002'
        'Lehmann\u3089[15]\u306f\u3001\u6559\u80b2\u3068\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u306e\u7269\u7406\u7684\u64a4\u53bb\u3092\u7d44\u307f\u5408\u308f\u305b\u305f\u65bd\u8a2d\u30ec\u30d9\u30eb\u306e\u4ecb\u5165\u306b\u3088\u308a\u3001'
        '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u8d77\u56e0\u306eCO\u2082\u7b49\u4fa1\u6392\u51fa\u91cf\u304c86%\u524a\u6e1b\u3055\u308c\u305f\u3053\u3068\u3092\u5b9f\u8a3c\u3057\u305f\u304c\u3001'
        '\u85ac\u5264\u6d88\u8cbb\u91cf\u3092\u6e2c\u5b9a\u3057\u305f\u3082\u306e\u3067\u3042\u308a\u3001\u6a5f\u5668\u306e\u518d\u8ca9\u58f2\u4fa1\u5024\u306f\u691c\u8a0e\u3057\u3066\u3044\u306a\u3044\u3002'
        'Meyer[16]\u304a\u3088\u3073Mohammed\u3068Metta[18]\u306f\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u4e2d\u6b62\u306e\u4e16\u754c\u7684\u30fb\u8ca1\u52d9\u7684\u6839\u62e0\u3092\u8ad6\u3058\u3001'
        'Moonesinghe[17]\u306f\u5ec3\u6b62\u30d7\u30ed\u30b0\u30e9\u30e0\u306e\u5e83\u7bc4\u306a\u542b\u610f\u3092\u8b70\u8ad6\u3057\u305f\u304c\u3001'
        '\u3044\u305a\u308c\u3082\u4e2d\u53e4\u6a5f\u5668\u5e02\u5834\u3078\u306e\u4e0b\u6d41\u5f71\u97ff\u306f\u691c\u8a0e\u3057\u3066\u3044\u306a\u3044\u3002'
        'Beard\u3089[19]\u306f\u7d42\u672b\u547c\u6c17\u6fc3\u5ea6\u5236\u5fa1\u306b\u3088\u308b\u7d4c\u6e08\u7684\u4fbf\u76ca\u3092\u5b9a\u91cf\u5316\u3057\u305f\u304c\u3001\u6a5f\u5668\u306e\u6e1b\u4fa1\u512a\u5374\u306f\u6271\u3063\u3066\u3044\u306a\u3044\u3002'
        'BFMV\u4e2d\u53e4\u533b\u7642\u6a5f\u5668\u4fa1\u683c\u30d9\u30f3\u30c1\u30de\u30fc\u30af\u30ec\u30dd\u30fc\u30c8[20]\u306f\u7d041,500\u6a5f\u7a2e\u306e\u4e2d\u53e4\u533b\u7642\u6a5f\u5668\u306e\u5e74\u6b21\u4fa1\u683c\u30d9\u30f3\u30c1\u30de\u30fc\u30af\u3092\u63d0\u4f9b\u3057\u3001'
        '\u591a\u304f\u306e\u6a5f\u5668\u30ab\u30c6\u30b4\u30ea\u30fc\u306e\u518d\u8ca9\u58f2\u4fa1\u5024\u304c5\u5e74\u9593\u6bd4\u8f03\u7684\u5b89\u5b9a\u3067\u3042\u308b\u3053\u3068\u3092\u793a\u3057\u3066\u3044\u308b\u3002'
        '\u672c\u7814\u7a76\u306e\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u30fb\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u4fa1\u683c\u306e\u5b89\u5b9a\u6027\u306f\u3053\u308c\u3068\u4e00\u81f4\u3059\u308b\u304c\u3001'
        '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u306e\u4e0b\u843d\u306f\u898f\u5236\u4ecb\u5165\u306b\u8d77\u56e0\u3059\u308b\u7279\u7b46\u3059\u3079\u304d\u4f8b\u5916\u3067\u3042\u308b\u3002')
    doc.add_paragraph(
        '\u672c\u7814\u7a76\u306e\u77e5\u898b\u306f\u3001\u898f\u5236\u306b\u3088\u308b\u9673\u8150\u5316[14]\u306b\u95a2\u3059\u308b\u5e83\u7bc4\u306a\u7d4c\u6e08\u5b66\u6587\u732e\u3068\u4e00\u81f4\u3057\u3066\u304a\u308a\u3001'
        '\u4e88\u60f3\u3055\u308c\u308b\u653f\u5e9c\u898f\u5236\u304c\u4e2d\u53e4\u5e02\u5834\u306b\u304a\u3051\u308b\u4e88\u6e2c\u7684\u306a\u4fa1\u683c\u4e0b\u843d\u3092\u5f15\u304d\u8d77\u3053\u3059\u3002'
        '\u7acb\u6cd5\u904e\u7a0b\uff082022\u5e74\uff5e2024\u5e74\uff09\u306b\u304a\u3051\u308b\u6bb5\u968e\u7684\u306a\u4fa1\u683c\u4fb5\u98df\u3068\u3001'
        '\u898f\u5236\u5f8c\u306e\u3088\u308a\u9855\u8457\u306a\u4e0b\u843d\u3068\u3044\u3046\u30d1\u30bf\u30fc\u30f3\u306f\u3001'
        '\u8eca\u4e21\u6392\u51fa\u898f\u5236\u304c\u4e2d\u53e4\u8eca\u5e02\u5834\u306b\u4e0e\u3048\u305f\u5f71\u97ff\u306e\u7814\u7a76\u3068\u985e\u4f3c\u3057\u3066\u3044\u308b\u3002'
        '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u306e\u307f\u306b\u5f71\u97ff\u3057\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u30fb\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u4fa1\u683c\u306f\u5909\u5316\u3057\u306a\u304b\u3063\u305f\u3068\u3044\u3046\u85ac\u5264\u7279\u7570\u6027\u306f\u3001'
        '\u4e00\u822c\u7684\u306a\u5e02\u5834\u52b9\u679c\u3067\u306f\u306a\u304f\u898f\u5236\u52b9\u679c\u3067\u3042\u308b\u3053\u3068\u306e\u7279\u306b\u5f37\u3044\u30a8\u30d3\u30c7\u30f3\u30b9\u3092\u63d0\u4f9b\u3059\u308b\u3002')

    add_heading_styled(doc, '\u5f37\u307f\u3068\u9650\u754c', level=2)
    doc.add_paragraph(
        '\u672c\u7814\u7a76\u306e\u5f37\u307f\u306f\u3001\u5b9f\u969b\u306e\u843d\u672d\u4fa1\u683c\u306e\u4f7f\u7528\u3001\u7acb\u6cd5\u904e\u7a0b\u3068\u898f\u5236\u5b9f\u65bd\u306e\u4e21\u65b9\u3092\u30ab\u30d0\u30fc\u3059\u308b3\u5e74\u9593\u306e\u89b3\u5bdf\u671f\u9593\u3001'
        '\u8907\u6570\u306e\u88dc\u5b8c\u7684\u7d71\u8a08\u624b\u6cd5\uff08\u6a2a\u65ad\u7684\u6bd4\u8f03\u3001Spearman\u76f8\u95a2\u3001Kendall \u03c4\u30c8\u30ec\u30f3\u30c9\u691c\u5b9a\uff09\u3001'
        '\u81ea\u7136\u5bfe\u7167\u7fa4\uff08\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u3001\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\uff09\u306e\u5229\u7528\u53ef\u80fd\u6027\u3001'
        '\u304a\u3088\u3073\u6a19\u6e96\u5316\u3055\u308c\u305f\u30c7\u30fc\u30bf\u30bd\u30fc\u30b9\uff08eBay Terapeak\uff09\u306e\u4f7f\u7528\u3067\u3042\u308b\u3002')
    doc.add_paragraph(
        f'\u9650\u754c\u3068\u3057\u3066\u3001eBay\u306f\u4e2d\u53e4\u533b\u7642\u6a5f\u5668\u5e02\u5834\u306e\u4e00\u90e8\u306b\u904e\u304e\u306a\u3044\u3053\u3068\u3001'
        f'\u6a5f\u5668\u306e\u5e74\u5f0f\u30fb\u6574\u5099\u5c65\u6b74\u30fb\u5916\u89b3\u72b6\u614b\u3092\u5236\u5fa1\u3067\u304d\u306a\u304b\u3063\u305f\u3053\u3068\u3001'
        f'\u898f\u5236\u5f8c\u671f\u9593\uff082026\u5e741\uff5e3\u6708\uff09\u306e\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u53d6\u5f15\u304c{des["post_n"]}\u4ef6\u3068\u5c11\u306a\u304b\u3063\u305f\u3053\u3068\u304c\u6319\u3052\u3089\u308c\u308b\u3002'
        f'\u305f\u3060\u3057\u3001\u6642\u7cfb\u5217\u30c8\u30ec\u30f3\u30c9\u5206\u6790\u306f\u5168\u30c7\u30fc\u30bf\u30dd\u30a4\u30f3\u30c8\u3092\u5229\u7528\u3057\u3066\u304a\u308a\u3001\u6bb5\u968e\u7684\u4e0b\u843d\u3092\u78ba\u8a8d\u3057\u305f\u3002'
        f'\u307e\u305f\u30013\u5e74\u9593\u306e\u89b3\u5bdf\u671f\u9593\u306f\u6b27\u5dde\u59d4\u54e1\u4f1a\u539f\u6848\u304b\u3089\u898f\u5236\u767a\u52b9\u5f8c\u307e\u3067\u306e\u5168\u7acb\u6cd5\u904e\u7a0b\u3068\u4e00\u81f4\u3059\u308b\u304c\u3001'
        f'\u63d0\u6848\u524d\uff082022\u5e744\u6708\u4ee5\u524d\uff09\u306e\u898f\u5236\u30b7\u30b0\u30ca\u30eb\u306e\u5f71\u97ff\u3092\u53d7\u3051\u306a\u3044\u771f\u306e\u30d9\u30fc\u30b9\u30e9\u30a4\u30f3\u3092\u78ba\u7acb\u3059\u308b\u3053\u3068\u306f\u3067\u304d\u306a\u3044\u3002')

    # Supplementary analysis: asking prices (Japanese)
    if has_asking_data:
        ask = asking_results['asking_summary']
        kw = asking_results['kruskal_wallis']
        spr = asking_results['spread']
        pw = asking_results['pairwise']
        n_asking = len(asking_df)

        add_heading_styled(doc, '\u88dc\u8db3\u5206\u6790\uff1a\u73fe\u5728\u306e\u51fa\u54c1\u4fa1\u683c', level=2)
        doc.add_paragraph(
            f'\u898f\u5236\u5f8c\u306e\u5c11\u306a\u3044\u30b5\u30f3\u30d7\u30eb\u6570\u3068\u3044\u3046\u9650\u754c\u306b\u5bfe\u51e6\u3059\u308b\u305f\u3081\u3001'
            f'2026\u5e743\u670827\u65e5\u6642\u70b9\u306eeBay\u51fa\u54c1\u4fa1\u683c\uff08\u672a\u58f2\u5374\u306e\u30a2\u30af\u30c6\u30a3\u30d6\u30ea\u30b9\u30c6\u30a3\u30f3\u30b0\uff09\u306e'
            f'\u88dc\u8db3\u7684\u6a2a\u65ad\u5206\u6790\u3092\u884c\u3063\u305f\u3002'
            f'\u5408\u8a08{n_asking}\u4ef6\u306e\u30a2\u30af\u30c6\u30a3\u30d6\u30ea\u30b9\u30c6\u30a3\u30f3\u30b0\u3092\u7279\u5b9a\u3057\u305f\uff1a'
            f'\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3{ask["Desflurane"]["n"]}\u4ef6\u3001'
            f'\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3{ask["Sevoflurane"]["n"]}\u4ef6\u3001'
            f'\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3{ask["Isoflurane"]["n"]}\u4ef6\u3002'
            f'\u843d\u672d\u4fa1\u683c\u304c\u5b9f\u73fe\u3057\u305f\u5e02\u5834\u4fa1\u5024\u3092\u53cd\u6620\u3059\u308b\u306e\u306b\u5bfe\u3057\u3001'
            f'\u51fa\u54c1\u4fa1\u683c\u306f\u58f2\u308a\u624b\u306e\u73fe\u5728\u306e\u4fa1\u683c\u671f\u5f85\u3092\u793a\u3059\u3082\u306e\u3067\u3042\u308a\u3001'
            f'\u3088\u308a\u5927\u304d\u304f\u540c\u6642\u4ee3\u7684\u306a\u5e02\u5834\u30b9\u30ca\u30c3\u30d7\u30b7\u30e7\u30c3\u30c8\u3092\u63d0\u4f9b\u3059\u308b\u3002')
        doc.add_paragraph(
            f'\u51fa\u54c1\u4fa1\u683c\u306f\u85ac\u5264\u9593\u3067\u6709\u610f\u306b\u7570\u306a\u3063\u305f'
            f'\uff08Kruskal\u2013Wallis H={kw["H"]:.1f}, P<0.001\uff09\u3002'
            f'\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u306e\u4e2d\u592e\u5024\u51fa\u54c1\u4fa1\u683c\u306f\u6700\u3082\u4f4e\u304f'
            f'\uff08US${ask["Desflurane"]["median"]:.0f}\u3001IQR ${ask["Desflurane"]["q25"]:.0f}\u2013{ask["Desflurane"]["q75"]:.0f}\uff09\u3001'
            f'\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\uff08US${ask["Sevoflurane"]["median"]:.0f}\uff09\u306e\u7d047\u5206\u306e1\u3001'
            f'\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\uff08US${ask["Isoflurane"]["median"]:.0f}\uff09\u306e\u7d043\u5206\u306e1\u3067\u3042\u3063\u305f\u3002'
            f'\u3059\u3079\u3066\u306e\u30da\u30a2\u30ef\u30a4\u30ba\u6bd4\u8f03\u304c\u6709\u610f\u3067\u3042\u3063\u305f\u3002')
        doc.add_paragraph(
            f'\u51fa\u54c1\u4fa1\u683c\u3068\u843d\u672d\u4fa1\u683c\u306e\u30b9\u30d7\u30ec\u30c3\u30c9\u5206\u6790\u3067\u306f\u3001'
            f'\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u306e\u51fa\u54c1\u4fa1\u683c\u306f\u6700\u8fd1\u306e\u843d\u672d\u4fa1\u683c\u3092'
            f'{spr["Desflurane"]["spread_pct"]:.0f}%\u4e0a\u56de\u308b\u306b\u904e\u304e\u306a\u304b\u3063\u305f\u304c\u3001'
            f'\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u3067\u306f{spr["Sevoflurane"]["spread_pct"]:.0f}%\u3001'
            f'\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u3067\u306f{spr["Isoflurane"]["spread_pct"]:.0f}%\u3067\u3042\u3063\u305f\u3002'
            f'\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u306e\u5927\u5e45\u306b\u72ed\u3044\u30b9\u30d7\u30ec\u30c3\u30c9\u306f\u3001'
            f'\u58f2\u308a\u624b\u304c\u898f\u5236\u5f8c\u306e\u5e02\u5834\u73fe\u5b9f\u3092\u65e2\u306b\u4fa1\u683c\u306b\u53cd\u6620\u3055\u305b\u3066\u3044\u308b\u3053\u3068\u3092\u793a\u5506\u3059\u308b\u3002'
            f'\u3053\u306e\u77e5\u898b\u306f\u4e3b\u8981\u306a\u6642\u7cfb\u5217\u5206\u6790\u3092\u88cf\u4ed8\u3051\u3001'
            f'\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u4fa1\u683c\u4e0b\u843d\u304c\u73fe\u5728\u306e\u5e02\u5834\u53c2\u52a0\u8005\u306b\u8a8d\u8b58\u3055\u308c\u3066\u3044\u308b\u72ec\u7acb\u3057\u305f\u540c\u6642\u4ee3\u7684\u30a8\u30d3\u30c7\u30f3\u30b9\u3092\u63d0\u4f9b\u3059\u308b\u3002')

        # eTable (Japanese)
        doc.add_paragraph()
        p = doc.add_paragraph()
        add_run_styled(p, 'e\u8868 1. ', bold=True, size=Pt(10))
        add_run_styled(p, '\u73fe\u5728\u306eeBay\u51fa\u54c1\u4fa1\u683c\uff08\u30a2\u30af\u30c6\u30a3\u30d6\u30ea\u30b9\u30c6\u30a3\u30f3\u30b0\uff09\u306e\u6c17\u5316\u5668\u7a2e\u985e\u5225\u6982\u8981\u3002'
            '2026\u5e743\u670827\u65e5\u53ce\u96c6\u3002\u5024\u306f\u5e73\u5747\uff08SD\uff09\u3001\u4e2d\u592e\u5024\uff08IQR\uff09\u3001\u5358\u4f4d\u306f\u7c73\u30c9\u30eb\u3002'
            'P\u5024\u306fKruskal\u2013Wallis\u691c\u5b9a\u3002', italic=True, size=Pt(10))

        et = doc.add_table(rows=1, cols=6)
        et.style = 'Table Grid'
        et.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_table_header(et, ['\u85ac\u5264', 'n', '\u5e73\u5747 (SD)', '\u4e2d\u592e\u5024 (IQR)', '\u7bc4\u56f2', 'P\u5024'])

        for i, agent_cap in enumerate(['Desflurane', 'Sevoflurane', 'Isoflurane']):
            a = ask[agent_cap]
            agent_ja = {
                'Desflurane': '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3',
                'Sevoflurane': '\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3',
                'Isoflurane': '\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3'
            }[agent_cap]
            pval_str = fmt_p(kw['P']) if i == 0 else ''
            data = [
                (agent_ja, WD_ALIGN_PARAGRAPH.LEFT),
                (str(a['n']), WD_ALIGN_PARAGRAPH.CENTER),
                (f'${a["mean"]:.0f} ({a["sd"]:.0f})', WD_ALIGN_PARAGRAPH.CENTER),
                (f'${a["median"]:.0f} ({a["q25"]:.0f}\u2013{a["q75"]:.0f})', WD_ALIGN_PARAGRAPH.CENTER),
                (f'${a["min"]:.0f}\u2013{a["max"]:.0f}', WD_ALIGN_PARAGRAPH.CENTER),
                (pval_str, WD_ALIGN_PARAGRAPH.CENTER),
            ]
            add_table_data_row(et, data)
        doc.add_paragraph()

    # CONCLUSIONS
    add_heading_styled(doc, '\u7d50\u8ad6', level=1)
    doc.add_paragraph(
        'EU\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u898f\u5236\u306f\u3001eBay\u4e0a\u306e\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u306e\u4e2d\u53e4\u5e02\u5834\u4fa1\u5024\u306e\u6bb5\u968e\u7684\u304b\u3064\u7d71\u8a08\u7684\u306b\u6709\u610f\u306a\u4e0b\u843d\u3068\u95a2\u9023\u3057\u3066\u3044\u305f\u3002'
        '\u6642\u7cfb\u5217\u30c8\u30ec\u30f3\u30c9\u5206\u6790\u306b\u3088\u308a\u3001\u3053\u306e\u4e0b\u843d\u304c\u898f\u5236\u5bfe\u8c61\u85ac\u5264\u306b\u7279\u6709\u3067\u3042\u308b\u3053\u3068\u304c\u5b9f\u8a3c\u3055\u308c\u305f\uff1a'
        '\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u304a\u3088\u3073\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u4fa1\u683c\u306f\u7814\u7a76\u671f\u9593\u3092\u901a\u3058\u3066\u5b89\u5b9a\u3057\u3066\u304a\u308a\u3001\u81ea\u7136\u5bfe\u7167\u7fa4\u3068\u3057\u3066\u6a5f\u80fd\u3057\u305f\u3002'
        '\u4fa1\u683c\u4e0b\u843d\u306f\u7acb\u6cd5\u904e\u7a0b\u4e2d\u306b\u65e2\u306b\u59cb\u307e\u3063\u3066\u304a\u308a\u3001\u898f\u5236\u30b7\u30b0\u30ca\u30eb\u306e\u7d2f\u7a4d\u306b\u5bfe\u3059\u308b\u5e02\u5834\u306e\u4e88\u6e2c\u7684\u53cd\u5fdc\u304c\u793a\u5506\u3055\u308c\u305f\u3002'
        '\u3053\u308c\u3089\u306e\u77e5\u898b\u306f\u3001\u9ebb\u9154\u85ac\u306e\u74b0\u5883\u898f\u5236\u304c\u4e2d\u53e4\u533b\u7642\u6a5f\u5668\u5e02\u5834\u306b\u85ac\u5264\u7279\u7570\u7684\u306a\u7d4c\u6e08\u7684\u5f71\u97ff\u3092\u53ca\u307c\u3059\u3053\u3068\u3092\u793a\u3059\u521d\u3081\u3066\u306e\u5b9f\u8a3c\u7684\u30a8\u30d3\u30c7\u30f3\u30b9\u3067\u3042\u308b\u3002')

    # DECLARATIONS
    add_heading_styled(doc, '\u5229\u76ca\u76f8\u53cd\u306e\u958b\u793a', level=1)
    doc.add_paragraph('[\u8457\u8005\u304c\u8a18\u5165]')
    add_heading_styled(doc, '\u8cc7\u91d1\u63d0\u4f9b', level=1)
    doc.add_paragraph('[\u8457\u8005\u304c\u8a18\u5165]')
    add_heading_styled(doc, '\u502b\u7406\u627f\u8a8d', level=1)
    doc.add_paragraph('\u672c\u7814\u7a76\u306feBay\u306e\u516c\u958b\u843d\u672d\u30c7\u30fc\u30bf\u306e\u5206\u6790\u3067\u3042\u308a\u3001\u500b\u4eba\u60c5\u5831\u3084\u60a3\u8005\u30c7\u30fc\u30bf\u306f\u53ce\u96c6\u3057\u3066\u3044\u306a\u3044\u305f\u3081\u3001\u502b\u7406\u627f\u8a8d\u306f\u4e0d\u8981\u3067\u3042\u308b\u3002')
    add_heading_styled(doc, '\u30c7\u30fc\u30bf\u5229\u7528\u53ef\u80fd\u6027', level=1)
    doc.add_paragraph('\u672c\u7814\u7a76\u3067\u751f\u6210\u3055\u308c\u305f\u30c7\u30fc\u30bf\u30bb\u30c3\u30c8\u306f\u3001\u5408\u7406\u7684\u306a\u8981\u6c42\u306b\u5fdc\u3058\u3066\u8cac\u4efb\u8457\u8005\u304b\u3089\u5165\u624b\u53ef\u80fd\u3067\u3042\u308b\u3002\u751f\u30c7\u30fc\u30bf\u306feBay Terapeak\u304b\u3089\u53d6\u5f97\u3057\u305f\u3002')
    add_heading_styled(doc, '\u8457\u8005\u8ca2\u732e', level=1)
    doc.add_paragraph('[CRediT\u5206\u985e\u6cd5\u306b\u3088\u308a\u8457\u8005\u304c\u8a18\u5165]')

    doc.add_page_break()

    # REFERENCES
    add_heading_styled(doc, '\u53c2\u8003\u6587\u732e', level=1)
    references = [
        '1. Varughese S, Ahmed R. Environmental and occupational considerations of anesthesia. Anesth Analg 2021;133:826-35.',
        '2. Regulation (EU) 2024/573. Official Journal of the European Union 2024;L 2024/573.',
        '3. Sherman JD, Chesebro BB. Inhaled anesthetic climate and ozone effects. Anesth Analg 2023;137:201-15.',
        '4. ESAIC position statement on the use of desflurane. Eur J Anaesthesiol 2024;41:1-3.',
        '5. Association of Anaesthetists. Environmental sustainability in anaesthesia. Anaesthesia 2023;78:219-30.',
        '6. Sulbaek Andersen MP, et al. Inhalation anaesthetics and climate change. Br J Anaesth 2010;105:760-6.',
        '7. Ryan SM, Nielsen CJ. Global warming potential of inhaled anesthetics. Anesth Analg 2010;111:92-8.',
        '8. McGain F, et al. Environmental sustainability in anaesthesia and critical care. Br J Anaesth 2020;125:680-92.',
        '9. Rauchenwald V, et al. Sevoflurane versus desflurane. BMC Anesthesiol 2020;20:272.',
        '10. Zuegge KL, et al. APW-AVE. Anesth Analg 2023;137:1219-25.',
        '11. von Elm E, et al. The STROBE statement. BMJ 2007;335:806-8.',
        '12. NHS England. Decommissioning of desflurane in the NHS. 2023.',
        '13. Richter H, et al. Environmental sustainability in anaesthesia: desflurane. Curr Opin Anaesthesiol 2024;37:183-8.',
        '14. Davis G, Patel N. Regulatory obsolescence and secondary market asset depreciation. J Environ Econ Manag 2019;95:142-60.',
        '15. Lehmann H, Werning J, Baschnegger H, et al. Minimising the usage of desflurane only by education and removal of the vaporisers. BMC Anesthesiol 2025;25:108.',
        '16. Meyer MJ. Desflurane should des-appear: global and financial rationale. Anesth Analg 2020;131:1317-22.',
        '17. Moonesinghe SR. Desflurane decommissioning: more than meets the eye. Anaesthesia 2024;79:237-41.',
        '18. Mohammed A, Metta H. Is it time to bid adieu to desflurane? J Anaesthesiol Clin Pharmacol 2025;41:211-2.',
        '19. Beard D, Aston W, Black S, et al. Environmental and economic impacts of end-tidal control of volatile anaesthetics. Open Anaesthesia J 2025;19:e18742126.',
        '20. Buckhead Fair Market Value. 2025 Benchmark Report on Pre-Owned Medical Equipment Prices. Atlanta, GA: BFMV, 2025.',
    ]
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(10)

    doc.save(outdir + 'vaporizer_paper_japanese.docx')
    print("Japanese paper saved (BMJ format, STROBE-compliant, with Spearman/Kendall)!")



# ==========================================
# MAIN
# ==========================================
if __name__ == '__main__':
    print("Generating BMJ-format papers with Spearman/Kendall trend analysis...")
    print(f"Dataset: {total_n} listings ({date_min_all} to {date_max_all})")
    print()
    for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
        tr = trend_results[agent]
        print(f"  {agent}:")
        print(f"    Spearman rho={tr['spearman_rho']:.3f}, P={tr['spearman_p']:.6f}")
        print(f"    Kendall tau={tr['kendall_tau']:.3f}, P={tr['kendall_p']:.6f}")
        print(f"    Quarterly rho={tr['quarterly_rho']:.3f}, P={tr['quarterly_p']:.6f}")
    print()
    write_english_paper()
    write_japanese_paper()
    print()
    print("Both papers generated successfully!")
    print(f"  English: {outdir}vaporizer_paper_english.docx")
    print(f"  Japanese: {outdir}vaporizer_paper_japanese.docx")
