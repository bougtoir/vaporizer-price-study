"""
Generate EJA (European Journal of Anaesthesiology) cover letter as editable .docx file.
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
out_dir = os.path.join(SCRIPT_DIR, 'papers')
os.makedirs(out_dir, exist_ok=True)


def setup_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing = 1.15
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    return doc


def add_para(doc, text, size=Pt(11), bold=False, italic=False,
             alignment=None, space_after=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    return p


def write_eja_cover_letter():
    doc = setup_doc()

    # Date
    add_para(doc, '[Date]', space_after=Pt(12))

    # Addressee
    add_para(doc, 'The Editor-in-Chief', bold=True, space_after=Pt(0))
    add_para(doc, 'European Journal of Anaesthesiology & Intensive Care', italic=True,
            space_after=Pt(0))
    add_para(doc, '', space_after=Pt(12))

    # Subject line
    p = doc.add_paragraph()
    run = p.add_run('Re: ')
    run.font.size = Pt(11)
    run.bold = True
    run = p.add_run(
        'Submission of original article \u2013 '
        '\u201cTargeted environmental regulation without collateral market damage: '
        'the EU desflurane ban and secondary market vaporiser prices\u201d')
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(12)

    # Salutation
    add_para(doc, 'Dear Editor,', space_after=Pt(12))

    # Body paragraphs
    add_para(doc,
        'We are pleased to submit the above manuscript for consideration for publication '
        'in the European Journal of Anaesthesiology & Intensive Care as an original article.',
        space_after=Pt(8))

    add_para(doc,
        'Environmental regulations targeting specific anaesthetic agents are increasing '
        '\u2014 the EU desflurane ban, ASA recommendations on nitrous oxide deactivation, '
        'and NHS decommissioning programmes are recent examples. A common concern is that '
        'restricting a single agent could destabilise the broader equipment market. Yet '
        'whether such targeted regulation actually produces collateral economic effects on '
        'non-regulated equipment has never been empirically examined.',
        space_after=Pt(8))

    add_para(doc,
        'In this study, we analysed 1,033 completed eBay sales of anaesthetic vaporisers '
        '(desflurane, sevoflurane and isoflurane) over three years, spanning the full legislative '
        'trajectory from the EC proposal through to post-ban implementation. Using complementary '
        'statistical approaches (Spearman rank correlation, Kendall \u03c4 trend test, '
        'Mann\u2013Whitney U and between-agent effect size comparison), we demonstrate that '
        'the EU desflurane ban achieved targeted economic effects: desflurane vaporiser prices '
        'declined progressively (P<0.001), while sevoflurane and isoflurane vaporiser prices '
        'remained completely stable. The between-agent effect size comparison confirmed this '
        'specificity (P=0.043), and the two non-regulated agents were indistinguishable from '
        'each other (P=0.17).',
        space_after=Pt(8))

    add_para(doc,
        'The central message is reassuring: the regulation operated surgically, affecting '
        'only its intended target without collateral damage to the wider anaesthetic equipment '
        'market. Additionally, the price decline began during the legislative process itself, '
        'suggesting that well-signalled regulation generates predictable and orderly market '
        'adjustments. Early compliance was associated with better cost recovery, potentially '
        'freeing capital for reinvestment in alternative equipment.',
        space_after=Pt(8))

    add_para(doc,
        'We believe this work is particularly well suited for EJAIC because: (1) it provides '
        'empirical evidence that ESAIC\u2019s advocacy for the desflurane ban produced '
        'precisely the targeted outcome intended\u2014without destabilising the wider '
        'market; (2) as further environmental restrictions are anticipated (e.g. nitrous '
        'oxide), these findings offer a reassuring precedent for the European anaesthesia '
        'community; (3) the actionable implications for equipment management during '
        'regulatory transitions are directly relevant to EJAIC\u2019s clinical readership; '
        'and (4) the natural experiment design, with non-regulated agents as controls, '
        'provides a rigorous methodological framework applicable to future regulatory '
        'impact assessments.',
        space_after=Pt(8))

    add_para(doc,
        'This manuscript has not been published elsewhere and is not under consideration by '
        'another journal. All authors have read and approved the final manuscript and meet '
        'ICMJE authorship criteria. There are no conflicts of interest to declare. '
        'No ethical approval was required as the study analysed publicly available, anonymised '
        'marketplace data with no human participant involvement.',
        space_after=Pt(8))

    add_para(doc,
        'The study is reported in accordance with the STROBE guidelines for cross-sectional '
        'studies. The completed STROBE checklist is provided as supplementary material.',
        space_after=Pt(8))

    add_para(doc,
        'Thank you for considering this manuscript. We look forward to your response.',
        space_after=Pt(12))

    # Closing
    add_para(doc, 'Yours sincerely,', space_after=Pt(24))

    add_para(doc, '[Corresponding author name]', bold=True, space_after=Pt(0))
    add_para(doc, '[Department]', space_after=Pt(0))
    add_para(doc, '[Institution]', space_after=Pt(0))
    add_para(doc, '[Address]', space_after=Pt(0))
    add_para(doc, '[Email]', space_after=Pt(0))

    path = os.path.join(out_dir, 'eja_cover_letter.docx')
    doc.save(path)
    print(f"EJA cover letter saved: {path}")
    return path


if __name__ == '__main__':
    write_eja_cover_letter()
