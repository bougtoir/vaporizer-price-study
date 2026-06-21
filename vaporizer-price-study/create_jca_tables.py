"""
Generate standalone editable .docx files for each table (JCA submission).
Adapted from create_eja_tables.py with American English spelling.
"""
import os
import pandas as pd
import numpy as np
from scipy import stats as sp_stats
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
out_dir = os.path.join(SCRIPT_DIR, 'papers')
os.makedirs(out_dir, exist_ok=True)

DATA_DIR = os.path.join(SCRIPT_DIR, 'data')


def fmt_p(p_val):
    if p_val < 0.001:
        return '<0.001'
    return f'{p_val:.3f}'


def setup_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    style.paragraph_format.line_spacing = 2.0
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    return doc


def add_run_styled(p, text, bold=False, italic=False, size=Pt(10)):
    run = p.add_run(text)
    run.font.size = size
    run.font.name = 'Times New Roman'
    run.bold = bold
    run.italic = italic
    return run


def add_table_header(table, headers):
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table_data_row(table, data):
    row = table.add_row()
    for i, (text, align) in enumerate(data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        p.alignment = align


BAN_DATE = pd.Timestamp('2026-01-01')
PROPOSAL_DATE = pd.Timestamp('2022-04-05')
AGREEMENT_DATE = pd.Timestamp('2023-10-05')
ADOPTION_DATE = pd.Timestamp('2024-02-07')


def classify_period(date):
    d = pd.Timestamp(date)
    if d < PROPOSAL_DATE:
        return 1
    elif d < AGREEMENT_DATE:
        return 2
    elif d < ADOPTION_DATE:
        return 3
    elif d < BAN_DATE:
        return 4
    else:
        return 5


def load_data():
    df = pd.read_csv(os.path.join(DATA_DIR, 'combined_cleaned.csv'))
    df['date_sold'] = pd.to_datetime(df['date_sold'])
    df['period'] = np.where(df['date_sold'] < BAN_DATE, 'pre', 'post')
    df['period_num'] = df['date_sold'].apply(classify_period)
    return df


def compute_trend(df):
    trend_results = {}
    for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
        ad = df[df['agent_type'] == agent].copy()
        ad['date_ord'] = ad['date_sold'].map(lambda d: d.toordinal())
        sp_rho, sp_p = sp_stats.spearmanr(ad['date_ord'], ad['price_usd'])
        kt_tau, kt_p = sp_stats.kendalltau(ad['period_num'], ad['price_usd'])
        ad['quarter'] = ad['date_sold'].dt.to_period('Q')
        qm = ad.groupby('quarter')['price_usd'].agg(['median', 'count'])
        qm = qm[qm['count'] >= 3]
        if len(qm) >= 4:
            q_rho, q_p = sp_stats.spearmanr(range(len(qm)), qm['median'].values)
        else:
            q_rho, q_p = float('nan'), float('nan')
        trend_results[agent] = {
            'spearman_rho': sp_rho, 'spearman_p': sp_p,
            'kendall_tau': kt_tau, 'kendall_p': kt_p,
            'quarterly_rho': q_rho, 'quarterly_p': q_p,
        }
    return trend_results


def compute_effect_sizes(df):
    stats_csv = pd.read_csv(os.path.join(DATA_DIR, 'statistics_summary.csv'), index_col=0)
    effect_sizes = {}
    for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
        ad = df[df['agent_type'] == agent]
        n1 = len(ad[ad['period'] == 'pre'])
        n2 = len(ad[ad['period'] == 'post'])
        d = float(stats_csv.loc[agent, 'cohens_d'])
        se_d = np.sqrt((n1 + n2) / (n1 * n2) + d**2 / (2 * (n1 + n2)))
        ci_lo = d - 1.96 * se_d
        ci_hi = d + 1.96 * se_d
        effect_sizes[agent] = {'d': d, 'se': se_d, 'ci_lo': ci_lo, 'ci_hi': ci_hi,
                                'n1': n1, 'n2': n2}
    return effect_sizes


def compute_es_comparisons(effect_sizes):
    es_comparisons = {}
    for a1, a2 in [('Desflurane', 'Sevoflurane'), ('Desflurane', 'Isoflurane'),
                   ('Sevoflurane', 'Isoflurane')]:
        e1, e2 = effect_sizes[a1], effect_sizes[a2]
        diff = e1['d'] - e2['d']
        se = np.sqrt(e1['se']**2 + e2['se']**2)
        z = diff / se if se > 0 else 0
        p = 2 * (1 - sp_stats.norm.cdf(abs(z)))
        es_comparisons[f'{a1}_vs_{a2}'] = {'diff': diff, 'se': se, 'z': z, 'p': p}
    return es_comparisons


def get_pval(df, agent, pval_type):
    ad = df[df['agent_type'] == agent]
    pre = ad[ad['period'] == 'pre']['price_usd']
    post = ad[ad['period'] == 'post']['price_usd']
    if pval_type == 'u_pval':
        _, p = sp_stats.mannwhitneyu(pre, post, alternative='two-sided')
    else:
        _, p = sp_stats.ttest_ind(pre, post, equal_var=False)
    return p


def write_table1(df):
    """Table 1: Time-series trend analysis."""
    doc = setup_doc()
    trend_results = compute_trend(df)

    p = doc.add_paragraph()
    add_run_styled(p, 'Table 1. ', bold=True, size=Pt(10))
    add_run_styled(p, ('Time-series trend analysis of vaporizer prices by agent type. Spearman rank '
                       'correlation tests monotonic association between sale date and price; '
                       'Kendall \u03c4 tests association between ordered regulatory phase and price.'),
                   italic=True, size=Pt(10))

    t1 = doc.add_table(rows=1, cols=7)
    t1.style = 'Table Grid'
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header(t1, ['Agent', 'Spearman \u03c1', 'P value', 'Kendall \u03c4', 'P value',
                          'Quarterly \u03c1', 'P value'])

    for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
        tr = trend_results[agent]
        data = [
            (agent, WD_ALIGN_PARAGRAPH.LEFT),
            (f'{tr["spearman_rho"]:.3f}', WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(tr['spearman_p']), WD_ALIGN_PARAGRAPH.CENTER),
            (f'{tr["kendall_tau"]:.3f}', WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(tr['kendall_p']), WD_ALIGN_PARAGRAPH.CENTER),
            (f'{tr["quarterly_rho"]:.3f}', WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(tr['quarterly_p']), WD_ALIGN_PARAGRAPH.CENTER),
        ]
        add_table_data_row(t1, data)

    path = os.path.join(out_dir, 'jca_table1.docx')
    doc.save(path)
    print(f"Table 1 saved: {path}")
    return path


def write_table2(df):
    """Table 2: Pre-/post-ban comparison with effect size analysis."""
    doc = setup_doc()

    summ = {}
    for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
        ad = df[df['agent_type'] == agent]
        pre = ad[ad['period'] == 'pre']['price_usd']
        post = ad[ad['period'] == 'post']['price_usd']
        summ[agent] = {
            'pre_n': len(pre), 'post_n': len(post),
            'pre_mean': pre.mean(), 'pre_sd': pre.std(),
            'post_mean': post.mean(), 'post_sd': post.std(),
        }

    effect_sizes = compute_effect_sizes(df)
    es_comparisons = compute_es_comparisons(effect_sizes)

    p = doc.add_paragraph()
    add_run_styled(p, 'Table 2. ', bold=True, size=Pt(10))
    add_run_styled(p, ('Pre- and post-ban vaporizer prices by agent type with between-agent '
                       'effect size comparison. Values are mean \u00b1 SD in US dollars.'),
                   italic=True, size=Pt(10))

    # Panel A
    p = doc.add_paragraph()
    add_run_styled(p, 'Panel A. ', bold=True, size=Pt(9))
    add_run_styled(p, 'Descriptive statistics and within-agent effect sizes (pre- vs post-ban)',
                   italic=True, size=Pt(9))

    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header(table, ['Agent', 'n (pre/post)', 'Pre-ban mean \u00b1 SD',
                             'Post-ban mean \u00b1 SD', '% change',
                             "Cohen\u2019s d", '95% CI'])

    for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
        s = summ[agent]
        es = effect_sizes[agent]
        pct = (s['post_mean'] - s['pre_mean']) / s['pre_mean'] * 100
        data = [
            (agent, WD_ALIGN_PARAGRAPH.LEFT),
            (f'{s["pre_n"]}/{s["post_n"]}', WD_ALIGN_PARAGRAPH.CENTER),
            (f'${s["pre_mean"]:.0f} \u00b1 {s["pre_sd"]:.0f}', WD_ALIGN_PARAGRAPH.CENTER),
            (f'${s["post_mean"]:.0f} \u00b1 {s["post_sd"]:.0f}', WD_ALIGN_PARAGRAPH.CENTER),
            (f'{pct:+.1f}%', WD_ALIGN_PARAGRAPH.CENTER),
            (f'{es["d"]:.2f}', WD_ALIGN_PARAGRAPH.CENTER),
            (f'{es["ci_lo"]:.2f} to {es["ci_hi"]:.2f}', WD_ALIGN_PARAGRAPH.CENTER),
        ]
        add_table_data_row(table, data)

    # Panel B
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_run_styled(p, 'Panel B. ', bold=True, size=Pt(9))
    add_run_styled(p, 'Between-agent comparison of effect sizes (z-test for independent '
                   "Cohen\u2019s d)",
                   italic=True, size=Pt(9))

    t2b = doc.add_table(rows=1, cols=5)
    t2b.style = 'Table Grid'
    t2b.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header(t2b, ['Comparison', '\u0394d', 'SE', 'z', 'P value'])

    for key, label in [('Desflurane_vs_Sevoflurane', 'Desflurane vs Sevoflurane'),
                        ('Desflurane_vs_Isoflurane', 'Desflurane vs Isoflurane'),
                        ('Sevoflurane_vs_Isoflurane', 'Sevoflurane vs Isoflurane')]:
        c = es_comparisons[key]
        data = [
            (label, WD_ALIGN_PARAGRAPH.LEFT),
            (f'{c["diff"]:.2f}', WD_ALIGN_PARAGRAPH.CENTER),
            (f'{c["se"]:.3f}', WD_ALIGN_PARAGRAPH.CENTER),
            (f'{c["z"]:.2f}', WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(c['p']), WD_ALIGN_PARAGRAPH.CENTER),
        ]
        add_table_data_row(t2b, data)

    # Footnote
    p = doc.add_paragraph()
    add_run_styled(p, 'Within-agent P values (Mann\u2013Whitney U): ', bold=False, italic=True,
                   size=Pt(8))
    footnote_parts = []
    for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
        u_p = fmt_p(get_pval(df, agent, 'u_pval'))
        t_p = fmt_p(get_pval(df, agent, 't_pval'))
        footnote_parts.append(f'{agent} U P={u_p}, t P={t_p}')
    add_run_styled(p, '; '.join(footnote_parts) + '.', italic=True, size=Pt(8))

    path = os.path.join(out_dir, 'jca_table2.docx')
    doc.save(path)
    print(f"Table 2 saved: {path}")
    return path


def write_table_s1(df):
    """Table S1: Current eBay asking prices (Supplementary)."""
    asking_path = os.path.join(DATA_DIR, 'ebay_asking_prices.csv')
    if not os.path.exists(asking_path):
        print("No asking price data found, skipping Table S1")
        return None

    ask_df = pd.read_csv(asking_path)
    doc = setup_doc()

    p = doc.add_paragraph()
    add_run_styled(p, 'Table S1. ', bold=True, size=Pt(10))
    add_run_styled(p, ('Current eBay asking prices (active listings) by vaporizer type, '
                       'collected 27 March 2026. Values are mean (SD), median (IQR) in US dollars. '
                       'P value from Kruskal\u2013Wallis test across three agent types.'),
                   italic=True, size=Pt(10))

    agents_map = {'Desflurane': 'desflurane', 'Sevoflurane': 'sevoflurane',
                   'Isoflurane': 'isoflurane'}
    agent_col = 'agent'
    price_col = 'price_usd'

    agents = ['Desflurane', 'Sevoflurane', 'Isoflurane']
    groups = []
    ask_summary = {}
    for agent in agents:
        agent_lower = agents_map[agent]
        ad = ask_df[ask_df[agent_col] == agent_lower][price_col].dropna()
        groups.append(ad)
        ask_summary[agent] = {
            'n': len(ad), 'mean': ad.mean(), 'sd': ad.std(),
            'median': ad.median(), 'q25': ad.quantile(0.25), 'q75': ad.quantile(0.75),
            'min': ad.min(), 'max': ad.max(),
        }

    kw_stat, kw_p = sp_stats.kruskal(*groups)

    et = doc.add_table(rows=1, cols=6)
    et.style = 'Table Grid'
    et.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header(et, ['Agent', 'n', 'Mean (SD)', 'Median (IQR)', 'Range', 'P value'])

    for i, agent in enumerate(agents):
        a = ask_summary[agent]
        pval_str = fmt_p(kw_p) if i == 0 else ''
        data = [
            (agent, WD_ALIGN_PARAGRAPH.LEFT),
            (str(a['n']), WD_ALIGN_PARAGRAPH.CENTER),
            (f'${a["mean"]:.0f} ({a["sd"]:.0f})', WD_ALIGN_PARAGRAPH.CENTER),
            (f'${a["median"]:.0f} ({a["q25"]:.0f}\u2013{a["q75"]:.0f})',
             WD_ALIGN_PARAGRAPH.CENTER),
            (f'${a["min"]:.0f}\u2013{a["max"]:.0f}', WD_ALIGN_PARAGRAPH.CENTER),
            (pval_str, WD_ALIGN_PARAGRAPH.CENTER),
        ]
        add_table_data_row(et, data)

    path = os.path.join(out_dir, 'jca_table_s1.docx')
    doc.save(path)
    print(f"Table S1 saved: {path}")
    return path


if __name__ == '__main__':
    df = load_data()
    write_table1(df)
    write_table2(df)
    write_table_s1(df)
