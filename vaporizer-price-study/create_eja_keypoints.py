"""
Generate EJA Key Points as a separate editable .docx file.
EJA requires 3-5 key bullet points summarising the article, each no longer than one sentence.
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
out_dir = os.path.join(SCRIPT_DIR, 'papers')
os.makedirs(out_dir, exist_ok=True)


def write_keypoints():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    for section in doc.sections:
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(3.0)

    # Heading
    p = doc.add_paragraph()
    run = p.add_run('Key Points')
    run.bold = True
    run.font.size = Pt(12)

    key_points = [
        'The EU desflurane ban was associated with a progressive, agent-specific decline in '
        'secondary market vaporiser prices, while sevoflurane and isoflurane prices remained '
        'stable throughout the study period.',

        'Sevoflurane and isoflurane vaporiser prices were unaffected by the desflurane '
        'regulation, indicating that the ban produced no collateral damage to the broader '
        'anaesthetic equipment market.',

        'Between-agent comparison confirmed that the desflurane effect size was significantly '
        'larger than that of sevoflurane (P=0.043), while the two non-regulated agents were '
        'indistinguishable from each other (P=0.17).',

        'The desflurane price decline began during the legislative process, well before the '
        'formal prohibition date, demonstrating that the regulation was well signalled and '
        'its economic consequences were predictable.',

        'These findings provide the first empirical evidence that targeted environmental '
        'regulation of a single anaesthetic agent can achieve its objectives without '
        'destabilising the wider equipment market\u2014a reassuring precedent as further '
        'environmental restrictions are anticipated.',
    ]

    for kp in key_points:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(kp)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'

    path = os.path.join(out_dir, 'eja_key_points.docx')
    doc.save(path)
    print(f"EJA Key Points saved: {path}")
    return path


if __name__ == '__main__':
    write_keypoints()
