"""
Generate Anaesthesia-format cover letters (English and Japanese) as .docx files.
Target journal: Anaesthesia (Association of Anaesthetists, Wiley)
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

outdir = '/home/ubuntu/vaporizer_research/papers/'


def setup_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing = 1.15
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    return doc


def add_para(doc, text, size=Pt(11), bold=False, italic=False,
             alignment=None, space_after=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    return p


def write_english_cover_letter():
    doc = setup_doc()

    # Date
    add_para(doc, '[Date]', space_after=Pt(12))

    # Addressee
    add_para(doc, 'The Editor', bold=True, space_after=Pt(0))
    add_para(doc, 'Anaesthesia', italic=True, space_after=Pt(0))
    add_para(doc, 'Association of Anaesthetists', space_after=Pt(12))

    # Subject line
    add_para(doc,
             'Re: Submission of Original Article \u2013 '
             '"Impact of the European Union desflurane regulation on secondary market '
             'prices of anaesthetic vaporisers: a cross-sectional time-series analysis '
             'of eBay sold listings"',
             bold=True, space_after=Pt(12))

    # Salutation
    add_para(doc, 'Dear Editor,', space_after=Pt(6))

    # Para 1 - Introduction
    add_para(doc,
             'We are pleased to submit the above manuscript for consideration as an '
             'Original Article in Anaesthesia. This study provides the first empirical '
             'evidence that environmental regulation of an inhaled anaesthetic agent '
             'produces agent-specific economic consequences in the secondary medical '
             'equipment market.',
             space_after=Pt(6))

    # Para 2 - Summary of findings
    add_para(doc,
             'Using three years of eBay sold-listing data (n = 1,033) obtained via '
             'Terapeak, we analysed price trends for desflurane, sevoflurane, and '
             'isoflurane vaporisers spanning the entire legislative timeline of the '
             'EU F-gas Regulation (EU) 2024/573. Desflurane vaporiser prices showed '
             'a statistically significant monotonic decline (Spearman \u03c1 = \u22120.28, '
             'p < 0.001), with progressive price erosion across successive regulatory '
             'milestones confirmed by Kendall \u03c4 (\u03c4 = \u22120.12, p = 0.049). '
             'Pre-to-post-ban mean prices fell by 31% (Cohen\u2019s d = 0.55). '
             'Crucially, sevoflurane (\u03c1 = 0.01, p = 0.862) and isoflurane '
             '(\u03c1 = \u22120.086, p = 0.044) vaporiser prices remained clinically '
             'stable throughout, serving as natural comparator groups and strengthening '
             'the inference that the desflurane price decline was regulation-specific '
             'rather than attributable to broader market forces.',
             space_after=Pt(6))

    # Para 3 - Why Anaesthesia
    add_para(doc,
             'We believe this manuscript is well suited to Anaesthesia for several '
             'reasons. First, Anaesthesia has been at the forefront of the environmental '
             'sustainability agenda in anaesthetic practice, including the landmark '
             'Association of Anaesthetists position on desflurane and the broader '
             'sustainability guidance. Our findings add an economic dimension to this '
             'discourse by quantifying the downstream financial consequences of '
             'regulation on equipment owners. Second, the concept of \u2018regulatory '
             'obsolescence\u2019 \u2013 where policy action erodes the residual value of '
             'agent-specific capital equipment \u2013 is directly relevant to '
             'anaesthetists, department heads, and hospital procurement teams planning '
             'capital expenditure. Third, the study design (STROBE-compliant '
             'cross-sectional analysis with built-in comparator agents) offers a '
             'novel, reproducible methodological approach for evaluating the economic '
             'impact of environmental health policy on medical equipment markets.',
             space_after=Pt(6))

    # Para 4 - Novelty
    add_para(doc,
             'A comprehensive literature review confirmed that no prior study has '
             'quantitatively analysed the impact of inhaled anaesthetic environmental '
             'regulation on secondary market equipment values. This study therefore '
             'opens a new line of inquiry at the intersection of sustainability policy '
             'and health economics.',
             space_after=Pt(6))

    # Para 5 - Ethics and admin
    add_para(doc,
             'This is an observational study of publicly available, anonymised '
             'marketplace data; no human participants or patient data were involved, '
             'and ethical approval was therefore not required. The manuscript has not '
             'been published elsewhere and is not under consideration by any other '
             'journal. All authors have reviewed and approved the manuscript and meet '
             'the ICMJE authorship criteria. There are no conflicts of interest to '
             'declare. No external funding was received.',
             space_after=Pt(6))

    # Para 6 - Reporting guidelines
    add_para(doc,
             'The manuscript is reported in accordance with the STROBE statement for '
             'cross-sectional studies. A completed STROBE checklist is included with '
             'this submission.',
             space_after=Pt(6))

    # Para 7 - Closing
    add_para(doc,
             'We would welcome the opportunity for peer review in Anaesthesia and look '
             'forward to hearing from the editorial team.',
             space_after=Pt(12))

    # Sign-off
    add_para(doc, 'Yours sincerely,', space_after=Pt(24))

    add_para(doc, '[Corresponding author name]', bold=True, space_after=Pt(0))
    add_para(doc, '[Title / Position]', space_after=Pt(0))
    add_para(doc, '[Institution]', space_after=Pt(0))
    add_para(doc, '[Address]', space_after=Pt(0))
    add_para(doc, '[Email address]', space_after=Pt(0))
    add_para(doc, '[Telephone number]', space_after=Pt(0))

    path = outdir + 'cover_letter_english.docx'
    doc.save(path)
    print(f"English cover letter saved: {path}")


def write_japanese_cover_letter():
    doc = setup_doc()

    # Date
    add_para(doc, '[\u65e5\u4ed8]', space_after=Pt(12))

    # Addressee
    add_para(doc, '\u7de8\u96c6\u90e8\u5fa1\u4e2d', bold=True, space_after=Pt(0))
    add_para(doc, 'Anaesthesia', italic=True, space_after=Pt(0))
    add_para(doc, 'Association of Anaesthetists', space_after=Pt(12))

    # Subject line
    add_para(doc,
             '\u4ef6\u540d: \u539f\u8457\u8ad6\u6587\u6295\u7a3f \u2013 '
             '\u300cEU\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u898f\u5236\u304c\u9ebb\u9154\u6c17\u5316\u5668'
             '\u306e\u4e2d\u53e4\u5e02\u5834\u4fa1\u683c\u306b\u4e0e\u3048\u308b\u5f71\u97ff\uff1a'
             'eBay\u843d\u672d\u30c7\u30fc\u30bf\u306e\u6a2a\u65ad\u7684\u6642\u7cfb\u5217\u5206\u6790\u300d',
             bold=True, space_after=Pt(12))

    # Salutation
    add_para(doc, '\u7de8\u96c6\u59d4\u54e1\u4f1a\u5fa1\u4e2d', space_after=Pt(6))

    # Para 1
    add_para(doc,
             '\u4e0a\u8a18\u306e\u8ad6\u6587\u3092Anaesthesia\u8a8c\u306eOriginal Article'
             '\u3068\u3057\u3066\u3054\u691c\u8a0e\u3044\u305f\u3060\u304d\u305f\u304f\u3001'
             '\u6295\u7a3f\u3044\u305f\u3057\u307e\u3059\u3002\u672c\u7814\u7a76\u306f\u3001'
             '\u5438\u5165\u9ebb\u9154\u85ac\u306e\u74b0\u5883\u898f\u5236\u304c\u4e2d\u53e4'
             '\u533b\u7642\u6a5f\u5668\u5e02\u5834\u306b\u85ac\u5264\u7279\u7570\u7684\u306a'
             '\u7d4c\u6e08\u7684\u5f71\u97ff\u3092\u53ca\u307c\u3059\u3053\u3068\u3092\u793a'
             '\u3059\u521d\u3081\u3066\u306e\u5b9f\u8a3c\u7684\u30a8\u30d3\u30c7\u30f3\u30b9'
             '\u3092\u63d0\u4f9b\u3057\u307e\u3059\u3002',
             space_after=Pt(6))

    # Para 2 - Key findings
    add_para(doc,
             'eBay\u306e3\u5e74\u5206\u306e\u843d\u672d\u30c7\u30fc\u30bf\uff08n = 1,033'
             '\uff09\u3092Terapeak\u7d4c\u7531\u3067\u53d6\u5f97\u3057\u3001EU F\u30ac\u30b9'
             '\u898f\u5247\uff08EU\uff09 2024/573\u306e\u5168\u7acb\u6cd5\u904e\u7a0b\u306b'
             '\u304a\u3051\u308b\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u3001\u30bb\u30dc\u30d5'
             '\u30eb\u30e9\u30f3\u3001\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668'
             '\u306e\u4fa1\u683c\u63a8\u79fb\u3092\u5206\u6790\u3057\u307e\u3057\u305f\u3002'
             '\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u4fa1\u683c\u306f\u7d71'
             '\u8a08\u7684\u306b\u6709\u610f\u306a\u5358\u8abf\u7684\u4e0b\u843d\u3092\u793a'
             '\u3057\uff08Spearman \u03c1 = \u22120.28, p < 0.001\uff09\u3001\u898f\u5236'
             '\u30de\u30a4\u30eb\u30b9\u30c8\u30fc\u30f3\u306b\u6cbf\u3063\u305f\u6bb5\u968e'
             '\u7684\u4e0b\u843d\u304cKendall \u03c4\u3067\u78ba\u8a8d\u3055\u308c\u307e\u3057'
             '\u305f\uff08\u03c4 = \u22120.12, p = 0.049\uff09\u3002\u898f\u5236\u524d\u5f8c'
             '\u306e\u5e73\u5747\u4fa1\u683c\u306f31%\u4e0b\u843d\u3057\u307e\u3057\u305f'
             '\uff08Cohen\u306ed = 0.55\uff09\u3002\u4e00\u65b9\u3001\u30bb\u30dc\u30d5\u30eb'
             '\u30e9\u30f3\uff08\u03c1 = 0.01, p = 0.862\uff09\u304a\u3088\u3073\u30a4\u30bd'
             '\u30d5\u30eb\u30e9\u30f3\uff08\u03c1 = \u22120.086, p = 0.044\uff09\u306e\u4fa1'
             '\u683c\u306f\u81e8\u5e8a\u7684\u306b\u5b89\u5b9a\u3057\u3066\u304a\u308a\u3001'
             '\u81ea\u7136\u5bfe\u7167\u7fa4\u3068\u3057\u3066\u6a5f\u80fd\u3057\u307e\u3057\u305f\u3002',
             space_after=Pt(6))

    # Para 3 - Why Anaesthesia
    add_para(doc,
             '\u672c\u8ad6\u6587\u304cAnaesthesia\u306b\u3075\u3055\u308f\u3057\u3044\u7406'
             '\u7531\u306f\u4ee5\u4e0b\u306e\u901a\u308a\u3067\u3059\u3002\u7b2c\u4e00\u306b'
             '\u3001Anaesthesia\u8a8c\u306f\u9ebb\u9154\u5b9f\u8df5\u306b\u304a\u3051\u308b'
             '\u74b0\u5883\u6301\u7d9a\u53ef\u80fd\u6027\u306e\u8b70\u8ad6\u3092\u30ea\u30fc'
             '\u30c9\u3057\u3066\u304a\u308a\u3001Association of Anaesthetists\u306e\u30c7'
             '\u30b9\u30d5\u30eb\u30e9\u30f3\u306b\u95a2\u3059\u308b\u7acb\u5834\u8868\u660e'
             '\u3084\u5e83\u7bc4\u306a\u6301\u7d9a\u53ef\u80fd\u6027\u30ac\u30a4\u30c0\u30f3'
             '\u30b9\u3092\u767a\u8868\u3057\u3066\u3044\u307e\u3059\u3002\u672c\u7814\u7a76'
             '\u306f\u3053\u306e\u8b70\u8ad6\u306b\u7d4c\u6e08\u7684\u5074\u9762\u3092\u52a0'
             '\u3048\u307e\u3059\u3002\u7b2c\u4e8c\u306b\u3001\u300c\u898f\u5236\u306b\u3088'
             '\u308b\u9673\u8150\u5316\u300d\u306e\u6982\u5ff5\u306f\u3001\u9ebb\u9154\u79d1'
             '\u533b\u3001\u90e8\u9580\u8cac\u4efb\u8005\u3001\u75c5\u9662\u8cfc\u8cb7\u90e8'
             '\u9580\u306b\u3068\u3063\u3066\u76f4\u63a5\u95a2\u4fc2\u304c\u3042\u308a\u307e'
             '\u3059\u3002\u7b2c\u4e09\u306b\u3001STROBE\u6e96\u62e0\u306e\u6a2a\u65ad\u7684'
             '\u5206\u6790\u3068\u3044\u3046\u7814\u7a76\u30c7\u30b6\u30a4\u30f3\u306f\u3001'
             '\u533b\u7642\u898f\u5236\u306e\u7d4c\u6e08\u7684\u5f71\u97ff\u3092\u7814\u7a76'
             '\u3059\u308b\u305f\u3081\u306e\u65b0\u898f\u304b\u3064\u518d\u73fe\u53ef\u80fd'
             '\u306a\u65b9\u6cd5\u8ad6\u7684\u30a2\u30d7\u30ed\u30fc\u30c1\u3092\u63d0\u793a'
             '\u3057\u3066\u3044\u307e\u3059\u3002',
             space_after=Pt(6))

    # Para 4 - Novelty
    add_para(doc,
             '\u5305\u62ec\u7684\u306a\u6587\u732e\u30ec\u30d3\u30e5\u30fc\u306e\u7d50\u679c'
             '\u3001\u5438\u5165\u9ebb\u9154\u85ac\u306e\u74b0\u5883\u898f\u5236\u304c\u95a2'
             '\u9023\u6a5f\u5668\u306e\u4e2d\u53e4\u5e02\u5834\u306b\u4e0e\u3048\u308b\u5f71'
             '\u97ff\u3092\u5b9a\u91cf\u7684\u306b\u5206\u6790\u3057\u305f\u5148\u884c\u7814'
             '\u7a76\u306f\u5b58\u5728\u3057\u306a\u3044\u3053\u3068\u304c\u78ba\u8a8d\u3055'
             '\u308c\u307e\u3057\u305f\u3002\u672c\u7814\u7a76\u306f\u3001\u6301\u7d9a\u53ef'
             '\u80fd\u6027\u653f\u7b56\u3068\u4fdd\u5065\u533b\u7642\u7d4c\u6e08\u5b66\u306e'
             '\u63a5\u70b9\u306b\u304a\u3051\u308b\u65b0\u305f\u306a\u7814\u7a76\u9818\u57df'
             '\u3092\u958b\u304f\u3082\u306e\u3067\u3059\u3002',
             space_after=Pt(6))

    # Para 5 - Ethics and admin
    add_para(doc,
             '\u672c\u7814\u7a76\u306f\u516c\u958b\u3055\u308c\u305f\u533f\u540d\u306e\u30de'
             '\u30fc\u30b1\u30c3\u30c8\u30d7\u30ec\u30a4\u30b9\u30c7\u30fc\u30bf\u306e\u307f'
             '\u3092\u4f7f\u7528\u3057\u305f\u89b3\u5bdf\u7814\u7a76\u3067\u3042\u308a\u3001'
             '\u30d2\u30c8\u3092\u5bfe\u8c61\u3068\u305b\u305a\u3001\u60a3\u8005\u30c7\u30fc'
             '\u30bf\u3084\u500b\u4eba\u3092\u7279\u5b9a\u3067\u304d\u308b\u60c5\u5831\u3082'
             '\u542b\u307e\u308c\u3066\u3044\u306a\u3044\u305f\u3081\u3001\u502b\u7406\u5be9'
             '\u67fb\u306f\u4e0d\u8981\u3067\u3059\u3002\u672c\u8ad6\u6587\u306f\u672a\u767a'
             '\u8868\u3067\u3042\u308a\u3001\u4ed6\u8a8c\u3067\u306e\u67fb\u8aad\u3082\u53d7'
             '\u3051\u3066\u3044\u307e\u305b\u3093\u3002\u3059\u3079\u3066\u306e\u8457\u8005'
             '\u306f\u539f\u7a3f\u3092\u78ba\u8a8d\u30fb\u627f\u8a8d\u3057\u3001ICMJE\u306e'
             '\u8457\u8005\u8cc7\u683c\u57fa\u6e96\u3092\u6e80\u305f\u3057\u3066\u3044\u307e'
             '\u3059\u3002\u5229\u76ca\u76f8\u53cd\u306f\u3042\u308a\u307e\u305b\u3093\u3002'
             '\u5916\u90e8\u8cc7\u91d1\u306f\u53d7\u3051\u3066\u3044\u307e\u305b\u3093\u3002',
             space_after=Pt(6))

    # Para 6 - Reporting guidelines
    add_para(doc,
             '\u672c\u8ad6\u6587\u306f\u6a2a\u65ad\u7814\u7a76\u306eSTROBE\u58f0\u660e\u306b'
             '\u6e96\u62e0\u3057\u3066\u4f5c\u6210\u3055\u308c\u3066\u3044\u307e\u3059\u3002'
             'STROBE\u30c1\u30a7\u30c3\u30af\u30ea\u30b9\u30c8\u3092\u540c\u5c01\u3057\u307e\u3059\u3002',
             space_after=Pt(6))

    # Para 7 - Closing
    add_para(doc,
             'Anaesthesia\u3067\u306e\u67fb\u8aad\u306e\u6a5f\u4f1a\u3092\u3044\u305f\u3060'
             '\u3051\u308c\u3070\u5e78\u3044\u3067\u3059\u3002\u7de8\u96c6\u90e8\u304b\u3089'
             '\u306e\u3054\u610f\u898b\u3092\u304a\u5f85\u3061\u3057\u3066\u304a\u308a\u307e\u3059\u3002',
             space_after=Pt(12))

    # Sign-off
    add_para(doc, '\u656c\u5177', space_after=Pt(24))

    add_para(doc, '[\u8cac\u4efb\u8457\u8005\u540d]', bold=True, space_after=Pt(0))
    add_para(doc, '[\u5f79\u8077\u30fb\u6240\u5c5e]', space_after=Pt(0))
    add_para(doc, '[\u6240\u5c5e\u6a5f\u95a2]', space_after=Pt(0))
    add_para(doc, '[\u4f4f\u6240]', space_after=Pt(0))
    add_para(doc, '[\u30e1\u30fc\u30eb\u30a2\u30c9\u30ec\u30b9]', space_after=Pt(0))
    add_para(doc, '[\u96fb\u8a71\u756a\u53f7]', space_after=Pt(0))

    path = outdir + 'cover_letter_japanese.docx'
    doc.save(path)
    print(f"Japanese cover letter saved: {path}")


if __name__ == '__main__':
    write_english_cover_letter()
    write_japanese_cover_letter()
    print("\nBoth cover letters generated successfully!")
