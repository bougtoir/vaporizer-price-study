"""
Generate JCA (Journal of Clinical Anesthesia) cover letter as editable .docx file.
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
out_dir = os.path.join(SCRIPT_DIR, 'papers')
os.makedirs(out_dir, exist_ok=True)


def write_cover_letter():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.15
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Date
    from datetime import date
    doc.add_paragraph(date.today().strftime('%B %d, %Y'))
    doc.add_paragraph()

    # Addressee
    doc.add_paragraph('Editor-in-Chief')
    doc.add_paragraph('Journal of Clinical Anesthesia')
    doc.add_paragraph()

    # Salutation
    doc.add_paragraph('Dear Editor,')
    doc.add_paragraph()

    # Body
    doc.add_paragraph(
        'We are pleased to submit our manuscript entitled "Targeted environmental regulation '
        'without observable collateral market damage: the EU desflurane ban and secondary '
        'market vaporizer prices" for consideration as an Original Contribution in the '
        'Journal of Clinical Anesthesia.')

    doc.add_paragraph(
        'Environmental regulation of anesthetic agents is accelerating, yet the downstream '
        'economic consequences for clinical equipment markets remain largely unknown. Using '
        'three years of eBay completed sale data (March 2023 to March 2026) spanning the full '
        'EU desflurane regulatory timeline, we provide the first empirical evidence on this '
        'question. We found that desflurane vaporizer prices showed a progressive, '
        'statistically significant decline (Spearman \u03c1=\u22120.28, P<0.001), while '
        'sevoflurane and isoflurane vaporizer prices remained stable throughout the study '
        'period. Between-agent effect size comparison confirmed the agent-specificity of this '
        'decline (P=0.043). These findings suggest that the EU\u2019s targeted regulation '
        'achieved its intended market-level effects without observable destabilization of the '
        'broader anesthetic equipment market.')

    doc.add_paragraph(
        'We believe this work is particularly relevant to the readership of the Journal of '
        'Clinical Anesthesia for several reasons. First, the economic impact of environmental '
        'regulation on anesthesia equipment is a practical concern for departments worldwide as '
        'similar regulations are anticipated for other agents and delivery systems. Second, the '
        'finding that secondary market values depreciate well before formal prohibition dates '
        'has direct implications for capital asset management and transition planning. Third, '
        'the study addresses a gap in the literature: while the environmental and clinical '
        'rationale for desflurane discontinuation is well established, the downstream equipment '
        'market effects have not been empirically examined.')

    doc.add_paragraph(
        'The manuscript has not been published previously and is not under consideration by '
        'any other journal. All authors have approved the manuscript and agree with its '
        'submission to the Journal of Clinical Anesthesia.')

    doc.add_paragraph(
        'We confirm that this study complies with the STROBE guidelines for observational '
        'studies. The completed STROBE checklist is provided as supplementary material. '
        'No ethical approval was required, as the study analyzed publicly available market '
        'data without involving human participants.')

    doc.add_paragraph()

    # Closing
    doc.add_paragraph('Sincerely,')
    doc.add_paragraph()
    doc.add_paragraph('[Corresponding author name]')
    doc.add_paragraph('[Department, Institution]')
    doc.add_paragraph('[Address]')
    doc.add_paragraph('[Email]')

    path = os.path.join(out_dir, 'jca_cover_letter.docx')
    doc.save(path)
    print(f"JCA cover letter saved: {path}")
    return path


if __name__ == '__main__':
    write_cover_letter()
