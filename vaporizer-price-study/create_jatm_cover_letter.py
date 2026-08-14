"""
Generate JATM (Journal of Anesthesia and Translational Medicine) cover letter
as editable .docx file.
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
out_dir = os.path.join(SCRIPT_DIR, 'papers')
os.makedirs(out_dir, exist_ok=True)


def write_cover_letter():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.15
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Date
    from datetime import date
    doc.add_paragraph(date.today().strftime('%B %d, %Y'))
    doc.add_paragraph()

    # Addressee
    doc.add_paragraph('Editor-in-Chief')
    doc.add_paragraph('Journal of Anesthesia and Translational Medicine')
    doc.add_paragraph()

    # Salutation
    doc.add_paragraph('Dear Editor,')
    doc.add_paragraph()

    # Body
    doc.add_paragraph(
        'We are resubmitting our revised manuscript entitled "Association between the EU '
        'desflurane phase-out and secondary market vaporizer prices: an observational '
        'time-series analysis" (Manuscript Number JATMED-D-26-00064R1) for reconsideration as '
        'an Article in the Journal of Anesthesia and Translational Medicine.')

    doc.add_paragraph(
        'We thank the Editor and the two reviewers for their constructive feedback. In this '
        'minor revision we have (1) further softened the causal language in the title and '
        'abstract; (2) clarified that the short post-restriction window reflects the Terapeak '
        'three-year historical window ending at data extraction in March 2026; (3) explicitly '
        'interpreted the non-significant transaction-level DiD post-restriction coefficient in '
        'light of a violated parallel-trend assumption; and (4) added an acknowledgment that '
        'trend tests across three agents raise a multiple-testing concern, with isoflurane '
        'reaching nominal significance in two transaction-level tests but not in the quarterly '
        'median analysis and with small effect sizes. A detailed point-by-point response to '
        'reviewers is enclosed.')

    doc.add_paragraph(
        'The manuscript has not been published previously and is not under consideration by '
        'any other journal. All authors have approved the revised manuscript and agree with '
        'its resubmission to the Journal of Anesthesia and Translational Medicine.')

    doc.add_paragraph(
        'We confirm that this study complies with the STROBE guidelines for observational '
        'studies. The completed STROBE checklist is provided as supplementary material. '
        'No ethical approval was required, as the study analyzed publicly available market '
        'data without involving human participants.')

    doc.add_paragraph()

    # Closing
    doc.add_paragraph('Sincerely,')
    doc.add_paragraph()
    doc.add_paragraph('Onishi Tatsuki')
    doc.add_paragraph('Data Science and AI Innovation Research Promotion Center, Shiga University')
    doc.add_paragraph('1-1-1, Bamba, Hikone, Shiga, 522-8522, Japan')
    doc.add_paragraph('E-mail: bougtoir@gmail.com')

    path = os.path.join(out_dir, 'jatm_cover_letter.docx')
    doc.save(path)
    print(f"JATM cover letter saved: {path}")
    return path


if __name__ == '__main__':
    write_cover_letter()
