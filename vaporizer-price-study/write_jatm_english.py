"""
Generate JATM (Journal of Anesthesia and Translational Medicine) format
English paper as editable .docx file with inline figures.

Target journal: Journal of Anesthesia and Translational Medicine (KeAi/Elsevier)
Official journal of the Chinese Society of Anesthesiology (CSA)

Key JATM format requirements:
  - Structured abstract (max 300 words): Background, Methods, Results, Conclusions
  - Introduction: max 500 words
  - Discussion: max 1500 words; Introduction + Discussion <= 2000 combined
  - Materials and Methods (with Statistical Analysis subsection at end)
  - Double-spaced, Times New Roman 12pt, 1-inch margins
  - Vancouver numbered references as superscript in order of first appearance
  - 3+ authors: list first 3 then et al.
  - STROBE checklist for observational studies
  - Figures/tables inline next to relevant text
"""
import pandas as pd
import numpy as np
from scipy import stats as sp_stats
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json
import re

# ==========================================
# Load analysis results
# ==========================================
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
data_dir = os.path.join(SCRIPT_DIR, 'data')
fig_dir = os.path.join(SCRIPT_DIR, 'figures')
out_dir = os.path.join(SCRIPT_DIR, 'papers')
os.makedirs(out_dir, exist_ok=True)

stats_df = pd.read_csv(os.path.join(data_dir, 'statistics_summary.csv'), index_col=0)
combined = pd.read_csv(os.path.join(data_dir, 'combined_cleaned.csv'))
combined['date_sold'] = pd.to_datetime(combined['date_sold'])

# Load asking price analysis results
try:
    with open(os.path.join(data_dir, 'asking_price_analysis.json'), 'r') as f:
        asking_results = json.load(f)
    asking_df = pd.read_csv(os.path.join(data_dir, 'ebay_asking_prices.csv'))
    has_asking_data = True
except FileNotFoundError:
    has_asking_data = False
    asking_results = None
    asking_df = None

# Load revision sensitivity analyses (bootstrap, power, ITS/DiD)
try:
    with open(os.path.join(data_dir, 'revision_sensitivity.json'), 'r') as f:
        revision_results = json.load(f)
    has_revision_sensitivity = True
except FileNotFoundError:
    revision_results = None
    has_revision_sensitivity = False

# Key dates
reg_date = pd.Timestamp('2026-01-01')
proposal_date = pd.Timestamp('2022-04-05')
agreement_date = pd.Timestamp('2023-10-05')
adoption_date = pd.Timestamp('2024-02-07')

# Compute summary statistics
summ = {}
for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
    sub = combined[combined['agent_type'] == agent]
    pre = sub[sub['date_sold'] < reg_date]['price_usd']
    post = sub[sub['date_sold'] >= reg_date]['price_usd']
    summ[agent] = {
        'total_n': len(sub),
        'pre_n': len(pre), 'post_n': len(post),
        'pre_mean': pre.mean() if len(pre) > 0 else float('nan'),
        'post_mean': post.mean() if len(post) > 0 else float('nan'),
        'pre_median': pre.median() if len(pre) > 0 else float('nan'),
        'post_median': post.median() if len(post) > 0 else float('nan'),
        'pre_sd': pre.std() if len(pre) > 0 else float('nan'),
        'post_sd': post.std() if len(post) > 0 else float('nan'),
    }

total_n = len(combined)
date_min_all = combined['date_sold'].min().strftime('%d %B %Y')
date_max_all = combined['date_sold'].max().strftime('%d %B %Y')


# ==========================================
# Compute trend statistics
# ==========================================
def classify_period(date):
    d = pd.Timestamp(date)
    if d < proposal_date:
        return 1
    elif d < agreement_date:
        return 2
    elif d < adoption_date:
        return 3
    elif d < reg_date:
        return 4
    else:
        return 5


combined['period_num'] = combined['date_sold'].apply(classify_period)

trend_results = {}
for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
    sub = combined[combined['agent_type'] == agent].copy()
    sub['days'] = (sub['date_sold'] - sub['date_sold'].min()).dt.days
    rho, rho_p = sp_stats.spearmanr(sub['days'], sub['price_usd'])
    tau, tau_p = sp_stats.kendalltau(sub['period_num'], sub['price_usd'])
    sub['quarter'] = sub['date_sold'].dt.to_period('Q')
    quarterly = sub.groupby('quarter')['price_usd'].agg(['median', 'count'])
    quarterly = quarterly[quarterly['count'] >= 3]
    q_nums = np.arange(len(quarterly))
    if len(quarterly) >= 4:
        q_rho, q_rho_p = sp_stats.spearmanr(q_nums, quarterly['median'])
    else:
        q_rho, q_rho_p = float('nan'), float('nan')
    trend_results[agent] = {
        'spearman_rho': rho, 'spearman_p': rho_p,
        'kendall_tau': tau, 'kendall_p': tau_p,
        'quarterly_rho': q_rho, 'quarterly_p': q_rho_p,
    }


def get_pval(agent, col='u_pval'):
    try:
        v = stats_df.loc[agent, col]
        if pd.notna(v):
            return float(v)
    except Exception:
        pass
    return float('nan')


def get_stat(agent, col):
    try:
        v = stats_df.loc[agent, col]
        if pd.notna(v):
            return float(v)
    except Exception:
        pass
    return float('nan')


def fmt_p(p):
    if np.isnan(p):
        return 'N/A'
    if p < 0.001:
        return '<0.001'
    return f'{p:.3f}'


def fmt_ci(lo, hi):
    return f'{lo:.0f} to {hi:.0f}'


def get_rev(key, sub=None, default=float('nan')):
    if not has_revision_sensitivity:
        return default
    try:
        v = revision_results[key]
        if sub:
            v = v[sub]
        return v
    except Exception:
        return default


def fmt_stat(val):
    if np.isnan(val):
        return 'N/A'
    return f'{val:.3f}'


# ==========================================
# Effect size comparison
# ==========================================
def var_cohens_d(n1, n2, d):
    return (n1 + n2) / (n1 * n2) + d**2 / (2 * (n1 + n2))


def se_cohens_d(n1, n2, d):
    return np.sqrt(var_cohens_d(n1, n2, d))


def ci_cohens_d(n1, n2, d, alpha=0.05):
    se = se_cohens_d(n1, n2, d)
    z_crit = sp_stats.norm.ppf(1 - alpha / 2)
    return d - z_crit * se, d + z_crit * se


def z_test_d_diff(d1, n1a, n1b, d2, n2a, n2b):
    diff = d1 - d2
    se = np.sqrt(var_cohens_d(n1a, n1b, d1) + var_cohens_d(n2a, n2b, d2))
    z = diff / se
    p = 2 * (1 - sp_stats.norm.cdf(abs(z)))
    return diff, se, z, p


effect_sizes = {}
for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
    d = get_stat(agent, 'cohens_d')
    n_pre = summ[agent]['pre_n']
    n_post = summ[agent]['post_n']
    se = se_cohens_d(n_pre, n_post, d)
    ci_lo, ci_hi = ci_cohens_d(n_pre, n_post, d)
    effect_sizes[agent] = {'d': d, 'se': se, 'ci_lo': ci_lo, 'ci_hi': ci_hi,
                           'n_pre': n_pre, 'n_post': n_post}

es_comparisons = {}
for a1, a2 in [('Desflurane', 'Sevoflurane'), ('Desflurane', 'Isoflurane'),
               ('Sevoflurane', 'Isoflurane')]:
    e1, e2 = effect_sizes[a1], effect_sizes[a2]
    diff, se, z, p = z_test_d_diff(e1['d'], e1['n_pre'], e1['n_post'],
                                    e2['d'], e2['n_pre'], e2['n_post'])
    es_comparisons[f'{a1}_vs_{a2}'] = {'diff': diff, 'se': se, 'z': z, 'p': p}


# ==========================================
# Figure mapping (figure number -> filename)
# ==========================================
FIGURE_MAP = [
    {'num': 1, 'file': 'fig1_price_timeseries.png',
     'caption': 'Time series of eBay completed sale prices for desflurane (red), '
                'sevoflurane (blue), and isoflurane (green) vaporizers over three years '
                '(March 2023 to March 2026). Vertical dashed lines indicate key EU regulatory '
                'milestones. Curved lines represent LOWESS trend estimates (fraction = 0.3). '
                'Data source: eBay Terapeak.'},
    {'num': 2, 'file': 'fig5_regulatory_timeline.png',
     'caption': 'Anesthetic vaporizer prices mapped against the EU regulatory timeline. '
                'Shaded regions indicate regulatory phases. Data source: eBay Terapeak.'},
    {'num': 3, 'file': 'fig6_quarterly_trends.png',
     'caption': 'Quarterly median price trends (upper panel) and sales volume (lower panel). '
                'Data source: eBay Terapeak.'},
    {'num': 4, 'file': 'fig2_boxplot_comparison.png',
     'caption': 'Box plot comparison of vaporizer prices before and after the EU desflurane '
                'ban (1 January 2026). Individual data points are shown as jittered dots. '
                'Data source: eBay Terapeak.'},
    {'num': 5, 'file': 'fig3_monthly_median.png',
     'caption': 'Monthly median prices of anesthetic vaporizers on eBay. Annotations '
                'indicate the number of transactions per month (n). Data source: eBay Terapeak.'},
    {'num': 6, 'file': 'fig4_histograms.png',
     'caption': 'Price distribution histograms for each vaporizer type, comparing pre-ban '
                '(solid fill) and post-ban (hatched) periods. Data source: eBay Terapeak.'},
]


# ==========================================
# Helper functions
# ==========================================
def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_run_styled(para, text, bold=False, italic=False, size=Pt(12)):
    run = para.add_run(text)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    return run


def add_superscript_text(para, text, size=Pt(12)):
    """Parse text with {ref} markers and create superscript runs."""
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            ref_text = part[1:-1]
            run = para.add_run(ref_text)
            run.font.size = size
            run.font.superscript = True
        else:
            run = para.add_run(part)
            run.font.size = size
    return para


def add_para(doc, text, size=Pt(12), bold=False, italic=False,
             alignment=None, space_after=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    return p


def add_para_with_refs(doc, text, size=Pt(12)):
    """Add paragraph with superscript citation references."""
    p = doc.add_paragraph()
    add_superscript_text(p, text, size=size)
    return p


def setup_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 2.0
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    return doc


def add_table_header(table, headers):
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(cell, 'D9E2F3')


def add_table_data_row(table, data):
    row = table.add_row()
    for i, (text, align) in enumerate(data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(str(text))
        run.font.size = Pt(9)
    return row


def insert_inline_figure(doc, fig_info):
    """Insert a PNG figure inline with caption."""
    fig_path = os.path.join(fig_dir, fig_info['file'])
    if not os.path.exists(fig_path):
        print(f"WARNING: {fig_path} not found, skipping")
        return

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(fig_path, width=Inches(5.5))

    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    label_run = cap.add_run(f'Fig. {fig_info["num"]}. ')
    label_run.bold = True
    label_run.font.size = Pt(10)
    text_run = cap.add_run(fig_info['caption'])
    text_run.font.size = Pt(10)
    text_run.italic = True
    doc.add_paragraph()


def count_words(text):
    return len(text.split())


# ==========================================
# JATM ENGLISH PAPER
# ==========================================
def write_jatm_paper():
    doc = setup_doc()
    des = summ['Desflurane']
    sevo = summ['Sevoflurane']
    iso = summ['Isoflurane']
    des_u_pval = get_pval('Desflurane', 'u_pval')
    des_t_pval = get_pval('Desflurane', 't_pval')
    sevo_u_pval = get_pval('Sevoflurane', 'u_pval')
    iso_u_pval = get_pval('Isoflurane', 'u_pval')
    des_d = get_stat('Desflurane', 'cohens_d')
    des_tr = trend_results['Desflurane']
    sevo_tr = trend_results['Sevoflurane']
    iso_tr = trend_results['Isoflurane']
    des_pct = abs((des['post_mean'] - des['pre_mean']) / des['pre_mean'] * 100)

    # ============================================================
    # TITLE PAGE
    # ============================================================
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(
        'Association between the EU desflurane phase-out and secondary market '
        'vaporizer prices: an observational time-series analysis')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_superscript_text(author_p, 'Onishi Tatsuki{1}, Tatsuyoshi Ikenoue{1,2}', size=Pt(12))
    aff1_p = doc.add_paragraph()
    aff1_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_superscript_text(aff1_p,
        '{1}Data Science and AI Innovation Research Promotion Center, Shiga University',
        size=Pt(12))
    aff2_p = doc.add_paragraph()
    aff2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_superscript_text(aff2_p,
        '{2}Department of Public Health, Faculty of Medicine, University of Miyazaki',
        size=Pt(12))
    doc.add_paragraph()

    # Corresponding author
    p = doc.add_paragraph()
    add_run_styled(p, 'Corresponding author: ', bold=True, size=Pt(12))
    add_run_styled(p, 'Onishi Tatsuki; Data Science and AI Innovation Research Promotion Center, '
                   'Shiga University; 1-1-1, Bamba, Hikone, Shiga, 522-8522, Japan; '
                   'Telephone: +81-749-27-1023; E-mail: bougtoir@gmail.com', size=Pt(12))

    doc.add_paragraph()

    # Word counts
    add_para(doc, 'Abstract word count: ~299 (max 300)', size=Pt(10))
    add_para(doc, 'Number of references: 22', size=Pt(10))
    add_para(doc, 'Number of tables: 2 (+1 supplementary)', size=Pt(10))
    add_para(doc, 'Number of figures: 6', size=Pt(10))

    doc.add_paragraph()

    # Keywords
    p = doc.add_paragraph()
    add_run_styled(p, 'Keywords: ', bold=True, size=Pt(12))
    add_run_styled(p, ('anesthetic vaporizer, desflurane, environmental regulation, '
                       'F-gas, secondary market, EU regulation, equipment management, '
                       'capital asset lifecycle'),
                   size=Pt(12))

    doc.add_page_break()

    # ============================================================
    # STRUCTURED ABSTRACT (JATM: max 300 words)
    # Background, Methods, Results, Conclusions
    # ============================================================
    add_heading_styled(doc, 'Abstract', level=1)

    # Background
    p = doc.add_paragraph()
    add_run_styled(p, 'Background: ', bold=True)
    add_run_styled(p,
        'The European Union restricted desflurane\u2014the volatile anesthetic with the '
        'highest commonly reported 100-year global warming potential (GWP100 \u2248 2540 '
        'CO\u2082 equivalents)\u2014for routine use from 1 January 2026 under Regulation (EU) '
        '2024/573, with exemptions for clinically necessary cases. Whether this targeted '
        'environmental regulation was associated with collateral economic effects on '
        'non-targeted anesthetic equipment markets has not been examined. We investigated '
        'whether the EU desflurane regulation was associated with changes in secondary '
        'market prices of anesthetic vaporizers, and whether any such changes were '
        'agent-specific or extended to non-regulated agents (sevoflurane, isoflurane).')

    # Methods
    p = doc.add_paragraph()
    add_run_styled(p, 'Methods: ', bold=True)
    add_run_styled(p,
        'We conducted a cross-sectional time-series analysis of completed eBay sales '
        'of desflurane, sevoflurane, and isoflurane vaporizers over three years '
        '(March 2023 to March 2026), spanning the full EU regulatory timeline. '
        'Temporal trends were assessed using Spearman rank correlation and Kendall \u03c4 '
        'across ordered regulatory phases. Pre-/post-ban comparison used the '
        'Mann\u2013Whitney U test with Cohen\u2019s d effect size. Sensitivity analyses '
        'included bootstrap resampling (10,000 iterations), post-hoc power simulation under '
        'observed lognormal parameters, and a comparative interrupted time-series (CITS) '
        'model of monthly median prices with sevoflurane and isoflurane as controls.')

    # Results
    p = doc.add_paragraph()
    add_run_styled(p, 'Results: ', bold=True)
    add_run_styled(p,
        f'{total_n} completed sales were analyzed ({des["total_n"]} desflurane, '
        f'{sevo["total_n"]} sevoflurane, {iso["total_n"]} isoflurane). '
        f'Desflurane vaporizer prices declined '
        f'(Spearman \u03c1={des_tr["spearman_rho"]:.2f}, P<0.001; '
        f'Kendall \u03c4={des_tr["kendall_tau"]:.2f}, '
        f'P={fmt_p(des_tr["kendall_p"])}), '
        f'with a {des_pct:.0f}% decline pre- to post-restriction '
        f'(Cohen\u2019s d={des_d:.2f}). '
        f'Sevoflurane showed no trend. Isoflurane reached '
        f'nominal significance in Spearman (\u03c1={iso_tr["spearman_rho"]:.2f}, '
        f'P={fmt_p(iso_tr["spearman_p"])}) and Kendall \u03c4 (\u03c4={iso_tr["kendall_tau"]:.2f}, '
        f'P={fmt_p(iso_tr["kendall_p"])}) tests, but the quarterly median trend was not '
        f'significant; these isolated P values, with small effect sizes, are interpreted cautiously.')

    # Conclusions
    p = doc.add_paragraph()
    add_run_styled(p, 'Conclusions: ', bold=True)
    add_run_styled(p,
        'Desflurane vaporizer prices on eBay were associated with an agent-specific '
        'decline; non-regulated agents remained comparatively stable. These observational '
        'findings are consistent with, but do not prove, an effect of the EU desflurane '
        'phase-out and should be interpreted as hypothesis-generating given the single-'
        'platform design, unmeasured confounders, small post-restriction sample, and '
        'multiple trend tests.')

    doc.add_page_break()

    # ============================================================
    # INTRODUCTION (JATM: max 500 words)
    # ============================================================
    add_heading_styled(doc, 'Introduction', level=1)

    add_para_with_refs(doc,
        'Environmental regulation of healthcare products is accelerating. The European Union '
        'restricted desflurane\u2014the volatile anesthetic with the highest commonly reported '
        '100-year global warming potential (GWP100 \u2248 2540 CO\u2082 equivalents){1,2,3\u20135}'
        '\u2014for routine use from 1 January 2026 under Regulation (EU) 2024/573, while '
        'permitting documented clinical exceptions.{2} The choice of GWP metric and time '
        'horizon remains debated for short-lived halogenated anesthetics such as desflurane '
        '(atmospheric lifetime \u2248 14 years), because GWP100 may not reflect their '
        'instantaneous radiative forcing or steady-state atmospheric concentrations.{6} '
        'Desflurane also has recognized pharmacokinetic advantages, including faster '
        'emergence and extubation than sevoflurane in selected surgical populations.{7} The '
        'American Society of Anesthesiologists has recommended deactivation of central '
        'nitrous oxide piping on environmental grounds.{8} NHS England and NHS Scotland '
        'have independently decommissioned desflurane.{9,10} Each of these measures targets '
        'a specific agent or delivery system, yet whether such targeted restrictions are '
        'associated with collateral economic effects on non-targeted equipment markets '
        'has not been empirically examined.')

    add_para_with_refs(doc,
        'The EU desflurane restriction provides an opportunity to address this question as a '
        'natural experiment. First, only a single agent is targeted for routine use; '
        'sevoflurane (GWP \u2248 130) and isoflurane (GWP \u2248 510){11,12} remain in '
        'unrestricted use and serve as natural controls. Second, the regulatory process '
        'advanced through clearly dated milestones \u2014 European Commission proposal (April '
        '2022), European Parliament plenary vote (March 2023), trilogue provisional agreement '
        '(October 2023), formal adoption (February 2024), and the 2026 restriction on '
        'routine use (January 2026) \u2014 enabling time-series analysis across successive '
        'phases. Third, anesthetic vaporizers are agent-specific capital assets with typical '
        'lifespans of 10\u201315 years, so the economic consequences of regulation may be '
        'reflected in secondary market values.')

    add_para_with_refs(doc,
        'Previous studies have addressed the financial rationale for discontinuing '
        'desflurane,{13} the clinical and policy implications of decommissioning '
        'programs,{14,15} the effectiveness of vaporizer removal at the institutional '
        'level,{16} and the cost savings from reduced volatile anesthetic '
        'consumption.{17,18} The secondary market for pre-owned medical equipment has been '
        'characterized for other device categories.{19} However, to our knowledge, no study '
        'has examined whether environmental regulation of a single anesthetic agent is '
        'associated with targeted secondary-market effects or whether it destabilizes the '
        'broader equipment market.')

    add_para_with_refs(doc,
        'We hypothesized that (1) the EU desflurane regulation would be associated with a '
        'progressive decline in secondary market prices of desflurane vaporizers, and '
        '(2) this decline would be agent-specific\u2014sevoflurane and isoflurane vaporizer '
        'prices would remain stable\u2014suggesting that any market-level changes were '
        'concentrated in the regulated agent.')

    # ============================================================
    # MATERIALS AND METHODS (JATM: "Materials and Methods")
    # ============================================================
    add_heading_styled(doc, 'Materials and Methods', level=1)

    # Ethics
    add_heading_styled(doc, 'Ethics', level=2)
    doc.add_paragraph(
        'Ethical approval was not required for this study. The study analyzed publicly available, '
        'anonymized completed sale data from an online marketplace (eBay). No individual-level, '
        'patient, or human participant data were collected.')

    # STROBE statement
    add_heading_styled(doc, 'Reporting guidelines', level=2)
    add_para_with_refs(doc,
        'This study is reported in accordance with the Strengthening the Reporting of '
        'Observational Studies in Epidemiology (STROBE) guidelines for cross-sectional '
        'studies.{20} The completed STROBE checklist is provided as supplementary material.')

    add_heading_styled(doc, 'Study design and data source', level=2)
    doc.add_paragraph(
        'We conducted a cross-sectional time-series analysis of anesthetic vaporizer prices using '
        'completed (sold) listings on eBay (www.ebay.com). '
        'Data were retrieved using Terapeak, eBay\u2019s official product research tool integrated '
        'within eBay Seller Hub. Terapeak provides access to up to three years of historical '
        'completed sale data, including item titles, sale prices, sale dates, and quantities sold. '
        'Data were collected in March 2026, covering the period from 28 March 2023 to 24 March 2026.')
    doc.add_paragraph(
        'Although the three-year window reflects the maximum retrievable period within Terapeak, '
        'this timeframe is analytically meaningful: it begins shortly after the European Parliament '
        'plenary vote approving the revised F-gas Regulation (March 2023) and captures the full '
        'legislative trajectory from the European Commission\u2019s original proposal (April 2022) '
        'through to the post-ban period, encompassing all key regulatory milestones. '
        'Because Terapeak\u2019s three-year window ends at the data extraction date (March 2026), '
        'the post-restriction observation period is limited to approximately three months after the '
        f'1 January 2026 effective date (n={des["post_n"]} desflurane transactions); a longer '
        'follow-up would require either prospective data collection or a later extraction, neither '
        'of which was possible here. '
        'We used a single marketplace (eBay) rather than integrating data from multiple '
        'platforms to avoid the risk of counting cross-listed items more than once.')

    add_heading_styled(doc, 'Eligibility criteria', level=2)
    doc.add_paragraph(
        'We searched Terapeak for completed sales using the search terms '
        '\u201cdesflurane vaporizer,\u201d \u201csevoflurane vaporizer,\u201d and '
        '\u201cisoflurane vaporizer\u201d with a three-year date range filter. Inclusion criteria '
        'were: (1) completed (sold) listings; (2) standalone anesthetic vaporizer units; and '
        '(3) valid sale price and date. Exclusion criteria were: (1) non-vaporizer items '
        '(keyed fillers, bottle adapters, accessories, pour-fill adapters, anti-spill caps); '
        '(2) veterinary-specific anesthesia systems or machines (rather than standalone vaporizers); '
        '(3) lot listings containing multiple heterogeneous items; and (4) listings with missing or '
        'implausible price data.')

    add_heading_styled(doc, 'Variables', level=2)
    doc.add_paragraph(
        'The primary outcome was sale price in US dollars. For each listing, we recorded: item title, '
        'sale price (USD), sale date, and quantity sold. The primary exposure variable was the regulatory '
        'period, classified relative to key milestones in the EU F-gas Regulation timeline. The primary '
        'comparison used 1 January 2026 (the desflurane prohibition effective date) as the cutpoint. '
        'A secondary multi-period classification divided the study period into four phases: '
        'post-proposal (after EC proposal, April 2022), post-agreement (after trilogue, October 2023), '
        'post-adoption (after formal adoption, February 2024), and post-ban (after 1 January 2026). '
        'These ordered phases were used for trend analysis.')

    # Statistical analysis (at end of Materials and Methods per JATM)
    add_heading_styled(doc, 'Statistical analysis', level=2)
    doc.add_paragraph(
        'The primary analysis was a time-series trend assessment using Spearman rank correlation '
        'between sale date and price for each agent type separately. This was complemented by '
        'Kendall \u03c4, computed between the ordered regulatory phase (1\u20135) and sale price, '
        'to test whether prices changed progressively across successive milestones. These trend '
        'tests, applied to each agent type independently, allowed direct comparison of temporal '
        'patterns between the regulated agent (desflurane) and the unregulated comparators '
        '(sevoflurane, isoflurane). Quarterly median prices were also assessed using Spearman '
        'correlation to evaluate the trend at an aggregated level.')
    doc.add_paragraph(
        'The pre-/post-ban comparison (Mann\u2013Whitney U test and Welch\u2019s t-test) was '
        'conducted as a secondary, exploratory analysis. Given the small post-ban sample size '
        f'(n={des["post_n"]} for desflurane), this comparison has limited statistical power, '
        'and its results should be interpreted with caution. Effect sizes were estimated '
        'using Cohen\u2019s d with 95% confidence intervals. To test whether the magnitude of the '
        'pre-/post-ban price change differed between agent types, pairwise z-tests for independent '
        'Cohen\u2019s d values were performed using the large-sample variance approximation.')
    doc.add_paragraph(
        'Because trend tests (Spearman, Kendall \u03c4, and quarterly Spearman) were applied '
        'to each of three agent types, the probability of nominally significant P '
        'values arising under the null is inflated. We did not apply a formal family-wise '
        'multiple-testing correction; instead, P values were interpreted alongside effect '
        'magnitude, direction, and consistency across tests. For isoflurane, both Spearman '
        f'(P={fmt_p(iso_tr["spearman_p"])}) and Kendall \u03c4 (P={fmt_p(iso_tr["kendall_p"])}) '
        'tests reached nominal significance at the transaction level, but the quarterly median '
        'trend was not significant and effects were small; these P values were therefore '
        'interpreted cautiously and not considered evidence of a true trend.')
    doc.add_paragraph(
        'Sensitivity analyses were performed to assess robustness and to address causal '
        'interpretation. First, we used bootstrap resampling (10,000 iterations) of the pre- and '
        'post-ban samples to derive non-parametric 95% confidence intervals for the difference in '
        'mean and median prices and to obtain bootstrap P values for the Mann\u2013Whitney U and '
        'Welch tests. Second, we estimated post-hoc statistical power for the pre-/post-ban '
        'Mann\u2013Whitney U test and Welch t-test by simulating 10,000 datasets from lognormal '
        'distributions parameterized by the observed pre- and post-ban means and standard '
        'deviations. Third, we fitted a comparative interrupted time-series (CITS) model to '
        'monthly median prices, with sevoflurane and isoflurane as concurrent controls, to '
        'estimate any desflurane-specific level or slope change at the 2026 restriction. A '
        'difference-in-differences (DiD) model on log-transformed individual transaction prices '
        'was also fitted as a further sensitivity check.')
    doc.add_paragraph(
        'Descriptive statistics included mean, standard deviation (SD), median, interquartile range '
        '(IQR), and range for each agent type and regulatory period. '
        'The Kruskal\u2013Wallis test was used for multi-period comparisons across regulatory phases. '
        'LOWESS (locally weighted scatterplot smoothing) trend lines were fitted to visualize '
        'price trajectories. Analyses were performed using Python 3.12 with pandas 2.2, '
        'scipy 1.14, and statsmodels 0.14. Statistical significance was set at P\u2009<\u20090.05 '
        '(two-sided). No a priori sample size calculation was performed, as this study aimed to '
        'capture all available transactions within the Terapeak data window.')

    # ============================================================
    # RESULTS
    # ============================================================
    add_heading_styled(doc, 'Results', level=1)
    doc.add_paragraph(
        f'A total of {total_n} completed eBay sales of anesthetic vaporizers were identified '
        f'and included in the analysis after applying exclusion criteria: '
        f'{des["total_n"]} desflurane, '
        f'{sevo["total_n"]} sevoflurane, and '
        f'{iso["total_n"]} isoflurane vaporizers. '
        f'The study period spanned from {date_min_all} to {date_max_all} (three years). '
        f'Desflurane vaporizers were predominantly Datex-Ohmeda/GE Tec 6 Plus and '
        f'Dr\u00e4ger D-Vapor models; '
        f'sevoflurane vaporizers included Dr\u00e4ger Vapor 2000, Penlon Sigma Delta, and Tec 7 '
        f'models; isoflurane vaporizers included Ohmeda Tec 3, Tec 5, Tec 7, and Dr\u00e4ger '
        f'Vapor 2000 models (Fig. 1).')

    # Insert Fig. 1 inline
    insert_inline_figure(doc, FIGURE_MAP[0])

    # Inline Table 1 (trend analysis)
    trend_data = compute_trend_data()
    p = doc.add_paragraph()
    add_run_styled(p, 'Table 1. ', bold=True, size=Pt(10))
    add_run_styled(p, ('Time-series trend analysis of vaporizer prices by agent type. Spearman rank '
                       'correlation tests monotonic association between sale date and price; '
                       'Kendall \u03c4 tests association between ordered regulatory phase and price.'),
                   italic=True, size=Pt(10))

    t1 = doc.add_table(rows=1, cols=7)
    t1.style = 'Table Grid'
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header(t1, ['Agent', 'Spearman \u03c1', 'P value', 'Kendall \u03c4', 'P value',
                          'Quarterly \u03c1', 'P value'])
    for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
        tr = trend_data[agent]
        data = [
            (agent, WD_ALIGN_PARAGRAPH.LEFT),
            (f'{tr["spearman_rho"]:.3f}', WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(tr['spearman_p']), WD_ALIGN_PARAGRAPH.CENTER),
            (f'{tr["kendall_tau"]:.3f}', WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(tr['kendall_p']), WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_stat(tr['quarterly_rho']), WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(tr['quarterly_p']), WD_ALIGN_PARAGRAPH.CENTER),
        ]
        add_table_data_row(t1, data)
    p_t1fn = doc.add_paragraph()
    add_run_styled(p_t1fn, 'Note: ', bold=True, size=Pt(9))
    add_run_styled(p_t1fn, 'The Spearman and Kendall \u03c4 tests for isoflurane reached '
        f'nominal significance (P={fmt_p(iso_tr["spearman_p"])} and '
        f'P={fmt_p(iso_tr["kendall_p"])}, respectively), but effect sizes were small and the '
        'quarterly median trend was not significant; this is therefore not interpreted '
        'as a clinically meaningful temporal trend.', size=Pt(9), italic=True)
    doc.add_paragraph()

    # Inline Table 2 (pre-/post-ban comparison)
    p = doc.add_paragraph()
    add_run_styled(p, 'Table 2. ', bold=True, size=Pt(10))
    add_run_styled(p, ('Pre- and post-ban vaporizer prices by agent type with between-agent '
                       'effect size comparison. Values are mean \u00b1 SD in US dollars.'),
                   italic=True, size=Pt(10))

    p2a = doc.add_paragraph()
    add_run_styled(p2a, 'Panel A. ', bold=True, size=Pt(9))
    add_run_styled(p2a, 'Descriptive statistics and within-agent effect sizes (pre- vs post-ban)',
                   italic=True, size=Pt(9))

    table2a = doc.add_table(rows=1, cols=7)
    table2a.style = 'Table Grid'
    table2a.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header(table2a, ['Agent', 'n (pre/post)', 'Pre-ban mean \u00b1 SD',
                               'Post-ban mean \u00b1 SD', '% change',
                               "Cohen\u2019s d", '95% CI'])
    for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
        s = summ[agent]
        es = effect_sizes[agent]
        pct = (s['post_mean'] - s['pre_mean']) / s['pre_mean'] * 100
        data = [
            (agent, WD_ALIGN_PARAGRAPH.LEFT),
            (f'{s["pre_n"]}/{s["post_n"]}', WD_ALIGN_PARAGRAPH.CENTER),
            (f'${s["pre_mean"]:.0f} \u00b1 {s["pre_sd"]:.0f}', WD_ALIGN_PARAGRAPH.CENTER),
            (f'${s["post_mean"]:.0f} \u00b1 {s["post_sd"]:.0f}', WD_ALIGN_PARAGRAPH.CENTER),
            (f'{pct:+.1f}%', WD_ALIGN_PARAGRAPH.CENTER),
            (f'{es["d"]:.2f}', WD_ALIGN_PARAGRAPH.CENTER),
            (f'{es["ci_lo"]:.2f} to {es["ci_hi"]:.2f}', WD_ALIGN_PARAGRAPH.CENTER),
        ]
        add_table_data_row(table2a, data)

    doc.add_paragraph()
    p2b = doc.add_paragraph()
    add_run_styled(p2b, 'Panel B. ', bold=True, size=Pt(9))
    add_run_styled(p2b, 'Between-agent comparison of effect sizes (z-test for independent '
                   "Cohen\u2019s d)",
                   italic=True, size=Pt(9))

    table2b = doc.add_table(rows=1, cols=5)
    table2b.style = 'Table Grid'
    table2b.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header(table2b, ['Comparison', '\u0394d', 'SE', 'z', 'P value'])
    for key, label in [('Desflurane_vs_Sevoflurane', 'Desflurane vs Sevoflurane'),
                        ('Desflurane_vs_Isoflurane', 'Desflurane vs Isoflurane'),
                        ('Sevoflurane_vs_Isoflurane', 'Sevoflurane vs Isoflurane')]:
        c = es_comparisons[key]
        data = [
            (label, WD_ALIGN_PARAGRAPH.LEFT),
            (f'{c["diff"]:.2f}', WD_ALIGN_PARAGRAPH.CENTER),
            (f'{c["se"]:.3f}', WD_ALIGN_PARAGRAPH.CENTER),
            (f'{c["z"]:.2f}', WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(c['p']), WD_ALIGN_PARAGRAPH.CENTER),
        ]
        add_table_data_row(table2b, data)

    # Within-agent P value footnote
    p_fn = doc.add_paragraph()
    add_run_styled(p_fn, 'Within-agent P values (Mann\u2013Whitney U): ', italic=True, size=Pt(8))
    footnote_parts = []
    for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
        u_p = fmt_p(get_pval(agent, 'u_pval'))
        t_p = fmt_p(get_pval(agent, 't_pval'))
        footnote_parts.append(f'{agent} U P={u_p}, t P={t_p}')
    add_run_styled(p_fn, '; '.join(footnote_parts) + '.', italic=True, size=Pt(8))
    doc.add_paragraph()

    # Bootstrap and power footnote
    if has_revision_sensitivity:
        des_boot = get_rev('bootstrap_pre_post', 'Desflurane')
        des_pow = get_rev('power_simulation', 'Desflurane')
        if isinstance(des_boot, dict):
            p_fn2 = doc.add_paragraph()
            add_run_styled(p_fn2, 'Sensitivity: ', bold=True, italic=True, size=Pt(8))
            mean_lo, mean_hi = des_boot['mean_diff_ci95']
            med_lo, med_hi = des_boot['median_diff_ci95']
            boot_text = (
                f"Bootstrap 95% CI for the desflurane pre-/post-ban mean difference: "
                f"US${mean_lo:.0f} to US${mean_hi:.0f}; median difference: "
                f"US${med_lo:.0f} to US${med_hi:.0f}. Post-hoc power (lognormal simulation, "
                f"n={des_pow['n_pre']}/{des_pow['n_post']}): Mann\u2013Whitney U "
                f"{des_pow['power_mannwhitney']*100:.1f}%, Welch t "
                f"{des_pow['power_welch_t']*100:.1f}%."
            )
            add_run_styled(p_fn2, boot_text, italic=True, size=Pt(8))
        doc.add_paragraph()

    # Results narrative - trend
    des_pct_val = (des['post_mean'] - des['pre_mean']) / des['pre_mean'] * 100
    doc.add_paragraph(
        f'Desflurane vaporizer prices showed a statistically significant downward trend over '
        f'the three-year study period (primary analysis). Spearman rank correlation demonstrated a '
        f'significant negative monotonic association between sale date and price '
        f'(\u03c1={des_tr["spearman_rho"]:.2f}, P<0.001). Kendall \u03c4 analysis '
        f'confirmed that prices decreased across successive regulatory phases '
        f'(\u03c4={des_tr["kendall_tau"]:.2f}, '
        f'P={fmt_p(des_tr["kendall_p"])}) (Fig. 2). '
        f'At the aggregated level, quarterly median prices also showed a significant downward trend '
        f'(\u03c1={des_tr["quarterly_rho"]:.2f}, '
        f'P={fmt_p(des_tr["quarterly_p"])}) (Table 1; Fig. 3).')

    # Insert Fig. 2 inline
    insert_inline_figure(doc, FIGURE_MAP[1])

    # Insert Fig. 3 inline
    insert_inline_figure(doc, FIGURE_MAP[2])

    doc.add_paragraph(
        f'In the exploratory pre-/post-ban comparison, the post-ban mean price '
        f'(US${des["post_mean"]:.0f} \u00b1 {des["post_sd"]:.0f}) was {abs(des_pct_val):.0f}% '
        f'lower than the pre-ban mean (US${des["pre_mean"]:.0f} \u00b1 {des["pre_sd"]:.0f}). '
        f'This difference was statistically significant on Welch\u2019s t-test '
        f'(P={fmt_p(des_t_pval)}) but did not reach significance on the '
        f'Mann\u2013Whitney U test (P={fmt_p(des_u_pval)}), likely reflecting the '
        f'small post-ban sample (n={des["post_n"]}). The effect size was medium '
        f'(Cohen\u2019s d={des_d:.2f}) (Table 2, Panel A; Fig. 4).')

    # Insert Fig. 4 inline
    insert_inline_figure(doc, FIGURE_MAP[3])

    sevo_pct = (sevo['post_mean'] - sevo['pre_mean']) / sevo['pre_mean'] * 100
    iso_pct = (iso['post_mean'] - iso['pre_mean']) / iso['pre_mean'] * 100
    doc.add_paragraph(
        f'In contrast, sevoflurane vaporizer prices showed no significant '
        f'temporal trend (Spearman \u03c1={sevo_tr["spearman_rho"]:.2f}, '
        f'P={fmt_p(sevo_tr["spearman_p"])}; '
        f'Kendall \u03c4={sevo_tr["kendall_tau"]:.2f}, '
        f'P={fmt_p(sevo_tr["kendall_p"])}). '
        f'Pre-/post-ban comparison showed a non-significant {abs(sevo_pct):.0f}% increase '
        f'(P={fmt_p(sevo_u_pval)}, Mann\u2013Whitney U) (Fig. 5).')

    # Insert Fig. 5 inline
    insert_inline_figure(doc, FIGURE_MAP[4])

    doc.add_paragraph(
        f'Isoflurane vaporizer prices were similarly stable. Spearman '
        f'(\u03c1={iso_tr["spearman_rho"]:.2f}, P={fmt_p(iso_tr["spearman_p"])}) and Kendall '
        f'\u03c4 (\u03c4={iso_tr["kendall_tau"]:.2f}, P={fmt_p(iso_tr["kendall_p"])}) tests reached '
        f'nominal significance at the transaction level, but the quarterly median trend was not '
        f'significant (\u03c1={iso_tr["quarterly_rho"]:.2f}, P={fmt_p(iso_tr["quarterly_p"])}) '
        f'and effect sizes were small. These isolated P values are expected under multiple testing '
        f'and do not indicate a consistent directional trend. The pre-/post-ban comparison showed '
        f'a non-significant {abs(iso_pct):.0f}% decline (P={fmt_p(iso_u_pval)}, '
        f'Mann\u2013Whitney U) (Fig. 6).')

    # Insert Fig. 6 inline
    insert_inline_figure(doc, FIGURE_MAP[5])

    # Between-agent effect size comparison
    des_vs_sevo = es_comparisons['Desflurane_vs_Sevoflurane']
    des_vs_iso = es_comparisons['Desflurane_vs_Isoflurane']
    des_es = effect_sizes['Desflurane']
    sevo_es = effect_sizes['Sevoflurane']
    iso_es = effect_sizes['Isoflurane']
    doc.add_paragraph(
        f'Between-agent comparison tested whether the magnitude of pre-/post-ban price change '
        f'differed across agent types (Table 2, Panel B). The effect size for desflurane '
        f'(d={des_es["d"]:.2f}; 95% CI {des_es["ci_lo"]:.2f} to {des_es["ci_hi"]:.2f}) '
        f'was significantly larger than that for sevoflurane '
        f'(d={sevo_es["d"]:.2f}; 95% CI {sevo_es["ci_lo"]:.2f} to {sevo_es["ci_hi"]:.2f}; '
        f'\u0394d={des_vs_sevo["diff"]:.2f}, z={des_vs_sevo["z"]:.2f}, '
        f'P={fmt_p(des_vs_sevo["p"])}). '
        f'The difference relative to isoflurane '
        f'(d={iso_es["d"]:.2f}; 95% CI {iso_es["ci_lo"]:.2f} to {iso_es["ci_hi"]:.2f}) '
        f'did not reach statistical significance '
        f'(\u0394d={des_vs_iso["diff"]:.2f}, z={des_vs_iso["z"]:.2f}, '
        f'P={fmt_p(des_vs_iso["p"])}).')

    # Sensitivity: DiD / CITS
    if has_revision_sensitivity:
        did_tx = get_rev('did_transaction_level')
        cits = get_rev('cits_monthly_medians')
        if isinstance(did_tx, dict) and 'error' not in did_tx:
            doc.add_paragraph(
                f'Sensitivity analyses with control series were consistent with the primary '
                f'time-series findings. In a comparative interrupted time-series model of '
                f'monthly median prices, desflurane showed a desflurane-specific level decline '
                f'at the restriction (log-price coefficient={cits["desflurane_level_change_coef"]:.2f}, '
                f'P<0.001) and a post-restriction slope difference '
                f'(coefficient={cits["desflurane_slope_change_coef"]:.2f}, P<0.001), whereas '
                f'isoflurane did not (level-change P={fmt_p(cits["isoflurane_level_change_p"])}). '
                f'In a transaction-level difference-in-differences model the desflurane '
                f'post-restriction coefficient was not significant '
                f'(log-price coefficient={did_tx["desflurane_post_coef"]:.3f}, '
                f'P={fmt_p(did_tx["desflurane_post_p"])}), partly because a significant '
                f'pre-existing trend difference was present (P={fmt_p(did_tx["desflurane_trend_p"])}), '
                f'violating the parallel-trend assumption. These models support '
                f'agent-specificity but do not establish causality.')

    # Supplementary analysis
    if has_asking_data:
        ask = asking_results['asking_summary']
        kw = asking_results['kruskal_wallis']
        spr = asking_results['spread']
        n_asking = len(asking_df)
        doc.add_paragraph(
            f'In a supplementary cross-sectional analysis of {n_asking} current eBay asking prices '
            f'(Table S1; active listings, 27 March 2026), desflurane vaporizers had the lowest '
            f'median asking price (US${ask["Desflurane"]["median"]:.0f}), '
            f'approximately one-seventh that of sevoflurane '
            f'(US${ask["Sevoflurane"]["median"]:.0f}) '
            f'and one-third that of isoflurane '
            f'(US${ask["Isoflurane"]["median"]:.0f}; '
            f'Kruskal\u2013Wallis H={kw["H"]:.1f}, P<0.001). '
            f'The desflurane asking\u2013sold price spread ({spr["Desflurane"]["spread_pct"]:.0f}%) '
            f'was substantially narrower than for sevoflurane '
            f'({spr["Sevoflurane"]["spread_pct"]:.0f}%) or isoflurane '
            f'({spr["Isoflurane"]["spread_pct"]:.0f}%), suggesting that sellers have '
            f'already adjusted their price expectations to reflect post-regulation market conditions.')

    # ============================================================
    # DISCUSSION (JATM: max 1500 words)
    # ============================================================
    add_heading_styled(doc, 'Discussion', level=1)
    doc.add_paragraph(
        'This observational study provides the first empirical evidence that environmental '
        'regulation of a single anesthetic agent is associated with agent-specific effects '
        'on the secondary equipment market. Using three years of eBay completed sale data and '
        'complementary statistical approaches, we found that desflurane vaporizer prices '
        'declined progressively over the study period. Sevoflurane and isoflurane vaporizer '
        'prices remained comparatively stable throughout\u2014despite being traded on the same '
        'marketplace and subject to the same macroeconomic conditions. The convergence of '
        'evidence from Spearman rank correlation, Kendall \u03c4, and the between-agent effect '
        f'size comparison (P={fmt_p(des_vs_sevo["p"])}) '
        'suggests a progressive and agent-specific desflurane price pattern. However, the '
        'single-platform, observational design precludes causal inference; the observed '
        'association may be partly or wholly attributable to concurrent secular trends.')
    if has_revision_sensitivity:
        did_disc = get_rev('did_transaction_level')
        if isinstance(did_disc, dict) and 'error' not in did_disc:
            doc.add_paragraph(
                f'The transaction-level difference-in-differences (DiD) model did not detect a '
                f'significant desflurane post-restriction coefficient '
                f'(coefficient={did_disc["desflurane_post_coef"]:.3f}, '
                f'P={fmt_p(did_disc["desflurane_post_p"])}). This null result should not be '
                f'interpreted as evidence of no effect. Instead, the DiD parallel-trend '
                f'assumption was violated '
                f'(P={fmt_p(did_disc["desflurane_trend_p"])} for the pre-restriction trend '
                f'difference), indicating that desflurane prices were already diverging from '
                f'controls before 1 January 2026. Because the regulation was proposed in 2022 and '
                f'adopted in 2024, market participants may have anticipated the 2026 effective date; '
                f'the DiD violation is consistent with a gradual, anticipation-driven decline '
                f'rather than a discrete shock. The DiD and CITS results are best interpreted as '
                f'sensitivity, not causal, analyses.')
    doc.add_paragraph(
        'We did not family-wise correct trend tests across three agents. For isoflurane, '
        f'Spearman (P={fmt_p(iso_tr["spearman_p"])}) and Kendall \u03c4 '
        f'(P={fmt_p(iso_tr["kendall_p"])}) tests were nominally significant at the transaction '
        'level, but the quarterly median trend was not significant and effect sizes were small; '
        'these isolated P values are expected by chance and are not interpreted as a true trend. '
        'The desflurane finding is supported by consistency across Spearman, Kendall \u03c4, '
        'and quarterly median analyses.')

    sevo_vs_iso = es_comparisons['Sevoflurane_vs_Isoflurane']
    doc.add_paragraph(
        'The stability of non-regulated agent prices is noteworthy, but it does not '
        'demonstrate that the EU regulation caused the desflurane decline. Several '
        'alternative explanations should be considered. NHS England and NHS Scotland '
        'decommissioned desflurane before the EU restriction took effect, so part of the '
        'price decline may reflect UK policy signals rather than EU regulation. A '
        'concurrent shift toward total intravenous anesthesia (TIVA) and the growing use of '
        'low-flow anesthesia and end-tidal control technologies reduce volatile anesthetic '
        'consumption generally, but should affect all agents rather than desflurane alone. '
        'Equipment-specific factors\u2014age, model year, service history, cosmetic condition, '
        'and calibration certificates\u2014are major determinants of vaporizer pricing and '
        'could produce compositional changes that mimic a regulatory effect. Macroeconomic '
        'trends and institutional stock reallocation\u2014whereby facilities redistributed '
        'existing non-desflurane vaporizers internally rather than purchasing on the '
        'secondary market\u2014may also have dampened market effects. The comparators '
        '(sevoflurane and isoflurane) are not perfect counterfactuals, but their stability '
        'on the same marketplace weakens broad-market explanations.')
    doc.add_paragraph(
        'Generalizability is also limited. eBay is a global consumer-to-consumer and small-business '
        'marketplace; transactions by large hospital systems are more likely to occur on '
        'specialized business-to-business platforms (e.g., DOTmed, Bimedis) or through private '
        'dealer networks, which we could not study. Because eBay does not disclose buyer or '
        'seller location, we could not determine whether observed transactions involved EU '
        'participants; if a substantial share were outside the EU, the price decline could '
        'represent a global reassessment of desflurane\u2019s long-term viability rather than a '
        'response confined to EU jurisdiction. Finally, the sample of post-restriction '
        'transactions was small, and the bootstrap median-difference confidence interval for '
        'desflurane included zero, underscoring the uncertainty of the pre-/post comparison.')

    add_para_with_refs(doc,
        'To our knowledge, no previous study has examined the secondary market impact of '
        'environmental regulation on anesthetic equipment. Lehmann et al.{16} demonstrated '
        'that combining education with physical removal of desflurane vaporizers reduced '
        'desflurane-attributable CO\u2082 equivalent emissions by 86%, but their study '
        'measured drug consumption rather than equipment resale values. Meyer{13} and Mohammed '
        'and Metta{15} articulated the global and financial rationale for desflurane '
        'discontinuation, while Moonesinghe{14} discussed the broader implications of '
        'decommissioning programs, but none examined downstream effects on the secondary '
        'equipment market. Our findings are consistent with the broader literature on '
        'environment-related stranded assets,{21} where anticipated regulatory restrictions '
        'are associated with anticipatory declines in asset values. The clinical and '
        'environmental trade-offs of restricting desflurane\u2014including its pharmacokinetic '
        'advantages, the contested suitability of GWP100 for short-lived anesthetics, and '
        'the risk of reducing anesthetic diversity amid recurring propofol and sevoflurane '
        'shortages\u2014 underscore the need for cautious, evidence-based policy and for '
        'market studies that do not rely on a single sales channel.{6,7,22}')

    doc.add_paragraph(
        'The timing of the price decline may have practical implications. A substantial '
        'proportion of the depreciation appears to have occurred before the 2026 '
        'restriction took effect, during the legislative process itself. Because NHS '
        'England and NHS Scotland had already announced decommissioning, the pre-2026 '
        'pattern may partly reflect UK policy signals rather than EU law alone. For '
        'anesthesia departments, this observation suggests that early planning for '
        'equipment transition\u2014rather than waiting for formal prohibition\u2014may '
        'improve cost recovery on the secondary market, although this inference is '
        'observational and cannot be translated directly to all jurisdictions.')

    doc.add_paragraph(
        'Strengths of this study include the use of actual completed sale prices (rather than '
        'asking prices), a three-year observation window spanning both the legislative process '
        'and restriction implementation, the use of multiple complementary statistical approaches '
        '(Spearman correlation, Kendall \u03c4 trend test, pre-/post-ban comparison, '
        'bootstrap resampling, post-hoc power estimation, and CITS/DiD sensitivity analyses), '
        'the availability of natural comparator groups (sevoflurane and isoflurane), '
        'and the use of a standardized data source (eBay Terapeak). '
        'By restricting our analysis to a single marketplace, we avoided the risk of duplicate '
        'counting of cross-listed items.')

    doc.add_paragraph(
        f'This study has several important limitations. First, this is an observational study '
        f'of secondary market data; no causal inference can be drawn. The association between '
        f'regulatory milestones and price changes may be confounded by unmeasured factors, '
        f'including changes in clinical practice patterns (e.g., increasing use of total '
        f'intravenous anesthesia), technological evolution of anesthesia delivery systems '
        f'(e.g., low-flow and end-tidal control systems), NHS decommissioning announcements '
        f'that pre-dated the 2026 EU restriction, or broader macroeconomic conditions. Second, '
        f'eBay represents only one segment of the secondary medical equipment market. Prices '
        f'on specialized platforms (e.g., DOTmed, Bimedis) or private dealer networks may '
        f'behave differently, and our findings may not generalize to those channels. Third, we '
        f'could not control for equipment age, model year, service history, cosmetic condition, '
        f'or the presence of manufacturer calibration certificates\u2014factors that '
        f'substantially influence vaporizer pricing and could create compositional effects. '
        f'Fourth, the post-restriction period (January\u2013March 2026) comprised only '
        f'{des["post_n"]} desflurane, {sevo["post_n"]} sevoflurane, and '
        f'{iso["post_n"]} isoflurane transactions. This small sample reflects the Terapeak '
        f'three-year window ending at data collection in March 2026, which constrains the '
        f'post-restriction follow-up to approximately three months after the 1 January 2026 '
        f'effective date and precludes extending the window without prospective data collection '
        f'or a later extraction. Bootstrap resampling and post-hoc power simulation confirmed '
        f'that the desflurane Mann\u2013Whitney U test had low power (approximately 28%) and '
        f'that the median-difference confidence interval included zero; the Welch t-test was '
        f'more powerful (approximately 58%) and the mean-difference confidence interval '
        f'excluded zero. Although the time-series trend analyses (which use all data points) '
        f'suggest a progressive decline, the pre-/post-ban comparison should be considered '
        f'exploratory. Fifth, eBay is a global marketplace; we could not distinguish between '
        f'EU and non-EU buyers or sellers, nor could we assess whether sellers were institutions '
        f'disposing of regulated equipment or private resellers. Finally, the three-year '
        f'observation period does not extend to the pre-proposal period (before April 2022), '
        f'limiting our ability to establish a true baseline unaffected by regulatory signals. The '
        f'DiD and CITS models support agent-specificity but do not satisfy the assumptions '
        f'required for causal identification; they should be interpreted as sensitivity, not '
        f'confirmatory, analyses. Additionally, multiple trend tests across three agents '
        f'inflate the chance of a nominally significant P value; isolated P values near 0.05 '
        f'(e.g., isoflurane Spearman P={fmt_p(iso_tr["spearman_p"])}) are interpreted cautiously '
        f'and in the context of effect size and consistency across tests.')

    add_para_with_refs(doc,
        'Looking ahead, environmental pressures are likely to prompt further regulatory '
        'interventions in anesthesia and healthcare more broadly. Nitrous oxide, for example, '
        'is already subject to emerging regulatory and institutional restrictions on '
        'environmental grounds.{8} Our findings provide preliminary, hypothesis-generating '
        'evidence that the EU desflurane restriction on routine use\u2014the first mandatory, '
        'agent-specific environmental restriction in anesthesia\u2014was associated with '
        'targeted secondary-market price patterns without observable destabilization of the '
        'wider equipment market. Future studies with larger post-restriction samples, multiple '
        'marketplaces, geographic transaction data, and controlled comparisons will be needed '
        'to confirm these findings and to determine whether they generalize to other '
        'regulatory contexts.')

    # ============================================================
    # CONCLUSION (JATM: singular)
    # ============================================================
    add_heading_styled(doc, 'Conclusion', level=1)
    doc.add_paragraph(
        'In this observational study, desflurane vaporizer prices on eBay showed a progressive, '
        'agent-specific decline during the EU regulatory process, while sevoflurane and '
        'isoflurane prices remained comparatively stable. These findings are consistent with '
        'the hypothesis that the EU desflurane phase-out was associated with targeted '
        'secondary-market effects, but they do not prove causation. They should be interpreted '
        'as hypothesis-generating, given the single-platform design, the inability to identify '
        'EU-specific transactions, the small post-restriction sample, and the possibility of '
        'unmeasured confounders such as NHS decommissioning and secular clinical trends.')

    # ============================================================
    # DECLARATIONS (JATM order)
    # ============================================================
    add_heading_styled(doc, 'CRediT authorship contribution statement', level=1)
    doc.add_paragraph('O.T.: Conceptualization, Methodology, Software, Formal analysis, '
                      'Writing - original draft, Writing - review & editing')
    doc.add_paragraph('T.I.: Writing - original draft, Writing - review & editing')

    add_heading_styled(doc, 'Disclosure statement', level=1)
    doc.add_paragraph('The authors declare that they have no known competing financial '
                      'interests or personal relationships that could have appeared to '
                      'influence the work reported in this paper.')

    add_heading_styled(doc, 'Ethical statement', level=1)
    doc.add_paragraph(
        'Ethical approval was not required for this study, which analyzed publicly available '
        'completed sale data from eBay. No individual-level or patient data were collected.')

    add_heading_styled(doc, 'Funding', level=1)
    doc.add_paragraph('This research did not receive any specific grant from funding '
                      'agencies in the public, commercial, or not-for-profit sectors.')

    add_heading_styled(doc, 'Data availability', level=1)
    doc.add_paragraph(
        'The datasets generated during this study are available from the corresponding author '
        'on reasonable request. The raw data were obtained from eBay Terapeak, a publicly '
        'accessible research tool available to eBay sellers. '
        'Analysis code is available at https://github.com/bougtoir/vaporizer-price-study.')

    add_heading_styled(doc, 'Declaration of competing interest', level=1)
    doc.add_paragraph('The authors declare that they have no known competing financial '
                      'interests or personal relationships that could have appeared to '
                      'influence the work reported in this paper.')

    add_heading_styled(doc, 'Acknowledgments', level=1)
    doc.add_paragraph('None.')

    add_heading_styled(doc, 'Declaration of Generative Artificial Intelligence (AI) in Scientific Writing', level=1)
    doc.add_paragraph(
        'During the preparation of this work the author used Devin (devin.ai) in order to format '
        'the text and choose words that suited the tone, and to help writing codes. After using '
        'this tool/service, the author reviewed and edited the content as needed and takes full '
        'responsibility for the content of the published article.')

    doc.add_page_break()

    # ============================================================
    # REFERENCES (JATM Vancouver style)
    # 3+ authors: list first 3, then et al.
    # Format: Author. Title. J Abbrev. Year;Vol:Pages. DOI
    # ============================================================
    add_heading_styled(doc, 'References', level=1)
    references = [
        # 1
        'Varughese S, Ahmed R. Environmental and occupational considerations of anesthesia: '
        'a narrative review and update. Anesth Analg. 2021;133:826\u201335.',
        # 2
        'Regulation (EU) 2024/573 of the European Parliament and of the Council of '
        '7 February 2024 on fluorinated greenhouse gases. Official Journal of the European '
        'Union. 2024;L 2024/573.',
        # 3
        'Sherman J, Le C, Lamers V, et al. Life cycle greenhouse gas emissions of '
        'anesthetic drugs. Anesth Analg. 2012;114:1086\u201390.',
        # 4
        'Hendrickx JFA, Nielsen OJ, De Hert S, et al. The science behind banning '
        'desflurane: a narrative review. Eur J Anaesthesiol. 2022;39:818\u201324.',
        # 5
        'White SM, Shelton CL, Gelb AW, et al. Principles of environmentally-sustainable '
        'anaesthesia: a global consensus statement from the World Federation of Societies of '
        'Anaesthesiologists. Anaesthesia. 2022;77:201\u201312.',
        # 6
        'Sulbaek Andersen MP, Nielsen OJ, Sherman JD. Assessing the potential climate impact '
        'of anaesthetic gases. Lancet Planet Health. 2023;7(7):e622\u2013e629. '
        'doi:10.1016/S2542-5196(23)00084-0.',
        # 7
        'Hariyanto H, Widiastuti M, Pandrya CO, et al. Comparison of desflurane and sevoflurane '
        'as maintenance inhalational anaesthetic agents for adult patients undergoing '
        'neurosurgeries: a systematic review and meta-analysis of randomised trials. '
        'Indian J Anaesth. 2025;69(1):65\u201377. doi:10.4103/ija.ija_1215_24.',
        # 8
        'American Society of Anesthesiologists Committee on Environmental Health. Statement on '
        'deactivating central piped nitrous oxide to mitigate avoidable health care pollution. '
        'Schaumburg, IL: ASA, 2024.',
        # 9
        'McGain F, Muret J, Guen CL, et al. Environmental sustainability in anaesthesia '
        'and critical care. Br J Anaesth. 2020;125:680\u201392.',
        # 10
        'NHS England. Decommissioning of desflurane in the NHS. 2023.',
        # 11
        'Sulbaek Andersen MP, Sander SP, Nielsen OJ, et al. Inhalation anaesthetics and '
        'climate change. Br J Anaesth. 2010;105:760\u20136.',
        # 12
        'Ryan SM, Nielsen CJ. Global warming potential of inhaled anesthetics: application '
        'to clinical use. Anesth Analg. 2010;111:92\u20138.',
        # 13
        'Meyer MJ. Desflurane should des-appear: global and financial rationale. Anesth Analg. '
        '2020;131:1317\u201322.',
        # 14
        'Moonesinghe SR. Desflurane decommissioning: more than meets the eye. Anaesthesia. '
        '2024;79:237\u201341.',
        # 15
        'Mohammed A, Metta H. Is it time to bid adieu to desflurane? J Anaesthesiol Clin '
        'Pharmacol. 2025;41:211\u20132.',
        # 16
        'Lehmann H, Werning J, Baschnegger H, et al. Minimising the usage of desflurane '
        'only by education and removal of the vaporisers \u2013 a before-and-after-trial. '
        'BMC Anesthesiol. 2025;25:108.',
        # 17
        'Feldman JM, Lo C, Hendrickx J. Estimating the impact of carbon dioxide absorbent '
        'performance differences on absorbent cost during low-flow anesthesia. Anesth Analg. '
        '2020;130(2):374\u201381. doi:10.1213/ANE.0000000000004059.',
        # 18
        'Beard J, Kennedy R, Philip J, et al. Environmental and economic impacts of end-tidal '
        'control of volatile anesthetics: a scoping review and analysis. Open Anesthesiol J. '
        '2025;19:e25896458355905.',
        # 19
        'Shukla S, Kalaiselvan V, Raghuvanshi RS. How to improve regulatory practices for '
        'refurbished medical devices. Bull World Health Organ. 2023;101:412\u20137.',
        # 20
        'von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of '
        'Observational Studies in Epidemiology (STROBE) statement: guidelines for reporting '
        'observational studies. BMJ. 2007;335:806\u20138.',
        # 21
        'Shimbar A. Environment-related stranded assets: what does the market think about the '
        'impact of collective climate action on the value of fossil fuel stocks? Energy Econ. '
        '2021;103:105579.',
        # 22
        'Kranke P, Kleinberg R, Meybohm P, et al. Anesthetic diversity: a pillar of healthcare '
        'system resilience and a strategic imperative in an era of uncertainty. Anaesthesiologie. '
        '2026. doi:10.1007/s00101-025-01639-x.',
    ]
    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        run_num = p.add_run(f'{i} ')
        run_num.font.size = Pt(10)
        run_num.font.superscript = True
        run_text = p.add_run(ref)
        run_text.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    # Supplementary table placeholder
    if has_asking_data:
        doc.add_page_break()
        add_heading_styled(doc, 'Supplementary material', level=1)
        p = doc.add_paragraph()
        add_run_styled(p, '[Table S1: Current eBay asking prices \u2014 uploaded as separate file]',
                       bold=True, italic=True, size=Pt(12))

    path = os.path.join(out_dir, 'jatm_manuscript_english.docx')
    doc.save(path)
    print(f"JATM English paper saved: {path}")
    return path


def compute_trend_data():
    """Compute trend data using same logic as create_jca_tables.py."""
    df = combined.copy()
    df['period_num'] = df['date_sold'].apply(classify_period)
    trend_results_local = {}
    for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
        ad = df[df['agent_type'] == agent].copy()
        ad['date_ord'] = ad['date_sold'].map(lambda d: d.toordinal())
        sp_rho, sp_p = sp_stats.spearmanr(ad['date_ord'], ad['price_usd'])
        kt_tau, kt_p = sp_stats.kendalltau(ad['period_num'], ad['price_usd'])
        ad['quarter'] = ad['date_sold'].dt.to_period('Q')
        qm = ad.groupby('quarter')['price_usd'].agg(['median', 'count'])
        qm = qm[qm['count'] >= 3]
        if len(qm) >= 4:
            q_rho, q_p = sp_stats.spearmanr(range(len(qm)), qm['median'].values)
        else:
            q_rho, q_p = float('nan'), float('nan')
        trend_results_local[agent] = {
            'spearman_rho': sp_rho, 'spearman_p': sp_p,
            'kendall_tau': kt_tau, 'kendall_p': kt_p,
            'quarterly_rho': q_rho, 'quarterly_p': q_p,
        }
    return trend_results_local


if __name__ == '__main__':
    write_jatm_paper()
