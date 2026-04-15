"""
Generate STROBE checklist for cross-sectional studies as editable .docx files.
English and Japanese versions.
Maps each STROBE item to the specific section/page of our manuscript.
"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

outdir = '/home/ubuntu/vaporizer_research/papers/'
os.makedirs(outdir, exist_ok=True)

def setup_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    pf = style.paragraph_format
    pf.line_spacing = 1.15
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.page_width = Inches(11.69)  # A4 landscape
        section.page_height = Inches(8.27)
    return doc

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


# STROBE checklist items for cross-sectional studies
# Format: (item_no, section, recommendation, manuscript_location_en, manuscript_location_ja)
STROBE_ITEMS = [
    # Title and Abstract
    ('1a', 'Title and abstract',
     'Indicate the study\u2019s design with a commonly used term in the title or the abstract',
     'Title: "...a cross-sectional time-series analysis of eBay sold listings";\nSummary: "Cross-sectional time-series analysis of completed (sold) listings"',
     '\u30bf\u30a4\u30c8\u30eb: \u300c\u2026\u6a2a\u65ad\u7684\u6642\u7cfb\u5217\u5206\u6790\u300d;\n\u8981\u65e8: \u300c\u58f2\u8cb7\u5b9f\u7e3e\u306e\u6a2a\u65ad\u7684\u6642\u7cfb\u5217\u5206\u6790\u300d'),

    ('1b', 'Title and abstract',
     'Provide in the abstract an informative and balanced summary of what was done and what was found',
     'Summary: Unstructured single-paragraph summary covering objectives, methods, key findings, and conclusions',
     '\u8981\u65e8: \u76ee\u7684\u3001\u65b9\u6cd5\u3001\u4e3b\u8981\u306a\u77e5\u898b\u3001\u7d50\u8ad6\u3092\u542b\u3080\u975e\u69cb\u9020\u5316\u5358\u4e00\u6bb5\u843d\u306e\u8981\u65e8'),

    # Introduction
    ('2', 'Background/rationale',
     'Explain the scientific background and rationale for the investigation being reported',
     'Introduction, paragraphs 1\u20134: GWP of anaesthetic agents, EU regulatory timeline, economic significance of vaporizers, gap in knowledge',
     '\u7dd2\u8a00, \u7b2c1\u20134\u6bb5\u843d: \u5438\u5165\u9ebb\u9154\u85ac\u306eGWP\u3001EU\u898f\u5236\u30bf\u30a4\u30e0\u30e9\u30a4\u30f3\u3001\u6c17\u5316\u5668\u306e\u7d4c\u6e08\u7684\u610f\u7fa9\u3001\u77e5\u8b58\u306e\u30ae\u30e3\u30c3\u30d7'),

    ('3', 'Objectives',
     'State specific objectives, including any prespecified hypotheses',
     'Introduction, paragraph 4: "We hypothesised that the EU desflurane regulation would be associated with a progressive decrease in secondary market prices for desflurane vaporizers specifically, while prices for sevoflurane and isoflurane vaporizers would remain stable."',
     '\u7dd2\u8a00, \u7b2c4\u6bb5\u843d: \u300c\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u6c17\u5316\u5668\u4fa1\u683c\u306e\u6bb5\u968e\u7684\u4e0b\u843d\u3068\u3001\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u30fb\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u4fa1\u683c\u306e\u5b89\u5b9a\u3092\u4eee\u8aac\u3068\u3057\u3066\u8a2d\u5b9a\u300d'),

    # Methods
    ('4', 'Study design',
     'Present key elements of study design early in the paper',
     'Methods, paragraph 1: "This study is reported following the STROBE guidelines..."\nStudy design and data source: "We conducted a cross-sectional time-series analysis..."',
     '\u65b9\u6cd5, \u7b2c1\u6bb5\u843d: \u300cSTROBE\u30ac\u30a4\u30c9\u30e9\u30a4\u30f3\u306b\u6e96\u62e0\u300d;\n\u7814\u7a76\u30c7\u30b6\u30a4\u30f3\u3068\u30c7\u30fc\u30bf\u30bd\u30fc\u30b9: \u300c\u6a2a\u65ad\u7684\u6642\u7cfb\u5217\u5206\u6790\u3092\u5b9f\u65bd\u300d'),

    ('5', 'Setting',
     'Describe the setting, locations, and relevant dates, including periods of recruitment, exposure, follow-up, and data collection',
     'Study design and data source: "eBay (www.ebay.com)...Data were collected in March 2026, covering the period from 28 March 2023 to 24 March 2026"',
     '\u7814\u7a76\u30c7\u30b6\u30a4\u30f3\u3068\u30c7\u30fc\u30bf\u30bd\u30fc\u30b9: \u300ceBay\u2026\u30c7\u30fc\u30bf\u306f2026\u5e743\u6708\u306b\u53ce\u96c6\u3001\u5bfe\u8c61\u671f\u9593\u306f2023\u5e743\u670828\u65e5\uff5e2026\u5e743\u670824\u65e5\u300d'),

    ('6', 'Participants',
     'Cross-sectional study\u2014Give the eligibility criteria, and the sources and methods of selection of participants',
     'Eligibility criteria: Inclusion (completed sales, standalone vaporizers, valid price/date) and exclusion criteria (non-vaporizer items, veterinary systems, lot listings, implausible data) fully described',
     '\u9069\u683c\u57fa\u6e96: \u7d44\u5165\u57fa\u6e96\uff08\u58f2\u8cb7\u5b8c\u4e86\u3001\u5358\u4f53\u6c17\u5316\u5668\u3001\u6709\u52b9\u306a\u4fa1\u683c/\u65e5\u4ed8\uff09\u304a\u3088\u3073\u9664\u5916\u57fa\u6e96\uff08\u975e\u6c17\u5316\u5668\u54c1\u3001\u7363\u533b\u7528\u3001\u30ed\u30c3\u30c8\u51fa\u54c1\u3001\u7570\u5e38\u30c7\u30fc\u30bf\uff09\u3092\u8a73\u8ff0'),

    ('7', 'Variables',
     'Clearly define all outcomes, exposures, predictors, potential confounders, and effect modifiers. Give diagnostic criteria, if applicable',
     'Variables: Primary outcome (sale price in USD), exposure (regulatory period classified by milestones), covariates (item title, sale date, quantity). Multi-period classification described.',
     '\u5909\u6570: \u4e3b\u8981\u30a2\u30a6\u30c8\u30ab\u30e0\uff08\u58f2\u5374\u4fa1\u683c USD\uff09\u3001\u66dd\u9732\uff08\u898f\u5236\u6bb5\u968e\u5225\u5206\u985e\uff09\u3001\u5171\u5909\u91cf\uff08\u30a2\u30a4\u30c6\u30e0\u540d\u3001\u58f2\u5374\u65e5\u3001\u6570\u91cf\uff09\u3002\u591a\u671f\u9593\u5206\u985e\u3092\u8a18\u8ff0'),

    ('8', 'Data sources/measurement',
     'For each variable of interest, give sources of data and details of methods of assessment (measurement). Describe comparability of assessment methods if there is more than one group',
     'Study design and data source: "Terapeak, eBay\u2019s official product research tool...provides access to up to three years of historical completed sale data, including item titles, sale prices, sale dates, and quantities sold"',
     '\u7814\u7a76\u30c7\u30b6\u30a4\u30f3\u3068\u30c7\u30fc\u30bf\u30bd\u30fc\u30b9: \u300cTerapeak\u306feBay\u516c\u5f0f\u306e\u88fd\u54c1\u30ea\u30b5\u30fc\u30c1\u30c4\u30fc\u30eb\u2026\u30a2\u30a4\u30c6\u30e0\u540d\u3001\u58f2\u5374\u4fa1\u683c\u3001\u58f2\u5374\u65e5\u3001\u6570\u91cf\u3092\u542b\u3080\u300d'),

    ('9', 'Bias',
     'Describe any efforts to address potential sources of bias',
     'Strengths and limitations: Single marketplace to avoid cross-listing duplicates; acknowledged inability to control for equipment age/condition; acknowledged global marketplace limitation (EU vs non-EU); acknowledged platform selection bias',
     '\u9577\u6240\u3068\u9650\u754c: \u30af\u30ed\u30b9\u30ea\u30b9\u30c6\u30a3\u30f3\u30b0\u91cd\u8907\u56de\u907f\u306e\u305f\u3081\u5358\u4e00\u30de\u30fc\u30b1\u30c3\u30c8\u30d7\u30ec\u30a4\u30b9\u306b\u9650\u5b9a\uff1b\u6a5f\u5668\u306e\u5e74\u6570/\u72b6\u614b\u306e\u5236\u5fa1\u4e0d\u80fd\u3001\u30b0\u30ed\u30fc\u30d0\u30eb\u5e02\u5834\u306e\u9650\u754c\u3001\u30d7\u30e9\u30c3\u30c8\u30d5\u30a9\u30fc\u30e0\u9078\u629e\u30d0\u30a4\u30a2\u30b9\u3092\u8a8d\u8b58'),

    ('10', 'Study size',
     'Explain how the study size was arrived at',
     'Statistical analysis: "No a priori sample size calculation was performed, as this study aimed to capture all available transactions within the Terapeak data window"',
     '\u7d71\u8a08\u5206\u6790: \u300c\u4e8b\u524d\u306e\u30b5\u30f3\u30d7\u30eb\u30b5\u30a4\u30ba\u8a08\u7b97\u306f\u5b9f\u65bd\u305b\u305a\u3001Terapeak\u30c7\u30fc\u30bf\u30a6\u30a3\u30f3\u30c9\u30a6\u5185\u306e\u5168\u53d6\u5f15\u3092\u53ce\u96c6\u300d'),

    ('11', 'Quantitative variables',
     'Explain how quantitative variables were handled in the analyses. If applicable, describe which groupings were chosen and why',
     'Variables: Regulatory period classification (4 phases based on milestones) described with rationale.\nStatistical analysis: LOWESS smoothing, quarterly aggregation described',
     '\u5909\u6570: \u898f\u5236\u6bb5\u968e\u5206\u985e\uff084\u6bb5\u968e\u3001\u30de\u30a4\u30eb\u30b9\u30c8\u30fc\u30f3\u57fa\u6e96\uff09\u3092\u8aac\u660e\u3002\n\u7d71\u8a08\u5206\u6790: LOWESS\u5e73\u6ed1\u5316\u3001\u56db\u534a\u671f\u96c6\u8a08\u3092\u8a18\u8ff0'),

    ('12a', 'Statistical methods',
     'Describe all statistical methods, including those used to control for confounding',
     'Statistical analysis: Mann-Whitney U, Welch\u2019s t-test, Cohen\u2019s d, Spearman rank correlation, Kendall \u03c4, Kruskal-Wallis, LOWESS. Software: Python 3.12, pandas 2.2, scipy 1.14, statsmodels 0.14',
     '\u7d71\u8a08\u5206\u6790: Mann-Whitney U\u3001Welch t\u691c\u5b9a\u3001Cohen d\u3001Spearman\u9806\u4f4d\u76f8\u95a2\u3001Kendall \u03c4\u3001Kruskal-Wallis\u3001LOWESS\u3002\u30bd\u30d5\u30c8: Python 3.12, pandas 2.2, scipy 1.14, statsmodels 0.14'),

    ('12b', 'Statistical methods',
     'Describe any methods used to examine subgroups and interactions',
     'Statistical analysis: Agent-specific analyses (desflurane, sevoflurane, isoflurane treated independently); multi-period analysis across regulatory phases',
     '\u7d71\u8a08\u5206\u6790: \u85ac\u5264\u5225\u5206\u6790\uff08\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u3001\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u3001\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u3092\u72ec\u7acb\u306b\u5206\u6790\uff09\uff1b\u898f\u5236\u6bb5\u968e\u5225\u306e\u591a\u671f\u9593\u5206\u6790'),

    ('12c', 'Statistical methods',
     'Explain how missing data were addressed',
     'Eligibility criteria: Exclusion of "listings with missing or implausible price data"',
     '\u9069\u683c\u57fa\u6e96: \u300c\u4fa1\u683c\u30c7\u30fc\u30bf\u306e\u6b20\u640d\u307e\u305f\u306f\u7570\u5e38\u5024\u306e\u3042\u308b\u51fa\u54c1\u3092\u9664\u5916\u300d'),

    ('12d', 'Statistical methods',
     'Cross-sectional study\u2014If applicable, describe analytical methods taking account of sampling strategy',
     'N/A \u2014 Complete enumeration of all available transactions (no sampling)',
     'N/A \u2014 \u5168\u53d6\u5f15\u306e\u5b8c\u5168\u679a\u6319\uff08\u30b5\u30f3\u30d7\u30ea\u30f3\u30b0\u306a\u3057\uff09'),

    ('12e', 'Statistical methods',
     'Describe any sensitivity analyses',
     'Statistical analysis: Welch\u2019s t-test as sensitivity analysis alongside Mann-Whitney U; quarterly median Spearman as aggregated-level sensitivity check; supplementary asking price analysis as robustness check',
     '\u7d71\u8a08\u5206\u6790: Mann-Whitney U\u306b\u52a0\u3048Welch t\u691c\u5b9a\u3092\u611f\u5ea6\u5206\u6790\u3068\u3057\u3066\u5b9f\u65bd\uff1b\u56db\u534a\u671f\u4e2d\u592e\u5024Spearman\u3092\u96c6\u8a08\u30ec\u30d9\u30eb\u306e\u611f\u5ea6\u691c\u8a3c\uff1b\u51fa\u54c1\u4fa1\u683c\u306e\u88dc\u8db3\u5206\u6790\u3092\u9811\u5065\u6027\u691c\u8a3c\u3068\u3057\u3066\u5b9f\u65bd'),

    # Results
    ('13a', 'Participants',
     'Report numbers of individuals at each stage of study\u2014e.g., numbers potentially eligible, examined for eligibility, confirmed eligible, included in the study, completing follow-up, and analysed',
     'Results, Study population: "A total of 1,033 completed eBay sales...after applying exclusion criteria: [n] desflurane, [n] sevoflurane, and [n] isoflurane"',
     '\u7d50\u679c, \u7814\u7a76\u5bfe\u8c61: \u300c\u9664\u5916\u57fa\u6e96\u9069\u7528\u5f8c\u3001\u5408\u8a081,033\u4ef6\u306eeBay\u58f2\u8cb7\u5b9f\u7e3e: \u30c7\u30b9\u30d5\u30eb\u30e9\u30f3[n]\u4ef6\u3001\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3[n]\u4ef6\u3001\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3[n]\u4ef6\u300d'),

    ('13b', 'Participants',
     'Give reasons for non-participation at each stage',
     'Eligibility criteria: Exclusion criteria specified (non-vaporizer items, veterinary systems, lot listings, implausible data)',
     '\u9069\u683c\u57fa\u6e96: \u9664\u5916\u7406\u7531\u3092\u660e\u793a\uff08\u975e\u6c17\u5316\u5668\u54c1\u3001\u7363\u533b\u7528\u30b7\u30b9\u30c6\u30e0\u3001\u30ed\u30c3\u30c8\u51fa\u54c1\u3001\u7570\u5e38\u30c7\u30fc\u30bf\uff09'),

    ('14', 'Descriptive data',
     'Give characteristics of study participants (e.g., demographic, clinical, social) and information on exposures and potential confounders',
     'Table 1: n, mean (SD), median (IQR), range by agent type and regulatory period.\nResults: Vaporizer models described (Tec 6, D-Vapor, Vapor 2000, Sigma Delta, etc.)',
     'Table 1: \u85ac\u5264\u578b\u30fb\u898f\u5236\u671f\u9593\u5225\u306en\u3001\u5e73\u5747(SD)\u3001\u4e2d\u592e\u5024(IQR)\u3001\u7bc4\u56f2\u3002\n\u7d50\u679c: \u6c17\u5316\u5668\u30e2\u30c7\u30eb\u3092\u8a18\u8ff0\uff08Tec 6, D-Vapor, Vapor 2000\u7b49\uff09'),

    ('15', 'Outcome data',
     'Cross-sectional study\u2014Report numbers of outcome events or summary measures',
     'Table 1: Mean, SD, median, IQR, range for each agent \u00d7 period.\nTable 2: Spearman \u03c1, Kendall \u03c4, quarterly \u03c1 with P values.\nFigures 1\u20136: Visual summaries',
     'Table 1: \u5404\u85ac\u5264\u00d7\u671f\u9593\u306e\u5e73\u5747\u3001SD\u3001\u4e2d\u592e\u5024\u3001IQR\u3001\u7bc4\u56f2\u3002\nTable 2: Spearman \u03c1\u3001Kendall \u03c4\u3001\u56db\u534a\u671f\u03c1\u3068P\u5024\u3002\nFigures 1\u20136: \u8996\u899a\u7684\u8981\u7d04'),

    ('16a', 'Main results',
     'Give unadjusted estimates and, if applicable, confounder-adjusted estimates and their precision (e.g., 95% CI). Make clear which confounders were adjusted for and why they were included',
     'Results: Spearman \u03c1, Kendall \u03c4, Mann-Whitney U P values, Welch\u2019s t-test P values, Cohen\u2019s d reported for each agent. No confounder adjustment (observational market data).',
     '\u7d50\u679c: Spearman \u03c1\u3001Kendall \u03c4\u3001Mann-Whitney U P\u5024\u3001Welch t P\u5024\u3001Cohen d\u3092\u5404\u85ac\u5264\u5225\u306b\u5831\u544a\u3002\u4ea4\u7d61\u56e0\u5b50\u8abf\u6574\u306a\u3057\uff08\u89b3\u5bdf\u7684\u5e02\u5834\u30c7\u30fc\u30bf\uff09'),

    ('16b', 'Main results',
     'Report category boundaries when continuous variables were categorized',
     'Variables: Regulatory phases defined by specific dates (EC proposal April 2022, trilogue October 2023, adoption February 2024, ban January 2026)',
     '\u5909\u6570: \u898f\u5236\u6bb5\u968e\u3092\u5177\u4f53\u7684\u306a\u65e5\u4ed8\u3067\u5b9a\u7fa9\uff08EC\u63d0\u6848 2022\u5e744\u6708\u3001\u30c8\u30ea\u30ed\u30fc\u30b0 2023\u5e7410\u6708\u3001\u63a1\u629e 2024\u5e742\u6708\u3001\u7981\u6b62 2026\u5e741\u6708\uff09'),

    ('17', 'Other analyses',
     'Report other analyses done\u2014e.g., analyses of subgroups and interactions, and sensitivity analyses',
     'Multi-period analysis: Kruskal-Wallis across 4 phases.\nSupplementary analysis: Asking price cross-sectional analysis (eTable 1).\nSensitivity: Welch\u2019s t-test, quarterly median trend',
     '\u591a\u671f\u9593\u5206\u6790: 4\u6bb5\u968e\u306eKruskal-Wallis\u3002\n\u88dc\u8db3\u5206\u6790: \u51fa\u54c1\u4fa1\u683c\u6a2a\u65ad\u5206\u6790\uff08eTable 1\uff09\u3002\n\u611f\u5ea6\u5206\u6790: Welch t\u691c\u5b9a\u3001\u56db\u534a\u671f\u4e2d\u592e\u5024\u30c8\u30ec\u30f3\u30c9'),

    # Discussion
    ('18', 'Key results',
     'Summarise key results with reference to study objectives',
     'Discussion, Principal findings: Agent-specific price decline for desflurane confirmed; sevoflurane and isoflurane stable as hypothesised',
     '\u8003\u5bdf, \u4e3b\u8981\u306a\u77e5\u898b: \u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u306e\u85ac\u5264\u7279\u7570\u7684\u4fa1\u683c\u4e0b\u843d\u3092\u78ba\u8a8d\uff1b\u30bb\u30dc\u30d5\u30eb\u30e9\u30f3\u30fb\u30a4\u30bd\u30d5\u30eb\u30e9\u30f3\u306f\u4eee\u8aac\u901a\u308a\u5b89\u5b9a'),

    ('19', 'Limitations',
     'Discuss limitations of the study, taking into account sources of potential bias or imprecision. Discuss both direction and magnitude of any potential bias',
     'Discussion, Strengths and limitations: Platform bias (eBay only), inability to control for equipment condition/age, small post-ban sample, global marketplace (EU/non-EU), limited pre-proposal baseline',
     '\u8003\u5bdf, \u9577\u6240\u3068\u9650\u754c: \u30d7\u30e9\u30c3\u30c8\u30d5\u30a9\u30fc\u30e0\u30d0\u30a4\u30a2\u30b9\uff08eBay\u306e\u307f\uff09\u3001\u6a5f\u5668\u72b6\u614b/\u5e74\u6570\u306e\u5236\u5fa1\u4e0d\u80fd\u3001\u7981\u6b62\u5f8c\u30b5\u30f3\u30d7\u30eb\u304c\u5c0f\u3001\u30b0\u30ed\u30fc\u30d0\u30eb\u5e02\u5834\u3001\u63d0\u6848\u524d\u30d9\u30fc\u30b9\u30e9\u30a4\u30f3\u306e\u6b20\u5982'),

    ('20', 'Interpretation',
     'Give a cautious overall interpretation of results considering objectives, limitations, multiplicity of analyses, results from similar studies, and other relevant evidence',
     'Discussion, Principal findings and Comparison with other studies: Contextualised with regulatory obsolescence literature, vehicle emission studies, BFMV benchmarks. Cautious language throughout.',
     '\u8003\u5bdf, \u4e3b\u8981\u306a\u77e5\u898b\u304a\u3088\u3073\u4ed6\u306e\u7814\u7a76\u3068\u306e\u6bd4\u8f03: \u898f\u5236\u9673\u8150\u5316\u6587\u732e\u3001\u8eca\u4e21\u6392\u51fa\u898f\u5236\u7814\u7a76\u3001BFMV\u30d9\u30f3\u30c1\u30de\u30fc\u30af\u3068\u6587\u8108\u5316\u3002\u614e\u91cd\u306a\u8868\u73fe\u3092\u4f7f\u7528'),

    ('21', 'Generalisability',
     'Discuss the generalisability (external validity) of the study results',
     'Discussion, Implications: "For healthcare facilities in jurisdictions considering similar regulations..."; acknowledged eBay as one segment of secondary market',
     '\u8003\u5bdf, \u793a\u5506: \u300c\u540c\u69d8\u306e\u898f\u5236\u3092\u691c\u8a0e\u4e2d\u306e\u7ba1\u8f44\u5730\u57df\u306e\u533b\u7642\u6a5f\u95a2\u306b\u3068\u3063\u3066\u2026\u300d\uff1beBay\u304c\u4e2d\u53e4\u5e02\u5834\u306e\u4e00\u90e8\u3067\u3042\u308b\u3053\u3068\u3092\u8a8d\u8b58'),

    # Other information
    ('22', 'Funding',
     'Give the source of funding and the role of the funders for the present study and, if applicable, for the original study on which the present article is based',
     'Transparency declaration: "No external funding was received for this study"',
     '\u900f\u660e\u6027\u5ba3\u8a00: \u300c\u672c\u7814\u7a76\u306b\u5bfe\u3059\u308b\u5916\u90e8\u8cc7\u91d1\u306f\u53d7\u3051\u3066\u3044\u306a\u3044\u300d'),
]


def write_english_checklist():
    doc = setup_doc()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('STROBE Statement\u2014Checklist of items that should be included in reports of cross-sectional studies')
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
           'Manuscript: Impact of the European Union desflurane regulation on secondary market '
           'prices of anaesthetic vaporisers: a cross-sectional time-series analysis of eBay sold listings')
    run.italic = True
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(12)

    # Create table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ['Item No', 'Recommendation', 'Reported on page/section', 'Comment']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(cell, '2E75B6')
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Section headers and items
    current_section = None
    for item_no, section, recommendation, location_en, location_ja in STROBE_ITEMS:
        # Add section header if new
        section_name = section.split('/')[0].strip() if '/' in section else section
        if section in ['Title and abstract'] and current_section != 'Title and abstract':
            row = table.add_row()
            cell = row.cells[0]
            # Merge all cells for section header
            cell_text = 'Title and abstract'
            cell.merge(row.cells[3])
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.bold = True
            run.font.size = Pt(9)
            set_cell_shading(cell, 'D6E4F0')
            current_section = 'Title and abstract'
        elif section == 'Background/rationale' and current_section != 'Introduction':
            row = table.add_row()
            cell = row.cells[0]
            cell.merge(row.cells[3])
            p = cell.paragraphs[0]
            run = p.add_run('Introduction')
            run.bold = True
            run.font.size = Pt(9)
            set_cell_shading(cell, 'D6E4F0')
            current_section = 'Introduction'
        elif section == 'Study design' and current_section != 'Methods':
            row = table.add_row()
            cell = row.cells[0]
            cell.merge(row.cells[3])
            p = cell.paragraphs[0]
            run = p.add_run('Methods')
            run.bold = True
            run.font.size = Pt(9)
            set_cell_shading(cell, 'D6E4F0')
            current_section = 'Methods'
        elif item_no == '13a' and current_section != 'Results':
            row = table.add_row()
            cell = row.cells[0]
            cell.merge(row.cells[3])
            p = cell.paragraphs[0]
            run = p.add_run('Results')
            run.bold = True
            run.font.size = Pt(9)
            set_cell_shading(cell, 'D6E4F0')
            current_section = 'Results'
        elif item_no == '18' and current_section != 'Discussion':
            row = table.add_row()
            cell = row.cells[0]
            cell.merge(row.cells[3])
            p = cell.paragraphs[0]
            run = p.add_run('Discussion')
            run.bold = True
            run.font.size = Pt(9)
            set_cell_shading(cell, 'D6E4F0')
            current_section = 'Discussion'
        elif item_no == '22' and current_section != 'Other':
            row = table.add_row()
            cell = row.cells[0]
            cell.merge(row.cells[3])
            p = cell.paragraphs[0]
            run = p.add_run('Other information')
            run.bold = True
            run.font.size = Pt(9)
            set_cell_shading(cell, 'D6E4F0')
            current_section = 'Other'

        # Add item row
        row = table.add_row()

        # Item number
        cell = row.cells[0]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item_no)
        run.font.size = Pt(9)
        run.bold = True

        # Recommendation
        cell = row.cells[1]
        p = cell.paragraphs[0]
        run = p.add_run(recommendation)
        run.font.size = Pt(9)

        # Location in manuscript
        cell = row.cells[2]
        p = cell.paragraphs[0]
        run = p.add_run(location_en)
        run.font.size = Pt(9)

        # Comment (checkmark)
        cell = row.cells[3]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Yes')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 128, 0)
        run.bold = True

    # Set column widths
    for row in table.rows:
        if len(row.cells) >= 4:
            set_cell_width(row.cells[0], 1.5)
            set_cell_width(row.cells[1], 8.0)
            set_cell_width(row.cells[2], 10.0)
            set_cell_width(row.cells[3], 1.5)

    # Footer note
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Note: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run(
        'An Explanation and Elaboration article discusses each checklist item and gives '
        'methodological background and published examples of transparent reporting. The STROBE '
        'checklist is best used in conjunction with this article (freely available at '
        'www.strobe-statement.org).')
    run.font.size = Pt(9)
    run.italic = True

    path = outdir + 'strobe_checklist_english.docx'
    doc.save(path)
    print(f"English STROBE checklist saved: {path}")


def write_japanese_checklist():
    doc = setup_doc()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('STROBE\u58f0\u660e\u2014\u6a2a\u65ad\u7814\u7a76\u306e\u5831\u544a\u306b\u542b\u3081\u308b\u3079\u304d\u9805\u76ee\u306e\u30c1\u30a7\u30c3\u30af\u30ea\u30b9\u30c8')
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        '\u8ad6\u6587: \u6b27\u5dde\u9023\u5408\u306e\u30c7\u30b9\u30d5\u30eb\u30e9\u30f3\u898f\u5236\u304c'
        '\u9ebb\u9154\u6c17\u5316\u5668\u306e\u4e2d\u53e4\u5e02\u5834\u4fa1\u683c\u306b\u4e0e\u3048\u308b'
        '\u5f71\u97ff\uff1aeBay\u58f2\u8cb7\u5b9f\u7e3e\u306e\u6a2a\u65ad\u7684\u6642\u7cfb\u5217\u5206\u6790')
    run.italic = True
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(12)

    # Create table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ['\u9805\u76ee\u756a\u53f7', '\u63a8\u5968\u4e8b\u9805', '\u8ad6\u6587\u4e2d\u306e\u8a72\u5f53\u7b87\u6240', '\u78ba\u8a8d']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(cell, '2E75B6')
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Section headers and items
    section_names_ja = {
        'Title and abstract': '\u30bf\u30a4\u30c8\u30eb\u3068\u62bd\u8c61',
        'Introduction': '\u7dd2\u8a00',
        'Methods': '\u65b9\u6cd5',
        'Results': '\u7d50\u679c',
        'Discussion': '\u8003\u5bdf',
        'Other': '\u305d\u306e\u4ed6\u306e\u60c5\u5831',
    }

    recommendation_ja = {
        '1a': '\u30bf\u30a4\u30c8\u30eb\u307e\u305f\u306f\u62bd\u8c61\u306b\u304a\u3044\u3066\u3001\u4e00\u822c\u7684\u306b\u7528\u3044\u3089\u308c\u308b\u7528\u8a9e\u3067\u7814\u7a76\u30c7\u30b6\u30a4\u30f3\u3092\u793a\u3059',
        '1b': '\u62bd\u8c61\u306b\u304a\u3044\u3066\u3001\u4f55\u3092\u884c\u3044\u4f55\u3092\u898b\u51fa\u3057\u305f\u304b\u306e\u60c5\u5831\u7684\u3067\u30d0\u30e9\u30f3\u30b9\u306e\u3068\u308c\u305f\u8981\u7d04\u3092\u793a\u3059',
        '2': '\u5831\u544a\u3055\u308c\u308b\u7814\u7a76\u306e\u79d1\u5b66\u7684\u80cc\u666f\u3068\u6839\u62e0\u3092\u8aac\u660e\u3059\u308b',
        '3': '\u4e8b\u524d\u306b\u8a2d\u5b9a\u3057\u305f\u4eee\u8aac\u3092\u542b\u3081\u3001\u5177\u4f53\u7684\u306a\u76ee\u7684\u3092\u8ff0\u3079\u308b',
        '4': '\u7814\u7a76\u30c7\u30b6\u30a4\u30f3\u306e\u4e3b\u8981\u7d20\u3092\u8ad6\u6587\u306e\u65e9\u3044\u6bb5\u968e\u3067\u793a\u3059',
        '5': '\u30c7\u30fc\u30bf\u53ce\u96c6\u671f\u9593\u3092\u542b\u3080\u3001\u8a2d\u5b9a\u3001\u5834\u6240\u3001\u95a2\u9023\u3059\u308b\u65e5\u4ed8\u3092\u8a18\u8ff0\u3059\u308b',
        '6': '\u6a2a\u65ad\u7814\u7a76\u2014\u9069\u683c\u57fa\u6e96\u3001\u304a\u3088\u3073\u53c2\u52a0\u8005\u306e\u9078\u629e\u306e\u30bd\u30fc\u30b9\u3068\u65b9\u6cd5\u3092\u793a\u3059',
        '7': '\u3059\u3079\u3066\u306e\u30a2\u30a6\u30c8\u30ab\u30e0\u3001\u66dd\u9732\u3001\u4ea4\u7d61\u56e0\u5b50\u3001\u52b9\u679c\u4fee\u98fe\u56e0\u5b50\u3092\u660e\u78ba\u306b\u5b9a\u7fa9\u3059\u308b',
        '8': '\u5404\u5909\u6570\u306b\u3064\u3044\u3066\u30c7\u30fc\u30bf\u30bd\u30fc\u30b9\u3068\u8a55\u4fa1\u65b9\u6cd5\u306e\u8a73\u7d30\u3092\u793a\u3059',
        '9': '\u30d0\u30a4\u30a2\u30b9\u306e\u6f5c\u5728\u7684\u306a\u30bd\u30fc\u30b9\u306b\u5bfe\u51e6\u3059\u308b\u53d6\u308a\u7d44\u307f\u3092\u8a18\u8ff0\u3059\u308b',
        '10': '\u7814\u7a76\u306e\u30b5\u30f3\u30d7\u30eb\u30b5\u30a4\u30ba\u306e\u6c7a\u5b9a\u65b9\u6cd5\u3092\u8aac\u660e\u3059\u308b',
        '11': '\u5b9a\u91cf\u5909\u6570\u306e\u5206\u6790\u3067\u306e\u53d6\u308a\u6271\u3044\u65b9\u3092\u8aac\u660e\u3059\u308b',
        '12a': '\u4ea4\u7d61\u56e0\u5b50\u306e\u5236\u5fa1\u306b\u7528\u3044\u305f\u65b9\u6cd5\u3092\u542b\u3080\u3001\u3059\u3079\u3066\u306e\u7d71\u8a08\u65b9\u6cd5\u3092\u8a18\u8ff0\u3059\u308b',
        '12b': '\u30b5\u30d6\u30b0\u30eb\u30fc\u30d7\u3068\u4ea4\u4e92\u4f5c\u7528\u306e\u691c\u8a0e\u65b9\u6cd5\u3092\u8a18\u8ff0\u3059\u308b',
        '12c': '\u6b20\u640d\u30c7\u30fc\u30bf\u306e\u5bfe\u51e6\u65b9\u6cd5\u3092\u8aac\u660e\u3059\u308b',
        '12d': '\u6a2a\u65ad\u7814\u7a76\u2014\u8a72\u5f53\u3059\u308b\u5834\u5408\u3001\u30b5\u30f3\u30d7\u30ea\u30f3\u30b0\u6226\u7565\u3092\u8003\u616e\u3057\u305f\u5206\u6790\u65b9\u6cd5\u3092\u8a18\u8ff0',
        '12e': '\u611f\u5ea6\u5206\u6790\u3092\u8a18\u8ff0\u3059\u308b',
        '13a': '\u7814\u7a76\u306e\u5404\u6bb5\u968e\u3067\u306e\u5bfe\u8c61\u8005\u6570\u3092\u5831\u544a\u3059\u308b',
        '13b': '\u5404\u6bb5\u968e\u3067\u306e\u975e\u53c2\u52a0\u306e\u7406\u7531\u3092\u793a\u3059',
        '14': '\u7814\u7a76\u53c2\u52a0\u8005\u306e\u7279\u5fb4\u3001\u66dd\u9732\u60c5\u5831\u3001\u4ea4\u7d61\u56e0\u5b50\u306b\u95a2\u3059\u308b\u60c5\u5831\u3092\u793a\u3059',
        '15': '\u6a2a\u65ad\u7814\u7a76\u2014\u30a2\u30a6\u30c8\u30ab\u30e0\u306e\u30a4\u30d9\u30f3\u30c8\u6570\u307e\u305f\u306f\u8981\u7d04\u7d71\u8a08\u91cf\u3092\u5831\u544a',
        '16a': '\u672a\u8abf\u6574\u306e\u63a8\u5b9a\u5024\u3001\u8a72\u5f53\u3059\u308b\u5834\u5408\u306f\u8abf\u6574\u6e08\u307f\u63a8\u5b9a\u5024\u3068\u305d\u306e\u7cbe\u5ea6\u3092\u793a\u3059',
        '16b': '\u9023\u7d9a\u5909\u6570\u3092\u30ab\u30c6\u30b4\u30ea\u5316\u3057\u305f\u5834\u5408\u3001\u305d\u306e\u5883\u754c\u5024\u3092\u5831\u544a\u3059\u308b',
        '17': '\u30b5\u30d6\u30b0\u30eb\u30fc\u30d7\u5206\u6790\u3001\u4ea4\u4e92\u4f5c\u7528\u3001\u611f\u5ea6\u5206\u6790\u306a\u3069\u3001\u305d\u306e\u4ed6\u306e\u5206\u6790\u3092\u5831\u544a',
        '18': '\u7814\u7a76\u76ee\u7684\u3092\u53c2\u7167\u3057\u3001\u4e3b\u8981\u306a\u7d50\u679c\u3092\u8981\u7d04\u3059\u308b',
        '19': '\u30d0\u30a4\u30a2\u30b9\u3084\u4e0d\u6b63\u78ba\u3055\u306e\u539f\u56e0\u3092\u8003\u616e\u3057\u3001\u7814\u7a76\u306e\u9650\u754c\u3092\u8b70\u8ad6\u3059\u308b',
        '20': '\u76ee\u7684\u3001\u9650\u754c\u3001\u5206\u6790\u306e\u591a\u91cd\u6027\u3001\u985e\u4f3c\u7814\u7a76\u306e\u7d50\u679c\u3092\u8003\u616e\u3057\u3001\u614e\u91cd\u306a\u5168\u4f53\u7684\u89e3\u91c8\u3092\u793a\u3059',
        '21': '\u7814\u7a76\u7d50\u679c\u306e\u4e00\u822c\u5316\u53ef\u80fd\u6027\uff08\u5916\u7684\u59a5\u5f53\u6027\uff09\u3092\u8b70\u8ad6\u3059\u308b',
        '22': '\u73fe\u5728\u306e\u7814\u7a76\u306e\u8cc7\u91d1\u6e90\u3068\u8cc7\u91d1\u63d0\u4f9b\u8005\u306e\u5f79\u5272\u3092\u793a\u3059',
    }

    current_section = None
    for item_no, section, recommendation, location_en, location_ja in STROBE_ITEMS:
        # Add section header if new
        if section in ['Title and abstract'] and current_section != 'Title and abstract':
            row = table.add_row()
            cell = row.cells[0]
            cell.merge(row.cells[3])
            p = cell.paragraphs[0]
            run = p.add_run(section_names_ja['Title and abstract'])
            run.bold = True
            run.font.size = Pt(9)
            set_cell_shading(cell, 'D6E4F0')
            current_section = 'Title and abstract'
        elif section == 'Background/rationale' and current_section != 'Introduction':
            row = table.add_row()
            cell = row.cells[0]
            cell.merge(row.cells[3])
            p = cell.paragraphs[0]
            run = p.add_run(section_names_ja['Introduction'])
            run.bold = True
            run.font.size = Pt(9)
            set_cell_shading(cell, 'D6E4F0')
            current_section = 'Introduction'
        elif section == 'Study design' and current_section != 'Methods':
            row = table.add_row()
            cell = row.cells[0]
            cell.merge(row.cells[3])
            p = cell.paragraphs[0]
            run = p.add_run(section_names_ja['Methods'])
            run.bold = True
            run.font.size = Pt(9)
            set_cell_shading(cell, 'D6E4F0')
            current_section = 'Methods'
        elif item_no == '13a' and current_section != 'Results':
            row = table.add_row()
            cell = row.cells[0]
            cell.merge(row.cells[3])
            p = cell.paragraphs[0]
            run = p.add_run(section_names_ja['Results'])
            run.bold = True
            run.font.size = Pt(9)
            set_cell_shading(cell, 'D6E4F0')
            current_section = 'Results'
        elif item_no == '18' and current_section != 'Discussion':
            row = table.add_row()
            cell = row.cells[0]
            cell.merge(row.cells[3])
            p = cell.paragraphs[0]
            run = p.add_run(section_names_ja['Discussion'])
            run.bold = True
            run.font.size = Pt(9)
            set_cell_shading(cell, 'D6E4F0')
            current_section = 'Discussion'
        elif item_no == '22' and current_section != 'Other':
            row = table.add_row()
            cell = row.cells[0]
            cell.merge(row.cells[3])
            p = cell.paragraphs[0]
            run = p.add_run(section_names_ja['Other'])
            run.bold = True
            run.font.size = Pt(9)
            set_cell_shading(cell, 'D6E4F0')
            current_section = 'Other'

        # Add item row
        row = table.add_row()

        # Item number
        cell = row.cells[0]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(item_no)
        run.font.size = Pt(9)
        run.bold = True

        # Recommendation (Japanese)
        cell = row.cells[1]
        p = cell.paragraphs[0]
        run = p.add_run(recommendation_ja.get(item_no, recommendation))
        run.font.size = Pt(9)

        # Location in manuscript (Japanese)
        cell = row.cells[2]
        p = cell.paragraphs[0]
        run = p.add_run(location_ja)
        run.font.size = Pt(9)

        # Comment (checkmark)
        cell = row.cells[3]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('\u6e08')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 128, 0)
        run.bold = True

    # Set column widths
    for row in table.rows:
        if len(row.cells) >= 4:
            set_cell_width(row.cells[0], 1.5)
            set_cell_width(row.cells[1], 8.0)
            set_cell_width(row.cells[2], 10.0)
            set_cell_width(row.cells[3], 1.5)

    # Footer note
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('\u6ce8: ')
    run.bold = True
    run.font.size = Pt(9)
    run = p.add_run(
        'STROBE\u30c1\u30a7\u30c3\u30af\u30ea\u30b9\u30c8\u306e\u8a73\u7d30\u306a\u8aac\u660e\u3068'
        '\u80cc\u666f\u306f\u3001Explanation and Elaboration\u8ad6\u6587\u3092\u53c2\u7167\u3057\u3066'
        '\u304f\u3060\u3055\u3044\uff08www.strobe-statement.org\u3067\u7121\u6599\u516c\u958b\uff09\u3002')
    run.font.size = Pt(9)
    run.italic = True

    path = outdir + 'strobe_checklist_japanese.docx'
    doc.save(path)
    print(f"Japanese STROBE checklist saved: {path}")


if __name__ == '__main__':
    write_english_checklist()
    write_japanese_checklist()
    print("\nBoth STROBE checklists generated successfully!")
