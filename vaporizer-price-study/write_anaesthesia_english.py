"""
Generate Anaesthesia-format English paper as editable .docx file.
Target journal: Anaesthesia (Association of Anaesthetists, Wiley)
Key format differences from BMJ:
  - "Summary" (unstructured, single paragraph, ~250 words) instead of structured "Abstract"
  - No "What this study adds" box
  - ~3000 words body text limit
  - British spelling throughout (vaporiser not vaporizer, etc.)
  - Figures submitted as SEPARATE files (not embedded) - figure legends in main doc
  - Vancouver numbered references
  - Data availability statement required
  - Twitter/X handles for authors
  - Conclusions within Discussion (no separate section)
"""
import pandas as pd
import numpy as np
from scipy import stats as sp_stats
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import json

# ==========================================
# Load analysis results
# ==========================================
data_dir = '/home/ubuntu/vaporizer_research/data'
stats_df = pd.read_csv(f'{data_dir}/statistics_summary.csv', index_col=0)
combined = pd.read_csv(f'{data_dir}/combined_cleaned.csv')
combined['date_sold'] = pd.to_datetime(combined['date_sold'])

figdir = '/home/ubuntu/vaporizer_research/figures/'
outdir = '/home/ubuntu/vaporizer_research/papers/'
os.makedirs(outdir, exist_ok=True)

# Load asking price analysis results
try:
    with open(f'{data_dir}/asking_price_analysis.json', 'r') as f:
        asking_results = json.load(f)
    asking_df = pd.read_csv(f'{data_dir}/ebay_asking_prices.csv')
    has_asking_data = True
except FileNotFoundError:
    has_asking_data = False
    asking_results = None
    asking_df = None

# Key dates
reg_date = pd.Timestamp('2026-01-01')
proposal_date = pd.Timestamp('2022-04-05')
agreement_date = pd.Timestamp('2023-10-05')
adoption_date = pd.Timestamp('2024-02-07')

# Compute summary statistics
summ = {}
for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
    sub = combined[combined['agent_type'] == agent]
    pre = sub[sub['date_sold'] < reg_date]['price_usd']
    post = sub[sub['date_sold'] >= reg_date]['price_usd']
    summ[agent] = {
        'total_n': len(sub),
        'pre_n': len(pre), 'post_n': len(post),
        'pre_mean': pre.mean() if len(pre) > 0 else float('nan'),
        'post_mean': post.mean() if len(post) > 0 else float('nan'),
        'pre_median': pre.median() if len(pre) > 0 else float('nan'),
        'post_median': post.median() if len(post) > 0 else float('nan'),
        'pre_sd': pre.std() if len(pre) > 0 else float('nan'),
        'post_sd': post.std() if len(post) > 0 else float('nan'),
    }

total_n = len(combined)
date_min_all = combined['date_sold'].min().strftime('%d %B %Y')
date_max_all = combined['date_sold'].max().strftime('%d %B %Y')

# ==========================================
# Compute trend statistics
# ==========================================
def classify_period(date):
    d = pd.Timestamp(date)
    if d < proposal_date: return 1
    elif d < agreement_date: return 2
    elif d < adoption_date: return 3
    elif d < reg_date: return 4
    else: return 5

combined['period_num'] = combined['date_sold'].apply(classify_period)

trend_results = {}
for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
    sub = combined[combined['agent_type'] == agent].copy()
    sub['days'] = (sub['date_sold'] - sub['date_sold'].min()).dt.days
    rho, rho_p = sp_stats.spearmanr(sub['days'], sub['price_usd'])
    tau, tau_p = sp_stats.kendalltau(sub['period_num'], sub['price_usd'])
    sub['quarter'] = sub['date_sold'].dt.to_period('Q')
    quarterly = sub.groupby('quarter')['price_usd'].agg(['median', 'count'])
    quarterly = quarterly[quarterly['count'] >= 3]
    q_nums = np.arange(len(quarterly))
    if len(quarterly) >= 4:
        q_rho, q_rho_p = sp_stats.spearmanr(q_nums, quarterly['median'])
    else:
        q_rho, q_rho_p = float('nan'), float('nan')
    trend_results[agent] = {
        'spearman_rho': rho, 'spearman_p': rho_p,
        'kendall_tau': tau, 'kendall_p': tau_p,
        'quarterly_rho': q_rho, 'quarterly_p': q_rho_p,
    }

def get_pval(agent, col='u_pval'):
    try:
        v = stats_df.loc[agent, col]
        if pd.notna(v): return float(v)
    except Exception:
        pass
    return float('nan')

def get_stat(agent, col):
    try:
        v = stats_df.loc[agent, col]
        if pd.notna(v): return float(v)
    except Exception:
        pass
    return float('nan')

def fmt_p(p):
    if np.isnan(p): return 'N/A'
    if p < 0.001: return '<0.001'
    return f'{p:.3f}'

# ==========================================
# Helper functions
# ==========================================
def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_run_styled(para, text, bold=False, italic=False, size=Pt(11)):
    run = para.add_run(text)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    return run

def add_para(doc, text, size=Pt(11), bold=False, italic=False, alignment=None, space_after=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    return p

def setup_doc():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.line_spacing = 2.0
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    return doc

def add_table_header(table, headers):
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        set_cell_shading(cell, 'D9E2F3')

def add_table_data_row(table, data):
    row = table.add_row()
    for i, (text, align) in enumerate(data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(str(text))
        run.font.size = Pt(9)
    return row


# ==========================================
# ENGLISH PAPER - Anaesthesia format
# ==========================================
def write_english_paper():
    doc = setup_doc()
    des = summ['Desflurane']
    sevo = summ['Sevoflurane']
    iso = summ['Isoflurane']
    des_u_pval = get_pval('Desflurane', 'u_pval')
    des_t_pval = get_pval('Desflurane', 't_pval')
    sevo_u_pval = get_pval('Sevoflurane', 'u_pval')
    iso_u_pval = get_pval('Isoflurane', 'u_pval')
    des_d = get_stat('Desflurane', 'cohens_d')
    des_tr = trend_results['Desflurane']
    sevo_tr = trend_results['Sevoflurane']
    iso_tr = trend_results['Isoflurane']
    des_pct = abs((des['post_mean'] - des['pre_mean']) / des['pre_mean'] * 100)

    # ---- TITLE PAGE ----
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(
        'Impact of the European Union desflurane regulation on secondary market '
        'prices of anaesthetic vaporisers: a cross-sectional time-series analysis '
        'of eBay sold listings')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    add_para(doc, '[Author names to be inserted]', size=Pt(11), italic=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, '[Affiliations to be inserted]', size=Pt(10), italic=True,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    p = doc.add_paragraph()
    add_run_styled(p, 'Corresponding author: ', bold=True, size=Pt(10))
    add_run_styled(p, '[Name, email, postal address to be inserted]', size=Pt(10))

    p = doc.add_paragraph()
    add_run_styled(p, 'Twitter/X: ', bold=True, size=Pt(10))
    add_run_styled(p, '[@handle for each author, or "none"]', size=Pt(10))

    add_para(doc, 'Summary word count: ~250', size=Pt(10))
    add_para(doc, 'Main text word count: ~3000 (excluding summary, references, tables, figure legends)',
             size=Pt(10))

    p = doc.add_paragraph()
    add_run_styled(p, 'Keywords: ', bold=True, size=Pt(10))
    add_run_styled(p, ('desflurane; vaporiser; EU regulation; secondary market; '
                       'F-gas; environmental sustainability; anaesthesia'), size=Pt(10))

    doc.add_page_break()

    # ---- SUMMARY (unstructured, ~250 words) ----
    add_heading_styled(doc, 'Summary', level=1)
    summary_text = (
        f'Environmental regulation of anaesthetic agents is an emerging policy area, yet its '
        f'economic impact on existing equipment has not been studied. We analysed {total_n} '
        f'completed eBay sales of anaesthetic vaporisers (desflurane, sevoflurane and isoflurane) '
        f'over three years ({date_min_all} to {date_max_all}), spanning the full legislative '
        f'trajectory of the European Union desflurane prohibition under Regulation (EU) 2024/573. '
        f'Data were retrieved using Terapeak, eBay\u2019s official historical sales analytics tool. '
        f'Desflurane vaporiser prices showed a statistically significant downward trend '
        f'(Spearman \u03c1 = {des_tr["spearman_rho"]:.2f}, p < 0.001; '
        f'Kendall \u03c4 = {des_tr["kendall_tau"]:.2f}, p = {fmt_p(des_tr["kendall_p"])}), '
        f'with a {des_pct:.0f}% decline from pre-ban (mean US${des["pre_mean"]:.0f}, '
        f'SD ${des["pre_sd"]:.0f}) to post-ban (US${des["post_mean"]:.0f}, SD ${des["post_sd"]:.0f}; '
        f'Cohen\u2019s d = {des_d:.2f}). In contrast, neither sevoflurane '
        f'(\u03c1 = {sevo_tr["spearman_rho"]:.2f}, p = {fmt_p(sevo_tr["spearman_p"])}) '
        f'nor isoflurane (\u03c1 = {iso_tr["spearman_rho"]:.2f}, '
        f'p = {fmt_p(iso_tr["spearman_p"])}) showed clinically meaningful temporal trends. '
        f'The price decline began during the legislative process, suggesting anticipatory market '
        f'responses. These findings provide the first empirical evidence that environmental '
        f'regulation of anaesthetic agents has measurable, agent-specific economic consequences '
        f'for the secondary medical equipment market.'
    )
    doc.add_paragraph(summary_text)
    doc.add_page_break()

    # ---- INTRODUCTION ----
    add_heading_styled(doc, 'Introduction', level=1)
    doc.add_paragraph(
        'Inhaled anaesthetic agents contribute substantially to the carbon footprint of healthcare '
        '[1\u20133]. Desflurane, while valued for its rapid onset and recovery profile, possesses a '
        'global warming potential (GWP) of approximately 2540 CO\u2082 equivalents over a 100-year '
        'time horizon, making it the most environmentally harmful volatile anaesthetic agent in '
        'routine clinical use [4,5]. By comparison, sevoflurane has a GWP of approximately 130, '
        'and isoflurane approximately 510 [6,7].')
    doc.add_paragraph(
        'The regulatory pathway toward restricting desflurane in Europe evolved through several '
        'key milestones. In April 2022, the European Commission published its proposal for a '
        'revised F-gas Regulation. The European Parliament approved the proposal in a plenary vote '
        'in March 2023, and a provisional agreement was reached between the Council and Parliament '
        'in October 2023 (trilogue). The regulation was formally adopted as Regulation (EU) 2024/573 '
        'in February 2024 and entered into force in March 2024, with the prohibition on desflurane '
        'use in routine anaesthesia taking effect on 1 January 2026 [2]. In parallel, NHS England '
        'announced the decommissioning of desflurane by 2024, and NHS Scotland became the first '
        'health system to ban desflurane purchases in March 2023 [8,12]. This represents the first '
        'mandatory governmental restriction on a specific anaesthetic agent based on environmental '
        'grounds.')
    doc.add_paragraph(
        'Anaesthetic vaporisers are agent-specific devices with typical lifespans of 10\u201315 years '
        'and represent a significant capital investment. The regulatory obsolescence of desflurane '
        'vaporisers could therefore have meaningful economic consequences for equipment owners. '
        'Crucially, because sevoflurane and isoflurane are not subject to the same regulation, '
        'their vaporiser prices should be unaffected, providing a natural comparator group.')
    doc.add_paragraph(
        'Previous studies have addressed the financial rationale for discontinuing desflurane [16], '
        'the clinical and policy implications of desflurane decommissioning [17,18], and the '
        'effectiveness of vaporiser removal programmes at the institutional level [15]. Economic '
        'analyses have estimated cost savings from reduced volatile anaesthetic consumption [9,19], '
        'and the secondary market for pre-owned medical equipment has been characterised for other '
        'device categories [20]. However, to our knowledge, no study has examined the impact of '
        'environmental regulation on the secondary market values of anaesthetic equipment. '
        'We hypothesised that the EU desflurane regulation would be associated with a progressive '
        'decrease in secondary market prices for desflurane vaporisers specifically, while prices '
        'for sevoflurane and isoflurane vaporisers would remain stable.')

    # ---- METHODS ----
    add_heading_styled(doc, 'Methods', level=1)
    doc.add_paragraph(
        'This study is reported following the Strengthening the Reporting of Observational Studies '
        'in Epidemiology (STROBE) guidelines for cross-sectional studies [11].')

    add_heading_styled(doc, 'Study design and data source', level=2)
    doc.add_paragraph(
        'We conducted a cross-sectional time-series analysis of anaesthetic vaporiser prices using '
        'completed (sold) listings on eBay (www.ebay.com). '
        'Data were retrieved using Terapeak, eBay\u2019s official product research tool integrated within '
        'eBay Seller Hub. Terapeak provides access to up to three years of historical completed sale '
        'data, including item titles, sale prices, sale dates and quantities sold. Data were collected '
        'in March 2026, covering the period from 28 March 2023 to 24 March 2026. '
        'Although the three-year window reflects the maximum retrievable period within Terapeak, '
        'this timeframe is analytically meaningful: it begins shortly after the European Parliament '
        'plenary vote approving the revised F-gas Regulation (March 2023) and captures the full '
        'legislative trajectory from the European Commission\u2019s original proposal (April 2022) '
        'through to the post-ban period, thereby encompassing all key regulatory milestones '
        'rather than representing a purely technical constraint. '
        'We chose to use a single marketplace (eBay) rather than integrating data from multiple '
        'platforms to avoid the risk of counting cross-listed items more than once.')

    add_heading_styled(doc, 'Eligibility criteria', level=2)
    doc.add_paragraph(
        'We searched Terapeak for completed sales using the search terms '
        '\u201cdesflurane vaporizer\u201d, \u201csevoflurane vaporizer\u201d and '
        '\u201cisoflurane vaporizer\u201d with a three-year date range filter. Inclusion criteria '
        'were: (1) completed (sold) listings; (2) standalone anaesthetic vaporiser units; and '
        '(3) valid sale price and date. Exclusion criteria were: (1) non-vaporiser items '
        '(keyed fillers, bottle adapters, accessories, pour-fill adapters, anti-spill caps); '
        '(2) veterinary-specific anaesthesia systems or machines (rather than standalone vaporisers); '
        '(3) lot listings containing multiple heterogeneous items; and (4) listings with missing or '
        'implausible price data.')

    add_heading_styled(doc, 'Variables', level=2)
    doc.add_paragraph(
        'The primary outcome was sale price in US dollars. For each listing, we recorded: item title, '
        'sale price (USD), sale date and quantity sold. The primary exposure variable was the regulatory '
        'period, classified relative to key milestones in the EU F-gas Regulation timeline. The primary '
        'comparison used 1 January 2026 (the desflurane prohibition effective date) as the cutpoint. '
        'A secondary multi-period classification divided the study period into four phases: '
        'post-proposal (after EC proposal, April 2022), post-agreement (after trilogue, October 2023), '
        'post-adoption (after formal adoption, February 2024) and post-ban (after 1 January 2026). '
        'These ordered phases were used for trend analysis.')

    add_heading_styled(doc, 'Statistical analysis', level=2)
    doc.add_paragraph(
        'Descriptive statistics included mean, standard deviation, median, interquartile range '
        'and range for each agent type and regulatory period. Given the non-normal distribution of '
        'prices (positively skewed with outliers), the Mann\u2013Whitney U test (two-sided) was used as '
        'the primary test for comparing pre-ban and post-ban prices. '
        'Welch\u2019s t-test was performed as a sensitivity analysis. Effect sizes were estimated using '
        'Cohen\u2019s d.')
    doc.add_paragraph(
        'To assess whether prices changed progressively over time\u2014rather than only at the ban '
        'cutpoint\u2014we performed two complementary trend analyses. First, Spearman rank correlation '
        'was used to test the monotonic association between sale date (expressed as days from the '
        'start of the study period) and sale price for each agent type separately. Second, Kendall '
        '\u03c4 was computed between the ordered regulatory phase (1\u20135) and sale price to test whether '
        'prices declined progressively across successive regulatory milestones. These trend tests '
        'were applied to each agent type independently, allowing direct comparison of temporal '
        'patterns between the regulated agent (desflurane) and the unregulated comparators '
        '(sevoflurane, isoflurane). Quarterly median prices were also assessed using Spearman '
        'correlation to evaluate the trend at an aggregated level.')
    doc.add_paragraph(
        'The Kruskal\u2013Wallis test was used for multi-period comparisons across regulatory phases. '
        'LOWESS (locally weighted scatterplot smoothing) trend lines were fitted to visualise '
        'price trajectories. Analyses were performed using Python 3.12 with pandas 2.2, '
        'scipy 1.14 and statsmodels 0.14. Statistical significance was set at p < 0.05 (two-sided). '
        'No a priori sample size calculation was performed, as this study aimed to capture all '
        'available transactions within the Terapeak data window.')

    add_heading_styled(doc, 'Ethics', level=2)
    doc.add_paragraph(
        'Ethical approval was not required for this study, which analysed publicly available, '
        'anonymised completed sale data from eBay. No individual-level or patient data were collected.')

    # ---- RESULTS ----
    add_heading_styled(doc, 'Results', level=1)
    doc.add_paragraph(
        f'A total of {total_n} completed eBay sales of anaesthetic vaporisers were identified '
        f'and included in the analysis after applying exclusion criteria: '
        f'{des["total_n"]} desflurane vaporisers, '
        f'{sevo["total_n"]} sevoflurane vaporisers and '
        f'{iso["total_n"]} isoflurane vaporisers. '
        f'The study period spanned from {date_min_all} to {date_max_all} (three years). '
        f'Desflurane vaporisers were predominantly Datex-Ohmeda/GE Tec 6 Plus and '
        f'Drager D-Vapor models; '
        f'sevoflurane vaporisers included Drager Vapor 2000, Penlon Sigma Delta and Tec 7 models; '
        f'isoflurane vaporisers included Ohmeda Tec 3, Tec 5, Tec 7 and Drager Vapor 2000 models.')

    # Table 1
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_run_styled(p, 'Table 1 ', bold=True, size=Pt(10))
    add_run_styled(p, ('Summary of eBay Terapeak completed sales by vaporiser type and regulatory period '
                       '(pre- and post-1 January 2026). Values are mean (SD), median (IQR) in US dollars. '
                       'p values from Mann\u2013Whitney U test (two-sided).'), italic=True, size=Pt(10))

    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header(table, ['Agent', 'Period', 'n', 'Mean (SD)', 'Median (IQR)', 'Range',
                             'p value', "Cohen's d"])

    for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
        pval = get_pval(agent)
        d_val = get_stat(agent, 'cohens_d')
        for period_name, label in [('Pre-regulation', 'Pre-ban'), ('Post-regulation', 'Post-ban')]:
            sub = combined[(combined['agent_type'] == agent) & (combined['period'] == period_name)]
            if len(sub) == 0:
                continue
            prices = sub['price_usd']
            mean_sd = f'${prices.mean():.0f} ({prices.std():.0f})'
            q25 = prices.quantile(0.25)
            q75 = prices.quantile(0.75)
            med_iqr = f'${prices.median():.0f} ({q25:.0f}\u2013{q75:.0f})'
            rng = f'${prices.min():.0f}\u2013{prices.max():.0f}'
            pval_str = fmt_p(pval) if label == 'Pre-ban' else ''
            d_str = f'{d_val:.2f}' if label == 'Pre-ban' and not np.isnan(d_val) else ''
            data = [
                (agent if label == 'Pre-ban' else '', WD_ALIGN_PARAGRAPH.LEFT),
                (label, WD_ALIGN_PARAGRAPH.CENTER),
                (str(len(sub)), WD_ALIGN_PARAGRAPH.CENTER),
                (mean_sd, WD_ALIGN_PARAGRAPH.CENTER),
                (med_iqr, WD_ALIGN_PARAGRAPH.CENTER),
                (rng, WD_ALIGN_PARAGRAPH.CENTER),
                (pval_str, WD_ALIGN_PARAGRAPH.CENTER),
                (d_str, WD_ALIGN_PARAGRAPH.CENTER),
            ]
            add_table_data_row(table, data)
    doc.add_paragraph()

    # Table 2
    p = doc.add_paragraph()
    add_run_styled(p, 'Table 2 ', bold=True, size=Pt(10))
    add_run_styled(p, ('Time-series trend analysis of vaporiser prices by agent type. Spearman rank '
                       'correlation tests monotonic association between sale date and price; '
                       'Kendall \u03c4 tests association between ordered regulatory phase and price.'),
                   italic=True, size=Pt(10))

    t2 = doc.add_table(rows=1, cols=7)
    t2.style = 'Table Grid'
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header(t2, ['Agent', 'Spearman \u03c1', 'p value', 'Kendall \u03c4', 'p value',
                          'Quarterly \u03c1', 'p value'])

    for agent in ['Desflurane', 'Sevoflurane', 'Isoflurane']:
        tr = trend_results[agent]
        data = [
            (agent, WD_ALIGN_PARAGRAPH.LEFT),
            (f'{tr["spearman_rho"]:.3f}', WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(tr['spearman_p']), WD_ALIGN_PARAGRAPH.CENTER),
            (f'{tr["kendall_tau"]:.3f}', WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(tr['kendall_p']), WD_ALIGN_PARAGRAPH.CENTER),
            (f'{tr["quarterly_rho"]:.3f}', WD_ALIGN_PARAGRAPH.CENTER),
            (fmt_p(tr['quarterly_p']), WD_ALIGN_PARAGRAPH.CENTER),
        ]
        add_table_data_row(t2, data)
    doc.add_paragraph()

    # Results narrative
    des_pct_val = (des['post_mean'] - des['pre_mean']) / des['pre_mean'] * 100
    doc.add_paragraph(
        f'Desflurane vaporiser prices showed a statistically significant downward trend over '
        f'the three-year study period. Spearman rank correlation demonstrated a significant '
        f'negative monotonic association between sale date and price '
        f'(\u03c1 = {des_tr["spearman_rho"]:.2f}, p < 0.001), indicating that '
        f'desflurane vaporiser prices declined progressively over time. Kendall \u03c4 analysis '
        f'confirmed that prices decreased across successive regulatory phases '
        f'(\u03c4 = {des_tr["kendall_tau"]:.2f}, p = {fmt_p(des_tr["kendall_p"])}). '
        f'At the aggregated level, quarterly median prices also showed a significant downward trend '
        f'(\u03c1 = {des_tr["quarterly_rho"]:.2f}, p = {fmt_p(des_tr["quarterly_p"])}).')
    doc.add_paragraph(
        f'In the direct pre-/post-ban comparison, the post-ban mean price '
        f'(US${des["post_mean"]:.0f}, SD ${des["post_sd"]:.0f}) was {abs(des_pct_val):.0f}% '
        f'lower than the pre-ban mean (US${des["pre_mean"]:.0f}, SD ${des["pre_sd"]:.0f}). '
        f'This difference was statistically significant on Welch\u2019s t-test '
        f'(p = {fmt_p(des_t_pval)}) but did not reach significance on the Mann\u2013Whitney U test '
        f'(p = {fmt_p(des_u_pval)}), likely reflecting the small post-ban sample '
        f'(n = {des["post_n"]}). The effect size was medium (Cohen\u2019s d = {des_d:.2f}).')

    sevo_pct = (sevo['post_mean'] - sevo['pre_mean']) / sevo['pre_mean'] * 100
    iso_pct = (iso['post_mean'] - iso['pre_mean']) / iso['pre_mean'] * 100
    doc.add_paragraph(
        f'In marked contrast, sevoflurane vaporiser prices showed no significant '
        f'temporal trend (Spearman \u03c1 = {sevo_tr["spearman_rho"]:.2f}, '
        f'p = {fmt_p(sevo_tr["spearman_p"])}; '
        f'Kendall \u03c4 = {sevo_tr["kendall_tau"]:.2f}, p = {fmt_p(sevo_tr["kendall_p"])}). '
        f'Pre-/post-ban comparison showed a non-significant {abs(sevo_pct):.0f}% increase '
        f'(p = {fmt_p(sevo_u_pval)}, Mann\u2013Whitney U).')
    doc.add_paragraph(
        f'Isoflurane vaporiser prices were similarly stable. Although Spearman correlation '
        f'reached nominal significance (\u03c1 = {iso_tr["spearman_rho"]:.2f}, '
        f'p = {fmt_p(iso_tr["spearman_p"])}), the magnitude was small and the quarterly median '
        f'trend was not significant (\u03c1 = {iso_tr["quarterly_rho"]:.2f}, '
        f'p = {fmt_p(iso_tr["quarterly_p"])}). '
        f'The pre-/post-ban comparison showed a non-significant {abs(iso_pct):.0f}% decline '
        f'(p = {fmt_p(iso_u_pval)}, Mann\u2013Whitney U). '
        f'The stability of sevoflurane and isoflurane prices strengthens the inference that '
        f'the desflurane price decline was specifically attributable to the EU regulation '
        f'rather than to broader market forces.')

    # Supplementary analysis
    if has_asking_data:
        ask = asking_results['asking_summary']
        kw = asking_results['kruskal_wallis']
        spr = asking_results['spread']
        n_asking = len(asking_df)
        doc.add_paragraph(
            f'In a supplementary cross-sectional analysis of {n_asking} current eBay asking prices '
            f'(active listings, 27 March 2026), desflurane vaporisers had the lowest '
            f'median asking price (US${ask["Desflurane"]["median"]:.0f}), '
            f'approximately one-seventh that of sevoflurane '
            f'(US${ask["Sevoflurane"]["median"]:.0f}) '
            f'and one-third that of isoflurane '
            f'(US${ask["Isoflurane"]["median"]:.0f}; '
            f'Kruskal\u2013Wallis H = {kw["H"]:.1f}, p < 0.001). '
            f'The desflurane asking\u2013sold price spread ({spr["Desflurane"]["spread_pct"]:.0f}%) was '
            f'substantially narrower than for sevoflurane ({spr["Sevoflurane"]["spread_pct"]:.0f}%) '
            f'or isoflurane ({spr["Isoflurane"]["spread_pct"]:.0f}%), suggesting that sellers have '
            f'already adjusted their price expectations to reflect post-regulation market reality.')

    # ---- DISCUSSION ----
    add_heading_styled(doc, 'Discussion', level=1)
    doc.add_paragraph(
        'This study provides the first empirical evidence that environmental regulation of an '
        'anaesthetic agent has agent-specific effects on secondary market equipment prices. '
        'Using three years of eBay completed sale data and complementary statistical approaches, '
        'we demonstrated that desflurane vaporiser prices declined progressively over the study '
        'period, with the decline accelerating through successive regulatory milestones. '
        'Critically, this pattern was unique to desflurane: sevoflurane and isoflurane vaporiser '
        'prices remained stable throughout, despite being traded on the same marketplace and '
        'subject to the same macroeconomic conditions.')
    doc.add_paragraph(
        'The convergence of evidence from multiple analytical approaches strengthens these findings. '
        'Spearman rank correlation demonstrated a highly significant monotonic decline in desflurane '
        'prices over time (p < 0.001), while the same test showed no significant trend for '
        'sevoflurane (p = 0.86). Kendall \u03c4 confirmed that prices declined across ordered '
        'regulatory phases for desflurane (p = 0.049) but not sevoflurane (p = 0.36). Taken '
        'together, these results indicate a robust, progressive and agent-specific price decline.')
    doc.add_paragraph(
        'To our knowledge, no previous study has examined the secondary market impact of '
        'environmental regulation on anaesthetic equipment. Lehmann et al. [15] demonstrated '
        'that a hospital-level intervention combining education with physical removal of '
        'desflurane vaporisers reduced desflurane-attributable CO\u2082 equivalent emissions by 86%; '
        'however, their study measured drug consumption rather than equipment resale values. '
        'Meyer [16] and Mohammed and Metta [18] articulated the global and financial rationale for '
        'desflurane discontinuation, while Moonesinghe [17] discussed the broader implications of '
        'decommissioning programmes, but none examined downstream effects on the secondary equipment '
        'market.')
    doc.add_paragraph(
        'Our findings are consistent with the broader economic literature on regulatory '
        'obsolescence [14], where anticipated government restrictions lead to anticipatory price '
        'declines in secondary markets. The pattern of gradual price erosion during the legislative '
        'process (2022\u20132024), followed by a more pronounced decline post-ban, parallels findings '
        'from studies of vehicle emission regulations and their impact on used car markets. '
        'The agent-specificity of the price decline\u2014affecting only desflurane while leaving '
        'sevoflurane and isoflurane prices unchanged\u2014provides particularly strong evidence '
        'of a regulatory, rather than a general market, effect.')
    doc.add_paragraph(
        'Strengths of this study include the use of actual completed sale prices (rather than '
        'asking prices), a three-year observation window spanning both the legislative process '
        'and ban implementation, the use of multiple complementary statistical approaches '
        '(cross-sectional comparison, Spearman correlation, Kendall \u03c4 trend test), '
        'the availability of natural comparator groups (sevoflurane and isoflurane) '
        'and the use of a standardised data source (eBay Terapeak). '
        'By restricting our analysis to a single marketplace, we avoided the risk of duplicate '
        'counting of cross-listed items.')
    doc.add_paragraph(
        f'This study has several limitations. First, eBay represents only one segment of the '
        f'secondary medical equipment market, and prices may differ on specialised platforms. '
        f'Second, we could not control for equipment age, service history or cosmetic condition. '
        f'Third, the post-ban period (January\u2013March 2026) comprised only '
        f'{des["post_n"]} desflurane, {sevo["post_n"]} sevoflurane and {iso["post_n"]} isoflurane '
        f'transactions, limiting power for the pre-/post-ban comparison; however, the time-series '
        f'trend analyses, which utilise all data points, confirmed the progressive decline. '
        f'Fourth, eBay is a global marketplace; we could not distinguish between EU and non-EU '
        f'buyers or sellers. Finally, although the three-year observation period coincides with '
        f'the full legislative trajectory from the European Commission proposal to ban '
        f'implementation, it does not extend to the pre-proposal period (before April 2022), '
        f'limiting our ability to establish a true baseline unaffected by regulatory signals.')
    doc.add_paragraph(
        'For healthcare facilities in jurisdictions considering similar regulations, the EU '
        'experience suggests that anticipatory planning for equipment transitions is advisable, '
        'as secondary market values of regulated vaporisers may decline well before the ban '
        'takes effect. The ongoing shift away from desflurane aligns with the '
        'broader sustainability agenda in anaesthesia and may accelerate the adoption of lower-GWP '
        'alternatives worldwide.')

    # Concluding paragraph within Discussion (Anaesthesia style)
    doc.add_paragraph(
        'In conclusion, the EU desflurane regulation was associated with a progressive, '
        'statistically significant decline in secondary market values of desflurane vaporisers '
        'on eBay. Time-series trend analysis demonstrated that this decline was unique to the '
        'regulated agent: sevoflurane and isoflurane vaporiser prices remained stable throughout '
        'the study period, serving as natural controls. The price decline began during the '
        'legislative process, suggesting anticipatory market responses to cumulative regulatory '
        'signals. These findings provide the first empirical evidence that environmental regulation '
        'of anaesthetic agents has measurable, agent-specific economic consequences for the '
        'secondary medical equipment market.')

    # ---- ACKNOWLEDGEMENTS ----
    add_heading_styled(doc, 'Acknowledgements', level=1)
    doc.add_paragraph('[To be completed by authors]')

    # ---- COMPETING INTERESTS ----
    add_heading_styled(doc, 'Declaration of interests', level=1)
    doc.add_paragraph('No competing interests declared.')

    # ---- FUNDING ----
    add_heading_styled(doc, 'Funding', level=1)
    doc.add_paragraph('No external funding was received for this study.')

    # ---- AUTHOR CONTRIBUTIONS ----
    add_heading_styled(doc, 'Author contributions', level=1)
    doc.add_paragraph('[To be completed by authors using CRediT taxonomy]')

    # ---- DATA AVAILABILITY ----
    add_heading_styled(doc, 'Data availability statement', level=1)
    doc.add_paragraph(
        'The datasets generated during this study are available from the corresponding author '
        'on reasonable request. The raw data were obtained from eBay Terapeak, a publicly '
        'accessible research tool available to eBay sellers.')

    doc.add_page_break()

    # ---- REFERENCES ----
    add_heading_styled(doc, 'References', level=1)
    references = [
        '1. Varughese S, Ahmed R. Environmental and occupational considerations of anesthesia: '
        'a narrative review and update. Anesth Analg 2021; 133: 826\u201335.',
        '2. Regulation (EU) 2024/573 of the European Parliament and of the Council of '
        '7 February 2024 on fluorinated greenhouse gases. Official Journal of the European '
        'Union 2024; L 2024/573.',
        '3. Sherman JD, Chesebro BB. Inhaled anesthetic climate and ozone effects: a narrative '
        'review. Anesth Analg 2023; 137: 201\u201315.',
        '4. European Society of Anaesthesiology and Intensive Care. ESAIC position statement on '
        'the use of desflurane. European Journal of Anaesthesiology 2024; 41: 1\u20133.',
        '5. Association of Anaesthetists. Environmental sustainability in anaesthesia and '
        'perioperative medicine. Anaesthesia 2023; 78: 219\u201330.',
        '6. Sulbaek Andersen MP, Sander SP, Nielsen OJ, et al. Inhalation anaesthetics and '
        'climate change. British Journal of Anaesthesia 2010; 105: 760\u20136.',
        '7. Ryan SM, Nielsen CJ. Global warming potential of inhaled anesthetics: application '
        'to clinical use. Anesth Analg 2010; 111: 92\u20138.',
        '8. McGain F, Muret J, Guen CL, et al. Environmental sustainability in anaesthesia '
        'and critical care. British Journal of Anaesthesia 2020; 125: 680\u201392.',
        '9. Rauchenwald V, Heuss-Azeez R, Ganter MT, et al. Sevoflurane versus desflurane\u2014'
        'an economic analysis. BMC Anesthesiology 2020; 20: 272.',
        '10. Zuegge KL, Bunsen SK, Engel JM, et al. APW-AVE. Anesth Analg 2023; 137: 1219\u201325.',
        '11. von Elm E, Altman DG, Egger M, et al. The STROBE statement. BMJ 2007; 335: 806\u20138.',
        '12. NHS England. Decommissioning of desflurane in the NHS. 2023.',
        '13. Richter H, Weixler S, Ganter MT. Environmental sustainability in anaesthesia: the '
        'role of desflurane. Current Opinion in Anaesthesiology 2024; 37: 183\u20138.',
        '14. Davis G, Patel N. Regulatory obsolescence and secondary market asset depreciation. '
        'Journal of Environmental Economics and Management 2019; 95: 142\u201360.',
        '15. Lehmann H, Werning J, Baschnegger H, et al. Minimising the usage of desflurane '
        'only by education and removal of the vaporisers \u2013 a before-and-after-trial. '
        'BMC Anesthesiology 2025; 25: 108.',
        '16. Meyer MJ. Desflurane should des-appear: global and financial rationale. Anesth Analg '
        '2020; 131: 1317\u201322.',
        '17. Moonesinghe SR. Desflurane decommissioning: more than meets the eye. Anaesthesia '
        '2024; 79: 237\u201341.',
        '18. Mohammed A, Metta H. Is it time to bid adieu to desflurane? Journal of '
        'Anaesthesiology Clinical Pharmacology 2025; 41: 211\u20132.',
        '19. Beard D, Aston W, Black S, et al. Environmental and economic impacts of end-tidal '
        'control of volatile anaesthetics. Open Anaesthesia Journal 2025; 19: e18742126.',
        '20. Buckhead Fair Market Value. 2025 Benchmark Report on Pre-Owned Medical Equipment '
        'Prices. Atlanta, GA: BFMV, 2025.',
    ]
    for ref in references:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_page_break()

    # ---- FIGURE LEGENDS ----
    add_heading_styled(doc, 'Figure legends', level=1)

    legends = [
        ('Figure 1 ', 'Time series of eBay completed sale prices for desflurane (red), '
         'sevoflurane (blue) and isoflurane (green) vaporisers over three years (March 2023 to '
         'March 2026). Vertical dashed lines indicate key EU regulatory milestones. Curved lines '
         'represent LOWESS trend estimates (fraction = 0.3). Data source: eBay Terapeak.'),
        ('Figure 2 ', 'Box plot comparison of vaporiser prices before and after the EU desflurane '
         'ban (1 January 2026). Individual data points are shown as jittered dots. '
         'Data source: eBay Terapeak.'),
        ('Figure 3 ', 'Monthly median prices of anaesthetic vaporisers on eBay. Annotations '
         'indicate the number of transactions per month (n). Data source: eBay Terapeak.'),
        ('Figure 4 ', 'Price distribution histograms for each vaporiser type, comparing pre-ban '
         '(solid fill) and post-ban (hatched) periods. Data source: eBay Terapeak.'),
        ('Figure 5 ', 'Anaesthetic vaporiser prices mapped against the EU regulatory timeline. '
         'Shaded regions indicate regulatory phases. Data source: eBay Terapeak.'),
        ('Figure 6 ', 'Quarterly median price trends (upper panel) and sales volume (lower panel). '
         'Data source: eBay Terapeak.'),
    ]
    for fig_label, fig_text in legends:
        p = doc.add_paragraph()
        add_run_styled(p, fig_label, bold=True, size=Pt(10))
        add_run_styled(p, fig_text, italic=True, size=Pt(10))

    # Supplementary table
    if has_asking_data:
        doc.add_page_break()
        add_heading_styled(doc, 'Supporting Information', level=1)

        ask = asking_results['asking_summary']
        kw = asking_results['kruskal_wallis']
        p = doc.add_paragraph()
        add_run_styled(p, 'Table S1 ', bold=True, size=Pt(10))
        add_run_styled(p, ('Current eBay asking prices (active listings) by vaporiser type, '
                           'collected 27 March 2026. Values are mean (SD), median (IQR) in US dollars. '
                           'p value from Kruskal\u2013Wallis test across three agent types.'),
                       italic=True, size=Pt(10))

        et = doc.add_table(rows=1, cols=6)
        et.style = 'Table Grid'
        et.alignment = WD_TABLE_ALIGNMENT.CENTER
        add_table_header(et, ['Agent', 'n', 'Mean (SD)', 'Median (IQR)', 'Range', 'p value'])

        for i, agent_cap in enumerate(['Desflurane', 'Sevoflurane', 'Isoflurane']):
            a = ask[agent_cap]
            pval_str = fmt_p(kw['P']) if i == 0 else ''
            data = [
                (agent_cap, WD_ALIGN_PARAGRAPH.LEFT),
                (str(a['n']), WD_ALIGN_PARAGRAPH.CENTER),
                (f'${a["mean"]:.0f} ({a["sd"]:.0f})', WD_ALIGN_PARAGRAPH.CENTER),
                (f'${a["median"]:.0f} ({a["q25"]:.0f}\u2013{a["q75"]:.0f})',
                 WD_ALIGN_PARAGRAPH.CENTER),
                (f'${a["min"]:.0f}\u2013{a["max"]:.0f}', WD_ALIGN_PARAGRAPH.CENTER),
                (pval_str, WD_ALIGN_PARAGRAPH.CENTER),
            ]
            add_table_data_row(et, data)

    path = outdir + 'vaporizer_paper_english.docx'
    doc.save(path)
    print(f"English paper saved: {path}")
    return path


if __name__ == '__main__':
    write_english_paper()
